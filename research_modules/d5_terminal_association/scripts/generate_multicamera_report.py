#!/usr/bin/env python3
"""生成 D5 多相机视觉关联报告、中文图表与稳定仿真图片副本。"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
import warnings
import zipfile
from collections import defaultdict
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
warnings.filterwarnings("ignore", message="Unable to import Axes3D.*")

import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from matplotlib.patches import Ellipse, FancyArrowPatch, FancyBboxPatch, Polygon
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MODULE_DIR = Path(__file__).resolve().parents[1]
REPO_ROOT = MODULE_DIR.parents[1]
DOCS_DIR = MODULE_DIR / "docs"
ASSET_DIR = DOCS_DIR / "assets" / "d5_multicamera_association"
DOCX_PATH = DOCS_DIR / "D5_MULTICAMERA_ASSOCIATION_REPORT_CN.docx"

FORMAL_ROOT = (
    REPO_ROOT
    / "research_modules"
    / "airsim_runtime"
    / "outputs"
    / "d5_cv_5v5_multicamera_formal_20260716"
)
DETECT_DIR = FORMAL_ROOT / "d5_cv_5v5_multicamera_formal_20260716_detect"
YOLO_DIR = FORMAL_ROOT / "d5_cv_5v5_multicamera_formal_20260716_yolo_bytetrack"

FORMAL_ASSETS = {
    "detect_t6_montage.png": (
        DETECT_DIR
        / "annotated_snapshots"
        / "frame_0024_t06.00s"
        / "montage_5primary_plus_recon.png"
    ),
    "yolo_t6_montage.png": (
        YOLO_DIR
        / "annotated_snapshots"
        / "frame_0024_t06.00s"
        / "montage_5primary_plus_recon.png"
    ),
    "detect_frame_metrics.csv": DETECT_DIR / "d5_multicamera_frame_metrics.csv",
    "yolo_frame_metrics.csv": YOLO_DIR / "d5_multicamera_frame_metrics.csv",
    "detect_candidates.csv": DETECT_DIR / "d5_multicamera_candidates.csv",
    "yolo_candidates.csv": YOLO_DIR / "d5_multicamera_candidates.csv",
}

REPORT_IMAGES = (
    "system_architecture.png",
    "time_alignment.png",
    "coordinate_projection.png",
    "uncertainty_gating.png",
    "principle_pipeline.png",
    "six_camera_fusion.png",
    "handoff_degradation.png",
    "backend_comparison.png",
    "detect_t6_montage.png",
    "yolo_t6_montage.png",
    "detect_coverage.png",
    "detect_timeline.png",
    "yolo_coverage.png",
    "yolo_timeline.png",
    "simulation_timeline.png",
)

SECTION_TITLES = (
    "1. 体系架构",
    "2. 关键技术",
    "3. 关联方案",
    "4. 实验结果",
    "5. 边界与计划",
)

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
TEAL = "2A7F86"
LIGHT_TEAL = "DDEFEF"
ORANGE = "D97924"
LIGHT_ORANGE = "FCE9D8"
GREEN = "548235"
LIGHT_GREEN = "E2F0D9"
RED = "C00000"
LIGHT_RED = "FCE4D6"
MID_GRAY = "66717E"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "1F1F1F"

FONT_REGULAR = FontProperties(
    fname="/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
)
FONT_BOLD = FontProperties(
    fname="/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
)
CHINESE_FONT = "Noto Sans CJK SC"

FORBIDDEN_REPORT_TERMS = (
    "领导汇报",
    "领导送阅",
    "仓库",
    "文件路径",
    "需要指出的是",
    "值得注意的是",
    "从本质上看",
    "更准确地说",
    "换句话说",
    "可以看到",
    "综上所述",
    "这充分说明了",
    "意义重大",
    "提供支撑",
    "全面提升",
)

PATH_MARKERS = (
    "/home/",
    "research_modules/",
    "docs/",
    "src/",
    ".py",
)


def sync_formal_assets() -> None:
    """显式同步正式仿真截图和绘图数据到模块稳定目录。"""
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    missing = [str(source) for source in FORMAL_ASSETS.values() if not source.is_file()]
    if missing:
        raise FileNotFoundError("缺少正式仿真产物：\n" + "\n".join(missing))
    for name, source in FORMAL_ASSETS.items():
        shutil.copy2(source, ASSET_DIR / name)


def _new_canvas(title: str, subtitle: str, *, height: float = 7.2):
    fig, ax = plt.subplots(figsize=(16, height), dpi=160)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 60)
    ax.axis("off")
    fig.patch.set_facecolor("white")
    ax.text(
        2,
        56,
        title,
        color="#17365d",
        fontsize=22,
        fontproperties=FONT_BOLD,
        va="center",
    )
    ax.text(
        2,
        52,
        subtitle,
        color="#66717e",
        fontsize=12,
        fontproperties=FONT_REGULAR,
        va="center",
    )
    return fig, ax


def _box(
    ax,
    x: float,
    y: float,
    width: float,
    height: float,
    text: str,
    *,
    face: str = "#eef5fb",
    edge: str = "#2f75b5",
    size: float = 12,
    bold: bool = False,
    radius: float = 0.7,
) -> None:
    patch = FancyBboxPatch(
        (x, y),
        width,
        height,
        boxstyle=f"round,pad=0.3,rounding_size={radius}",
        linewidth=1.7,
        facecolor=face,
        edgecolor=edge,
    )
    ax.add_patch(patch)
    ax.text(
        x + width / 2,
        y + height / 2,
        text,
        ha="center",
        va="center",
        color="#1f2933",
        fontsize=size,
        fontproperties=FONT_BOLD if bold else FONT_REGULAR,
        linespacing=1.3,
    )


def _arrow(
    ax,
    start: tuple[float, float],
    end: tuple[float, float],
    *,
    color: str = "#46637f",
    width: float = 1.7,
    curve: float = 0,
) -> None:
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=width,
            color=color,
            connectionstyle=f"arc3,rad={curve}",
        )
    )


def _save_figure(fig, name: str) -> None:
    fig.tight_layout(pad=0.7)
    fig.savefig(ASSET_DIR / name, bbox_inches="tight")
    plt.close(fig)


def generate_system_architecture() -> None:
    fig, ax = _new_canvas(
        "D5 在体系信息流中的位置",
        "中心航迹和版本化分配进入相机几何链路；一致性证据分别供降级判断和视觉导引切换使用",
        height=7.6,
    )
    _box(
        ax,
        3,
        35,
        17,
        10,
        "D1 / D2\n中心全局航迹\n位置、速度、协方差",
        face="#e7f0fa",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        3,
        19,
        17,
        10,
        "D3\n版本化目标分配\n资源、目标、有效期",
        face="#e7f0fa",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        29,
        27,
        24,
        13,
        "D5 多相机视觉关联\n按图像时刻预测并投影\n汇总局部视觉一致性证据",
        face="#dcebf7",
        edge="#17365d",
        size=13,
        bold=True,
    )
    _box(
        ax,
        62,
        35,
        17,
        10,
        "D4\n主动降级判断\n中心 / 二级 / 分布式",
        face="#fce9d8",
        edge="#d97924",
        bold=True,
    )
    _box(
        ax,
        62,
        19,
        17,
        10,
        "D7\n视觉导引切换\n独立检查安全门控",
        face="#e2f0d9",
        edge="#548235",
        bold=True,
    )
    _box(
        ax,
        84,
        27,
        13,
        9,
        "执行侧\n继续中心模式\n降级或允许切换",
        face="#f2f4f7",
        edge="#66717e",
        bold=True,
    )

    _arrow(ax, (20, 40), (29, 35))
    _arrow(ax, (20, 24), (29, 31))
    _arrow(ax, (53, 35), (62, 40))
    _arrow(ax, (53, 31), (62, 24))
    _arrow(ax, (79, 40), (84, 35))
    _arrow(ax, (79, 24), (84, 31))
    ax.plot([20, 27, 56], [43, 49, 49], color="#8c6d31", linewidth=1.2)
    _arrow(ax, (56, 49), (62, 43), color="#8c6d31", width=1.2)
    ax.plot([20, 25, 57], [22, 13, 13], color="#8c6d31", linewidth=1.2)
    _arrow(ax, (57, 13), (62, 37), color="#8c6d31", width=1.2)
    ax.text(
        41,
        50,
        "上游不确定度",
        ha="center",
        color="#8c6d31",
        fontsize=10.5,
        fontproperties=FONT_BOLD,
    )
    ax.text(
        41,
        14,
        "计划有效性",
        ha="center",
        color="#8c6d31",
        fontsize=10.5,
        fontproperties=FONT_BOLD,
    )

    ax.text(
        50,
        9,
        "D4 同时参考上游不确定度、计划有效性和 D5 一致性；D7 在切换前检查计划版本、时效、相机能力、机动能力和安全门。",
        ha="center",
        color="#4f5b66",
        fontsize=11.3,
        fontproperties=FONT_REGULAR,
    )
    ax.text(
        50,
        3,
        "D5 只给出视觉关联证据，不重新分配目标，不改变中心持有的目标编号，不输出控制量。",
        ha="center",
        color="#9c2525",
        fontsize=11.8,
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "system_architecture.png")


def generate_time_alignment() -> None:
    fig, ax = _new_canvas(
        "时间对齐与证据新鲜度",
        "几何预测使用量测时刻；证据授权使用到达时刻与量测时刻的差值",
        height=7.0,
    )
    y = 30
    ax.plot([8, 92], [y, y], color="#46637f", linewidth=2)
    points = ((18, "航迹状态时刻\n$t_0$"), (50, "图像量测时刻\n$t_m$"), (82, "证据到达时刻\n$t_a$"))
    for x, label in points:
        ax.plot([x, x], [y - 2, y + 2], color="#17365d", linewidth=2.2)
        ax.text(
            x,
            y + 5,
            label,
            ha="center",
            va="bottom",
            fontsize=12,
            fontproperties=FONT_BOLD,
            color="#17365d",
        )

    _box(
        ax,
        24,
        37,
        52,
        8,
        r"预测到图像时刻：  $\mathbf{p}(t_m)=\mathbf{p}(t_0)+\mathbf{v}(t_0)(t_m-t_0)$",
        face="#eef5fb",
        edge="#2f75b5",
        size=13,
        bold=True,
    )
    _box(
        ax,
        56,
        16,
        35,
        8,
        r"证据年龄：  $\Delta t_{\rm age}=t_a-t_m$",
        face="#e0f0ef",
        edge="#2a7f86",
        size=13,
        bold=True,
    )
    _arrow(ax, (18, 32), (50, 32), color="#2f75b5")
    _arrow(ax, (50, 28), (82, 28), color="#2a7f86")
    ax.text(
        34,
        25,
        "只补偿目标在采样前后的运动",
        ha="center",
        fontsize=10.8,
        color="#46637f",
        fontproperties=FONT_REGULAR,
    )
    ax.text(
        66,
        25,
        "传输、排队和推理延迟",
        ha="center",
        fontsize=10.8,
        color="#2a7f86",
        fontproperties=FONT_REGULAR,
    )
    ax.text(
        50,
        8,
        "延迟增大时，预测协方差随时间增长。几何门可相应放宽，陈旧证据的授权权重必须降低；超过时效上限时进入暂停执行或受限重获取。",
        ha="center",
        fontsize=11.3,
        color="#9c2525",
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "time_alignment.png")


def generate_coordinate_projection() -> None:
    fig, ax = _new_canvas(
        "全局坐标到像素平面",
        "外参决定相机位置和镜头指向，内参决定焦距、主点和像素尺度",
        height=7.3,
    )
    ax.plot([8, 19], [29, 29], color="#2f75b5", linewidth=2)
    ax.plot([8, 8], [29, 41], color="#2a7f86", linewidth=2)
    ax.plot([8, 14], [29, 22], color="#d97924", linewidth=2)
    ax.text(20, 29, "北", color="#2f75b5", fontproperties=FONT_BOLD, fontsize=11)
    ax.text(7, 43, "东", color="#2a7f86", fontproperties=FONT_BOLD, fontsize=11)
    ax.text(15, 20, "地", color="#d97924", fontproperties=FONT_BOLD, fontsize=11)
    ax.scatter([18], [39], s=90, color="#c00000")
    ax.text(18, 42, "预测目标点", ha="center", fontproperties=FONT_BOLD, fontsize=11)

    _box(
        ax,
        26,
        30,
        19,
        11,
        r"相机外参" "\n" r"$\mathbf{P}_c=\mathbf{R}\mathbf{P}_{ned}+\mathbf{t}$",
        face="#eef5fb",
        edge="#2f75b5",
        size=13,
        bold=True,
    )
    _arrow(ax, (20, 36), (26, 36))

    camera = Polygon(
        [[51, 29], [58, 34], [58, 24]],
        closed=True,
        facecolor="#fce9d8",
        edgecolor="#d97924",
        linewidth=1.8,
    )
    ax.add_patch(camera)
    ax.text(54, 20, "相机光心与镜头指向", ha="center", fontproperties=FONT_BOLD, fontsize=10.5)
    ax.plot([58, 81], [34, 43], color="#d97924", linewidth=1.5)
    ax.plot([58, 81], [24, 15], color="#d97924", linewidth=1.5)
    ax.fill_between([58, 81], [24, 15], [34, 43], color="#fce9d8", alpha=0.35)
    ax.text(69, 31, "视场边界", ha="center", color="#9c5a18", fontproperties=FONT_BOLD, fontsize=11)

    ax.plot([82, 82], [12, 46], color="#17365d", linewidth=2.5)
    ax.scatter([82], [37], s=90, facecolors="white", edgecolors="#c00000", linewidths=2)
    ax.text(86, 38, "预测像素", fontproperties=FONT_BOLD, fontsize=11)
    _box(
        ax,
        59,
        4,
        37,
        8,
        r"针孔投影：  $u=f_x X_c/Z_c+c_x$，$v=f_y Y_c/Z_c+c_y$",
        face="#e0f0ef",
        edge="#2a7f86",
        size=13,
        bold=True,
    )
    ax.text(
        29,
        8,
        "目标在镜头后方、投影越出图像边界或深度无效时，本帧不进入正常匹配。",
        ha="center",
        color="#9c2525",
        fontsize=11.2,
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "coordinate_projection.png")


def generate_uncertainty_gating() -> None:
    fig, ax = _new_canvas(
        "不确定性传播与像素门控",
        "三维航迹误差、检测框中心误差、标定误差和局部跟踪误差共同决定像素门形状",
        height=7.2,
    )
    _box(
        ax,
        3,
        31,
        18,
        10,
        "三维预测状态\n位置协方差 P",
        face="#e7f0fa",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        26,
        31,
        18,
        10,
        "投影雅可比 J\n局部线性传播",
        face="#eef5fb",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        49,
        31,
        22,
        10,
        "像素协方差\n" r"$\Sigma_{px}=\mathbf{J}\mathbf{P}\mathbf{J}^{T}+\mathbf{R}$",
        face="#e0f0ef",
        edge="#2a7f86",
        bold=True,
        size=13,
    )
    _box(
        ax,
        76,
        31,
        21,
        10,
        "马氏距离门\n按总不确定度归一化",
        face="#fce9d8",
        edge="#d97924",
        bold=True,
    )
    _arrow(ax, (21, 36), (26, 36))
    _arrow(ax, (44, 36), (49, 36))
    _arrow(ax, (71, 36), (76, 36))

    ax.scatter([24], [17], s=90, color="#c00000", label="预测像素")
    ax.scatter([32], [20], s=90, marker="x", color="#17365d", linewidths=2, label="检测中心")
    ax.add_patch(
        Ellipse(
            (24, 17),
            width=28,
            height=13,
            angle=18,
            fill=False,
            linewidth=2,
            edgecolor="#2f75b5",
        )
    )
    ax.text(14, 8, "像素协方差椭圆", color="#2f75b5", fontproperties=FONT_BOLD, fontsize=11)
    _box(
        ax,
        45,
        12,
        49,
        10,
        r"$d^2=(\mathbf{z}-\hat{\mathbf{z}})^{T}\Sigma_{px}^{-1}"
        r"(\mathbf{z}-\hat{\mathbf{z}})$",
        face="#f2f4f7",
        edge="#66717e",
        size=13,
        bold=True,
    )
    ax.text(
        70,
        6,
        "R 汇总检测框中心、相机标定和局部跟踪误差。门外候选被拒绝；门内候选进入关联代价。",
        ha="center",
        color="#4f5b66",
        fontsize=11.2,
        fontproperties=FONT_REGULAR,
    )
    _save_figure(fig, "uncertainty_gating.png")


def generate_association_pipeline() -> None:
    fig, ax = _new_canvas(
        "单相机关联与时序稳定",
        "候选代价先在单个相机内部完成一对一匹配，再由 3 帧窗口过滤单帧跳变",
        height=7.8,
    )
    labels = (
        ("投影误差", 3, "#eef5fb", "#2f75b5"),
        ("角速度一致性", 19, "#eef5fb", "#2f75b5"),
        ("类别一致性", 35, "#eef5fb", "#2f75b5"),
        ("局部航迹质量", 51, "#e0f0ef", "#2a7f86"),
        ("身份冲突", 67, "#fce9d8", "#d97924"),
        ("侦察提示", 83, "#fce9d8", "#d97924"),
    )
    for text, x, face, edge in labels:
        _box(ax, x, 41, 14, 7, text, face=face, edge=edge, bold=True, size=11.5)
        _arrow(ax, (x + 7, 41), (31, 33), curve=0.06, width=1.2)

    _box(
        ax,
        20,
        25,
        22,
        8,
        "候选代价矩阵\n全局目标 × 本地轨迹",
        face="#e7f0fa",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        48,
        25,
        19,
        8,
        "匈牙利算法\n相机内一对一匹配",
        face="#fce9d8",
        edge="#d97924",
        bold=True,
    )
    _box(
        ax,
        73,
        25,
        23,
        8,
        "3 帧稳定窗口\n至少 2 帧有效匹配",
        face="#e2f0d9",
        edge="#548235",
        bold=True,
    )
    _arrow(ax, (42, 29), (48, 29))
    _arrow(ax, (67, 29), (73, 29))

    states = (
        ("稳定关联", 3, LIGHT_GREEN, GREEN),
        ("候选模糊", 27, LIGHT_ORANGE, ORANGE),
        ("暂停执行", 51, LIGHT_RED, RED),
        ("受限重获取", 75, LIGHT_BLUE, BLUE),
    )
    for text, x, face, edge in states:
        _box(ax, x, 9, 20, 7, text, face=f"#{face.lower()}", edge=f"#{edge.lower()}", bold=True)
        _arrow(ax, (84, 25), (x + 10, 16), curve=0.08, width=1.2)
    ax.text(
        50,
        3,
        "漏检、遮挡、检测框抖动或投影偏差都可能产生单帧候选。单帧结果不授权末端交接。",
        ha="center",
        color="#9c2525",
        fontsize=11.5,
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "principle_pipeline.png")


def generate_multicamera_fusion() -> None:
    fig, ax = _new_canvas(
        "六路相机证据汇总",
        "五路本地相机独立配准；一路宽视场侦察相机提供广域提示；证据按中心目标编号汇总",
        height=8.0,
    )
    _box(
        ax,
        4,
        43,
        92,
        8,
        "中心持有的全局航迹：目标 1 · 目标 2 · 目标 3 · 目标 4 · 目标 5",
        face="#e7f0fa",
        edge="#2f75b5",
        bold=True,
        size=13,
    )
    camera_labels = ("本地相机 1", "本地相机 2", "本地相机 3", "本地相机 4", "本地相机 5", "侦察相机")
    starts = (4, 20, 36, 52, 68, 84)
    for index, (label, x) in enumerate(zip(camera_labels, starts)):
        is_recon = index == 5
        _box(
            ax,
            x,
            29,
            13,
            8,
            label + ("\n广域提示" if is_recon else "\n局部目标子集"),
            face="#e0f0ef" if is_recon else "#eef5fb",
            edge="#2a7f86" if is_recon else "#2f75b5",
            bold=True,
            size=11.5,
        )
        _arrow(ax, (x + 6.5, 43), (x + 6.5, 37), width=1.2)

    _box(
        ax,
        5,
        15,
        26,
        8,
        "各相机独立几何配准\n局部编号按资源 / 相机隔离",
        face="#fce9d8",
        edge="#d97924",
        bold=True,
    )
    _box(
        ax,
        37,
        15,
        25,
        8,
        "稳定窗口\n保留当前实测证据",
        face="#e2f0d9",
        edge="#548235",
        bold=True,
    )
    _box(
        ax,
        68,
        15,
        27,
        8,
        "按同一中心目标编号汇总\n形成单视角或多视角支持",
        face="#e0f0ef",
        edge="#2a7f86",
        bold=True,
    )
    for x in starts:
        _arrow(ax, (x + 6.5, 29), (18, 23), curve=0.05, width=1.05)
    _arrow(ax, (31, 19), (37, 19))
    _arrow(ax, (62, 19), (68, 19))
    ax.text(
        50,
        8,
        "单个本地相机无需看到全部目标。侦察提示不能替代当前局部实测，也不能绕过版本、身份和稳定门。",
        ha="center",
        color="#4f5b66",
        fontsize=11.3,
        fontproperties=FONT_REGULAR,
    )
    ax.text(
        50,
        3,
        "当前未实现多视线三维联合优化和跨相机外观重识别。",
        ha="center",
        color="#9c2525",
        fontsize=11.5,
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "six_camera_fusion.png")


def generate_handoff_degradation() -> None:
    fig, ax = _new_canvas(
        "视觉交接与主动降级",
        "D5 证据分别进入 D7 切换门和 D4 降级判断；两个决策保持独立",
        height=7.5,
    )
    _box(
        ax,
        3,
        32,
        22,
        11,
        "D5 视觉一致性证据\n稳定关联 / 候选模糊\n暂停执行 / 受限重获取",
        face="#dcebf7",
        edge="#17365d",
        bold=True,
    )
    _box(
        ax,
        34,
        35,
        27,
        10,
        "D7 视觉导引切换门\n计划版本 · 证据时效 · 相机能力\n机动能力 · 安全门控",
        face="#e2f0d9",
        edge="#548235",
        bold=True,
    )
    _box(
        ax,
        70,
        35,
        26,
        10,
        "切换判定\n保持当前导引\n或允许视觉导引",
        face="#f2f4f7",
        edge="#66717e",
        bold=True,
    )
    _arrow(ax, (25, 39), (34, 40))
    _arrow(ax, (61, 40), (70, 40))

    _box(
        ax,
        3,
        13,
        22,
        9,
        "D1 / D2 不确定度\nD3 计划有效性",
        face="#eef5fb",
        edge="#2f75b5",
        bold=True,
    )
    _box(
        ax,
        34,
        12,
        27,
        11,
        "D4 主动降级判断\n中心链路、二级节点和\n分布式模式的可用性",
        face="#fce9d8",
        edge="#d97924",
        bold=True,
    )
    _box(
        ax,
        70,
        13,
        26,
        9,
        "降级判定\n继续中心模式\n或进入受控降级",
        face="#f2f4f7",
        edge="#66717e",
        bold=True,
    )
    _arrow(ax, (25, 18), (34, 18))
    _arrow(ax, (25, 35), (34, 22), curve=0.1)
    _arrow(ax, (61, 18), (70, 18))
    ax.text(
        50,
        5,
        "已验证友方证据触发暂停执行。身份未知保持未知，不能推定为敌方。",
        ha="center",
        color="#9c2525",
        fontsize=11.8,
        fontproperties=FONT_BOLD,
    )
    _save_figure(fig, "handoff_degradation.png")


def _camera_label(raw: str) -> str:
    match = re.search(r"Primary_(\d+)", raw)
    if match:
        return f"本地相机{match.group(1)}"
    return "侦察相机"


def _read_frame_metrics(path: Path):
    by_camera: dict[str, list[tuple[float, int]]] = defaultdict(list)
    totals: dict[float, int] = defaultdict(int)
    with path.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            timestamp = float(row["timestamp"])
            count = int(row["online_detection_count"])
            label = _camera_label(row["camera_id"])
            by_camera[label].append((timestamp, count))
            totals[timestamp] += count
    return by_camera, dict(sorted(totals.items()))


def _read_candidate_timeline(path: Path):
    selected: dict[float, int] = defaultdict(int)
    stable: dict[float, int] = defaultdict(int)
    with path.open(newline="", encoding="utf-8") as handle:
        for row in csv.DictReader(handle):
            timestamp = float(row["timestamp"])
            if row["selected"] == "True":
                selected[timestamp] += 1
                if row["stable_cross_view_support"] == "True":
                    stable[timestamp] += 1
    timestamps = sorted(set(selected) | set(stable))
    return (
        timestamps,
        [selected[timestamp] for timestamp in timestamps],
        [stable[timestamp] for timestamp in timestamps],
    )


def generate_backend_status() -> None:
    fig, ax = plt.subplots(figsize=(16, 8.2), dpi=160)
    ax.axis("off")
    fig.patch.set_facecolor("white")
    ax.text(
        0.01,
        0.96,
        "单种子仿真指标与分支阈值",
        transform=ax.transAxes,
        fontsize=22,
        color="#17365d",
        fontproperties=FONT_BOLD,
        va="top",
    )
    ax.text(
        0.01,
        0.90,
        "绿色表示达到本分支建议阈值，橙色表示未达到；结果只适用于 seed 7 的固定仿真场景",
        transform=ax.transAxes,
        fontsize=12,
        color="#66717e",
        fontproperties=FONT_REGULAR,
        va="top",
    )
    rows = (
        ("检测召回率", "≥0.95", "1.000", "达到", "≥0.90", "0.622", "未达到"),
        ("严格关联准确率", "≥0.95", "1.000", "达到", "≥0.95", "0.966", "达到"),
        ("稳定注册率", "≥0.90", "0.975", "达到", "≥0.90", "0.955", "达到"),
        ("局部联合覆盖率", "≥0.95", "1.000", "达到", "≥0.95", "1.000", "达到"),
        ("侦察全视野率", "≥0.90", "0.918", "达到", "≥0.90", "0.878", "未达到"),
        ("局部身份切换", "≤0", "0", "达到", "≤5", "25", "未达到"),
        ("在线真值身份使用", "=0", "0", "达到", "=0", "0", "达到"),
        ("中心目标编号改写", "=0", "0", "达到", "=0", "0", "达到"),
    )
    columns = (
        "指标",
        "detect 阈值",
        "detect 实测",
        "detect 判定",
        "YOLO 路线阈值",
        "YOLO 路线实测",
        "YOLO 路线判定",
    )
    table = ax.table(
        cellText=rows,
        colLabels=columns,
        cellLoc="center",
        colLoc="center",
        bbox=[0.01, 0.04, 0.98, 0.80],
        colWidths=[0.20, 0.12, 0.12, 0.12, 0.14, 0.14, 0.14],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(11.5)
    for (row, col), cell in table.get_celld().items():
        cell.get_text().set_fontproperties(FONT_BOLD if row == 0 else FONT_REGULAR)
        cell.set_edgecolor("#9aa6b2")
        if row == 0:
            cell.set_facecolor("#17365d")
            cell.get_text().set_color("white")
        elif col in (3, 6):
            value = rows[row - 1][col]
            if value == "达到":
                cell.set_facecolor("#e2f0d9")
                cell.get_text().set_color("#3f6b28")
            else:
                cell.set_facecolor("#fce9d8")
                cell.get_text().set_color("#a65313")
        elif row % 2 == 0:
            cell.set_facecolor("#f6f8fa")
    _save_figure(fig, "backend_comparison.png")


def generate_coverage_plot(filename: str, data_filename: str, title: str) -> None:
    path = ASSET_DIR / data_filename
    if not path.is_file():
        raise FileNotFoundError(path)
    by_camera, _ = _read_frame_metrics(path)
    fig, ax = plt.subplots(figsize=(14, 7), dpi=160)
    for label in sorted(by_camera):
        pairs = sorted(by_camera[label])
        ax.plot(
            [timestamp for timestamp, _ in pairs],
            [count for _, count in pairs],
            marker="o",
            markersize=2.8,
            linewidth=1.5,
            label=label,
        )
    ax.set_title(title, fontproperties=FONT_BOLD, fontsize=17, color="#17365d")
    ax.set_xlabel("仿真时间（秒）", fontproperties=FONT_REGULAR, fontsize=11)
    ax.set_ylabel("在线局部轨迹数", fontproperties=FONT_REGULAR, fontsize=11)
    ax.set_ylim(-0.2, 6.4)
    ax.grid(alpha=0.25)
    legend = ax.legend(prop=FONT_REGULAR, ncol=3, loc="lower center")
    for text in legend.get_texts():
        text.set_fontproperties(FONT_REGULAR)
    fig.tight_layout()
    fig.savefig(ASSET_DIR / filename, bbox_inches="tight")
    plt.close(fig)


def generate_timeline_plot(filename: str, data_filename: str, title: str) -> None:
    path = ASSET_DIR / data_filename
    if not path.is_file():
        raise FileNotFoundError(path)
    timestamps, selected, stable = _read_candidate_timeline(path)
    fig, ax = plt.subplots(figsize=(14, 6.3), dpi=160)
    ax.plot(timestamps, selected, linewidth=2, color="#2f75b5", label="即时选择")
    ax.plot(timestamps, stable, linewidth=2, color="#d97924", label="稳定注册")
    ax.set_title(title, fontproperties=FONT_BOLD, fontsize=17, color="#17365d")
    ax.set_xlabel("仿真时间（秒）", fontproperties=FONT_REGULAR, fontsize=11)
    ax.set_ylabel("每帧注册数", fontproperties=FONT_REGULAR, fontsize=11)
    ax.grid(alpha=0.25)
    legend = ax.legend(prop=FONT_REGULAR)
    for text in legend.get_texts():
        text.set_fontproperties(FONT_REGULAR)
    fig.tight_layout()
    fig.savefig(ASSET_DIR / filename, bbox_inches="tight")
    plt.close(fig)


def generate_simulation_timeline() -> None:
    detect_frame = ASSET_DIR / "detect_frame_metrics.csv"
    yolo_frame = ASSET_DIR / "yolo_frame_metrics.csv"
    detect_candidates = ASSET_DIR / "detect_candidates.csv"
    yolo_candidates = ASSET_DIR / "yolo_candidates.csv"
    for path in (detect_frame, yolo_frame, detect_candidates, yolo_candidates):
        if not path.is_file():
            raise FileNotFoundError(path)

    _, detect_totals = _read_frame_metrics(detect_frame)
    _, yolo_totals = _read_frame_metrics(yolo_frame)
    d_time, _, d_stable = _read_candidate_timeline(detect_candidates)
    y_time, _, y_stable = _read_candidate_timeline(yolo_candidates)

    fig, axes = plt.subplots(2, 1, figsize=(14, 9), dpi=160, sharex=True)
    axes[0].plot(
        list(detect_totals),
        list(detect_totals.values()),
        color="#2f75b5",
        linewidth=2,
        label="AirSim detect",
    )
    axes[0].plot(
        list(yolo_totals),
        list(yolo_totals.values()),
        color="#d97924",
        linewidth=2,
        label="YOLOv8 + ByteTrack",
    )
    axes[0].set_ylabel("六路在线检测总数", fontproperties=FONT_REGULAR, fontsize=11)
    axes[0].set_title(
        "单种子仿真检测与稳定注册时间序列",
        fontproperties=FONT_BOLD,
        fontsize=17,
        color="#17365d",
    )
    axes[0].grid(alpha=0.25)

    axes[1].plot(d_time, d_stable, color="#2f75b5", linewidth=2, label="detect 稳定注册")
    axes[1].plot(y_time, y_stable, color="#d97924", linewidth=2, label="YOLO 路线稳定注册")
    axes[1].set_xlabel("仿真时间（秒）", fontproperties=FONT_REGULAR, fontsize=11)
    axes[1].set_ylabel("每帧稳定注册数", fontproperties=FONT_REGULAR, fontsize=11)
    axes[1].grid(alpha=0.25)
    for ax in axes:
        legend = ax.legend(prop=FONT_REGULAR)
        for text in legend.get_texts():
            text.set_fontproperties(FONT_REGULAR)
    fig.tight_layout()
    fig.savefig(ASSET_DIR / "simulation_timeline.png", bbox_inches="tight")
    plt.close(fig)


def generate_all_figures() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    generate_system_architecture()
    generate_time_alignment()
    generate_coordinate_projection()
    generate_uncertainty_gating()
    generate_association_pipeline()
    generate_multicamera_fusion()
    generate_handoff_degradation()
    generate_backend_status()
    generate_coverage_plot(
        "detect_coverage.png",
        "detect_frame_metrics.csv",
        "AirSim detect：六路相机在线局部轨迹",
    )
    generate_timeline_plot(
        "detect_timeline.png",
        "detect_candidates.csv",
        "AirSim detect：即时选择与稳定注册",
    )
    generate_coverage_plot(
        "yolo_coverage.png",
        "yolo_frame_metrics.csv",
        "YOLOv8 + ByteTrack：六路相机在线局部轨迹",
    )
    generate_timeline_plot(
        "yolo_timeline.png",
        "yolo_candidates.csv",
        "YOLOv8 + ByteTrack：即时选择与稳定注册",
    )
    generate_simulation_timeline()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell,
    *,
    top: int = 90,
    start: int = 100,
    bottom: int = 90,
    end: int = 100,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _set_run_font(
    run,
    size: float = 10.5,
    *,
    bold: bool = False,
    color: str = BLACK,
) -> None:
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)


def _set_body_format(paragraph, *, first_line: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(5)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def _add_body(document: Document, text: str, *, keep_with_next: bool = False):
    paragraph = document.add_paragraph()
    _set_body_format(paragraph)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    _set_run_font(run)
    return paragraph


def _add_bullet(document: Document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.45)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.3
    run = paragraph.add_run(text)
    _set_run_font(run)
    return paragraph


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_field(paragraph, instruction: str, result: str = "—") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = result
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))
    _set_run_font(run, size=9)


def _add_heading(
    document: Document,
    text: str,
    level: int,
    *,
    bookmark: str | None = None,
):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    _set_run_font(
        run,
        size=17 if level == 1 else 12.5,
        bold=True,
        color=NAVY if level == 1 else BLUE,
    )
    if bookmark:
        _add_bookmark(paragraph, bookmark, 100 + int(bookmark.removeprefix("sec")))
    return paragraph


def _add_table(
    document: Document,
    headers: tuple[str, ...],
    rows: tuple[tuple[str, ...], ...] | list[tuple[str, ...]],
    *,
    widths_cm: tuple[float, ...],
    status_column: int | None = None,
):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    _set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        cell.width = Cm(widths_cm[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_shading(cell, NAVY)
        _set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        _set_run_font(run, size=9.3, bold=True, color=WHITE)

    for values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        for index, text in enumerate(values):
            cell = row.cells[index]
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            color = BLACK
            bold = False
            if status_column is not None and index == status_column:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bold = True
                if text in ("已实现", "达到", "通过"):
                    _set_cell_shading(cell, LIGHT_GREEN)
                    color = GREEN
                elif text in ("待验证", "未达到", "未实现"):
                    _set_cell_shading(cell, LIGHT_RED)
                    color = RED
                else:
                    _set_cell_shading(cell, LIGHT_ORANGE)
                    color = ORANGE
            run = paragraph.add_run(text)
            _set_run_font(run, size=9.1, bold=bold, color=color)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def _add_callout(
    document: Document,
    title: str,
    text: str,
    *,
    fill: str = PALE_BLUE,
    accent: str = BLUE,
) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(16.2)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title + "  ")
    _set_run_font(run, size=10.3, bold=True, color=accent)
    run = paragraph.add_run(text)
    _set_run_font(run, size=10.3)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def _add_picture(
    document: Document,
    name: str,
    caption: str,
    *,
    width_cm: float = 16.0,
) -> None:
    path = ASSET_DIR / name
    if not path.is_file():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(7)
    caption_paragraph.paragraph_format.keep_with_next = False
    run = caption_paragraph.add_run(caption)
    _set_run_font(run, size=9, color=MID_GRAY)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)

    for style_name, size, color in (
        ("Heading 1", 17, NAVY),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 11, TEAL),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)

    if "Report Kicker" not in [style.name for style in document.styles]:
        style = document.styles.add_style("Report Kicker", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)

    caption = document.styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CHINESE_FONT)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("MSM｜D5 多相机视觉关联报告")
    _set_run_font(run, size=8.5, color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("技术报告  ·  第 ")
    _set_run_font(run, size=8.5, color=MID_GRAY)
    _add_field(footer, "PAGE", "1")
    run = footer.add_run(" 页")
    _set_run_font(run, size=8.5, color=MID_GRAY)

    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _add_cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(48)
    kicker = document.add_paragraph(style="Report Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MSM 反无人机多机协同研究系统")
    _set_run_font(run, size=12, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("D5 多相机视觉关联报告")
    _set_run_font(run, size=27, bold=True, color=NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("架构、方法与单种子仿真证据")
    _set_run_font(run, size=15, bold=True, color=TEAL)

    _add_callout(
        document,
        "文档范围",
        "说明 D5 多相机视觉关联的体系位置、时间与几何方法、局部检测跟踪、跨视角稳定注册、"
        "正式实验结果和标定计划。AirSim 数据属于固定场景单种子仿真，不代表实飞、"
        "物理拦截或跨场景泛化结果。",
        fill=PALE_BLUE,
        accent=NAVY,
    )

    document.add_paragraph().paragraph_format.space_after = Pt(22)
    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = (
        ("文档类型", "学术工程报告"),
        ("当前阶段", "几何关联基线已实现；单种子仿真已完成"),
        ("验证日期", "2026-07-16"),
        ("验证边界", "不评价实飞、控制效果或物理拦截结果"),
    )
    for row, values in zip(table.rows, rows):
        _prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = Cm(4 if index == 0 else 11.8)
            _set_cell_margins(cell, top=120, start=140, bottom=120, end=140)
            if index == 0:
                _set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            _set_run_font(
                run,
                size=10,
                bold=index == 0,
                color=NAVY if index == 0 else BLACK,
            )
    document.add_paragraph().paragraph_format.space_after = Pt(18)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第五研究模块（D5）｜2026 年 7 月")
    _set_run_font(run, size=10.5, color=MID_GRAY)
    document.add_page_break()


def _add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run("目录")
    _set_run_font(run, size=21, bold=True, color=NAVY)

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, title in enumerate(SECTION_TITLES, start=1):
        row = table.add_row()
        _prevent_row_split(row)
        left, right = row.cells
        left.width = Cm(14.6)
        right.width = Cm(1.2)
        for cell in (left, right):
            _set_cell_margins(cell, top=130, start=80, bottom=130, end=80)
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        run = left_p.add_run(title)
        _set_run_font(run, size=11.3, bold=True, color=NAVY)
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_after = Pt(0)
        _add_field(right_p, f"PAGEREF sec{index} \\h", "—")
    document.add_page_break()


def _start_chapter(
    document: Document,
    number: int,
    title: str,
    *,
    page_break: bool = True,
) -> None:
    heading = _add_heading(
        document,
        f"{number}. {title}",
        1,
        bookmark=f"sec{number}",
    )
    if page_break:
        heading.paragraph_format.page_break_before = True


def _build_report_body(document: Document) -> None:
    _start_chapter(document, 1, "体系架构", page_break=False)
    _add_body(
        document,
        "第五研究模块（D5）位于中心航迹、版本化任务分配、主动降级和末端视觉导引之间。"
        "第一、第二研究模块提供中心全局航迹对象（GlobalTrack），第三研究模块提供带版本、"
        "有效期和资源范围的目标分配。D5 判断当前本地视觉证据是否支持既有绑定，不重新分配目标。"
    )
    _add_body(
        document,
        "图 1 按中心状态、视觉证据和运行决策三个层次组织信息流。中间的多相机链路只输出"
        "关联置信度、歧义、时效与冲突依据；右侧主动降级和视觉导引切换继续由各自模块独立判断。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "system_architecture.png",
        "图 1  D5 在 D1—D7 体系信息流中的位置",
        width_cm=16.2,
    )
    _add_body(
        document,
        "图中两条右向分支共享视觉一致性证据，但不会合并成一个控制开关。"
        "第四研究模块综合上游不确定度、计划有效性和 D5 证据判断是否降级；"
        "第七研究模块在视觉导引切换前继续检查计划版本、证据时效、相机能力、机动能力和安全条件。"
    )
    _add_heading(document, "1.1 信息流", 2)
    _add_body(
        document,
        "D5 将已分配目标预测到各相机的图像量测时刻，再按每个相机的内参、外参、图像尺寸"
        "和镜头指向完成像素投影。局部检测与跟踪结果只在本资源、本相机作用域内参与匹配。"
        "即时选择通过时序稳定窗口后，才按中心持有的同一目标标识汇总。"
    )
    _add_body(
        document,
        "每个阶段都保留可审计中间量，包括预测像素、像素协方差、马氏距离、候选代价、"
        "唯一匹配结果、稳定计数和拒绝原因。故障可被定位为时间对齐、投影无效、门控拒绝、"
        "相机内竞争或稳定窗口不足，避免把全部失败压缩为一个未锁定状态。"
    )
    _add_heading(document, "1.2 职责边界", 2)
    _add_table(
        document,
        ("事项", "D5 处理", "边界"),
        (
            ("目标身份", "只读使用中心全局航迹标识（global_track_id）", "不创建、不改写、不换绑中心标识"),
            ("视觉关联", "输出末端关联对象（TerminalAssociation）及四态判断", "状态只表示证据质量"),
            ("侦察提示", "降低合适候选的关联代价", "不能替代本地实测或越过安全门"),
            ("交接与降级", "提供一致性和冲突证据", "不决定导引律，不输出控制量"),
            ("真值标签", "只允许离线评分使用", "actor 名称和 object_id 不进入在线选择"),
        ),
        widths_cm=(3.0, 7.0, 6.0),
    )
    _add_body(
        document,
        "中心标识在输入对象和输出对象中均按只读合同传递。本地多目标跟踪编号按资源和相机隔离；"
        "两个相机产生相同数字编号时，仍表示两条独立局部轨迹。末端重获取只能围绕原分配目标搜索，"
        "局部像素距离不能触发中心身份换绑。"
    )
    _add_callout(
        document,
        "身份原则",
        "已验证友方证据与候选重叠时进入暂停执行。身份未知保持未知，不能推定为敌方。",
        fill=LIGHT_RED,
        accent=RED,
    )

    _start_chapter(document, 2, "关键技术")
    _add_heading(document, "2.1 时间对齐", 2)
    _add_body(
        document,
        "量测时间戳（measurement_timestamp）表示图像实际采样时间，到达时间戳"
        "（arrival_timestamp）表示检测或跟踪结果进入 D5 的时间。"
        "几何预测必须对齐量测时刻。若使用到达时刻投影，传输、排队和推理延迟会被误认为目标运动。"
    )
    _add_body(
        document,
        "图 2 分别标出航迹状态时刻、图像量测时刻和证据到达时刻。上方关系用于目标运动补偿，"
        "下方关系用于证据新鲜度；两个时间差进入不同计算，不能互换。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "time_alignment.png",
        "图 2  量测时刻、到达时刻与状态预测",
        width_cm=15.8,
    )
    _add_body(
        document,
        "对航迹状态时刻 t₀ 和图像量测时刻 tₘ，基线采用常速度预测："
        "p(tₘ)=p(t₀)+v(t₀)(tₘ−t₀)。到达时刻减量测时刻得到证据年龄。"
        "延迟增大时，预测协方差随时间增长；陈旧证据的授权权重必须降低，超过时效上限时进入"
        "暂停执行或受限重获取。正式专项采用逐相机批次量测时刻，没有用 episode 末时刻覆盖整段观测。"
    )

    _add_heading(document, "2.2 坐标投影", 2)
    _add_body(
        document,
        "中心航迹使用北东地坐标系（North-East-Down，NED）。相机外参给出相机在该坐标系中的位置和姿态，"
        "用于把全局目标点变换到相机坐标。相机内参给出水平和垂直焦距、主点及像素尺度。"
    )
    _add_body(
        document,
        "图 3 展示从 NED 世界点、相机外参到像素平面的三步关系。AirSim 的前、右、下轴需要"
        "转换到 OpenCV 的右、下、前光学轴；轴约定或相机姿态错误会对多个目标产生方向一致的系统残差。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "coordinate_projection.png",
        "图 3  北东地坐标、相机外参与针孔投影",
        width_cm=15.8,
    )
    _add_body(
        document,
        "相机坐标记为 (Xc, Yc, Zc)，像素坐标满足 "
        "u=fxXc/Zc+cx、v=fyYc/Zc+cy。镜头指向决定目标是否位于相机前方，"
        "水平视场、垂直视场和图像尺寸共同决定有效像素边界。越界、非有限或位于镜头后方的投影"
        "不进入正常匹配。当前实现可消费畸变系数；本轮没有独立标定板误差证据。"
    )

    _add_heading(document, "2.3 不确定性传播", 2)
    _add_body(
        document,
        "三维航迹协方差不能直接作为像素门限。D5 使用投影雅可比 J 将三维协方差 P 传播到像素平面："
        "Σpx=J P Jᵀ+R。R 汇总相机量测、检测框中心和局部跟踪误差，并加入数值正则项。"
    )
    _add_body(
        document,
        "批量注册在边界框面积可用时还按框尺度和图像分辨率形成附加像素协方差；"
        "缺少框面积时使用相机批次给定的协方差。图 4 中椭圆的方向与尺度因此会随航迹、"
        "相机几何和本地量测共同变化。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "uncertainty_gating.png",
        "图 4  三维协方差传播、像素协方差与马氏门",
        width_cm=15.8,
    )
    _add_body(
        document,
        "检测中心 z 与预测像素 ẑ 的平方马氏距离为 "
        "d²=(z−ẑ)ᵀΣpx⁻¹(z−ẑ)。该距离按当前总不确定度归一化，"
        "同样的像素偏差在不同距离、标定质量和航迹协方差下会得到不同判定。默认平方门限为 9.21；"
        "门外候选被置为不可选大代价。像素误差与马氏距离需同时保留，前者便于直观检查，后者负责门控。"
    )

    _add_heading(document, "2.4 关联代价", 2)
    _add_body(
        document,
        "门内候选的总代价由投影误差、角速度一致性、类别一致性、局部航迹质量、身份冲突和侦察提示组成。"
        "身份冲突提高代价并可触发暂停执行。有效侦察提示只降低符合时效和作用域的候选代价。"
    )
    _add_body(
        document,
        "图 5 将候选代价、一对一匹配、稳定窗口和末端四态放在同一处理链中。"
        "图中候选到稳定注册属于几何注册层，稳定关联等四态属于末端决策层，"
        "一次当前帧选择不能直接解释为末端锁定。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "principle_pipeline.png",
        "图 5  关联代价、一对一匹配与时序稳定",
        width_cm=16.2,
    )
    _add_body(
        document,
        "每个相机内部构建“中心目标—本地轨迹”代价矩阵，并使用匈牙利算法执行一对一匹配。"
        "一条本地轨迹在同一相机批次内最多支持一个中心目标。正式 5+1 批量注册用门内马氏距离"
        "作为矩阵代价，以隔离验证时间和几何链；通用末端决策才消费完整多项代价。"
        "缺少 SciPy 时使用确定性贪心唯一匹配，该回退不宣称取得匈牙利全局最优结果。"
    )
    _add_body(
        document,
        "门内未选边仍保留候选概率，便于分析竞争关系。联合概率数据关联"
        "（Joint Probabilistic Data Association，JPDA）当前没有在线实现；未来实现仍需遵守"
        "上游绑定、友方门、计划版本、授权和稳定窗口。"
    )

    _add_heading(document, "2.5 时序稳定", 2)
    _add_body(
        document,
        "当前基线使用 3 帧窗口，要求同一资源、相机、本地轨迹和中心目标组合至少 2 帧有效匹配。"
        "第一次有效选择只保留为 candidate；满足累计条件后，当前选择才标记为 stable 并进入稳定跨视角汇总。"
        "资源、相机、本地编号或中心绑定发生变化时，稳定历史重新建立。"
    )
    _add_table(
        document,
        ("状态", "判定", "处置"),
        (
            ("稳定关联（locked）", "唯一实测候选通过几何、身份、版本和稳定门", "作为 D7 的一项前置证据"),
            ("候选模糊（ambiguous）", "有候选，但代价间隔、质量或稳定性不足", "继续观测或请求侦察提示"),
            ("暂停执行（hold）", "友方冲突、计划失效、时效超限或安全条件不满足", "阻止当前视觉交接"),
            ("受限重获取（reacquire）", "投影无效、无门内候选或检测丢失", "围绕原中心目标编号重获取"),
        ),
        widths_cm=(2.5, 8.0, 5.5),
    )
    _add_body(
        document,
        "只有 measured 实测轨迹可以通过正常几何门；predicted 和 lost 轨迹不能直接形成稳定关联。"
        "重获取恢复后还要重新满足实测和稳定条件。已验证友方重叠触发暂停，未知身份保持未知。"
    )

    _add_heading(document, "2.6 跨相机汇总", 2)
    _add_body(
        document,
        "各相机使用自己的内参、外参、图像尺寸和量测时刻独立完成几何配准。"
        "本地编号按资源和相机隔离。两个相机都出现编号 1 时，二者仍属于不同局部轨迹。"
    )
    _add_body(
        document,
        "图 6 中五路本地相机承担当前局部实测，一路宽视场侦察相机提供广域提示和附加视角。"
        "所有证据沿既有中心标识汇总；侦察框不填补局部空批次，也不绕过稳定窗口。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "six_camera_fusion.png",
        "图 6  五路本地相机与一路侦察相机的证据汇总",
        width_cm=16.2,
    )
    _add_body(
        document,
        "单个本地相机只需覆盖其当前视场内的目标子集。侦察相机提供广域提示和附加视角，"
        "但不能替代本地相机当前实测。当前实现按中心持有的同一目标编号汇总证据，"
        "尚未实现多视线三维联合优化和跨相机外观重识别。多个资源支持同一目标时还需核对联盟成员、"
        "所需资源数和授权作用域。"
    )

    _add_heading(document, "2.7 交接与降级", 2)
    _add_body(
        document,
        "图 7 将 D5 证据通往视觉导引切换和主动降级的两条支路并列展示。"
        "两条支路各自保留版本、时效和安全拒绝原因，视觉稳定状态不会直接触发控制动作。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "handoff_degradation.png",
        "图 7  D5 证据进入 D7 交接门和 D4 降级判断",
        width_cm=16.0,
    )
    _add_body(
        document,
        "D5 的稳定关联只确认本地视觉证据与既有中心目标一致。"
        "D7 仍需检查计划版本、证据时效、相机能力、机动能力和安全门控。"
        "D4 使用上游不确定度、计划有效性和 D5 一致性判断主动降级。相机几何不可用、"
        "计划版本落后、证据过期或联盟范围不符时，下游仍应拒绝交接。"
    )

    _start_chapter(document, 3, "关联方案")
    _add_heading(document, "3.1 数据合同", 2)
    _add_body(
        document,
        "中心航迹、相机本地轨迹批、本地视觉轨迹和末端关联构成四个主要合同。"
        "它们分别约束中心状态所有权、相机作用域、检测与跟踪证据，以及下游可以消费的保守判断。"
    )
    _add_table(
        document,
        ("对象", "核心内容", "使用边界"),
        (
            ("中心全局航迹", "标识、NED 位置与速度、三维协方差、时间、版本", "只读预测与投影"),
            ("相机本地轨迹批", "资源、相机模型、图像尺寸、本地轨迹、量测与到达时间", "每路相机独立处理"),
            ("本地视觉轨迹", "本地编号、框与中心、质量、历史、量测状态、后端来源", "仅实测状态进入正常门控"),
            ("末端关联", "中心标识、本地候选、置信度、歧义、冲突、四态和计划证据", "只报告关联，不授权控制"),
        ),
        widths_cm=(3.0, 7.4, 5.6),
    )
    _add_body(
        document,
        "相机空批次只表示该相机当前没有本地量测，其他相机的框不能填入。末端关联保留计划版本、"
        "授权状态、联盟角色、候选代价和拒绝原因，使下游能够复核证据成立的条件。"
    )

    _add_heading(document, "3.2 基线选择", 2)
    _add_body(
        document,
        "当前方案采用中心目标锚定、相机独立配准、时序稳定确认和跨视角证据汇总。"
        "该结构保留中心身份所有权，并能逐相机审计几何残差、时效和拒绝原因。"
    )
    _add_table(
        document,
        ("环节", "当前方案", "取舍"),
        (
            ("中心输入", "带协方差的中心航迹和版本化分配", "依赖上游航迹质量与计划有效性"),
            ("几何关联", "逐相机投影、马氏门和匈牙利匹配", "需要可靠标定与时间对齐"),
            ("时序确认", "3 帧窗口至少 2 次有效", "抑制单帧跳变，增加有限确认延迟"),
            ("跨视角", "按同一中心标识汇总稳定证据", "暂不生成联合三维状态"),
            ("安全决策", "友方冲突暂停，未知身份保持未知", "优先限制错误交接"),
        ),
        widths_cm=(3.0, 7.2, 5.8),
    )

    _add_heading(document, "3.3 局部检测与跟踪", 2)
    _add_body(
        document,
        "局部视觉链先从单帧图像生成检测框，再由多目标跟踪"
        "（Multi-Object Tracking，MOT）维护相机内连续编号。"
        "本地轨迹区分 measured、predicted 和 lost；预测或丢失状态不能直接形成稳定注册，"
        "重新出现的实测框需再次通过几何门和稳定窗口。"
    )
    _add_body(
        document,
        "正式 detect episode 使用匿名 AirSim detect 轨迹验证几何链。另一个 episode 的五路本地相机"
        "使用第八版“只看一次”目标检测器（YOLOv8）和 ByteTrack，侦察相机继续使用 AirSim detect。"
        "后端只负责产生本地框与本地编号，不能读取或生成中心标识。"
    )
    _add_body(
        document,
        "质量分数、MOT 历史长度、框裁切状态、检测来源和跟踪状态用于解释候选可靠性。"
        "正式 5+1 注册用几何代价隔离验证投影与唯一匹配，因此本轮数据没有覆盖所有质量项和身份项的组合验收。"
    )

    _add_heading(document, "3.4 后端路线", 2)
    _add_table(
        document,
        ("路线", "定位", "当前状态"),
        (
            ("AirSim detect 几何基线", "验证时间、几何、唯一匹配和稳定窗口", "单种子仿真达到分支阈值"),
            ("YOLOv8 + ByteTrack", "评估实际检测与局部多目标跟踪链路", "可选路线，单种子仿真未达到全部阈值"),
            ("联合概率数据关联", "处理密集候选和联合事件", "待评估，在线未实现"),
            ("多视线三维优化", "利用多相机射线联合估计", "待评估"),
            ("跨相机外观重识别", "增强跨视角身份连续性", "待评估"),
        ),
        widths_cm=(4.0, 7.5, 4.5),
        status_column=2,
    )
    _add_body(
        document,
        "AirSim detect 的通过结论属于时间—几何基线。YOLOv8 + ByteTrack 的状态取决于召回、"
        "严格关联准确率、本地身份连续性、侦察覆盖和时延的联合结果。"
        "JPDA、多视线三维优化和跨相机外观重识别均未计入当前能力。"
    )

    _add_heading(document, "3.5 约束条件", 2)
    for text in (
        "相机内外参、镜头指向和图像量测时刻必须可用；失效时进入保守状态。",
        "中心全局航迹标识由中心持有；本地检测编号和离线真值标签不能替代中心标识。",
        "侦察提示必须满足作用域和时效；过期提示不进入当前交接判断。",
        "稳定关联不能单独触发视觉导引；D7 和 D4 的独立门控保持有效。",
    ):
        _add_bullet(document, text)

    _start_chapter(document, 4, "实验结果")
    _add_heading(document, "4.1 验证边界", 2)
    _add_body(
        document,
        "2026-07-16 的正式专项包含两个固定场景 episode，只使用 seed 7。"
        "两个 episode 共享相机布局、目标初始位置和运动时长，分别评估 AirSim detect 几何基线"
        "与 YOLOv8 + ByteTrack 可选路线。所有数值均来自这批正式指标和候选记录。"
    )
    _add_table(
        document,
        ("边界", "本轮设置", "结论限制"),
        (
            ("随机性", "单个 seed，49 帧", "无均值、方差和置信区间"),
            ("中心航迹", "未在线运行 D1/D2，由 actor 真值运动生成夹具", "不证明上游在线融合性能"),
            ("在线关联", "代价、匈牙利选择和稳定窗口不读取真值身份", "只证明被审计环节真值隔离"),
            ("离线评分", "真值框用于召回、错配和身份切换计算", "评分不回流在线选择"),
            ("系统范围", "只验证 D5 视觉关联", "不评价控制、拦截距离、实飞或物理拦截"),
        ),
        widths_cm=(3.0, 7.5, 5.5),
    )
    _add_body(
        document,
        "中心航迹夹具使几何注册可以隔离验证，但不能代替完整上游链。"
        "在线真值身份使用计数为 0 只约束候选代价、唯一匹配和稳定窗口，"
        "不表示夹具生成和离线评分没有使用真值。"
    )

    _add_heading(document, "4.2 实验条件", 2)
    _add_body(
        document,
        "AirSim Blocks 运行在计算机视觉（ComputerVision）模式。五路本地相机允许只看到目标子集，"
        "一路高位侦察相机提供更宽视场。侦察相机在两个 episode 中均使用 AirSim detect，"
        "因此 YOLO 路线的侦察覆盖仍受同一侦察几何和 detect 可见性约束。"
    )
    _add_table(
        document,
        ("项目", "配置"),
        (
            ("相机", "5 个本地相机，1920×1080，水平视场 60°；1 个侦察相机，3840×2160，75°"),
            ("目标", "5 个移动 Quadrotor1 actor；初距约 30 m；目标间距 8 m；目标尺度 2.0"),
            ("高度", "本地相机与目标高度约 50 m；侦察相机再高约 50 m"),
            ("时序", "12.0 s；0.25 s 间隔；49 帧；seed 7"),
            ("目标运动", "12 s 内前向 7.2—9.6 m，横向绝对位移 0—6.0 m，高度不变"),
            ("中心航迹", "专项未在线运行 D1/D2；系统编排侧根据 actor 真值运动合成中心航迹夹具"),
            ("真值用途", "只用于离线召回、错配和身份切换评分，不回流在线关联"),
        ),
        widths_cm=(3.2, 12.8),
    )
    _add_body(
        document,
        "目标在 12 s 内发生可观运动，逐相机量测时间因此属于必要输入。"
        "若将早期帧统一投影到末时刻，预测会提前移动 7.2—9.6 m，并产生最多 6.0 m 的横向偏移。"
        "正式结果采用逐相机批次量测时刻；修复前没有保留同配置完整对照，故不量化修复增益。"
    )

    _add_heading(document, "4.3 总体结果", 2)
    _add_body(
        document,
        "detect 产生 1254 个在线检测，全部与离线真值框匹配；1220 个已选注册全部可评分，错配为 0。"
        "YOLO 路线产生 822 个在线检测，其中 781 个与真值框匹配；794 个注册被选中，"
        "770 个可评分，24 个无法匹配真值框。"
    )
    _add_table(
        document,
        ("指标", "AirSim detect", "YOLOv8 + ByteTrack"),
        (
            ("在线检测数", "1254", "822"),
            ("检测召回率", "1.000", "0.622"),
            ("关联准确率", "1.000", "0.996"),
            ("严格关联准确率", "1.000", "0.966"),
            ("稳定注册率", "0.975", "0.955"),
            ("本地相机联合覆盖率", "1.000", "1.000"),
            ("侦察全视野率", "0.918", "0.878"),
            ("局部身份切换", "0", "25"),
            ("在线真值身份使用", "0", "0"),
            ("中心目标编号改写", "0", "0"),
            ("推理时延 P50 / P95", "不适用", "10.42 / 12.37 ms"),
        ),
        widths_cm=(7.0, 4.0, 5.0),
    )
    _add_body(
        document,
        "第 50 和第 95 百分位分别记为 P50 和 P95。严格关联准确率把已选但无法与离线真值框匹配"
        "的结果计为错误。本地身份切换计数（Identity Switch，IDSW）由离线真值评分，"
        "不参与在线候选选择。"
    )
    _add_body(
        document,
        "图 8 将实测值与本专项分支门限并列。绿色只表示该项在 seed 7 固定场景中达到门限，"
        "不能解释为跨场景统计通过。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "backend_comparison.png",
        "图 8  两条后端路线的单种子仿真指标与分支阈值",
        width_cm=16.2,
    )
    _add_callout(
        document,
        "当前判定",
        "AirSim detect 几何基线达到本次 seed 7 分支阈值。"
        "YOLOv8 + ByteTrack 的召回率、侦察全视野率和局部身份切换未达到阈值，继续作为可选路线。",
        fill=LIGHT_ORANGE,
        accent=ORANGE,
    )
    _add_body(
        document,
        "YOLO 路线有 3 个可评分错配：G-102→G-103 一次，G-105→G-104 两次。"
        "其严格准确率和稳定注册率达到门限，但召回、侦察全覆盖和 IDSW 未通过。"
        "较高稳定率不能抵消漏检和局部身份连续性问题。"
    )

    _add_heading(document, "4.4 逐相机结果", 2)
    _add_body(
        document,
        "逐相机统计用于定位错误集中视角。下表准确率以可评分结果为分母；"
        "无真值匹配结果单列，并已在总体严格准确率中按错误处理。"
    )
    _add_table(
        document,
        ("相机", "detect 已选/错配/无真值", "YOLO 已选/错配/无真值", "YOLO 准确率"),
        (
            ("本地相机 1", "167 / 0 / 0", "110 / 0 / 2", "1.000"),
            ("本地相机 2", "210 / 0 / 0", "149 / 3 / 11", "0.978"),
            ("本地相机 3", "240 / 0 / 0", "139 / 0 / 9", "1.000"),
            ("本地相机 4", "205 / 0 / 0", "101 / 0 / 2", "1.000"),
            ("本地相机 5", "160 / 0 / 0", "56 / 0 / 0", "1.000"),
            ("侦察相机", "238 / 0 / 0", "239 / 0 / 0", "1.000"),
        ),
        widths_cm=(3.0, 4.8, 5.0, 3.2),
    )
    _add_body(
        document,
        "YOLO 的 3 个可评分错配和 11 个无真值匹配结果集中在第二路本地相机；"
        "第三路另有 9 个无真值匹配。侦察相机在两次运行中都使用 detect，"
        "因此该对照更直接反映局部检测与跟踪后端差异。"
    )
    _add_table(
        document,
        ("相机", "detect 均值/P50/P95 px", "YOLO 均值/P50/P95 px"),
        (
            ("本地相机 1", "10.12 / 10.68 / 14.74", "9.44 / 8.50 / 14.79"),
            ("本地相机 2", "9.31 / 10.04 / 13.09", "17.94 / 12.52 / 15.13"),
            ("本地相机 3", "9.57 / 10.95 / 13.56", "11.72 / 13.19 / 17.56"),
            ("本地相机 4", "11.49 / 11.00 / 15.49", "9.28 / 9.33 / 17.29"),
            ("本地相机 5", "10.07 / 7.10 / 16.02", "9.65 / 8.98 / 15.05"),
            ("侦察相机", "6.25 / 6.49 / 10.08", "6.24 / 6.51 / 10.08"),
        ),
        widths_cm=(3.0, 6.4, 6.6),
    )
    _add_body(
        document,
        "YOLO 第二路本地相机包含一个 284.93 px 的所选结果，使均值升至 17.94 px，"
        "而 P95 为 15.13 px。该尾部异常与同相机的 3 个错配同时出现；现有数据尚不能把原因"
        "唯一归结为检测框、局部编号变化或标定偏差。"
    )

    _add_heading(document, "4.5 六路画面", 2)
    _add_body(
        document,
        "t=6 s 拼图用于核对局部视场子集和侦察广域视角。图 9 对应 detect 几何基线，"
        "五路本地相机检测数依次为 4、5、5、4、3，侦察相机检测到 5 个目标。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "detect_t6_montage.png",
        "图 9  AirSim detect 分支在 t=6 s 的六路仿真画面",
        width_cm=16.4,
    )
    _add_body(
        document,
        "各本地画面覆盖不同目标子集，侦察俯视画面提供全局覆盖。"
        "拼图只承担视觉核对，召回和关联结论仍以逐帧正式指标为准。"
    )
    _add_body(
        document,
        "图 10 对应 YOLOv8 + ByteTrack。五路本地相机检测数依次为 4、2、3、2、1；"
        "侦察相机继续使用 AirSim detect 并检测到 5 个目标。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "yolo_t6_montage.png",
        "图 10  YOLOv8 + ByteTrack 分支在 t=6 s 的六路仿真画面",
        width_cm=16.4,
    )
    _add_body(
        document,
        "YOLO 局部画面的漏检比 detect 明显，与总体召回率 0.622 一致。"
        "侦察画面全目标可见不能补齐局部实测；稳定注册仍依赖本地相机自己的当前框。"
    )

    _add_heading(document, "4.6 覆盖与时间线", 2)
    _add_body(
        document,
        "图 11 给出 detect 六路相机随时间的在线局部轨迹数，用于区分单路子集覆盖与五路联合覆盖。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "detect_coverage.png",
        "图 11  AirSim detect 六路相机在线局部轨迹数",
        width_cm=15.8,
    )
    _add_body(
        document,
        "detect 的五路本地相机联合覆盖率为 1.000，侦察全覆盖帧率为 0.918。"
        "联合覆盖表示局部视场并集覆盖全部目标，不要求单路相机在每帧都看到五个目标。"
    )
    _add_body(
        document,
        "图 12 比较 detect 每帧即时选择和通过 3 帧 2 次稳定窗口的注册数。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "detect_timeline.png",
        "图 12  AirSim detect 即时选择与稳定注册",
        width_cm=15.8,
    )
    _add_body(
        document,
        "detect 共选择 1220 个注册，其中 1189 个形成稳定支持，31 个未满足稳定窗口，"
        "稳定注册率为 0.975。几何门拒绝数为 0 只适用于本次协方差、目标间距和相机设置。"
    )
    _add_body(
        document,
        "图 13 给出 YOLO 路线的六路在线局部轨迹数；侦察曲线仍来自 AirSim detect。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "yolo_coverage.png",
        "图 13  YOLOv8 + ByteTrack 六路相机在线局部轨迹数",
        width_cm=15.8,
    )
    _add_body(
        document,
        "YOLO 路线的局部相机联合覆盖仍为 1.000，侦察全覆盖为 0.878。"
        "视场并集覆盖不能抵消逐框召回不足，也不能解释本地身份切换。"
    )
    _add_body(
        document,
        "图 14 比较 YOLO 路线的即时选择与稳定注册。局部漏检和编号变化会使稳定键重新累计。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "yolo_timeline.png",
        "图 14  YOLOv8 + ByteTrack 即时选择与稳定注册",
        width_cm=15.8,
    )
    _add_body(
        document,
        "YOLO 路线共选择 794 个注册，其中 758 个形成稳定支持，36 个未通过稳定窗口，"
        "稳定注册率为 0.955。该比例需与召回、严格准确率和 IDSW 联合判读。"
    )
    _add_body(
        document,
        "图 15 将两个 episode 的六路检测总数和每帧稳定注册数置于同一时间轴。"
        "它用于观察固定场景中的持续差异，不提供统计显著性。",
        keep_with_next=True,
    )
    _add_picture(
        document,
        "simulation_timeline.png",
        "图 15  两条后端路线的检测与稳定注册时间序列",
        width_cm=15.8,
    )
    _add_body(
        document,
        "YOLO 路线的检测总数长期低于 detect，稳定注册数随本地检测供应变化。"
        "两个 episode 的在线真值身份使用和中心标识改写均为 0；中心航迹夹具和离线评分仍使用真值。"
    )

    _start_chapter(document, 5, "边界与计划")
    _add_heading(document, "5.1 能力分层", 2)
    _add_table(
        document,
        ("分层", "内容", "状态"),
        (
            ("算法基线", "双时间戳、NED 投影、协方差传播、马氏门和相机内唯一匹配", "已实现"),
            ("稳定与汇总", "3 帧 2 次稳定窗口、局部编号隔离和跨视角证据汇总", "已实现"),
            ("安全边界", "友方冲突暂停、未知身份保守、中心标识只读", "已实现"),
            ("固定场景", "2026-07-16，5+1 相机，5 个目标，seed 7，49 帧", "单种子已验证"),
            ("YOLO 路线", "召回、侦察全覆盖和本地 IDSW 未通过本专项门限", "可选"),
            ("统计稳定性", "多 seed、遮挡、交叉、标定和时间扰动", "待验证"),
            ("增强路线", "在线 JPDA、多视线三维优化和跨相机外观重识别", "未实现"),
        ),
        widths_cm=(3.1, 9.4, 3.5),
        status_column=2,
    )
    _add_body(
        document,
        "“已实现”描述当前代码能力，“单种子已验证”描述正式实验覆盖，"
        "“可选”表示存在实现但未通过全部分支门限。建议值与待验证项不应写成当前性能。"
    )

    _add_heading(document, "5.2 失败边界", 2)
    _add_body(
        document,
        "正式场景没有注入内参误差、外参旋转和平移偏差、时钟偏移或网络抖动。"
        "几何门拒绝为 0 只说明本次相机模型、协方差和目标间距没有产生门外选择，"
        "不能证明同一门限在标定漂移下仍然有效。"
    )
    _add_body(
        document,
        "本轮也没有遮挡分级、近距离交叉、背景变化、目标尺度变化、强机动或长时间丢失。"
        "YOLO 的 25 次 IDSW 与 3 次错配揭示局部连续性问题，单次运行无法估计发生概率和置信区间。"
    )
    _add_table(
        document,
        ("环节", "真值身份使用", "判断"),
        (
            ("D1/D2 在线链路", "未运行", "不形成上游在线性能结论"),
            ("中心航迹夹具", "使用 actor 真值运动学", "只形成隔离测试输入"),
            ("像素门与匈牙利选择", "0 次", "使用几何、协方差和本地质量"),
            ("稳定窗口", "0 次", "使用作用域、中心绑定和时序历史"),
            ("召回、错配与 IDSW", "离线使用", "评分不回流在线关联"),
            ("中心标识改写", "0 次", "两个 episode 均满足只读约束"),
        ),
        widths_cm=(4.0, 4.0, 8.0),
    )
    _add_body(
        document,
        "候选记录中的离线真值列只用于运行后评分。在线真值身份使用计数为 0，"
        "不能解释为整个专项没有真值；它只证明被审计的在线关联环节未读取真值身份。"
    )

    _add_heading(document, "5.3 标定计划", 2)
    _add_body(
        document,
        "下一轮应先建立可追溯的相机内参、畸变、相机到机体外参、机体到 NED 姿态和时钟同步基线。"
        "每次标定需记录来源、时间、适用分辨率和有效期，并用独立数据给出重投影误差分布。"
        "本报告的目标关联像素误差不能替代标定板或独立场景的标定残差。"
    )
    _add_body(
        document,
        "外参实验应分别注入旋转和平移偏差，同时固定检测框与目标轨迹。"
        "输出至少包括像素误差 P50/P95、马氏门通过率、错配率、稳定窗口失败数和恢复时间。"
        "量测时间偏差与到达延迟需要分开注入，前者影响几何预测，后者影响证据年龄和授权。"
    )

    _add_heading(document, "5.4 后续实验", 2)
    _add_table(
        document,
        ("优先级", "建议工作", "报告指标"),
        (
            ("P1", "同一相机配置和目标轨迹下运行至少 10 个配对 seeds", "均值、标准差、置信区间、最差 seed"),
            ("P1", "提升 YOLOv8 召回并降低 ByteTrack 身份切换", "召回、严格准确率、IDSW、P95 时延"),
            ("P1", "分级注入遮挡、目标交叉、外参与量测时间偏差", "拒绝率、错配率、恢复时间、门控余量"),
            ("P1", "改变侦察视场、姿态和提示时效", "全覆盖、提示有效率、过期拒绝"),
            ("P2", "离线评估 JPDA、多视线三维优化和外观重识别", "收益、计算预算、安全合同影响"),
        ),
        widths_cm=(2.0, 8.4, 5.6),
    )
    _add_body(
        document,
        "至少 10 个配对 seeds 属于下一阶段建议值，当前尚未执行。detect 与 YOLO 路线应在相同 seed、"
        "相机设置和 actor 轨迹下运行，减少场景差异。YOLO 后端需在召回、严格准确率、IDSW、"
        "侦察覆盖和时延同时满足后续正式准入条件后，再评估默认状态。"
    )

    _add_heading(document, "5.5 结论", 2)
    for text in (
        "D5 已实现中心航迹锚定的多相机几何关联基线，覆盖时间、投影、门控、唯一匹配、时序稳定和跨视角汇总；D5 只输出证据，不授权控制。",
        "seed 7 固定场景中，AirSim detect 通过全部专项门限；YOLOv8 + ByteTrack 未通过召回、侦察覆盖和 IDSW 门限，保持可选。结论不外推至实飞或物理拦截。",
    ):
        _add_bullet(document, text)


def build_docx() -> None:
    missing = [name for name in REPORT_IMAGES if not (ASSET_DIR / name).is_file()]
    if missing:
        raise FileNotFoundError("缺少报告图片：\n" + "\n".join(missing))
    document = Document()
    _configure_document(document)
    _add_cover(document)
    _add_toc(document)
    _build_report_body(document)

    properties = document.core_properties
    properties.title = "D5 多相机视觉关联报告"
    properties.subject = "体系架构、关键技术、关联方案与 AirSim 仿真结果"
    properties.author = "MSM 第五研究模块（D5）"
    properties.keywords = "D5, 多相机, 视觉关联, AirSim, 仿真"
    properties.comments = "区分已实现、单种子仿真验证、建议指标和待验证内容。"
    temporary_path = DOCX_PATH.with_suffix(".tmp.docx")
    document.save(temporary_path)
    temporary_path.replace(DOCX_PATH)


IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\(([^)]+)\)")


def validate_markdown_images() -> list[str]:
    failures: list[str] = []
    for markdown in MODULE_DIR.rglob("*.md"):
        text = markdown.read_text(encoding="utf-8")
        for raw_target in IMAGE_PATTERN.findall(text):
            target = raw_target.strip().split(maxsplit=1)[0].strip("<>")
            if target.startswith(("http://", "https://", "data:")):
                continue
            if not (markdown.parent / target).resolve().is_file():
                failures.append(f"{markdown.relative_to(MODULE_DIR)} -> {raw_target}")
    return failures


def validate_report_markdown() -> dict[str, int]:
    report_path = DOCS_DIR / "D5_MULTICAMERA_ASSOCIATION_REPORT_CN.md"
    text = report_path.read_text(encoding="utf-8")
    forbidden = [term for term in FORBIDDEN_REPORT_TERMS if term in text]
    if forbidden:
        raise RuntimeError("Markdown 正文残留禁用词：" + "、".join(forbidden))
    path_markers = [marker for marker in PATH_MARKERS if marker in text]
    if path_markers:
        raise RuntimeError("Markdown 正文残留文件路径标记：" + "、".join(path_markers))
    image_targets = IMAGE_PATTERN.findall(text)
    expected_targets = {
        f"assets/d5_multicamera_association/{name}" for name in REPORT_IMAGES
    }
    actual_targets = {
        raw_target.strip().split(maxsplit=1)[0].strip("<>")
        for raw_target in image_targets
    }
    missing_targets = sorted(expected_targets - actual_targets)
    if missing_targets:
        raise RuntimeError("Markdown 缺少报告图片引用：" + "、".join(missing_targets))
    return {
        "headings": len(re.findall(r"^#{1,3} ", text, flags=re.MULTILINE)),
        "images": len(image_targets),
        "characters": len(text),
    }


def _all_document_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def validate_docx() -> dict[str, int]:
    if not DOCX_PATH.is_file():
        raise FileNotFoundError(DOCX_PATH)
    with zipfile.ZipFile(DOCX_PATH) as archive:
        bad_member = archive.testzip()
        if bad_member is not None:
            raise RuntimeError(f"DOCX 压缩成员校验失败：{bad_member}")
        media = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
    document = Document(DOCX_PATH)
    text = _all_document_text(document)
    forbidden = [term for term in FORBIDDEN_REPORT_TERMS if term in text]
    if forbidden:
        raise RuntimeError("DOCX 正文残留禁用词：" + "、".join(forbidden))
    path_markers = [marker for marker in PATH_MARKERS if marker in text]
    if path_markers:
        raise RuntimeError("DOCX 正文残留文件路径或代码标记：" + "、".join(path_markers))
    for title in SECTION_TITLES:
        if title not in text:
            raise RuntimeError(f"DOCX 缺少固定章节：{title}")
    if len(document.inline_shapes) != len(REPORT_IMAGES) or len(media) != len(REPORT_IMAGES):
        raise RuntimeError(
            f"DOCX 图片数量异常：media={len(media)}, inline={len(document.inline_shapes)}"
        )
    captions = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Caption"
    ]
    expected_captions = [f"图 {index}" for index in range(1, len(REPORT_IMAGES) + 1)]
    if len(captions) != len(REPORT_IMAGES):
        raise RuntimeError(
            f"DOCX 图题数量异常：captions={len(captions)}, expected={len(REPORT_IMAGES)}"
        )
    for expected, caption in zip(expected_captions, captions):
        if not caption.startswith(expected + "  "):
            raise RuntimeError(f"DOCX 图题顺序异常：期望 {expected}，实际 {caption}")
    if CHINESE_FONT not in document_xml or CHINESE_FONT not in styles_xml:
        raise RuntimeError(f"DOCX 未完整声明中文字体：{CHINESE_FONT}")
    page_break_before_count = document_xml.count("<w:pageBreakBefore")
    if page_break_before_count < len(SECTION_TITLES) - 1:
        raise RuntimeError(
            "DOCX 章节分页不足："
            f"pageBreakBefore={page_break_before_count}, expected>={len(SECTION_TITLES) - 1}"
        )
    keep_next_count = document_xml.count("<w:keepNext")
    if keep_next_count < len(REPORT_IMAGES):
        raise RuntimeError(
            f"DOCX 图文分页保护不足：keepNext={keep_next_count}, images={len(REPORT_IMAGES)}"
        )
    metadata = "\n".join(
        (
            document.core_properties.title or "",
            document.core_properties.subject or "",
            document.core_properties.comments or "",
        )
    )
    forbidden_metadata = [
        term for term in ("领导汇报", "领导送阅", "仓库") if term in metadata
    ]
    if forbidden_metadata:
        raise RuntimeError("DOCX 元数据残留禁用词：" + "、".join(forbidden_metadata))
    return {
        "embedded_images": len(media),
        "inline_shapes": len(document.inline_shapes),
        "captions": len(captions),
        "page_break_before": page_break_before_count,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--sync-formal-assets",
        action="store_true",
        help="从正式 AirSim 产物显式同步截图和绘图数据",
    )
    parser.add_argument(
        "--check-only",
        action="store_true",
        help="只检查现有 Markdown 图片和 DOCX 结构，不重新生成",
    )
    args = parser.parse_args()

    if not args.check_only:
        if args.sync_formal_assets:
            sync_formal_assets()
        generate_all_figures()
        build_docx()

    markdown_failures = validate_markdown_images()
    if markdown_failures:
        raise FileNotFoundError(
            "Markdown 本地图片链接失效：\n" + "\n".join(markdown_failures)
        )
    markdown_stats = validate_report_markdown()
    stats = validate_docx()
    print(
        "DOCX validation: "
        + ", ".join(f"{key}={value}" for key, value in stats.items())
    )
    print(
        "Markdown validation: "
        + ", ".join(f"{key}={value}" for key, value in markdown_stats.items())
    )
    print("Markdown image validation: all local image links exist")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
