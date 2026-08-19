#!/usr/bin/env python3
"""Build a dual-optical multi-target Markdown report as a Word document."""

from __future__ import annotations

import argparse
import re
import tempfile
import warnings
from pathlib import Path
from zipfile import ZipFile

warnings.filterwarnings(
    "ignore",
    message="Unable to import Axes3D.*",
    category=UserWarning,
)
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MODULE_DIR = Path(__file__).resolve().parent
DEFAULT_SOURCE = (
    MODULE_DIR
    / "outputs"
    / "airsim_seed_20260811_epipolar_mht_run01"
    / "DUAL_OPTICAL_40TARGET_AIRSIM_REPORT_CN.md"
)

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
MONO_FONT = "Consolas"

BLUE = "1F4E78"
TEAL = "176B73"
INK = "202833"
MUTED = "5F6B78"
WHITE = "FFFFFF"
ROW_FILL = "F4F7FA"

INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\$[^$]+?\$)")
TABLE_SEPARATOR_RE = re.compile(r"^:?-{3,}:?$")


def set_run_font(
    run,
    *,
    size: float = 10.5,
    bold: bool | None = None,
    color: str = INK,
    east_asia: str = BODY_FONT,
    latin: str = LATIN_FONT,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(4)

    style_specs = {
        "Title": (24, BLUE, 0, 14),
        "Heading 1": (16, BLUE, 14, 7),
        "Heading 2": (13.5, TEAL, 10, 5),
        "Heading 3": (11.5, INK, 8, 4),
    }
    for style_name, (size, color, before, after) in style_specs.items():
        style = document.styles[style_name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def apply_page_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    set_run_font(run, size=8.5, color=MUTED)


def set_page_number_start(section, value: int) -> None:
    section_properties = section._sectPr
    page_number = section_properties.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section_properties.append(page_number)
    page_number.set(qn("w:start"), str(value))


def configure_header_footer(section, title: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(title)
    set_run_font(run, size=8.5, color=MUTED, east_asia=HEADING_FONT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    add_page_field(footer)


def add_cover(document: Document, title: str, target_count: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(125)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_run_font(run, size=26, bold=True, color=BLUE, east_asia=HEADING_FONT)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("AirSim独立试验报告")
    set_run_font(run, size=16, bold=True, color=TEAL, east_asia=HEADING_FONT)

    scenario = document.add_paragraph()
    scenario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scenario.paragraph_format.space_before = Pt(42)
    scenario.paragraph_format.first_line_indent = Cm(0)
    run = scenario.add_run(
        f"双固定光电节点  |  {target_count}个移动目标  |  轨迹关联"
    )
    set_run_font(run, size=11.5, color=INK, east_asia=HEADING_FONT)

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.paragraph_format.space_before = Pt(175)
    date_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = date_paragraph.add_run("2026年8月11日")
    set_run_font(run, size=11, color=MUTED, east_asia=HEADING_FONT)


def add_table_of_contents(document: Document) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.first_line_indent = Cm(0)
    run = heading.add_run("目录")
    set_run_font(run, size=16, bold=True, color=BLUE, east_asia=HEADING_FONT)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后更新目录字段"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    set_run_font(run, size=10.5, color=MUTED)


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def math_to_plain(text: str) -> str:
    result = text.strip().strip("$")
    replacements = {
        r"\alpha": "α",
        r"\beta": "β",
        r"\theta": "θ",
        r"\lambda": "λ",
        r"\Delta": "Δ",
        r"\kappa": "κ",
        r"\rho": "ρ",
        r"\psi": "ψ",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\times": "×",
        r"\quad": "  ",
        r"\,": " ",
    }
    for source, replacement in replacements.items():
        result = result.replace(source, replacement)
    result = re.sub(r"\\mathbf\s+([A-Za-z0-9])", r"\1", result)
    result = re.sub(r"\\(?:mathbf|mathrm|text)\{([^{}]*)\}", r"\1", result)
    result = result.replace(r"\|", "‖").replace("^T", "ᵀ")
    return result.replace("{", "").replace("}", "")


def add_inline(paragraph, text: str, *, size: float = 10.5) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size - 0.5, color=TEAL, latin=MONO_FONT)
        else:
            run = paragraph.add_run(math_to_plain(token))
            set_run_font(run, size=size, latin=MATH_FONT)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    first = lines[index].strip()
    separator = lines[index + 1].strip()
    if not (first.startswith("|") and first.endswith("|")):
        return False
    cells = split_table_row(separator)
    return bool(cells) and all(TABLE_SEPARATOR_RE.fullmatch(cell) for cell in cells)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = [split_table_row(lines[index])]
    index += 2
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        rows.append(split_table_row(line))
        index += 1
    return rows, index


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_table(document: Document, rows: list[list[str]]) -> None:
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    font_size = 9.0 if column_count <= 3 else 8.2

    for row_index, row in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        prevent_split = OxmlElement("w:cantSplit")
        row_properties.append(prevent_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            row_properties.append(repeat)
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.1
            paragraph.paragraph_format.space_after = Pt(0)
            content = row[column_index] if column_index < len(row) else ""
            if row_index == 0:
                set_cell_shading(cell, BLUE)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(content)
                set_run_font(
                    run,
                    size=font_size,
                    bold=True,
                    color=WHITE,
                    east_asia=HEADING_FONT,
                )
            else:
                if row_index % 2 == 0:
                    set_cell_shading(cell, ROW_FILL)
                add_inline(paragraph, content, size=font_size)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_image(document: Document, source: Path, alt: str, relative: str, number: int) -> None:
    image_path = (source.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"missing report image: {image_path}")
    with Image.open(image_path) as raster:
        ratio = raster.width / max(raster.height, 1)
    max_width_cm = 16.4
    max_height_cm = 18.3
    width_cm = min(max_width_cm, max_height_cm * ratio)
    height_cm = width_cm / ratio

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(
        str(image_path), width=Cm(width_cm), height=Cm(height_cm)
    )

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run(f"图{number}  {alt}")
    set_run_font(run, size=9, color=MUTED, east_asia=BODY_FONT)


def render_equation(tex: str, path: Path) -> None:
    normalized = re.sub(r"\\mathbf\s+([A-Za-z0-9])", r"\\mathbf{\1}", tex)
    normalized = normalized.replace(r"\|", r"\Vert")
    figure = plt.figure(figsize=(12.0, 0.9), facecolor="white")
    figure.text(
        0.5,
        0.5,
        f"${normalized}$",
        ha="center",
        va="center",
        fontsize=17,
        color="#202833",
    )
    figure.savefig(
        path,
        dpi=220,
        bbox_inches="tight",
        pad_inches=0.08,
        facecolor="white",
    )
    plt.close(figure)


def add_equation(document: Document, tex: str, image_path: Path) -> None:
    render_equation(tex, image_path)
    with Image.open(image_path) as raster:
        natural_width_cm = raster.width / 220.0 * 2.54
        ratio = raster.width / max(raster.height, 1)
    width_cm = min(15.5, max(6.0, natural_width_cm))
    height_cm = width_cm / ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run().add_picture(
        str(image_path), width=Cm(width_cm), height=Cm(height_cm)
    )


def extract_display_equation(lines: list[str], index: int) -> tuple[str, int]:
    line = lines[index].strip()
    if line.startswith("$$") and line.endswith("$$") and len(line) > 4:
        return line[2:-2].strip(), index + 1
    parts = [line.removeprefix("$$")]
    index += 1
    while index < len(lines):
        item = lines[index].strip()
        if item.endswith("$$"):
            parts.append(item.removesuffix("$$"))
            return " ".join(parts).strip(), index + 1
        parts.append(item)
        index += 1
    raise RuntimeError("unterminated display equation")


def add_heading(document: Document, level: int, text: str) -> None:
    style = {2: "Heading 1", 3: "Heading 2"}.get(level, "Heading 3")
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    sizes = {2: 16, 3: 13.5}
    colors = {2: BLUE, 3: TEAL}
    set_run_font(
        run,
        size=sizes.get(level, 11.5),
        bold=True,
        color=colors.get(level, INK),
        east_asia=HEADING_FONT,
    )


def add_list_item(document: Document, marker: str, content: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(-0.55)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(2)
    if marker in {"-", "*"} and content.startswith("[ ] "):
        prefix = "☐"
        content = content[4:]
    elif marker in {"-", "*"}:
        prefix = "•"
    else:
        prefix = marker
    run = paragraph.add_run(f"{prefix}  ")
    set_run_font(run, size=10.0, bold=True, color=BLUE, east_asia=HEADING_FONT)
    add_inline(paragraph, content, size=10.0)


def build_document(source: Path, output: Path) -> dict[str, int]:
    lines = source.read_text(encoding="utf-8").splitlines()
    title_line = next((line for line in lines if line.startswith("# ")), None)
    if title_line is None:
        raise RuntimeError("report title not found")
    title = title_line[2:].strip()
    target_match = re.search(r"(\d+)目标", title)
    if target_match is None:
        target_match = re.search(r"DUAL_OPTICAL_(\d+)TARGET", source.name)
    if target_match is None:
        raise RuntimeError("target count is missing from report title and filename")
    target_count = int(target_match.group(1))

    document = Document()
    apply_page_layout(document.sections[0])
    configure_styles(document)
    add_cover(document, title, target_count)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(content_section)
    set_page_number_start(content_section, 1)
    configure_header_footer(content_section, title)
    add_table_of_contents(document)
    document.add_page_break()

    figure_count = 0
    equation_count = 0
    table_count = 0
    index = 0
    with tempfile.TemporaryDirectory(prefix="dual-optical-word-") as temporary:
        temporary_dir = Path(temporary)
        while index < len(lines):
            line = lines[index].strip()
            if not line:
                index += 1
                continue
            if line.startswith("# "):
                index += 1
                continue
            if line.startswith("$$"):
                equation, index = extract_display_equation(lines, index)
                equation_count += 1
                add_equation(
                    document,
                    equation,
                    temporary_dir / f"equation_{equation_count:02d}.png",
                )
                continue
            image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if image_match:
                figure_count += 1
                add_image(
                    document,
                    source,
                    image_match.group(1),
                    image_match.group(2),
                    figure_count,
                )
                index += 1
                continue
            if is_table_start(lines, index):
                rows, index = parse_table(lines, index)
                add_table(document, rows)
                table_count += 1
                continue
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                add_heading(document, level, line[level:].strip())
                index += 1
                continue
            list_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", line)
            if list_match:
                add_list_item(document, list_match.group(1), list_match.group(2))
                index += 1
                continue

            paragraph_lines = [line]
            index += 1
            while index < len(lines):
                candidate = lines[index].strip()
                if not candidate:
                    break
                if candidate.startswith(("#", "$$", "![")):
                    break
                if is_table_start(lines, index):
                    break
                if re.match(r"^([-*]|\d+\.)\s+", candidate):
                    break
                paragraph_lines.append(candidate)
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.widow_control = True
            add_inline(paragraph, " ".join(paragraph_lines))

    properties = document.core_properties
    properties.title = title
    properties.subject = "双光电多目标轨迹关联独立试验"
    properties.author = "MSM项目组"
    properties.keywords = (
        f"AirSim, 双光电, {target_count}目标, 轨迹关联, 匈牙利算法"
    )
    properties.comments = "由中文Markdown实验报告生成"
    set_update_fields(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    metrics = validate_document(
        output,
        expected_figures=figure_count,
        expected_tables=table_count,
    )
    metrics["equations"] = equation_count
    return metrics


def validate_document(
    path: Path, *, expected_figures: int, expected_tables: int
) -> dict[str, int]:
    reopened = Document(path)
    text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    if "待办事项" not in text or "取消对检测框面积的强依赖" not in text:
        raise RuntimeError("Word report is missing the TODO section")
    if len(reopened.tables) != expected_tables:
        raise RuntimeError(f"unexpected table count: {len(reopened.tables)}")
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        if not required.issubset(names):
            raise RuntimeError("Word package is incomplete")
        media = [name for name in names if name.startswith("word/media/")]
    if len(media) < expected_figures:
        raise RuntimeError(
            f"embedded image count {len(media)} is below figure count {expected_figures}"
        )
    return {
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "figures": expected_figures,
        "media": len(media),
        "bytes": path.stat().st_size,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def main() -> None:
    arguments = parse_args()
    source = arguments.source.resolve()
    output = (
        arguments.output.resolve()
        if arguments.output is not None
        else source.with_suffix(".docx")
    )
    metrics = build_document(source, output)
    print(f"output={output}")
    print(
        "paragraphs={paragraphs} tables={tables} figures={figures} "
        "equations={equations} media={media} bytes={bytes}".format(**metrics)
    )


if __name__ == "__main__":
    main()
