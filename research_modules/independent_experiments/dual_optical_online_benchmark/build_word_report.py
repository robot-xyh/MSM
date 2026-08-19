#!/usr/bin/env python3
"""Build a dual-optical benchmark Markdown report as a formatted Word document."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path
from zipfile import ZipFile

from matplotlib.font_manager import FontProperties
from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from dual_optical_40target import build_word_report as base


ROOT = Path(__file__).resolve().parent
DEFAULT_SOURCE = ROOT / "DUAL_OPTICAL_SCALE_FUNNEL_DETAILED_REPORT_CN.md"
CHINESE_PLOT_FONT = Path(
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Medium.ttc"
)


def add_cover(document: Document, title: str) -> None:
    is_s180 = "S180" in title
    subtitle_text = (
        "AirSim双光电S180专项试验"
        if is_s180
        else "AirSim双光电独立试验"
    )
    scenario_text = (
        "双固定光电节点  |  20、40、60目标  |  1秒单程180度扫描"
        if is_s180
        else "双固定光电节点  |  20至100目标  |  连续周扫轨迹配准"
    )
    date_text = "2026年8月17日" if is_s180 else "2026年8月"

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(120)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.first_line_indent = Cm(0)
    run = title_paragraph.add_run(title)
    base.set_run_font(
        run,
        size=25,
        bold=True,
        color=base.BLUE,
        east_asia=base.HEADING_FONT,
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Cm(0)
    subtitle.paragraph_format.space_before = Pt(20)
    run = subtitle.add_run(subtitle_text)
    base.set_run_font(
        run,
        size=16,
        bold=True,
        color=base.TEAL,
        east_asia=base.HEADING_FONT,
    )

    scenario = document.add_paragraph()
    scenario.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scenario.paragraph_format.first_line_indent = Cm(0)
    scenario.paragraph_format.space_before = Pt(42)
    run = scenario.add_run(scenario_text)
    base.set_run_font(
        run,
        size=11.5,
        color=base.INK,
        east_asia=base.HEADING_FONT,
    )

    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.paragraph_format.first_line_indent = Cm(0)
    date_paragraph.paragraph_format.space_before = Pt(175)
    run = date_paragraph.add_run(date_text)
    base.set_run_font(
        run,
        size=11,
        color=base.MUTED,
        east_asia=base.HEADING_FONT,
    )


def render_flowchart(path: Path) -> None:
    dot = r'''
digraph flow {
  graph [rankdir=TB, bgcolor="white", nodesep=0.40, ranksep=0.34,
         pad=0.12, fontname="Noto Sans CJK SC"];
  node [shape=box, style="rounded,filled", fillcolor="#F4F7FA",
        color="#356A8A", penwidth=1.2, fontname="Noto Sans CJK SC",
        fontsize=11, margin="0.14,0.09"];
  edge [color="#5F6B78", penwidth=1.15, arrowsize=0.7,
        fontname="Noto Sans CJK SC", fontsize=9];

  A [label="双站检测框"];
  B [label="像素转单位视线"];
  C [label="单圈重复检测合并"];
  D [label="单站连续跟踪"];
  E [label="跨站候选筛选"];
  F1 [label="增强几何", fillcolor="#EAF2F8"];
  F2 [label="轻量几何", fillcolor="#FCEFE7"];
  F3 [label="图神经网络", fillcolor="#EAF5EE"];
  G [label="一一关系选择"];
  H [label="连续多圈确认"];
  I [label="计算是否不超过1秒", shape=diamond, fillcolor="#FFF8E1"];
  J [label="输出匿名对应关系", fillcolor="#EAF5EE"];
  K [label="本圈不输出", fillcolor="#FBE9E7"];
  L [label="试验结束后统计结果"];

  A -> B -> C -> D -> E;
  {rank=same; F1; F2; F3;}
  E -> F1; E -> F2; E -> F3;
  F1 -> G; F2 -> G; F3 -> G;
  G -> H -> I;
  {rank=same; J; K;}
  I -> J [label="满足"];
  I -> K [label="超时"];
  J -> L;
}
'''
    subprocess.run(
        ["dot", "-Tpng", "-Gdpi=180", "-o", str(path)],
        input=dot,
        text=True,
        check=True,
    )


def add_flowchart(document: Document, image_path: Path) -> None:
    render_flowchart(image_path)
    with Image.open(image_path) as raster:
        ratio = raster.width / max(raster.height, 1)
    width_cm = min(16.4, 13.0 * ratio)
    height_cm = width_cm / ratio

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(
        str(image_path), width=Cm(width_cm), height=Cm(height_cm)
    )

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("算法总体流程")
    base.set_run_font(run, size=9, color=base.MUTED, east_asia=base.BODY_FONT)


def extract_fenced_block(
    lines: list[str], index: int
) -> tuple[str, list[str], int]:
    opener = lines[index].strip()
    language = opener.removeprefix("```").strip()
    index += 1
    content: list[str] = []
    while index < len(lines) and lines[index].strip() != "```":
        content.append(lines[index])
        index += 1
    if index >= len(lines):
        raise RuntimeError("unterminated fenced block")
    return language, content, index + 1


def add_code_block(document: Document, content: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("\n".join(content))
    base.set_run_font(
        run,
        size=8.5,
        color=base.INK,
        east_asia=base.BODY_FONT,
        latin=base.MONO_FONT,
    )


def normalize_mathtext_equation(tex: str) -> str:
    """Convert common LaTeX environments to Matplotlib mathtext syntax."""

    normalized = re.sub(r"\\mathcal\s+([A-Za-z0-9])", r"\\mathcal{\1}", tex)

    def replace_matrix(match: re.Match[str]) -> str:
        body = match.group(1).strip().replace("&", r"\quad")
        return r"\left[\substack{" + body + r"}\right]"

    normalized = re.sub(
        r"\\begin\{bmatrix\}(.*?)\\end\{bmatrix\}",
        replace_matrix,
        normalized,
        flags=re.DOTALL,
    )

    def replace_aligned(match: re.Match[str]) -> str:
        body = match.group(1).strip().replace("&", "")
        return r"\substack{" + body + "}"

    return re.sub(
        r"\\begin\{aligned\}(.*?)\\end\{aligned\}",
        replace_aligned,
        normalized,
        flags=re.DOTALL,
    )


def render_chinese_cases_equation(tex: str, path: Path) -> None:
    """Render a two-row cases expression while retaining its Chinese labels."""

    body_match = re.search(
        r"(.*?)\\begin\{cases\}(.*?)\\end\{cases\}", tex, flags=re.DOTALL
    )
    if body_match is None:
        raise RuntimeError("invalid cases equation")
    left = body_match.group(1).strip().replace("y_{ij}", "y(i,j)")
    rows = re.findall(r"([^,&]+),?&\\text\{([^}]*)\}", body_match.group(2))
    if len(rows) != 2:
        raise RuntimeError("only two-row Chinese cases equations are supported")

    first_value, first_label = rows[0]
    second_value, second_label = rows[1]
    display = (
        f"{left}  {{  {first_value.strip()}，{first_label}\n"
        f"          {second_value.strip()}，{second_label}"
    )
    font = FontProperties(fname=str(CHINESE_PLOT_FONT), size=16)
    figure = base.plt.figure(figsize=(12.0, 1.25), facecolor="white")
    figure.text(
        0.5,
        0.5,
        display,
        ha="center",
        va="center",
        fontproperties=font,
        color="#202833",
        linespacing=1.5,
    )
    figure.savefig(
        path,
        dpi=220,
        bbox_inches="tight",
        pad_inches=0.08,
        facecolor="white",
    )
    base.plt.close(figure)


def add_report_equation(
    document: Document, tex: str, image_path: Path
) -> None:
    if r"\begin{cases}" not in tex:
        base.add_equation(document, normalize_mathtext_equation(tex), image_path)
        return

    render_chinese_cases_equation(tex, image_path)
    with Image.open(image_path) as raster:
        natural_width_cm = raster.width / 220.0 * 2.54
        ratio = raster.width / max(raster.height, 1)
    width_cm = min(15.5, max(6.0, natural_width_cm))
    height_cm = width_cm / ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run().add_picture(
        str(image_path), width=Cm(width_cm), height=Cm(height_cm)
    )


def validate_document(
    path: Path,
    *,
    expected_figures: int,
    expected_tables: int,
    expected_equations: int,
    expected_flowcharts: int,
    required_text: tuple[str, ...],
) -> dict[str, int]:
    reopened = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )
    text = f"{paragraph_text}\n{table_text}"
    missing = [item for item in required_text if item not in text]
    if missing:
        raise RuntimeError(f"Word report is missing required text: {missing}")
    if len(reopened.tables) != expected_tables:
        raise RuntimeError(f"unexpected table count: {len(reopened.tables)}")

    with ZipFile(path) as archive:
        names = set(archive.namelist())
        required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
        if not required.issubset(names):
            raise RuntimeError("Word package is incomplete")
        media = [name for name in names if name.startswith("word/media/")]

    expected_media = expected_figures + expected_equations + expected_flowcharts
    if len(media) < expected_media:
        raise RuntimeError(
            f"embedded image count {len(media)} is below expected {expected_media}"
        )
    return {
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "figures": expected_figures,
        "equations": expected_equations,
        "flowcharts": expected_flowcharts,
        "media": len(media),
        "bytes": path.stat().st_size,
    }


def build_document(source: Path, output: Path) -> dict[str, int]:
    lines = source.read_text(encoding="utf-8").splitlines()
    title_line = next((line for line in lines if line.startswith("# ")), None)
    if title_line is None:
        raise RuntimeError("report title not found")
    title = title_line[2:].strip()
    if source.resolve() == DEFAULT_SOURCE.resolve():
        required_text = (
            "1. 结论",
            "6. 试验结果",
            "6.5 100目标",
            "图神经网络（离线测试）",
            "10. 验证与证据",
        )
    else:
        required_text = (title,) + tuple(
            line[3:].strip() for line in lines if line.startswith("## ")
        )

    document = Document()
    base.apply_page_layout(document.sections[0])
    base.configure_styles(document)
    add_cover(document, title)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    base.apply_page_layout(content_section)
    base.set_page_number_start(content_section, 1)
    base.configure_header_footer(content_section, title)
    base.add_table_of_contents(document)
    document.add_page_break()

    figure_count = 0
    equation_count = 0
    flowchart_count = 0
    table_count = 0
    index = 0
    with tempfile.TemporaryDirectory(prefix="dual-optical-scale-word-") as temporary:
        temporary_dir = Path(temporary)
        while index < len(lines):
            line = lines[index].strip()
            if not line:
                index += 1
                continue
            if line.startswith("# "):
                index += 1
                continue
            if line.startswith("```"):
                language, content, index = extract_fenced_block(lines, index)
                if language == "mermaid":
                    flowchart_count += 1
                    add_flowchart(
                        document,
                        temporary_dir / f"flowchart_{flowchart_count:02d}.png",
                    )
                else:
                    add_code_block(document, content)
                continue
            if line.startswith("$$"):
                equation, index = base.extract_display_equation(lines, index)
                equation_count += 1
                add_report_equation(
                    document,
                    equation,
                    temporary_dir / f"equation_{equation_count:02d}.png",
                )
                continue
            image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
            if image_match:
                figure_count += 1
                base.add_image(
                    document,
                    source,
                    image_match.group(1),
                    image_match.group(2),
                    figure_count,
                )
                index += 1
                continue
            if base.is_table_start(lines, index):
                rows, index = base.parse_table(lines, index)
                base.add_table(document, rows)
                table_count += 1
                continue
            if line.startswith("#"):
                level = len(line) - len(line.lstrip("#"))
                base.add_heading(document, level, line[level:].strip())
                index += 1
                continue
            list_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", line)
            if list_match:
                base.add_list_item(document, list_match.group(1), list_match.group(2))
                index += 1
                continue

            paragraph_lines = [line]
            index += 1
            while index < len(lines):
                candidate = lines[index].strip()
                if not candidate:
                    break
                if candidate.startswith(("#", "$$", "![", "```")):
                    break
                if base.is_table_start(lines, index):
                    break
                if re.match(r"^([-*]|\d+\.)\s+", candidate):
                    break
                paragraph_lines.append(candidate)
                index += 1
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.widow_control = True
            base.add_inline(paragraph, " ".join(paragraph_lines))

    properties = document.core_properties
    properties.title = title
    is_s180 = "S180" in title
    properties.subject = (
        "双光电S180扫描多目标轨迹配准试验"
        if is_s180
        else "双光电连续周扫多目标轨迹配准试验"
    )
    properties.author = "MSM项目组"
    properties.keywords = (
        "AirSim, 双光电, S180扫描, 轨迹配准, 图神经网络"
        if is_s180
        else "AirSim, 双光电, 连续周扫, 轨迹配准, 图神经网络"
    )
    properties.comments = "由中文Markdown试验报告生成"
    base.set_update_fields(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    return validate_document(
        output,
        expected_figures=figure_count,
        expected_tables=table_count,
        expected_equations=equation_count,
        expected_flowcharts=flowchart_count,
        required_text=required_text,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def main() -> None:
    arguments = parse_args()
    source = arguments.source.resolve()
    output = (
        arguments.output.resolve()
        if arguments.output is not None
        else source.with_suffix(".docx")
    )
    metrics = build_document(source, output)
    print(f"output={output}")
    print(" ".join(f"{key}={value}" for key, value in metrics.items()))


if __name__ == "__main__":
    main()
