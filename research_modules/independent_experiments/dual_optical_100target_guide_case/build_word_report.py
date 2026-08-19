#!/usr/bin/env python3
"""Build an A4 Word document from the generated Chinese Markdown report."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import re
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_SOURCE = Path(__file__).resolve().parent / "outputs" / "airsim_seed_20260812_guide_run01" / "DUAL_OPTICAL_100TARGET_GUIDE_AIRSIM_REPORT_CN.md"


def _set_font(run, name: str = "SimSun", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)
    for style_name, size, bold in (
        ("Normal", 10.5, False),
        ("Title", 22.0, True),
        ("Heading 1", 16.0, True),
        ("Heading 2", 14.0, True),
        ("Heading 3", 12.0, True),
    ):
        style = document.styles[style_name]
        style.font.name = "SimSun"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\$[^$]+\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_font(run, name="Consolas", size=9.0)
        elif part.startswith("$") and part.endswith("$"):
            run = paragraph.add_run(part[1:-1])
            _set_font(run, name="Cambria Math", size=10.0)
        else:
            run = paragraph.add_run(part)
            _set_font(run)


def _parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build_document(source: Path, output: Path) -> dict[str, int]:
    lines = source.read_text(encoding="utf-8").splitlines()
    title = next(line[2:] for line in lines if line.startswith("# "))
    document = Document()
    _configure(document)
    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(150)
    run = cover.add_run(title)
    _set_font(run, size=22.0, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("独立AirSim案例")
    _set_font(run, size=14.0)
    document.add_page_break()

    figure_count = 0
    table_count = 0
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line or line.startswith("# "):
            index += 1
            continue
        image = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", line)
        if image:
            image_path = (source.parent / image.group(2)).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(image_path)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
            caption = document.add_paragraph(f"图{figure_count + 1} {image.group(1)}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption.runs:
                _set_font(run, size=9.0)
            figure_count += 1
            index += 1
            continue
        if line.startswith("$$"):
            equation_lines = [line.removeprefix("$$")]
            index += 1
            while index < len(lines) and not lines[index].strip().endswith("$$"):
                equation_lines.append(lines[index].strip())
                index += 1
            if index < len(lines):
                equation_lines.append(lines[index].strip().removesuffix("$$"))
                index += 1
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(" ".join(item for item in equation_lines if item))
            _set_font(run, name="Cambria Math", size=10.0)
            continue
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].strip().startswith("|"):
            rows, index = _parse_table(lines, index)
            if rows:
                table = document.add_table(rows=len(rows), cols=max(map(len, rows)))
                table.style = "Table Grid"
                for row_index, values in enumerate(rows):
                    for column_index, value in enumerate(values):
                        cell = table.cell(row_index, column_index)
                        cell.text = value
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                _set_font(run, size=8.5, bold=row_index == 0)
                table_count += 1
            continue
        if line.startswith("#"):
            level = min(len(line) - len(line.lstrip("#")), 3)
            document.add_heading(line[level:].strip(), level=level)
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, re.sub(r"^[-*]\s+", "", line))
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate or candidate.startswith(("#", "![", "|", "$$", "- ", "* ")):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.35
        _add_inline(paragraph, " ".join(paragraph_lines))

    properties = document.core_properties
    properties.title = title
    properties.subject = "双站光电100目标轨迹配准独立试验"
    properties.author = "MSM项目组"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    with ZipFile(output) as archive:
        media_count = len(
            [name for name in archive.namelist() if name.startswith("word/media/")]
        )
    if media_count < figure_count:
        raise RuntimeError("Word images are incomplete")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": table_count,
        "figures": figure_count,
        "media": media_count,
        "bytes": output.stat().st_size,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    source = args.source.resolve()
    output = args.output.resolve() if args.output else source.with_suffix(".docx")
    metrics = build_document(source, output)
    print(f"output={output}")
    print(json.dumps(metrics, ensure_ascii=False))


if __name__ == "__main__":
    main()
