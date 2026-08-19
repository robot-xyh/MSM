"""Convert the experiment Markdown report to an A4 Word document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
HEADING_ONE_RE = re.compile(r"^#(?!#)\s+(.+?)\s*$")
DEFAULT_DOCUMENT_TITLE = "文档报告"


def _document_title(lines: list[str]) -> str:
    fence: str | None = None
    for raw_line in lines:
        line = raw_line.strip()
        if line.startswith(("```", "~~~")):
            marker = line[:3]
            if fence is None:
                fence = marker
            elif fence == marker:
                fence = None
            continue
        if fence is not None:
            continue
        match = HEADING_ONE_RE.match(line)
        if not match:
            continue
        title = re.sub(r"\s+#+\s*$", "", match.group(1)).strip()
        if title:
            return title
    return DEFAULT_DOCUMENT_TITLE


def _font(run, *, size: float = 11.0, bold: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _add_inline(paragraph, text: str, *, size: float = 11.0) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _font(run, size=size)
        else:
            run = paragraph.add_run(part)
            _font(run, size=size)


def build_word_report(source: str | Path, output: str | Path | None = None) -> Path:
    source = Path(source).resolve()
    output = Path(output).resolve() if output else source.with_suffix(".docx")
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)

    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        image_match = IMAGE_RE.match(line)
        if image_match:
            image_path = (source.parent / image_match.group(2)).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(f"report image not found: {image_path}")
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(image_path), width=Cm(15.2))
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(image_match.group(1))
            _font(run, size=9.5)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and set(lines[index + 1].replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            rows = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            if len(rows) >= 2:
                rows.pop(1)
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, values in enumerate(rows):
                for column_index, value in enumerate(values):
                    cell = table.cell(row_index, column_index)
                    if row_index == 0:
                        _cell_shading(cell, "D9E7F0")
                    paragraph = cell.paragraphs[0]
                    _add_inline(paragraph, value, size=9.5)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            paragraph = document.add_paragraph()
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(30)
                paragraph.paragraph_format.space_after = Pt(20)
                run = paragraph.add_run(text)
                _font(run, size=20, bold=True)
            else:
                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(5)
                run = paragraph.add_run(text)
                _font(run, size=15 if level == 2 else 12, bold=True)
            index += 1
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, line[2:])
            index += 1
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        _add_inline(paragraph, line)
        index += 1

    core = document.core_properties
    core.title = _document_title(lines)
    core.subject = "独立AirSim算法验证"
    document.save(output)
    return output


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True)
    parser.add_argument("--output")
    args = parser.parse_args()
    print(build_word_report(args.source, args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
