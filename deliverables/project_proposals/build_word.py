#!/usr/bin/env python3
"""Build the three Chinese project proposals as Word documents."""

from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCES = (
    HERE / "无人机协同探测项目建议书.md",
    HERE / "无人机协同拦截项目建议书.md",
    HERE / "大模型辅助评估项目建议书.md",
)

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
MONO_FONT = "Consolas"

BLUE = "1F4E78"
TEAL = "176B73"
INK = "202833"
MUTED = "5F6B78"
WHITE = "FFFFFF"
LIGHT_BLUE = "EAF2F8"
LIGHT_AMBER = "FFF5D9"
ROW_FILL = "F5F7FA"


MATH_REPLACEMENTS = {
    r"\sigma": "σ",
    r"\theta": "θ",
    r"\lambda": "λ",
    r"\Delta": "Δ",
    r"\pi": "π",
    r"\tau": "τ",
    r"\approx": "≈",
    r"\propto": "∝",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\le": "≤",
    r"\ge": "≥",
    r"\in": "∈",
    r"\sum": "Σ",
    r"\min": "min",
    r"\max": "max",
    r"\cdot": "·",
    r"\rightarrow": "→",
    r"\quad": "  ",
    r"\;": " ",
    r"\,": " ",
}

INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|\\\(.+?\\\)|\$[^$]+?\$)"
)
TABLE_SEPARATOR_RE = re.compile(r"^:?-{3,}:?$")


def set_run_font(run, *, size: float = 11, bold: bool | None = None,
                 color: str = INK, east_asia: str = BODY_FONT,
                 latin: str = LATIN_FONT) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 55, start: int = 65,
                     bottom: int = 55, end: int = 65) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def math_to_plain(text: str) -> str:
    result = text.strip().strip("$")
    if result.startswith(r"\(") and result.endswith(r"\)"):
        result = result[2:-2]
    for old, new in MATH_REPLACEMENTS.items():
        result = result.replace(old, new)
    result = result.replace(r"\lvert", "|").replace(r"\rvert", "|")
    result = result.replace(r"\times", "×").replace(r"\top", "ᵀ")
    result = re.sub(r"\\text\{([^{}]*)\}", r"\1", result)
    result = result.replace("{", "").replace("}", "")
    return result


def add_hyperlink(paragraph, label: str, url: str, size: float) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend((fonts, color, underline, sz))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((r_pr, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, size: float = 11,
               color: str = INK, bold: bool = False) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=size, bold=bold, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(size - 0.7, 7), color=TEAL, east_asia=BODY_FONT, latin=MONO_FONT)
        elif token.startswith("["):
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2), size)
        else:
            run = paragraph.add_run(math_to_plain(token))
            set_run_font(run, size=size, color=color, east_asia=BODY_FONT, latin="Cambria Math")
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, bold=bold, color=color)


def apply_page_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.15)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, size=8.5, color=MUTED, east_asia=HEADING_FONT)


def set_page_number_start(section, value: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(value))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    specs = {
        "Title": (26, BLUE, 0, 14),
        "Heading 1": (16, BLUE, 14, 8),
        "Heading 2": (14, TEAL, 11, 6),
        "Heading 3": (12, INK, 9, 5),
        "Heading 4": (11, INK, 7, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True


def add_header_footer(section, title: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(header, title, size=8.3, color=MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "项目建议书  ·  ", size=8.3, color=MUTED)
    add_page_field(footer)


def add_cover(doc: Document, title: str, cutoff: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kind = doc.add_paragraph()
    kind.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(kind, "项目建议书", size=18, color=TEAL, bold=True)
    doc.add_paragraph()
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(heading, title.removesuffix("项目建议书"), size=29, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(subtitle, "研究论证与技术方案", size=15, color=BLUE, bold=True)

    rule = doc.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    rule.autofit = False
    cell = rule.cell(0, 0)
    cell.width = Cm(11.5)
    set_cell_shading(cell, LIGHT_BLUE)
    cell_p = cell.paragraphs[0]
    cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_p.paragraph_format.first_line_indent = Cm(0)
    add_inline(cell_p, "立项依据  ·  研究目标  ·  研究任务  ·  技术方案", size=11, color=BLUE, bold=True)

    for _ in range(8):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(meta, "MSM 项目组", size=12, color=INK)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(date, cutoff or "证据截止日期：2026年7月19日", size=10.5, color=MUTED)
    boundary = doc.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(boundary, "科研论证、仿真与受控试验用途", size=9.5, color=TEAL)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.page_break_before = False
    add_inline(p, "目录", size=16, color=BLUE, bold=True)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = 'TOC \\o "1-4" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在 Word 中更新域以生成目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    set_run_font(run, size=10, color=MUTED)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    separator = lines[index + 1].strip()
    if not (current.startswith("|") and current.endswith("|") and separator.startswith("|")):
        return False
    cells = split_table_row(separator)
    return bool(cells) and all(TABLE_SEPARATOR_RE.fullmatch(cell) for cell in cells)


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    rows = [split_table_row(lines[index])]
    index += 2
    while index < len(lines):
        line = lines[index].strip()
        if not (line.startswith("|") and line.endswith("|")):
            break
        rows.append(split_table_row(line))
        index += 1
    return rows, index


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    font_size = 8.8 if columns <= 3 else 8.1 if columns <= 5 else 7.2

    for row_index, row in enumerate(rows):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        no_split = OxmlElement("w:cantSplit")
        tr_pr.append(no_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.paragraph_format.space_after = Pt(0)
            content = row[col_index] if col_index < len(row) else ""
            if row_index == 0:
                set_cell_shading(cell, BLUE)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, content, size=font_size, color=WHITE, bold=True)
            else:
                if row_index % 2 == 0:
                    set_cell_shading(cell, ROW_FILL)
                add_inline(paragraph, content, size=font_size, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_display_equation(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(math_to_plain(" ".join(lines)))
    set_run_font(run, size=11, color=INK, east_asia=BODY_FONT, latin="Cambria Math")


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)


def add_image(doc: Document, source: Path, alt: str, relative: str, figure_number: int) -> None:
    image_path = (source.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"Missing image referenced by {source.name}: {image_path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.1))

    caption = alt if re.match(r"^图\s*\d", alt) else f"图4-{figure_number}  {alt}"
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.first_line_indent = Cm(0)
    caption_p.paragraph_format.space_after = Pt(6)
    add_inline(caption_p, caption, size=9, color=MUTED)


def add_quote(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.25)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_AMBER)
    p_pr.append(shading)
    add_inline(paragraph, text, size=9.5, color=INK)


def build_document(source: Path) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    title_line = next(line for line in lines if line.startswith("# "))
    title = title_line[2:].strip()
    cutoff_match = re.search(r"证据截止日期[：:]?\*?\*?[：:]?\s*([^\n。]+)", "\n".join(lines[:15]))
    cutoff = f"证据截止日期：{cutoff_match.group(1).strip()}" if cutoff_match else "证据截止日期：2026年7月19日"

    doc = Document()
    apply_page_layout(doc.sections[0])
    configure_styles(doc)
    add_cover(doc, title, cutoff)

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(content_section)
    set_page_number_start(content_section, 1)
    add_header_footer(content_section, title)
    add_toc(doc)
    doc.add_page_break()

    figure_number = 0
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line == "---":
            add_horizontal_rule(doc)
            index += 1
            continue
        if line.startswith(r"\["):
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith(r"\]"):
                equation_lines.append(lines[index].strip())
                index += 1
            add_display_equation(doc, equation_lines)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if image_match:
            figure_number += 1
            add_image(doc, source, image_match.group(1), image_match.group(2), figure_number)
            index += 1
            continue
        if is_table_start(lines, index):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            heading = line[level:].strip()
            style_map = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3", 5: "Heading 4"}
            paragraph = doc.add_paragraph(style=style_map.get(level, "Heading 4"))
            paragraph.paragraph_format.first_line_indent = Cm(0)
            sizes = {2: 16, 3: 14, 4: 12, 5: 11}
            colors = {2: BLUE, 3: TEAL, 4: INK, 5: INK}
            add_inline(paragraph, heading, size=sizes.get(level, 11), color=colors.get(level, INK), bold=True)
            index += 1
            continue
        if line.startswith(">"):
            add_quote(doc, line.lstrip("> "))
            index += 1
            continue
        list_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", line)
        if list_match:
            marker, content = list_match.groups()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.78)
            paragraph.paragraph_format.first_line_indent = Cm(-0.55)
            paragraph.paragraph_format.space_after = Pt(2)
            prefix = "•" if marker in {"-", "*"} else marker
            prefix_run = paragraph.add_run(f"{prefix}  ")
            set_run_font(prefix_run, size=10.5, bold=True, color=BLUE, east_asia=HEADING_FONT)
            add_inline(paragraph, content, size=10.5, color=INK)
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate or candidate == "---" or candidate.startswith(("#", ">", r"\[", "![")):
                break
            if is_table_start(lines, lookahead):
                break
            if re.match(r"^([-*]|\d+\.)\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, " ".join(parts), size=11, color=INK)
        index = lookahead

    properties = doc.core_properties
    properties.title = title
    properties.subject = "项目建议书"
    properties.author = "MSM 项目组"
    properties.keywords = "无人机, 协同探测, 协同拦截, 大模型辅助评估"
    properties.comments = "由项目建议书Markdown源文档生成"

    destination = source.with_suffix(".docx")
    doc.save(destination)
    return destination


def validate_document(path: Path) -> dict[str, int]:
    reopened = Document(path)
    with ZipFile(path) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
    if not images:
        raise RuntimeError(f"No embedded image found in {path}")
    if len(reopened.paragraphs) < 100:
        raise RuntimeError(f"Unexpectedly short document: {path}")
    return {
        "paragraphs": len(reopened.paragraphs),
        "tables": len(reopened.tables),
        "images": len(images),
        "bytes": path.stat().st_size,
    }


def main() -> None:
    for source in SOURCES:
        if not source.exists():
            raise FileNotFoundError(source)
        output = build_document(source)
        metrics = validate_document(output)
        print(
            f"{output.name}: paragraphs={metrics['paragraphs']}, "
            f"tables={metrics['tables']}, images={metrics['images']}, "
            f"bytes={metrics['bytes']}"
        )


if __name__ == "__main__":
    main()
