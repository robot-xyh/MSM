#!/usr/bin/env python3
"""Build the 200-v-200 full-flow Chinese project report as a Word document."""

from __future__ import annotations

import re
import warnings
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from PIL import Image

warnings.filterwarnings(
    "ignore",
    message="Unable to import Axes3D.*",
    category=UserWarning,
)

from matplotlib.mathtext import math_to_image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "200V200_三维全流程技术报告.md"
OUTPUT = HERE / "200V200_三维全流程技术报告.docx"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
MONO_FONT = "Consolas"
MATH_FONT = "Cambria Math"

BLUE = "17365D"
TEAL = "365F78"
INK = "1F1F1F"
MUTED = "66727D"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F3F2"
ROW_FILL = "F6F8FA"
WHITE = "FFFFFF"

INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\\\(.+?\\\)|\$[^$]+?\$)"
)
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")
TABLE_SEPARATOR_RE = re.compile(r"^:?-{3,}:?$")

MATH_REPLACEMENTS = {
    r"\mathbf": "",
    r"\boldsymbol": "",
    r"\mathsf": "",
    r"\mathrm": "",
    r"\operatorname": "",
    r"\beginbmatrix": "[",
    r"\endbmatrix": "]",
    r"\begin{bmatrix}": "[",
    r"\end{bmatrix}": "]",
    r"\begin{aligned}": "",
    r"\end{aligned}": "",
    r"\frac": "frac",
    r"\sqrt": "sqrt",
    r"\left": "",
    r"\right": "",
    r"\qquad": "    ",
    r"\quad": "  ",
    r"\;": " ",
    r"\,": " ",
    r"\times": "×",
    r"\cdot": "·",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\le": "≤",
    r"\ge": "≥",
    r"\approx": "≈",
    r"\in": "∈",
    r"\sum": "Σ",
    r"\min": "min",
    r"\max": "max",
    r"\tanh": "tanh",
    r"\operatorname": "",
    r"\mathcal": "",
    r"\Delta": "Δ",
    r"\sigma": "σ",
    r"\theta": "θ",
    r"\tau": "τ",
    r"\omega": "ω",
    r"\boldsymbol": "",
    r"\mathsf T": "ᵀ",
    r"^\mathsf{T}": "ᵀ",
    r"^\mathsf T": "ᵀ",
    r"^{\mathsf T}": "ᵀ",
}


def set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool | None = None,
    color: str = INK,
    east_asia: str = BODY_FONT,
    latin: str = LATIN_FONT,
    italic: bool | None = None,
) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, top: int = 55, start: int = 70, bottom: int = 55, end: int = 70
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    elif not value and keep is not None:
        p_pr.remove(keep)


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    set_run_font(run, size=8.5, color=MUTED)


def add_toc(paragraph) -> None:
    entries = (
        ("报告要点", 3),
        ("第一章  项目背景与建设目标", 4),
        ("第二章  总体方案", 6),
        ("第三章  分系统方案", 10),
        ("第四章  当前仿真进展", 16),
        ("第五章  学习算法增强", 26),
        ("第六章  主要问题与后续安排", 31),
        ("技术附件", 34),
    )
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.8
    for index, (title, page) in enumerate(entries):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        dots = "·" * max(4, 26 - len(title))
        run = paragraph.add_run(f"{title}  {dots}  {page}")
        set_run_font(run, size=11, color=INK)


def configure_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def apply_page_layout(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.15)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.75)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.widow_control = True

    specs = {
        "Title": (24, BLUE, 0, 14),
        "Heading 1": (16, BLUE, 14, 8),
        "Heading 2": (13.5, TEAL, 11, 6),
        "Heading 3": (12, INK, 9, 5),
        "Heading 4": (11, INK, 7, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
        if name == "Heading 1":
            style.paragraph_format.page_break_before = True

    caption = doc.styles["Caption"]
    caption.font.name = LATIN_FONT
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False


def add_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("大规模无人机群协同探测与拦截仿真方案")
    set_run_font(run, size=8.5, color=MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("项目技术报告  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(footer)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    category = doc.add_paragraph()
    category.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = category.add_run("项目技术报告")
    set_run_font(run, size=15, bold=True, color=TEAL, east_asia=HEADING_FONT)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("大规模无人机群协同探测\n与拦截仿真方案")
    set_run_font(run, size=26, bold=True, color=BLUE, east_asia=HEADING_FONT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run("200目标与200拦截资源三维全流程验证")
    set_run_font(run, size=14, color=INK, east_asia=HEADING_FONT)

    for _ in range(6):
        doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(12)
    run = line.add_run("资料截止：2026年7月31日")
    set_run_font(run, size=11.5, color=MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run("资料范围：软件测试、三维质点仿真与AirSim受控试验")
    set_run_font(run, size=10.5, color=MUTED)


def math_to_plain(text: str) -> str:
    value = text.strip()
    for old, new in MATH_REPLACEMENTS.items():
        value = value.replace(old, new)
    value = re.sub(r"\\text\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\([A-Za-z]+)", r"\1", value)
    value = value.replace("{", "").replace("}", "")
    value = value.replace("&", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def add_inline(paragraph, text: str, *, size: float = 11, color: str = INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                size=max(size - 0.6, 8),
                color=TEAL,
                east_asia=BODY_FONT,
                latin=MONO_FONT,
            )
        else:
            value = token
            if value.startswith(r"\(") and value.endswith(r"\)"):
                value = value[2:-2]
            value = value.strip("$")
            run = paragraph.add_run(math_to_plain(value))
            set_run_font(
                run, size=size, color=color, east_asia=BODY_FONT, latin=MATH_FONT
            )
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=size, color=color)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(paragraph, text)


def add_list_item(doc: Document, text: str, marker: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.5)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(f"{marker} ")
    set_run_font(run, size=11, bold=True, color=TEAL)
    add_inline(paragraph, text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F6F8")
    set_cell_margins(cell, top=90, start=130, bottom=90, end=130)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        if index:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        set_run_font(
            run,
            size=8.7,
            color=INK,
            east_asia=BODY_FONT,
            latin=MONO_FONT,
        )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_equation(doc: Document, lines: list[str]) -> None:
    source = " ".join(line.strip() for line in lines if line.strip())
    math_source = source.replace(r"\mathsf T", r"\mathsf{T}").replace(
        r"\|", r"\Vert "
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    try:
        stream = BytesIO()
        math_to_image(
            f"${math_source}$",
            stream,
            dpi=300,
            format="png",
            color=f"#{INK}",
        )
        stream.seek(0)
        with Image.open(stream) as image:
            width_px, height_px = image.size
        stream.seek(0)
        ratio = width_px / max(height_px, 1)
        width_cm = min(15.5, max(4.0, width_px / 300 * 2.54))
        height_cm = width_cm / ratio
        if height_cm > 2.2:
            height_cm = 2.2
            width_cm = height_cm * ratio
        run = paragraph.add_run()
        run.add_picture(stream, width=Cm(width_cm), height=Cm(height_cm))
    except Exception:
        run = paragraph.add_run(math_to_plain(source))
        set_run_font(
            run,
            size=10.5,
            color=INK,
            east_asia=BODY_FONT,
            latin=MATH_FONT,
        )


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    font_size = 8.3 if columns >= 5 else 9.0
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        set_cant_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index in range(columns):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, ROW_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if row_index == 0
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            value = values[column_index] if column_index < len(values) else ""
            add_inline(
                paragraph,
                value,
                size=font_size,
                color=WHITE if row_index == 0 else INK,
            )
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_picture(doc: Document, alt: str, target: str) -> None:
    path = (SOURCE.parent / target).resolve()
    if not path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"[图片缺失：{alt}]")
        set_run_font(run, size=9, color="B00020")
        return

    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / max(height_px, 1)
    max_width_cm = 16.2
    max_height_cm = 13.2
    width_cm = max_width_cm
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    set_keep_with_next(paragraph)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    value = text.strip()
    if value.startswith("*") and value.endswith("*"):
        value = value[1:-1]
    run = paragraph.add_run(value)
    set_run_font(run, size=9, color=MUTED, italic=False)


def add_heading(doc: Document, level: int, text: str) -> None:
    style_level = min(max(level, 1), 4)
    paragraph = doc.add_paragraph(style=f"Heading {style_level}")
    paragraph.paragraph_format.first_line_indent = Cm(0)
    add_inline(
        paragraph,
        text,
        size={1: 17, 2: 14, 3: 12, 4: 11}[style_level],
        color={1: BLUE, 2: TEAL, 3: INK, 4: INK}[style_level],
    )
    for run in paragraph.runs:
        run.bold = True
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)


def render_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0

    # The first title and version are represented on the cover.
    if lines and lines[0].startswith("# "):
        index = 1
    while index < len(lines) and (
        not lines[index].strip() or lines[index].startswith("**版本日期")
    ):
        index += 1

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if not line:
            index += 1
            continue

        if line.startswith("```"):
            block: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            add_code_block(doc, block)
            index += 1
            continue

        if line == "$$":
            equation: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                equation.append(lines[index])
                index += 1
            add_equation(doc, equation)
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            add_heading(doc, len(heading.group(1)), heading.group(2))
            index += 1
            continue

        image = IMAGE_RE.match(line)
        if image:
            add_picture(doc, image.group(1), image.group(2))
            if index + 1 < len(lines):
                next_line = lines[index + 1].strip()
                if next_line.startswith("*图") and next_line.endswith("*"):
                    add_caption(doc, next_line)
                    index += 1
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines):
            separator = parse_table_row(lines[index + 1])
            if separator and all(TABLE_SEPARATOR_RE.fullmatch(cell) for cell in separator):
                rows = [parse_table_row(line)]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    rows.append(parse_table_row(lines[index]))
                    index += 1
                add_table(doc, rows)
                continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            add_list_item(doc, numbered.group(2), f"{numbered.group(1)}.")
            index += 1
            continue

        if line.startswith("- "):
            add_list_item(doc, line[2:], "•")
            index += 1
            continue

        if line.startswith("*") and line.endswith("*"):
            add_caption(doc, line)
            index += 1
            continue

        add_body_paragraph(doc, line)
        index += 1


def validate_docx(path: Path) -> None:
    required = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
    with ZipFile(path) as archive:
        names = set(archive.namelist())
    missing = required - names
    if missing:
        raise RuntimeError(f"Invalid DOCX, missing: {sorted(missing)}")


def build() -> Path:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_update_fields(doc)
    apply_page_layout(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True

    properties = doc.core_properties
    properties.title = "大规模无人机群协同探测与拦截仿真方案"
    properties.subject = "200目标与200拦截资源三维全流程验证"
    properties.author = "MSM项目组"
    properties.comments = "报告结论限于软件测试、三维质点仿真和AirSim受控试验。"

    add_cover(doc)
    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(content_section)
    add_header_footer(content_section)

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(16)
    run = toc_title.add_run("目录")
    set_run_font(run, size=20, bold=True, color=BLUE, east_asia=HEADING_FONT)
    toc = doc.add_paragraph()
    add_toc(toc)
    doc.add_page_break()

    render_markdown(doc, markdown)
    doc.save(OUTPUT)
    validate_docx(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
