#!/usr/bin/env python3
"""Build the UAV-interceptor technology review as a Word document.

Reuses the typography of build_200v200_full_flow_report.py and adds three
things that document did not need: blockquote front matter, real hyperlinks,
and inline math disabled (the review contains "$18M" twice on one line, which
the base inline regex would otherwise swallow as a math span).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

import build_200v200_full_flow_report as base  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.opc.constants import RELATIONSHIP_TYPE as RT  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Pt  # noqa: E402
from docx.text.run import Run  # noqa: E402
from docx.image.image import Image as DocxImage  # noqa: E402
from PIL import Image  # noqa: E402

base_add_picture = base.add_picture

SOURCE = HERE / "无人机拦截无人机技术现状调研.md"
OUTPUT = HERE / "无人机拦截无人机技术现状调研.docx"
CACHE_DIR = HERE / "figures" / "uav_interceptor_technology_review" / "_docx"


def docx_can_read(path: Path) -> bool:
    try:
        DocxImage.from_file(str(path))
    except Exception:
        return False
    return True

TITLE = "无人机拦截无人机技术现状调研"
SUBTITLE = "单拦单、多拦单、多拦多与集群对抗的公开证据分级"
CUTOFF = "调研截止：2026年8月3日"
LINK = "1155CC"

# Link forms first so a URL is consumed before anything inside it can match.
# Bold before italic so "**" is never split into two "*" spans.
# Inline math is deliberately absent.
INLINE_RE = re.compile(
    r"(\[[^\]]+\]\(<?[^)\s]+>?\)"      # [text](url)
    r"|<https?://[^>\s]+>"             # <url>
    r"|\*\*.+?\*\*"                    # bold
    r"|\*[^*\n]+?\*"                   # italic, used for source titles
    r"|`.+?`)"                         # code
)
MD_LINK_RE = re.compile(r"^\[([^\]]+)\]\(<?([^)\s]+?)>?\)$")


def add_hyperlink(paragraph, url: str, text: str, size: float) -> None:
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    node = OxmlElement("w:hyperlink")
    node.set(qn("r:id"), r_id)
    run_node = OxmlElement("w:r")
    node.append(run_node)
    paragraph._p.append(node)
    run = Run(run_node, paragraph)
    run.text = text
    base.set_run_font(run, size=size, color=LINK)
    run.underline = True


def add_inline(paragraph, text: str, *, size: float = 11, color: str = base.INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            base.set_run_font(run, size=size, color=color)
        token = match.group(0)
        link = MD_LINK_RE.match(token)
        if link:
            add_hyperlink(paragraph, link.group(2), link.group(1), size)
        elif token.startswith("<http"):
            url = token[1:-1]
            add_hyperlink(paragraph, url, url, max(size - 1.0, 8))
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            base.set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            base.set_run_font(run, size=size, color=color, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            base.set_run_font(
                run,
                size=max(size - 0.6, 8),
                color=base.TEAL,
                east_asia=base.BODY_FONT,
                latin=base.MONO_FONT,
            )
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        base.set_run_font(run, size=size, color=color)


def add_picture(doc: Document, alt: str, target: str) -> None:
    """Same as the base version, but transcode images python-docx cannot read.

    Two cases show up in this figure set: a WebP downloaded from the original
    publisher, and Wikimedia thumbnails whose JPEG stream carries no JFIF or
    Exif marker. python-docx sniffs the header itself and rejects both, so probe
    first and re-encode to PNG only when the probe fails.
    """
    path = (SOURCE.parent / target).resolve()
    if path.exists() and not docx_can_read(path):
        CACHE_DIR.mkdir(parents=True, exist_ok=True)
        converted = CACHE_DIR / (path.stem + ".png")
        if not converted.exists() or converted.stat().st_mtime < path.stat().st_mtime:
            with Image.open(path) as image:
                image.convert("RGB").save(converted, "PNG")
        target = str(converted.relative_to(SOURCE.parent))
    base_add_picture(doc, alt, target)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    category = doc.add_paragraph()
    category.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = category.add_run("技术调研报告")
    base.set_run_font(run, size=15, bold=True, color=base.TEAL, east_asia=base.HEADING_FONT)

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    base.set_run_font(run, size=26, bold=True, color=base.BLUE, east_asia=base.HEADING_FONT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(18)
    run = subtitle.add_run(SUBTITLE)
    base.set_run_font(run, size=14, color=base.INK, east_asia=base.HEADING_FONT)

    for _ in range(6):
        doc.add_paragraph()

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Pt(12)
    run = line.add_run(CUTOFF)
    base.set_run_font(run, size=11.5, color=base.MUTED)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run("公开证据分级：政府通报、厂商演示、同行评议与仿真分别标注")
    base.set_run_font(run, size=10.5, color=base.MUTED)


def add_header_footer(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("技术调研报告  ·  ")
    base.set_run_font(run, size=8.5, color=base.MUTED)
    base.add_page_field(footer)


def prepare_markdown(text: str) -> str:
    """Turn the leading blockquote into a normal 调研说明 section."""
    lines = text.splitlines()
    out: list[str] = []
    quote: list[str] = []
    for line in lines:
        if line.startswith(">"):
            value = line.lstrip(">").strip()
            if value:
                quote.append(value)
            continue
        if quote:
            out.append("## 调研说明")
            out.append("")
            for item in quote:
                out.append(item)
                out.append("")
            quote = []
        out.append(line)
    return "\n".join(out)


def build() -> Path:
    base.SOURCE = SOURCE
    base.add_inline = add_inline
    base.add_picture = add_picture

    markdown = prepare_markdown(SOURCE.read_text(encoding="utf-8"))

    doc = Document()
    base.configure_styles(doc)
    base.configure_update_fields(doc)
    base.apply_page_layout(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True

    properties = doc.core_properties
    properties.title = TITLE
    properties.subject = SUBTITLE
    properties.author = "MSM项目组"
    properties.comments = (
        "报告只判断公开成熟度。厂商指标、战时通报和百科转引参数均保留来源限定，"
        "未经独立核实的宣称已在正文标注。"
    )

    add_cover(doc)
    content = doc.add_section(WD_SECTION.NEW_PAGE)
    base.apply_page_layout(content)
    add_header_footer(content)

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(16)
    run = toc_title.add_run("目录")
    base.set_run_font(run, size=20, bold=True, color=base.BLUE, east_asia=base.HEADING_FONT)
    base.add_toc(doc.add_paragraph())
    doc.add_page_break()

    base.render_markdown(doc, markdown)
    doc.save(OUTPUT)
    base.validate_docx(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
