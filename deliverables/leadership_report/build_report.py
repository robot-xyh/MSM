#!/usr/bin/env python3
"""Build the MSM leadership report figures, DOCX, and PDF.

The source of truth is C_UAS_PROJECT_LEADERSHIP_REPORT_CN.md.  This script is
kept beside the deliverable so the office documents can be regenerated without
touching any D1-D7 module-owned file.
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path
from shutil import copy2

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib import font_manager
from matplotlib.patches import Arc, Ellipse, FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parent.parent
ASSETS = HERE / "assets"
SOURCE = HERE / "C_UAS_PROJECT_LEADERSHIP_REPORT_CN.md"
DOCX_PATH = HERE / "C_UAS_PROJECT_LEADERSHIP_REPORT_CN.docx"
PDF_PATH = HERE / "C_UAS_PROJECT_LEADERSHIP_REPORT_CN.pdf"

BLUE = "1F4E78"
MID_BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
TEAL = "0F6B78"
GREEN = "2E7D32"
LIGHT_GREEN = "E2F0D9"
AMBER = "C47F00"
LIGHT_AMBER = "FFF2CC"
RED = "B03A2E"
LIGHT_RED = "FCE4D6"
INK = "1F2937"
MUTED = "5B6573"
PAPER = "F7F9FC"
WHITE = "FFFFFF"

FIG_FONT_PATH = "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
# Matplotlib registers this TTC under the JP family name; the collection still
# contains the Simplified Chinese glyphs used by this report.
FIG_FONT = "Noto Sans CJK JP"
BODY_FONT = "Noto Serif CJK SC"
HEADING_FONT = "Noto Sans CJK SC"


def _hex(value: str) -> str:
    return f"#{value}"


def setup_plot() -> None:
    font_manager.fontManager.addfont(FIG_FONT_PATH)
    plt.rcParams.update(
        {
            "font.family": FIG_FONT,
            "font.sans-serif": [FIG_FONT, "DejaVu Sans"],
            "mathtext.fontset": "dejavusans",
            "axes.unicode_minus": False,
            "figure.facecolor": "white",
            "savefig.facecolor": "white",
        }
    )


def add_card(ax, xy, width, height, title, body="", color=BLUE, face=WHITE, title_size=14, body_size=10):
    x, y = xy
    box = FancyBboxPatch(
        (x, y), width, height,
        boxstyle="round,pad=0.018,rounding_size=0.025",
        linewidth=1.5, edgecolor=_hex(color), facecolor=_hex(face),
    )
    ax.add_patch(box)
    ax.text(x + width / 2, y + height * 0.68, title, ha="center", va="center",
            fontsize=title_size, weight="bold", color=_hex(color))
    if body:
        ax.text(x + width / 2, y + height * 0.30, body, ha="center", va="center",
                fontsize=body_size, color=_hex(INK), linespacing=1.5)
    return box


def add_arrow(ax, start, end, color=MID_BLUE, connectionstyle="arc3,rad=0"):
    arrow = FancyArrowPatch(
        start, end, arrowstyle="-|>", mutation_scale=16, linewidth=1.8,
        color=_hex(color), connectionstyle=connectionstyle,
    )
    ax.add_patch(arrow)


def draw_drone_icon(ax, x, y, color=MID_BLUE, scale=1.0, zorder=7):
    """Draw a compact top-view quadrotor icon in the axis data coordinates."""
    rotor_offset = 1.15 * scale
    rotor_radius = 0.42 * scale
    body_radius = 0.48 * scale
    for sx, sy in [(-1, -1), (-1, 1), (1, -1), (1, 1)]:
        rx = x + sx * rotor_offset
        ry = y + sy * rotor_offset
        ax.plot([x, rx], [y, ry], color=_hex(color), linewidth=1.7, zorder=zorder)
        ax.add_patch(plt.Circle((rx, ry), rotor_radius, edgecolor=_hex(color),
                                facecolor="white", linewidth=1.4, zorder=zorder + 1))
    ax.add_patch(plt.Circle((x, y), body_radius, edgecolor="white",
                            facecolor=_hex(color), linewidth=1.0, zorder=zorder + 2))
    ax.plot([x, x + 0.95 * scale], [y, y], color=_hex(color), linewidth=2.0,
            solid_capstyle="round", zorder=zorder + 2)


def finish_figure(fig, path: Path) -> None:
    fig.savefig(path, dpi=220, bbox_inches="tight", pad_inches=0.18)
    plt.close(fig)


def prepare_project_figures() -> None:
    """Refresh figures when source evidence exists, otherwise use the archive."""
    sources = {
        REPO_ROOT / "research_modules/d1_sensor_fusion/reports/tracks_xy.png":
            ASSETS / "d1_fusion_tracks.png",
        REPO_ROOT / (
            "research_modules/airsim_runtime/outputs/p1_terminal_handoff_tuned_002/"
            "episode_006_full_flow/airsim_3d_intercept_trajectories.png"
        ): ASSETS / "airsim_intercept_3d.png",
    }
    for source, destination in sources.items():
        if source.exists():
            copy2(source, destination)
        elif not destination.exists():
            raise FileNotFoundError(
                f"Missing both project evidence and archived figure: {source}"
            )

    panorama = ASSETS / "battlefield_panorama.png"
    if not panorama.exists():
        raise FileNotFoundError(f"Missing bundled panorama: {panorama}")


def build_system_closed_loop() -> None:
    fig, ax = plt.subplots(figsize=(15.2, 4.4))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    labels = [
        ("多源探测", "雷达 / 可见光 / 红外"),
        ("可信融合", "时间 + 坐标 + 协方差"),
        ("身份关联", "稳定全局航迹 ID"),
        ("动态规划", "威胁 + 资源 + 时间窗"),
        ("韧性协同", "中心 → 二级 → 分布式"),
        ("末端确认", "视觉身份与计划一致"),
        ("导引执行", "中段 PN / 末段 PNG"),
        ("评估迭代", "日志、指标、再规划"),
    ]
    colors = [MID_BLUE, MID_BLUE, TEAL, TEAL, AMBER, AMBER, RED, GREEN]
    w, h, gap = 0.105, 0.48, 0.017
    x0 = 0.015
    for idx, ((title, body), color) in enumerate(zip(labels, colors)):
        x = x0 + idx * (w + gap)
        add_card(ax, (x, 0.31), w, h, title, body, color=color, face=PAPER, title_size=12.2, body_size=8.7)
        if idx < len(labels) - 1:
            add_arrow(ax, (x + w + 0.002, 0.55), (x + w + gap - 0.003, 0.55), color=MUTED)
    add_arrow(ax, (0.93, 0.27), (0.075, 0.23), color=GREEN, connectionstyle="arc3,rad=-0.06")
    ax.text(0.50, 0.10, "执行结果回流，驱动航迹、计划和策略持续更新", ha="center", va="center",
            fontsize=11, color=_hex(GREEN), weight="bold")
    ax.text(0.5, 0.93, "反无人机多无人机协同拦截闭环", ha="center", va="center",
            fontsize=18, color=_hex(INK), weight="bold")
    finish_figure(fig, ASSETS / "system_closed_loop.png")


def build_five_capabilities() -> None:
    fig, ax = plt.subplots(figsize=(14.8, 6.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.94, "五项关键技术围绕同一条任务闭环协同工作", ha="center",
            va="center", fontsize=18, weight="bold", color=_hex(INK))
    cards = [
        ("01", "可信航迹", "异步融合\nNED 坐标\n状态 + 协方差", MID_BLUE),
        ("02", "身份连续", "统计门控\n全局匹配\n稳定目标编号", TEAL),
        ("03", "滚动规划", "动态分配\n计划版本\n主用 + 备用", AMBER),
        ("04", "韧性接管", "中心统筹\n二级接管\n分布式保底", RED),
        ("05", "中末段协同", "雷达 PN\n视觉确认\n受控 PNG", GREEN),
    ]
    w, gap, x0 = 0.17, 0.025, 0.03
    for i, (num, title, body, color) in enumerate(cards):
        x = x0 + i * (w + gap)
        box = FancyBboxPatch((x, 0.20), w, 0.61, boxstyle="round,pad=0.02,rounding_size=0.03",
                             linewidth=1.5, edgecolor=_hex(color), facecolor=_hex(PAPER))
        ax.add_patch(box)
        ax.text(x + w / 2, 0.73, num, ha="center", va="center", fontsize=25, weight="bold", color=_hex(color))
        ax.text(x + w / 2, 0.58, title, ha="center", va="center", fontsize=15, weight="bold", color=_hex(INK))
        ax.text(x + w / 2, 0.37, body, ha="center", va="center", fontsize=11, color=_hex(MUTED), linespacing=1.6)
    ax.text(0.5, 0.09, "先把目标看准、任务分清、接管做稳，再进入末端受控执行",
            ha="center", va="center", fontsize=12, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "five_capabilities.png")


def build_identity_crossing_scene() -> None:
    """Qualitative dense-crossing scene without metric bars or IDSW charts."""
    rng = np.random.default_rng(20260714)
    t = np.linspace(0.0, 1.0, 90)
    paths = [
        (-115 + 230 * t, -48 + 98 * t + 8 * np.sin(np.pi * t), "G-101", MID_BLUE),
        (-115 + 230 * t, 52 - 102 * t - 7 * np.sin(np.pi * t), "G-102", TEAL),
        (-92 + 184 * t, -8 + 20 * np.sin(2 * np.pi * t), "G-103", AMBER),
        (92 - 184 * t, 10 - 22 * np.sin(2 * np.pi * t), "G-104", RED),
    ]
    fig, ax = plt.subplots(figsize=(14.8, 7.0))
    ax.axvspan(-16, 16, color=_hex(LIGHT_AMBER), alpha=0.55, label="密集交叉区")
    for x, y, label, color in paths:
        noise_x = rng.normal(0.0, 2.2, size=t.size)
        noise_y = rng.normal(0.0, 2.2, size=t.size)
        ax.scatter(x[::3] + noise_x[::3], y[::3] + noise_y[::3], s=11,
                   color=_hex(color), alpha=0.20, edgecolors="none")
        ax.plot(x, y, color=_hex(color), linewidth=2.6)
        ax.scatter([x[0]], [y[0]], s=55, marker="o", facecolor="white",
                   edgecolor=_hex(color), linewidth=1.8, zorder=5)
        ax.scatter([x[-1]], [y[-1]], s=58, marker=">", color=_hex(color), zorder=5)
        ax.text(x[0] - 4, y[0] + 6, f"{label} 起点", fontsize=9.2,
                color=_hex(color), ha="left", weight="bold")
        ax.text(x[-1] + (3 if x[-1] > 0 else -3), y[-1] + 5, f"{label} 保持",
                fontsize=9.2, color=_hex(color), ha="left" if x[-1] > 0 else "right", weight="bold")
    ax.add_patch(Ellipse((0, 1), 31, 19, edgecolor=_hex(AMBER), facecolor="none",
                         linestyle="--", linewidth=1.6))
    ax.annotate("位置靠近时，按预测状态与协方差联合判别\n不因最近点变化直接改写目标身份",
                xy=(0, 3), xytext=(35, 66), fontsize=10.5, color=_hex(INK),
                bbox=dict(boxstyle="round,pad=0.45", fc="white", ec=_hex(AMBER)),
                arrowprops=dict(arrowstyle="->", color=_hex(AMBER), linewidth=1.5))
    ax.text(-114, -72, "淡色散点：带噪观测    实线：关联后的稳定身份航迹    箭头：运动方向",
            fontsize=10, color=_hex(MUTED))
    ax.set_xlim(-125, 125)
    ax.set_ylim(-80, 82)
    ax.set_xlabel("东向位置 / m", fontsize=10.5)
    ax.set_ylabel("北向位置 / m", fontsize=10.5)
    ax.set_title("多目标交叉条件下的身份连续性（机制场景示意）",
                 fontsize=17, weight="bold", color=_hex(INK), pad=15)
    ax.grid(color="#D9DEE6", linewidth=0.75, alpha=0.75)
    ax.set_axisbelow(True)
    for side in ["top", "right"]:
        ax.spines[side].set_visible(False)
    fig.tight_layout()
    finish_figure(fig, ASSETS / "identity_crossing_scene.png")


def build_rolling_assignment_scene() -> None:
    """Top-down M5N2 assignment snapshot focused on geometry and roles."""
    fig, ax = plt.subplots(figsize=(14.8, 7.2))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 68)
    ax.set_aspect("equal", adjustable="box")
    ax.set_facecolor(_hex(PAPER))
    ax.grid(color="#DCE3EB", linewidth=0.7, alpha=0.75)
    ax.set_xlabel("任务区东向 / km", fontsize=10)
    ax.set_ylabel("任务区北向 / km", fontsize=10)
    ax.set_title("M5N2 滚动资源规划",
                 fontsize=17, weight="bold", color=_hex(INK), pad=14)

    target_data = {
        "T-01 高威胁": ((75, 49), (92, 59), RED),
        "T-02": ((74, 19), (94, 14), AMBER),
    }
    for name, (position, future, color) in target_data.items():
        ax.add_patch(Ellipse(position, 11, 7, angle=18, edgecolor=_hex(color),
                             facecolor=_hex(LIGHT_RED if color == RED else LIGHT_AMBER),
                             alpha=0.62, linewidth=1.4))
        draw_drone_icon(ax, *position, color=color, scale=1.35)
        ax.annotate("", xy=future, xytext=position,
                    arrowprops=dict(arrowstyle="-|>", color=_hex(color), linewidth=2.2))
        ax.text(position[0] + 2, position[1] + 5.5, name, color=_hex(color),
                fontsize=10.5, weight="bold")

    resources = {
        "R-01 主用": ((15, 57), MID_BLUE),
        "R-02 主用": ((14, 39), TEAL),
        "R-03 备用": ((34, 54), AMBER),
        "R-04": ((18, 16), GREEN),
        "R-05 待命": ((42, 21), MUTED),
    }
    for name, (position, color) in resources.items():
        draw_drone_icon(ax, *position, color=color, scale=1.45)
        ax.text(position[0], position[1] - 5, name, ha="center", fontsize=9.5,
                color=_hex(color), weight="bold")

    assignments = [
        ((15, 57), (75, 49), MID_BLUE, "主用 A", "-"),
        ((14, 39), (75, 49), TEAL, "主用 B", "-"),
        ((34, 54), (75, 49), AMBER, "备用", "--"),
        ((18, 16), (74, 19), GREEN, "主用", "-"),
    ]
    for start, end, color, role, style in assignments:
        ax.annotate("", xy=end, xytext=start,
                    arrowprops=dict(arrowstyle="-|>", color=_hex(color), linewidth=1.9,
                                    linestyle=style, alpha=0.9))
        mid = ((start[0] + end[0]) / 2, (start[1] + end[1]) / 2)
        ax.text(mid[0], mid[1] + 1.8, role, color=_hex(color), fontsize=8.8,
                ha="center", bbox=dict(fc="white", ec="none", alpha=0.75, pad=1.5))

    info = "当前计划 v17\nT-01：2 主用 + 1 备用\nT-02：1 主用\nR-05：保留机动余量"
    ax.text(66, 4.2, info, fontsize=9.6, color=_hex(INK), linespacing=1.55,
            bbox=dict(boxstyle="round,pad=0.55", fc="white", ec=_hex(BLUE), lw=1.2))
    ax.text(2.5, 2.7, "椭圆表示位置不确定范围；新观测或资源状态变化后发布下一版本",
            fontsize=9.5, color=_hex(MUTED))
    for side in ["top", "right"]:
        ax.spines[side].set_visible(False)
    fig.tight_layout()
    finish_figure(fig, ASSETS / "rolling_assignment_scene.png")


def build_d5_terminal_association_scene() -> None:
    """Illustrate D5 projection, local matching, and conservative decision states."""
    fig, ax = plt.subplots(figsize=(15.0, 6.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.95, "D5 末端视觉关联", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))

    # Global track and camera geometry.
    ax.add_patch(FancyBboxPatch((0.025, 0.17), 0.27, 0.66,
                                boxstyle="round,pad=0.018,rounding_size=0.025",
                                linewidth=1.3, edgecolor=_hex(MID_BLUE), facecolor=_hex(PAPER)))
    ax.text(0.16, 0.77, "全局航迹与相机", ha="center", fontsize=13,
            weight="bold", color=_hex(MID_BLUE))
    draw_drone_icon(ax, 0.085, 0.29, color=MID_BLUE, scale=0.015)
    ax.text(0.085, 0.23, "拦截资源", ha="center", fontsize=9, color=_hex(MID_BLUE))
    ax.fill([0.10, 0.255, 0.255, 0.10], [0.31, 0.66, 0.52, 0.27],
            color=_hex(LIGHT_BLUE), alpha=0.45, edgecolor=_hex(MID_BLUE), linewidth=1.1)
    ax.add_patch(Ellipse((0.235, 0.60), 0.085, 0.12, angle=22,
                         edgecolor=_hex(TEAL), facecolor="white", linewidth=1.8, linestyle="--"))
    draw_drone_icon(ax, 0.235, 0.60, color=TEAL, scale=0.012)
    ax.text(0.235, 0.69, "中心航迹 G-102", ha="center", fontsize=9.3,
            color=_hex(TEAL), weight="bold")
    ax.text(0.16, 0.12, "量测时刻 · 相机内外参 · 航迹协方差",
            ha="center", fontsize=8.8, color=_hex(MUTED))
    add_arrow(ax, (0.30, 0.50), (0.345, 0.50), color=MID_BLUE)

    # Camera image with projected uncertainty and competing local tracks.
    frame = FancyBboxPatch((0.35, 0.17), 0.34, 0.66,
                           boxstyle="round,pad=0.015,rounding_size=0.02",
                           linewidth=1.4, edgecolor=_hex(INK), facecolor="#EEF3F7")
    ax.add_patch(frame)
    ax.text(0.52, 0.77, "相机像面候选", ha="center", fontsize=13,
            weight="bold", color=_hex(INK))
    ax.plot([0.37, 0.67], [0.50, 0.50], color="#CBD4DE", linewidth=0.8)
    ax.plot([0.52, 0.52], [0.20, 0.75], color="#CBD4DE", linewidth=0.8)
    ax.add_patch(Ellipse((0.505, 0.54), 0.145, 0.22, angle=-10,
                         edgecolor=_hex(TEAL), facecolor=_hex(LIGHT_BLUE), alpha=0.45,
                         linewidth=1.8, linestyle="--"))
    ax.text(0.475, 0.68, "投影区域", fontsize=8.8, color=_hex(TEAL), weight="bold")
    # Stable local track history.
    for idx, (cx, cy, alpha) in enumerate([(0.478, 0.525, 0.28), (0.492, 0.535, 0.50), (0.505, 0.545, 1.0)]):
        rect = FancyBboxPatch((cx - 0.026, cy - 0.036), 0.052, 0.072,
                              boxstyle="round,pad=0.003,rounding_size=0.006",
                              linewidth=1.5, edgecolor=_hex(GREEN), facecolor="none", alpha=alpha)
        ax.add_patch(rect)
    ax.text(0.505, 0.49, "L-07  连续 3 帧", ha="center", fontsize=9.2,
            color=_hex(GREEN), weight="bold")
    # Competing candidate and verified friendly object.
    ax.add_patch(FancyBboxPatch((0.59, 0.38), 0.054, 0.075,
                                boxstyle="round,pad=0.003,rounding_size=0.006",
                                linewidth=1.5, edgecolor=_hex(AMBER), facecolor="none"))
    ax.text(0.617, 0.35, "L-11", ha="center", fontsize=8.8, color=_hex(AMBER))
    ax.add_patch(FancyBboxPatch((0.405, 0.30), 0.052, 0.072,
                                boxstyle="round,pad=0.003,rounding_size=0.006",
                                linewidth=1.5, edgecolor=_hex(MID_BLUE), facecolor="none"))
    ax.text(0.431, 0.27, "友方声明", ha="center", fontsize=8.6, color=_hex(MID_BLUE))
    ax.text(0.52, 0.12, "马氏门控 · 候选唯一性 · 友方冲突 · 证据时效",
            ha="center", fontsize=8.8, color=_hex(MUTED))
    add_arrow(ax, (0.695, 0.50), (0.735, 0.50), color=TEAL)

    # Conservative output states.
    ax.add_patch(FancyBboxPatch((0.74, 0.17), 0.235, 0.66,
                                boxstyle="round,pad=0.018,rounding_size=0.025",
                                linewidth=1.3, edgecolor=_hex(GREEN), facecolor=_hex(PAPER)))
    ax.text(0.858, 0.77, "关联决策", ha="center", fontsize=13,
            weight="bold", color=_hex(GREEN))
    add_card(ax, (0.775, 0.57), 0.165, 0.12, "锁定", "G-102 ↔ L-07", GREEN, LIGHT_GREEN, 12, 8.8)
    add_card(ax, (0.755, 0.38), 0.085, 0.105, "歧义", "候选接近", AMBER, LIGHT_AMBER, 9.2, 7.4)
    add_card(ax, (0.87, 0.38), 0.07, 0.105, "保持", "合同冲突", RED, LIGHT_RED, 9.2, 7.4)
    add_card(ax, (0.775, 0.22), 0.165, 0.105, "重捕", "丢测后重新门控", MID_BLUE, LIGHT_BLUE, 10, 7.6)
    ax.text(0.5, 0.045, "D5 只核对中心既有全局身份，不创建、不改写、不换绑全局航迹编号",
            ha="center", fontsize=10.2, color=_hex(RED), weight="bold")
    finish_figure(fig, ASSETS / "d5_terminal_association_scene.png")


def build_d7_proportional_guidance_scene() -> None:
    """Illustrate PN geometry and the visual LOS-rate input used in terminal PNG."""
    fig, ax = plt.subplots(figsize=(15.0, 6.4))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 66)
    ax.set_aspect("equal", adjustable="box")
    ax.set_facecolor(_hex(PAPER))
    ax.grid(color="#DCE3EB", linewidth=0.7, alpha=0.65)
    ax.set_xlabel("东向相对位置", fontsize=10)
    ax.set_ylabel("北向相对位置", fontsize=10)
    ax.set_title("D7 比例导引几何", fontsize=18, weight="bold", color=_hex(INK), pad=14)

    # Curved interceptor path and target motion.
    path_x = np.array([8, 18, 29, 41, 53, 64])
    path_y = np.array([8, 13, 21, 32, 42, 49])
    ax.plot(path_x, path_y, color=_hex(MID_BLUE), linewidth=2.8)
    ax.annotate("", xy=(64, 49), xytext=(53, 42),
                arrowprops=dict(arrowstyle="-|>", color=_hex(MID_BLUE), linewidth=2.5))
    ax.plot([61, 91], [53, 57], color=_hex(RED), linewidth=2.4, linestyle="--")
    ax.annotate("", xy=(91, 57), xytext=(78, 55.3),
                arrowprops=dict(arrowstyle="-|>", color=_hex(RED), linewidth=2.2))
    draw_drone_icon(ax, 41, 32, color=MID_BLUE, scale=1.8)
    draw_drone_icon(ax, 78, 55.3, color=RED, scale=1.7)
    ax.text(36, 26.5, "拦截资源", fontsize=9.5, color=_hex(MID_BLUE), weight="bold")
    ax.text(79, 61, "运动目标", fontsize=9.5, color=_hex(RED), weight="bold")

    # LOS at two instants and the core quantities.
    ax.plot([29, 69], [21, 54.1], color=_hex(MUTED), linestyle=":", linewidth=1.5)
    ax.plot([41, 78], [32, 55.3], color=_hex(TEAL), linestyle="--", linewidth=2.0)
    ax.text(47, 48.5, "LOS(t−Δt)", fontsize=8.8, color=_hex(MUTED), rotation=37)
    ax.text(56, 45.5, "LOS(t)", fontsize=9.3, color=_hex(TEAL), rotation=32, weight="bold")
    ax.add_patch(Arc((41, 32), 18, 18, angle=0, theta1=22, theta2=40,
                     color=_hex(AMBER), linewidth=2.0))
    ax.text(49, 39.2, "视线角变化", fontsize=9, color=_hex(AMBER), weight="bold")
    add_arrow(ax, (47, 35.8), (58, 42.7), color=TEAL)
    ax.text(53, 33.8, "闭合速度  Vc", fontsize=9.2, color=_hex(TEAL))
    add_arrow(ax, (41, 32), (32, 46), color=RED)
    ax.text(20, 49.5, "横向指令  a_n", fontsize=9.2, color=_hex(RED), weight="bold")

    # Camera inset showing terminal LOS-rate measurement.
    inset = FancyBboxPatch((70, 5), 27, 22, boxstyle="round,pad=0.7,rounding_size=1.1",
                           linewidth=1.3, edgecolor=_hex(INK), facecolor="white")
    ax.add_patch(inset)
    ax.plot([83.5, 83.5], [7, 25], color="#D1D8E0", linewidth=0.8)
    ax.plot([72, 95], [16, 16], color="#D1D8E0", linewidth=0.8)
    for cx, alpha in [(79.0, 0.28), (81.0, 0.55), (83.0, 1.0)]:
        ax.add_patch(FancyBboxPatch((cx - 2.0, 13.2), 4.0, 4.8,
                                    boxstyle="round,pad=0.1,rounding_size=0.25",
                                    linewidth=1.5, edgecolor=_hex(GREEN), facecolor="none", alpha=alpha))
    ax.annotate("像面中心变化", xy=(83, 19), xytext=(86, 22), fontsize=8.5,
                color=_hex(GREEN), arrowprops=dict(arrowstyle="->", color=_hex(GREEN)))
    ax.text(83.5, 8.5, "末端视觉给出 LOS 角速度", ha="center",
            fontsize=8.8, color=_hex(INK), weight="bold")

    ax.text(8, 60, r"$a_n=N V_c \dot{\lambda}$", fontsize=18, color=_hex(INK),
            bbox=dict(boxstyle="round,pad=0.45", fc="white", ec=_hex(MID_BLUE), lw=1.3))
    ax.text(8, 55, "目标是压低视线角速度，在机动约束内形成提前量",
            fontsize=9.5, color=_hex(MUTED))
    for side in ["top", "right"]:
        ax.spines[side].set_visible(False)
    fig.tight_layout()
    finish_figure(fig, ASSETS / "d7_proportional_guidance_scene.png")


def build_resilient_takeover_scene() -> None:
    """Before/after scene for center failure and secondary-node takeover."""
    fig, axes = plt.subplots(1, 2, figsize=(15.0, 6.2))
    panels = [
        (axes[0], "中心正常", False),
        (axes[1], "中心失联，二级节点接管区域任务", True),
    ]
    for ax, title, failed in panels:
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis("off")
        ax.set_title(title, fontsize=14, weight="bold",
                     color=_hex(RED if failed else BLUE), pad=10)
        c2_face = "E6E8EB" if failed else LIGHT_BLUE
        c2_color = MUTED if failed else MID_BLUE
        add_card(ax, (0.36, 0.77), 0.28, 0.12, "中心 C2", "全局航迹与计划", c2_color, c2_face, 12, 8.5)
        add_card(ax, (0.32, 0.50), 0.36, 0.13, "二级节点", "区域态势 · 缓存计划 · 新任期", TEAL,
                 "E2F4F5", 12, 8.5)
        if failed:
            ax.plot([0.39, 0.61], [0.79, 0.89], color=_hex(RED), linewidth=3)
            ax.plot([0.39, 0.61], [0.89, 0.79], color=_hex(RED), linewidth=3)
            ax.text(0.73, 0.82, "心跳超时", fontsize=9.2, color=_hex(RED), weight="bold")
        else:
            add_arrow(ax, (0.50, 0.77), (0.50, 0.64), color=MID_BLUE)
        xs = [0.10, 0.34, 0.58, 0.82]
        for idx, x in enumerate(xs, 1):
            add_card(ax, (x - 0.075, 0.18), 0.15, 0.13, f"资源 {idx}", "执行/上报",
                     AMBER, LIGHT_AMBER, 9.8, 7.8)
            origin = (0.50, 0.50) if failed else (0.50, 0.77)
            add_arrow(ax, origin, (x, 0.32), color=TEAL if failed else MID_BLUE,
                      connectionstyle=f"arc3,rad={0.10 * (idx - 2.5):.2f}")
        if failed:
            ax.text(0.50, 0.40, "新计划所有者 / 任期 / 租约 / 版本",
                    ha="center", fontsize=8.8, color=_hex(TEAL), weight="bold")
            ax.text(0.50, 0.09, "无法形成唯一接管者时保持安全状态",
                    ha="center", fontsize=9.5, color=_hex(RED), weight="bold")
        else:
            ax.text(0.50, 0.09, "中心是正常状态下的计划事实源",
                    ha="center", fontsize=9.5, color=_hex(BLUE), weight="bold")
    fig.suptitle("中心—二级—分布式的韧性接管", fontsize=17, weight="bold",
                 color=_hex(INK), y=0.99)
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    finish_figure(fig, ASSETS / "resilient_takeover_scene.png")


def build_three_layer_architecture() -> None:
    fig, ax = plt.subplots(figsize=(14.5, 7.6))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.95, "中心最优、二级连续、分布式保底", ha="center", va="center",
            fontsize=19, weight="bold", color=_hex(INK))
    add_card(ax, (0.31, 0.72), 0.38, 0.16, "中心 C2", "融合航迹 · 全局规划 · 授权 · 评估", MID_BLUE, LIGHT_BLUE, 15, 10.5)
    add_card(ax, (0.20, 0.43), 0.60, 0.17, "高空侦察 / 区域协调节点", "高视角光电 · 局部补盲 · 区域接管 · 通信中继", TEAL, "E2F4F5", 15, 10.5)
    xs = [0.05, 0.29, 0.53, 0.77]
    for idx, x in enumerate(xs, 1):
        add_card(ax, (x, 0.12), 0.18, 0.17, f"资源节点 {idx}", "机载视觉\n状态上报\n受控导引", AMBER, LIGHT_AMBER, 12.5, 9.5)
        add_arrow(ax, (0.5, 0.43), (x + 0.09, 0.30), color=TEAL)
    add_arrow(ax, (0.50, 0.72), (0.50, 0.61), color=MID_BLUE)
    add_arrow(ax, (0.45, 0.61), (0.45, 0.71), color=MID_BLUE)
    ax.text(0.71, 0.65, "正常：中心下发版本化计划", fontsize=10.5, color=_hex(MID_BLUE))
    ax.text(0.71, 0.39, "退化：二级节点在新所有者和任期下接管", fontsize=10.5, color=_hex(TEAL))
    ax.text(0.5, 0.045, "所有执行均受目标身份、计划版本、成员角色、租约和授权门控",
            ha="center", va="center", fontsize=11, color=_hex(RED), weight="bold")
    finish_figure(fig, ASSETS / "three_layer_architecture.png")


def build_sensor_roles() -> None:
    fig, ax = plt.subplots(figsize=(14.5, 6.7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.93, "雷达与三类光电分工接力", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    items = [
        (0.03, "地面雷达", "全局航迹骨架", "距离 / 方位 / 俯仰 / 速度\n全天候、宽域、需协方差", MID_BLUE, LIGHT_BLUE),
        (0.275, "地面光电", "识别与复核", "广角搜索 + 变焦凝视\n可见光纹理 + 红外弱光", TEAL, "E2F4F5"),
        (0.52, "高空光电", "补盲与区域态势", "俯视减遮挡、区域引导线索\n高度增加会降低目标像素尺度", AMBER, LIGHT_AMBER),
        (0.765, "机载相机", "末端身份与视线率", "全局快门、已知相机参数\n本地编号对齐全局编号", RED, LIGHT_RED),
    ]
    for x, title, subtitle, body, color, face in items:
        box = FancyBboxPatch((x, 0.22), 0.205, 0.55, boxstyle="round,pad=0.018,rounding_size=0.025",
                             linewidth=1.5, edgecolor=_hex(color), facecolor=_hex(face))
        ax.add_patch(box)
        ax.text(x + 0.1025, 0.67, title, ha="center", va="center", fontsize=14, weight="bold", color=_hex(color))
        ax.text(x + 0.1025, 0.54, subtitle, ha="center", va="center", fontsize=12, weight="bold", color=_hex(INK))
        ax.text(x + 0.1025, 0.37, body, ha="center", va="center", fontsize=9.7, color=_hex(MUTED), linespacing=1.6)
    for x in [0.245, 0.49, 0.735]:
        add_arrow(ax, (x - 0.012, 0.50), (x + 0.012, 0.50), color=MUTED)
    ax.text(0.5, 0.10, "交接基础：统一时钟、NED 坐标、相机内外参、协方差和稳定目标身份",
            ha="center", va="center", fontsize=11.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "sensor_roles.png")


def build_sensor_handover_flow() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 5.4))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.93, "雷达—光电—机载视觉分段接力", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    steps = [
        (0.03, "远距发现", "雷达搜索\n输出量测与协方差", MID_BLUE, LIGHT_BLUE),
        (0.25, "融合跟踪", "时间对齐 + NED\n形成全局航迹", TEAL, "E2F4F5"),
        (0.47, "中段执行", "版本化计划\n雷达比例导引", AMBER, LIGHT_AMBER),
        (0.69, "末段交接门", "身份一致\n计划有效\n视觉稳定", RED, LIGHT_RED),
        (0.86, "视觉末段", "D5 锁定\nD7 视觉比例导航", GREEN, LIGHT_GREEN),
    ]
    widths = [0.16, 0.16, 0.16, 0.13, 0.11]
    for i, ((x, title, body, color, face), width) in enumerate(zip(steps, widths)):
        add_card(ax, (x, 0.43), width, 0.32, title, body, color=color, face=face,
                 title_size=12.5, body_size=9.4)
        if i < len(steps) - 1:
            next_x = steps[i + 1][0]
            add_arrow(ax, (x + width + 0.005, 0.59), (next_x - 0.008, 0.59), color=MUTED)
    ax.text(0.755, 0.37, "未通过", fontsize=9.5, color=_hex(RED), weight="bold")
    add_arrow(ax, (0.755, 0.42), (0.755, 0.24), color=RED)
    add_card(ax, (0.53, 0.08), 0.45, 0.14, "门控不通过：继续雷达中段 / 保持 / 重捕 / 请求重规划",
             "不能用最近检测框替代目标身份", color=RED, face=PAPER, title_size=10.5, body_size=8.8)
    ax.text(0.5, 0.025, "交接条件综合误差、视场、身份、版本和证据新鲜度",
            ha="center", va="center", fontsize=10.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "sensor_handover_flow.png")


def build_fusion_estimation_flow() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 5.6))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.94, "从异步观测到可信全局航迹", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    input_items = [
        ("雷达", "距离/角度/速度", MID_BLUE, LIGHT_BLUE),
        ("地面/高空光电", "像素框/角度线索", TEAL, "E2F4F5"),
        ("资源状态", "位姿/速度/健康", AMBER, LIGHT_AMBER),
    ]
    ys = [0.67, 0.43, 0.19]
    for (title, body, color, face), y in zip(input_items, ys):
        add_card(ax, (0.025, y), 0.17, 0.16, title, body, color=color, face=face,
                 title_size=11.5, body_size=8.5)
        add_arrow(ax, (0.20, y + 0.08), (0.25, 0.50), color=MUTED)
    stages = [
        (0.25, "时间治理", "量测时刻\n到达时刻\n乱序/过期", MID_BLUE, LIGHT_BLUE),
        (0.43, "空间统一", "外参变换\nNED 坐标\n误差传播", TEAL, "E2F4F5"),
        (0.61, "状态估计", "预测 + EKF 更新\n状态均值 + 协方差", AMBER, LIGHT_AMBER),
        (0.79, "身份关联", "马氏门控\nHungarian\n生命周期", RED, LIGHT_RED),
    ]
    for i, (x, title, body, color, face) in enumerate(stages):
        add_card(ax, (x, 0.34), 0.145, 0.32, title, body, color=color, face=face,
                 title_size=11.5, body_size=8.8)
        if i < len(stages) - 1:
            add_arrow(ax, (x + 0.15, 0.50), (stages[i + 1][0] - 0.005, 0.50), color=MUTED)
    add_card(ax, (0.925, 0.34), 0.065, 0.32, "输出", "全局\n航迹", GREEN, LIGHT_GREEN, 10.5, 8.5)
    add_arrow(ax, (0.94, 0.31), (0.69, 0.19), color=GREEN, connectionstyle="arc3,rad=0.15")
    ax.text(0.70, 0.12, "协方差、残差和 ID 风险继续进入分配与末端门控", ha="center",
            va="center", fontsize=10.5, color=_hex(GREEN), weight="bold")
    finish_figure(fig, ASSETS / "fusion_estimation_flow.png")


def build_allocation_closed_loop() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 5.6))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.94, "资源—目标分配滚动闭环", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    stages = [
        (0.025, "可信态势", "全局航迹\n协方差/威胁", MID_BLUE, LIGHT_BLUE),
        (0.205, "资源状态", "速度/健康\n视场/余量", TEAL, "E2F4F5"),
        (0.385, "构造代价", "时间窗 + 风险\n冲突 + 迟滞", AMBER, LIGHT_AMBER),
        (0.565, "约束求解", "Hungarian\n需求槽/联盟", RED, LIGHT_RED),
        (0.745, "发布计划", "所有者/版本\n角色/时间窗", GREEN, LIGHT_GREEN),
    ]
    for i, (x, title, body, color, face) in enumerate(stages):
        add_card(ax, (x, 0.45), 0.15, 0.28, title, body, color=color, face=face,
                 title_size=11.5, body_size=8.8)
        if i < len(stages) - 1:
            add_arrow(ax, (x + 0.155, 0.59), (stages[i + 1][0] - 0.005, 0.59), color=MUTED)
    add_card(ax, (0.91, 0.45), 0.075, 0.28, "执行", "D4/D5/D7\n门控", BLUE, PAPER, 10.5, 8.2)
    add_arrow(ax, (0.95, 0.42), (0.70, 0.23), color=GREEN, connectionstyle="arc3,rad=0.12")
    add_card(ax, (0.32, 0.08), 0.50, 0.16, "执行反馈", "目标更新 · 资源损失 · 视觉歧义 · 计划过期 → 新版本重规划",
             GREEN, PAPER, 11.5, 9.2)
    add_arrow(ax, (0.32, 0.16), (0.10, 0.43), color=GREEN, connectionstyle="arc3,rad=-0.18")
    ax.text(0.5, 0.025, "迟滞抑制微小波动；新身份、新硬风险或窗口关闭则立即触发重评估",
            ha="center", va="center", fontsize=10.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "allocation_closed_loop.png")


def build_didi_vs_cuas_flow() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 6.3))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.95, "同样使用匹配算法，前置条件和闭环难度完全不同", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    ax.text(0.035, 0.77, "网约车抽象", fontsize=13, weight="bold", color=_hex(MID_BLUE))
    upper = [
        (0.18, "合作式定位", "司机/乘客身份已知"),
        (0.40, "确定性代价", "距离/ETA"),
        (0.62, "一次匹配", "Hungarian/派单"),
        (0.84, "服务执行", "异常再派单"),
    ]
    for i, (x, title, body) in enumerate(upper):
        add_card(ax, (x, 0.68), 0.15, 0.19, title, body, MID_BLUE, LIGHT_BLUE, 11, 8.7)
        if i < len(upper) - 1:
            add_arrow(ax, (x + 0.155, 0.775), (upper[i + 1][0] - 0.005, 0.775), color=MID_BLUE)
    ax.text(0.035, 0.40, "空中目标", fontsize=13, weight="bold", color=_hex(RED))
    lower = [
        (0.18, "间接观测", "噪声/漏检/延迟"),
        (0.36, "状态估计", "均值 + 协方差"),
        (0.54, "身份关联", "IDSW/遮挡"),
        (0.72, "约束分配", "窗口/联盟/安全"),
        (0.88, "执行反馈", "再观测/再规划"),
    ]
    for i, (x, title, body) in enumerate(lower):
        width = 0.13 if i < 4 else 0.10
        add_card(ax, (x, 0.30), width, 0.20, title, body, RED if i in {0, 2} else AMBER,
                 LIGHT_RED if i in {0, 2} else LIGHT_AMBER, 10.5, 8.2)
        if i < len(lower) - 1:
            add_arrow(ax, (x + width + 0.004, 0.40), (lower[i + 1][0] - 0.005, 0.40), color=MUTED)
    add_arrow(ax, (0.93, 0.28), (0.43, 0.17), color=GREEN, connectionstyle="arc3,rad=0.16")
    add_arrow(ax, (0.43, 0.17), (0.245, 0.28), color=GREEN, connectionstyle="arc3,rad=0.12")
    ax.text(0.59, 0.10, "执行改变观测几何，必须持续回流", fontsize=10.5,
            color=_hex(GREEN), weight="bold", ha="center")
    ax.text(0.5, 0.025, "空中问题 = 部分可观测状态估计 + 身份管理 + 滚动随机规划",
            ha="center", va="center", fontsize=11, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "didi_vs_cuas_flow.png")


def build_coalition_2plus1() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 5.8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.94, "高威胁目标  2 个主用资源 + 1 个备用资源", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    target = plt.Circle((0.78, 0.52), 0.10, edgecolor=_hex(RED), facecolor=_hex(LIGHT_RED), linewidth=2)
    ax.add_patch(target)
    ax.text(0.78, 0.55, "高威胁目标", ha="center", va="center", fontsize=12, weight="bold", color=_hex(RED))
    ax.text(0.78, 0.48, "同一全局航迹编号", ha="center", va="center", fontsize=9, color=_hex(MUTED))
    resources = [
        (0.10, 0.72, "主用资源 A", "第一扇区\n当前版本已激活", MID_BLUE, LIGHT_BLUE),
        (0.10, 0.35, "主用资源 B", "第二扇区\n当前版本已激活", TEAL, "E2F4F5"),
        (0.38, 0.08, "备用资源", "安全待命\n未激活不得进入末端", AMBER, LIGHT_AMBER),
    ]
    for x, y, title, body, color, face in resources:
        add_card(ax, (x, y), 0.19, 0.18, title, body, color, face, 11.5, 8.8)
    add_arrow(ax, (0.30, 0.81), (0.67, 0.60), color=MID_BLUE)
    add_arrow(ax, (0.30, 0.44), (0.67, 0.48), color=TEAL)
    add_arrow(ax, (0.57, 0.17), (0.71, 0.40), color=AMBER, connectionstyle="arc3,rad=-0.08")
    ax.text(0.52, 0.28, "仅在新计划版本激活后补位", fontsize=9.5, color=_hex(AMBER), weight="bold")
    add_card(ax, (0.89, 0.35), 0.095, 0.34, "D3/D4", "角色\n时间窗\n版本\n确认/租约", GREEN, LIGHT_GREEN, 10.5, 8.4)
    add_arrow(ax, (0.89, 0.52), (0.885, 0.52), color=GREEN)
    ax.text(0.5, 0.025, "联盟完成必须按目标和联盟统计；单个主用资源进入 5 m 不能推导整个联盟完成",
            ha="center", va="center", fontsize=10.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "coalition_2plus1.png")


def build_guidance_gate_flow() -> None:
    fig, ax = plt.subplots(figsize=(15.0, 5.8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.94, "雷达 PN 切换到视觉 PNG 的联合门控", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    add_card(ax, (0.03, 0.45), 0.15, 0.26, "雷达中段 PN", "持续预测目标状态\n保持当前计划", MID_BLUE, LIGHT_BLUE, 12, 9)
    gates = [
        (0.23, "计划门", "所有者/版本\n角色/窗口"),
        (0.40, "仲裁门", "D4 允许继续\n无降级冲突"),
        (0.57, "身份门", "D5 已锁定\n同一全局编号"),
        (0.74, "几何门", "视场/时延\n机动余量"),
    ]
    for i, (x, title, body) in enumerate(gates):
        add_card(ax, (x, 0.45), 0.13, 0.26, title, body, AMBER if i < 2 else RED,
                 LIGHT_AMBER if i < 2 else LIGHT_RED, 11, 8.5)
        if i == 0:
            add_arrow(ax, (0.185, 0.58), (x - 0.005, 0.58), color=MUTED)
        if i < len(gates) - 1:
            add_arrow(ax, (x + 0.135, 0.58), (gates[i + 1][0] - 0.005, 0.58), color=MUTED)
    add_card(ax, (0.90, 0.45), 0.085, 0.26, "视觉末段", "PNG\n受控执行", GREEN, LIGHT_GREEN, 10.5, 8.5)
    add_arrow(ax, (0.875, 0.58), (0.895, 0.58), color=GREEN)
    for x in [0.295, 0.465, 0.635, 0.805]:
        add_arrow(ax, (x, 0.44), (x, 0.27), color=RED)
    add_card(ax, (0.22, 0.08), 0.66, 0.16, "任一门不通过", "继续雷达比例导引 / 保持 / 重捕 / 请求重规划；过期计划和备用越权必须保守拒绝",
             RED, PAPER, 11.5, 9.2)
    ax.text(0.5, 0.025, "切换依据由计划、仲裁、身份和几何证据共同构成",
            ha="center", va="center", fontsize=10.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "guidance_gate_flow.png")


def build_evidence_overview() -> None:
    fig, ax = plt.subplots(figsize=(13.6, 6.7))
    labels = [
        "D4 故障安全结果",
        "5v5 YOLO 离线召回",
        "M5N2 最佳联盟完成",
        "D2 候选 IDSW 降幅",
        "原生 MOT 正式准入",
    ]
    values = [1.00, 0.678, 0.50, 0.546, 0.0]
    colors = [_hex(GREEN), _hex(AMBER), _hex(RED), _hex(AMBER), _hex(RED)]
    notes = ["60/60", "recall=0.678", "5/10（门限 8/10）", "1.3583 → 0.6167，未晋级", "0/18"]
    y = list(range(len(labels)))[::-1]
    ax.barh(y, values, color=colors, height=0.52)
    ax.set_xlim(0, 1.08)
    ax.set_yticks(y, labels, fontsize=11)
    ax.set_xticks([0, 0.2, 0.4, 0.6, 0.8, 1.0], ["0", "20%", "40%", "60%", "80%", "100%"])
    ax.grid(axis="x", color="#D9DEE6", linewidth=0.8, alpha=0.8)
    ax.set_axisbelow(True)
    for yi, val, note in zip(y, values, notes):
        xpos = min(max(val + 0.025, 0.03), 0.92)
        ax.text(xpos, yi, note, va="center", fontsize=10.5, color=_hex(INK), weight="bold")
    ax.axvline(0.8, color=_hex(MID_BLUE), linestyle="--", linewidth=1.5)
    ax.text(0.805, 4.45, "示意验收参考线", fontsize=9, color=_hex(MID_BLUE))
    ax.set_title("阶段证据概览：成功项与开放缺口同时呈现", fontsize=17, weight="bold", color=_hex(INK), pad=18)
    ax.set_xlabel("归一化比例（不同指标不可相互替代）", fontsize=10.5, color=_hex(MUTED))
    for side in ["top", "right", "left"]:
        ax.spines[side].set_visible(False)
    ax.spines["bottom"].set_color("#BBC3CE")
    fig.tight_layout()
    finish_figure(fig, ASSETS / "evidence_overview.png")


def build_roadmap() -> None:
    fig, ax = plt.subplots(figsize=(14.8, 6.5))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    ax.text(0.5, 0.93, "下一阶段：先闭合真实数据和物理链路，再晋级高阶算法", ha="center", va="center",
            fontsize=18, weight="bold", color=_hex(INK))
    phases = [
        ("阶段一", "真实数据基线", "实飞日志\n平台包线\n雷达/相机标定", MID_BLUE, LIGHT_BLUE),
        ("阶段二", "物理闭环 P1", "第二主用资源\n30/50 m 视觉\nM5N2 ≥ 8/10", RED, LIGHT_RED),
        ("阶段三", "半实物与网络", "飞控台架\n真实时延/丢包\n非破坏性缩比飞行", AMBER, LIGHT_AMBER),
        ("阶段四", "体系化验证", "多天气/多规模\n多故障/多随机种子\n高级算法同场对照", GREEN, LIGHT_GREEN),
    ]
    xs = [0.035, 0.285, 0.535, 0.785]
    for i, (phase, title, body, color, face) in enumerate(phases):
        x = xs[i]
        box = FancyBboxPatch((x, 0.23), 0.18, 0.55, boxstyle="round,pad=0.02,rounding_size=0.025",
                             linewidth=1.5, edgecolor=_hex(color), facecolor=_hex(face))
        ax.add_patch(box)
        ax.text(x + 0.09, 0.70, phase, ha="center", va="center", fontsize=11, color=_hex(color), weight="bold")
        ax.text(x + 0.09, 0.58, title, ha="center", va="center", fontsize=14, color=_hex(INK), weight="bold")
        ax.text(x + 0.09, 0.39, body, ha="center", va="center", fontsize=10, color=_hex(MUTED), linespacing=1.6)
        if i < 3:
            add_arrow(ax, (x + 0.19, 0.50), (xs[i + 1] - 0.01, 0.50), color=MUTED)
    ax.text(0.5, 0.105, "统一验收原则：相同场景版本、相同初始几何、配对随机种子、证据可回放",
            ha="center", va="center", fontsize=11.5, color=_hex(BLUE), weight="bold")
    finish_figure(fig, ASSETS / "roadmap.png")


def build_figures() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    setup_plot()
    prepare_project_figures()
    build_system_closed_loop()
    build_five_capabilities()
    build_identity_crossing_scene()
    build_rolling_assignment_scene()
    build_resilient_takeover_scene()
    build_three_layer_architecture()
    build_sensor_roles()
    build_sensor_handover_flow()
    build_fusion_estimation_flow()
    build_allocation_closed_loop()
    build_didi_vs_cuas_flow()
    build_coalition_2plus1()
    build_d5_terminal_association_scene()
    build_d7_proportional_guidance_scene()
    build_guidance_gate_flow()
    build_roadmap()


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


MATH_REPLACEMENTS = {
    r"\sigma": "σ", r"\theta": "θ", r"\lambda": "λ", r"\Delta": "Δ",
    r"\pi": "π", r"\tau": "τ", r"\approx": "≈", r"\propto": "∝",
    r"\le": "≤", r"\ge": "≥", r"\in": "∈", r"\sum": "Σ",
    r"\min": "min", r"\max": "max", r"\cdot": "·", r"\rightarrow": "→",
    r"\qquad": "    ", r"\;": " ", r"\,": " ",
}


def math_to_plain(text: str) -> str:
    result = text.strip().strip("$")
    for old, new in MATH_REPLACEMENTS.items():
        result = result.replace(old, new)
    result = result.replace(r"\hat{x}", "x̂").replace(r"\frac", "frac")
    result = result.replace("{", "").replace("}", "")
    return result


INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\$[^$]+?\$)")


def add_inline(paragraph, text: str, default_size=10.5, default_color=INK) -> None:
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_run_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True, color=default_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="DejaVu Sans Mono", size=max(default_size - 0.6, 8), color=TEAL)
        else:
            run = paragraph.add_run(math_to_plain(token))
            set_run_font(run, name="Cambria Math", size=default_size, color=default_color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=default_size, color=default_color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, name=HEADING_FONT, size=8.5, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    heading_specs = {
        "Title": (24, BLUE, 0, 18),
        "Heading 1": (18, BLUE, 12, 8),
        "Heading 2": (14, TEAL, 9, 5),
        "Heading 3": (11.5, INK, 7, 3),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = HEADING_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "DejaVu Sans Mono"
        code_style.font.size = Pt(8.8)
        code_style.paragraph_format.left_indent = Cm(0.5)
        code_style.paragraph_format.right_indent = Cm(0.5)
        code_style.paragraph_format.space_before = Pt(2)
        code_style.paragraph_format.space_after = Pt(2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(header, "MSM 反无人机多无人机协同拦截体系", 8.3, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "MSM 项目组  ·  ", 8.3, MUTED)
    add_page_number(footer)

    props = doc.core_properties
    props.title = "MSM 反无人机多无人机协同拦截体系"
    props.subject = "体系架构、关键技术、指控与资源规划、当前进展"
    props.author = "MSM 项目组"
    props.keywords = "C-UAS, 多无人机, 目标分配, AirSim, 体系方案"


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MSM 反无人机多无人机协同拦截体系")
    set_run_font(run, name=HEADING_FONT, size=25, bold=True, color=BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("体系方案与阶段进展")
    set_run_font(r2, name=HEADING_FONT, size=30, bold=True, color=INK)
    doc.add_paragraph()
    rule = doc.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    rule.autofit = False
    cell = rule.cell(0, 0)
    cell.width = Cm(12)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run("体系架构  ·  关键技术  ·  主要方案  ·  当前进展")
    set_run_font(cr, name=HEADING_FONT, size=12, bold=True, color=BLUE)
    for _ in range(7):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(meta, "阶段报告", 13, TEAL)
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(date, "2026 年 7 月", 12, MUTED)
    boundary = doc.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(boundary, "科研仿真与阶段评估用途", 9.5, RED)
    doc.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(rows):
        tr_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if r_idx == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.12
            paragraph.paragraph_format.space_after = Pt(0)
            if r_idx == 0:
                set_cell_shading(cell, BLUE)
                add_inline(paragraph, text, 8.3, WHITE)
                for run in paragraph.runs:
                    run.bold = True
            else:
                if r_idx % 2 == 0:
                    set_cell_shading(cell, "F3F6F9")
                add_inline(paragraph, text, 8.1, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_equation(doc: Document, tex: str, equation_index: int) -> None:
    tex = tex.strip()
    image_path = ASSETS / f"equation_{equation_index:02d}.png"
    try:
        # Equations are supporting evidence, not page-filling illustrations.
        # Short formulas use a compact 6.2 cm width; longer state/covariance
        # expressions grow only as far as 11 cm.
        if len(tex) < 34:
            figure_width, image_width = 4.8, 6.2
        elif len(tex) < 72:
            figure_width, image_width = 7.0, 8.6
        else:
            figure_width, image_width = 9.0, 11.0
        fig = plt.figure(figsize=(figure_width, 0.48))
        fig.text(0.5, 0.5, f"${tex}$", ha="center", va="center", fontsize=14, color=_hex(INK))
        fig.savefig(image_path, dpi=220, bbox_inches="tight", pad_inches=0.08, transparent=True)
        plt.close(fig)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run().add_picture(str(image_path), width=Cm(image_width))
    except Exception:
        plt.close("all")
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run(math_to_plain(tex))
        set_run_font(run, name="Cambria Math", size=11.5, color=INK)


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    paragraph.paragraph_format.first_line_indent = Cm(0)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF2F6")
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(code_lines))
    set_run_font(run, name="DejaVu Sans Mono", size=8.6, color=INK)


def build_docx() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    # The Markdown title block is represented by the designed cover.
    first_break = next(i for i, line in enumerate(lines) if line.strip() == "<!-- PAGEBREAK -->")
    idx = first_break + 1
    equation_index = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue
        if line == "<!-- PAGEBREAK -->":
            doc.add_page_break()
            idx += 1
            continue
        if line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                caption, relative = match.groups()
                image_path = (HERE / relative).resolve()
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.add_run().add_picture(str(image_path), width=Cm(16.0))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.first_line_indent = Cm(0)
                run = cp.add_run(f"图  {caption}")
                set_run_font(run, name=HEADING_FONT, size=8.8, color=MUTED)
            idx += 1
            continue
        if line == "$$":
            eq_lines = []
            idx += 1
            while idx < len(lines) and lines[idx].strip() != "$$":
                eq_lines.append(lines[idx].strip())
                idx += 1
            equation_index += 1
            add_equation(doc, " ".join(eq_lines), equation_index)
            idx += 1
            continue
        if line.startswith("```"):
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(doc, code_lines)
            idx += 1
            continue
        if line.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
            paragraph = doc.add_paragraph(style=style)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, title, {1: 18, 2: 14, 3: 11.5}.get(level, 10.5),
                       BLUE if level == 1 else TEAL if level == 2 else INK)
            idx += 1
            continue
        if line.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.right_indent = Cm(0.3)
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(4)
            p_pr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_AMBER)
            p_pr.append(shd)
            add_inline(paragraph, line.lstrip("> "), 9.4, INK)
            idx += 1
            continue
        list_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", line)
        if list_match:
            marker, content = list_match.groups()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.55)
            paragraph.paragraph_format.space_after = Pt(1.5)
            prefix = "•" if marker in {"-", "*"} else marker
            prefix_run = paragraph.add_run(f"{prefix}  ")
            set_run_font(prefix_run, name=HEADING_FONT, size=10.1, bold=True, color=BLUE)
            add_inline(paragraph, content, 10.1, INK)
            idx += 1
            continue
        if line == "---":
            paragraph = doc.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), MID_BLUE)
            borders.append(bottom)
            p_pr.append(borders)
            idx += 1
            continue

        # Join adjacent plain lines into one paragraph. Markdown uses blank lines
        # between paragraphs, so this primarily protects wrapped source text.
        parts = [line]
        look = idx + 1
        while look < len(lines) and lines[look].strip() and not re.match(
            r"^(#|>|\||```|\$\$|!\[|<!-- PAGEBREAK -->|[-*]\s+|\d+\.\s+|---$)", lines[look].strip()
        ):
            parts.append(lines[look].strip())
            look += 1
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, " ".join(parts), 10.5, INK)
        idx = look

    doc.save(DOCX_PATH)


def build_pdf() -> None:
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(HERE), str(DOCX_PATH)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    if not PDF_PATH.exists():
        raise RuntimeError(f"PDF was not generated: {PDF_PATH}")


def main() -> None:
    build_figures()
    build_docx()
    build_pdf()
    print(f"Markdown: {SOURCE}")
    print(f"DOCX: {DOCX_PATH}")
    print(f"PDF: {PDF_PATH}")


if __name__ == "__main__":
    main()
