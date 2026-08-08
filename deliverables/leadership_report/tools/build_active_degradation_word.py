#!/usr/bin/env python3
"""Build the active-degradation leadership report as a Word document."""

from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from build_search_visual_assignment_word import (
    BLUE,
    BODY_FONT,
    HEADING_FONT,
    INK,
    MUTED,
    TEAL,
    _add_inline as add_inline,
    _add_page_number as add_page_number,
    _configure_section as configure_section,
    _configure_styles as configure_styles,
    _set_run_font as set_run_font,
)


REPORT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = REPORT_ROOT / "ACTIVE_DEGRADATION_SECONDARY_DISTRIBUTED_ASSIGNMENT_REPORT_CN.md"
OUTPUT = REPORT_ROOT / "ACTIVE_DEGRADATION_SECONDARY_DISTRIBUTED_ASSIGNMENT_REPORT_CN.docx"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def add_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(header, "主动降级与分级目标分配", size=8.5, color=MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "MSM 项目组  ·  ", size=8.5, color=MUTED)
    add_page_number(footer)


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        title.add_run("主动降级与分级目标分配方案"),
        size=27,
        bold=True,
        color=INK,
        east_asia=HEADING_FONT,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        subtitle.add_run("中心节点、二级侦察节点与完全分布式协同"),
        size=15,
        bold=True,
        color=BLUE,
        east_asia=HEADING_FONT,
    )
    for _ in range(9):
        document.add_paragraph()
    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(owner.add_run("MSM 项目组"), size=12, color=INK)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("2026 年 8 月"), size=11, color=MUTED)
    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(boundary.add_run("科研仿真与技术论证材料"), size=9.5, color=TEAL)


def add_image(document: Document, alt: str, relative: str, number: int) -> None:
    image_path = (SOURCE.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.1))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(5)
    caption_text = re.sub(r"^图\s*\d+\s*", "", alt).strip()
    set_run_font(caption.add_run(f"图 {number}  {caption_text}"), size=9, color=MUTED)


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    add_cover(document)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(content_section)
    add_header_footer(content_section)

    image_number = 0
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## ") or line.startswith("### "):
            level = 1 if line.startswith("## ") and not line.startswith("### ") else 2
            title = line[3:].strip() if level == 1 else line[4:].strip()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            if title == "九、后续工作":
                paragraph.paragraph_format.page_break_before = True
            add_inline(
                paragraph,
                title,
                size=16 if level == 1 else 13.5,
                color=BLUE if level == 1 else TEAL,
            )
            index += 1
            continue
        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            image_number += 1
            add_image(document, image_match.group(1), image_match.group(2), image_number)
            index += 1
            continue
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            note = document.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note.paragraph_format.first_line_indent = Cm(0)
            note.paragraph_format.space_after = Pt(6)
            add_inline(note, line[1:-1], size=9, color=MUTED)
            index += 1
            continue
        list_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if list_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.55)
            paragraph.paragraph_format.space_after = Pt(4)
            set_run_font(
                paragraph.add_run(f"{list_match.group(1)}.  "),
                size=11,
                bold=True,
                color=BLUE,
                east_asia=HEADING_FONT,
            )
            add_inline(paragraph, list_match.group(2), size=11)
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate or candidate.startswith(("#", "![", "*")):
                break
            if re.match(r"^\d+\.\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, " ".join(parts), size=11)
        index = lookahead

    properties = document.core_properties
    properties.title = "主动降级与分级目标分配方案"
    properties.subject = "中心节点、二级侦察节点和完全分布式协同方案"
    properties.author = "MSM 项目组"
    properties.keywords = "主动降级, 二级节点, 目标重分配, 分布式协同"
    document.save(OUTPUT)


def validate_document() -> dict[str, int]:
    document = Document(OUTPUT)
    with ZipFile(OUTPUT) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("generated DOCX archive is damaged")
        images = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
    if len(images) != 5:
        raise RuntimeError(f"expected 5 embedded images, found {len(images)}")
    if len(document.paragraphs) < 50:
        raise RuntimeError("generated document is unexpectedly short")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for required in (
        "首次主动降级",
        "二级节点目标重新分配",
        "再次主动降级",
        "分布式分配",
        "降级后的算法与技术路径",
        "二级节点分配算法",
        "分布式协商算法",
        "当前进展",
    ):
        if required not in text:
            raise RuntimeError(f"missing section: {required}")
    return {
        "paragraphs": len(document.paragraphs),
        "images": len(images),
        "bytes": OUTPUT.stat().st_size,
    }


def main() -> None:
    build_document()
    metrics = validate_document()
    print(
        f"{OUTPUT.name}: paragraphs={metrics['paragraphs']}, "
        f"images={metrics['images']}, bytes={metrics['bytes']}"
    )


if __name__ == "__main__":
    main()
