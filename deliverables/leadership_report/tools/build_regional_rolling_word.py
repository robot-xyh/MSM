#!/usr/bin/env python3
"""Build the regional rolling scheduling report as a styled Word document."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from build_visual_registration_word import (
    BLUE,
    BODY_FONT,
    HEADING_FONT,
    INK,
    MUTED,
    TEAL,
    add_page_number,
    configure_section,
    configure_styles,
    set_run_font,
)


REPORT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = REPORT_ROOT / "REGIONAL_ROLLING_SCHEDULING_REINFORCEMENT_LEARNING_SOLUTION_CN.md"
OUTPUT = REPORT_ROOT / "REGIONAL_ROLLING_SCHEDULING_REINFORCEMENT_LEARNING_SOLUTION_CN.docx"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)")
TABLE_DIVIDER_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")


def add_inline(paragraph, value: str, *, size: float = 11, color: str = INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(value):
        if match.start() > position:
            set_run_font(paragraph.add_run(value[position : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color, east_asia=HEADING_FONT)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color)
            run.italic = True
        position = match.end()
    if position < len(value):
        set_run_font(paragraph.add_run(value[position:]), size=size, color=color)


def add_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(header, "强化学习区域滚动调度", size=8.5, color=MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "MSM 项目组  ·  ", size=8.5, color=MUTED)
    add_page_number(footer)


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        title.add_run("来袭预警后的区域滚动调度方案"),
        size=27,
        bold=True,
        color=INK,
        east_asia=HEADING_FONT,
    )
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        subtitle.add_run("大规模高动态来袭条件下的强化学习资源调度"),
        size=15,
        bold=True,
        color=BLUE,
        east_asia=HEADING_FONT,
    )
    for _ in range(9):
        document.add_paragraph()
    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(owner.add_run("MSM 项目组"), size=12, color=INK)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("2026 年 8 月"), size=11, color=MUTED)
    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(boundary.add_run("科研仿真与技术论证材料"), size=9.5, color=TEAL)


def add_image(document: Document, alt: str, relative: str, number: int) -> None:
    image_path = (SOURCE.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.0))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(5)
    set_run_font(caption.add_run(f"图 {number}  {alt}"), size=9, color=MUTED)


def _equation_text(raw: str) -> str:
    if "D_r(H)" in raw:
        return r"$D_r(H)=\sum_j p_{jr}(H)\,p_j^{\mathrm{exist}}\,k_j\,s_j$"
    if "R_t=" in raw:
        return r"$R_t=w_hH_t+w_cC_t+w_bB_t-w_gG_t-w_eE_t-w_sS_t-w_fF_t$"
    return f"${raw.strip()}$"


def add_equation(document: Document, raw: str) -> None:
    figure = plt.figure(figsize=(7.5, 0.7), dpi=220)
    figure.patch.set_alpha(0)
    figure.text(0.5, 0.5, _equation_text(raw), ha="center", va="center", fontsize=18, color="#18212B")
    stream = BytesIO()
    figure.savefig(stream, format="png", dpi=220, transparent=True, bbox_inches="tight", pad_inches=0.08)
    plt.close(figure)
    stream.seek(0)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run().add_picture(stream, width=Cm(12.0))


def _table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_table(document: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            text = values[column_index] if column_index < len(values) else ""
            run = paragraph.add_run(text)
            set_run_font(
                run,
                size=9.2,
                bold=row_index == 0,
                color="FFFFFF" if row_index == 0 else INK,
                east_asia=HEADING_FONT if row_index == 0 else BODY_FONT,
            )
            _shade_cell(cell, "1F5F99" if row_index == 0 else ("F4F7FA" if row_index % 2 == 0 else "FFFFFF"))
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    add_cover(document)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(content_section)
    add_header_footer(content_section)

    image_number = 0
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## ") or line.startswith("### "):
            level = 1 if line.startswith("## ") and not line.startswith("### ") else 2
            title = line[3:].strip() if level == 1 else line[4:].strip()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            if title in {"六、训练与验证", "九、当前基础与需协调事项"}:
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, title, size=16 if level == 1 else 13.5, color=BLUE if level == 1 else TEAL)
            index += 1
            continue

        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            image_number += 1
            add_image(document, image_match.group(1), image_match.group(2), image_number)
            index += 1
            continue

        if line == r"\[":
            equation_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != r"\]":
                equation_lines.append(lines[index].strip())
                index += 1
            if index >= len(lines):
                raise RuntimeError("unterminated display equation")
            add_equation(document, " ".join(equation_lines))
            index += 1
            continue

        if line.startswith("|") and index + 1 < len(lines) and TABLE_DIVIDER_RE.fullmatch(lines[index + 1].strip()):
            rows = [_table_cells(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_table_cells(lines[index]))
                index += 1
            add_table(document, rows)
            continue

        if line.startswith("**") and line.endswith("**") and line.count("**") == 2:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, line, size=11, color=BLUE)
            index += 1
            continue

        list_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        bullet_match = re.match(r"^-\s+(.*)$", line)
        if list_match or bullet_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.72)
            paragraph.paragraph_format.first_line_indent = Cm(-0.50)
            paragraph.paragraph_format.space_after = Pt(3)
            marker = f"{list_match.group(1)}.  " if list_match else "•  "
            content = list_match.group(2) if list_match else bullet_match.group(1)
            set_run_font(paragraph.add_run(marker), size=10.5, bold=True, color=BLUE, east_asia=HEADING_FONT)
            add_inline(paragraph, content, size=10.5)
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate or candidate.startswith(("#", "![", "|", r"\[")):
                break
            if re.match(r"^(?:\d+\.|-)\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, " ".join(parts), size=11)
        index = lookahead

    properties = document.core_properties
    properties.title = "来袭预警后的区域滚动调度方案"
    properties.subject = "大规模高动态来袭条件下的强化学习资源调度"
    properties.author = "MSM 项目组"
    properties.keywords = "区域滚动调度, 强化学习, 图神经网络, 资源调度"
    document.save(OUTPUT)


def validate_document() -> dict[str, int]:
    document = Document(OUTPUT)
    with ZipFile(OUTPUT) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("generated DOCX archive is damaged")
        images = [name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/")]
    if len(images) != 3:
        raise RuntimeError(f"expected 3 figures, found {len(images)} images")
    if len(document.tables) != 2:
        raise RuntimeError(f"expected 2 tables, found {len(document.tables)}")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for required in (
        "大规模高动态来袭条件下的强化学习资源调度",
        "总体方案",
        "强化学习区域调度",
        "算法结构",
        "安全与降级",
        "训练与验证",
        "需协调事项",
    ):
        if required not in text:
            raise RuntimeError(f"missing content: {required}")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(images),
        "bytes": OUTPUT.stat().st_size,
    }


def main() -> None:
    build_document()
    metrics = validate_document()
    print(
        f"{OUTPUT.name}: paragraphs={metrics['paragraphs']}, tables={metrics['tables']}, "
        f"images={metrics['images']}, bytes={metrics['bytes']}"
    )


if __name__ == "__main__":
    main()
