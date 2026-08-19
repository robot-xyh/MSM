#!/usr/bin/env python3
"""Integrate the reviewed 4.1, 4.4 and 4.5 material into the scheme template."""

from __future__ import annotations

import argparse
import copy
import hashlib
import os
import shutil
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
DEFAULT_DOCUMENT = ROOT / "deliverables/leadership_report/方案模板.docx"
ASSET_DIR = ROOT / "deliverables/leadership_report/assets/scheme_template_material"
DUAL_OPTICAL_REPORT = ROOT / "deliverables/leadership_report/双光电多目标轨迹配准与交汇定位试验报告_CN.md"
DUAL_OPTICAL_EVIDENCE = ROOT / "deliverables/leadership_report/双光电多目标轨迹配准与交汇定位试验报告_EVIDENCE.json"
COOPERATIVE_SEARCH_REPORT = ROOT / "deliverables/leadership_report/协同搜索试验报告_CN.md"
TERMINAL_ASSOCIATION_REPORT = ROOT / "deliverables/leadership_report/末端目标配准试验报告_CN.md"
DUAL_OPTICAL_REPORT_ASSET_DIR = ROOT / "deliverables/leadership_report/assets/dual_optical_registration_report"
CENTER_TERMINAL_REPORT_ASSET_DIR = ROOT / "deliverables/leadership_report/assets/center_terminal_split_reports"
DUAL_OPTICAL_DIR = ROOT / "research_modules/independent_experiments/dual_optical_online_benchmark"
DUAL_OPTICAL_RAW_DIR = (
    ROOT
    / "research_modules/independent_experiments/dual_optical_40target/outputs/airsim_seed_20260810_run11/figures"
)
CENTER_TERMINAL_DIR = (
    ROOT / "research_modules/independent_experiments/center_terminal_cv_campaign/outputs"
)

BODY_FONT = "仿宋"
HEADING_FONT = "黑体"
CAPTION_FONT = "楷体"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def verify_reviewed_sources() -> None:
    """Refuse to build from missing or superseded source reports."""

    required_markers = {
        DUAL_OPTICAL_REPORT: (
            "### 3.1 360度理想单站条件",
            "### 3.2 360度实际单站航迹",
            "### 3.3 180度扫描",
            "360度理想单站条件下20、40、60目标的三算法矩阵没有现成记录",
        ),
        COOPERATIVE_SEARCH_REPORT: (
            "### 1.2 搜索单元",
            "### 1.3 搜索收益和一一分配",
            "20目标/8机",
            "40目标/50机",
        ),
        TERMINAL_ASSOCIATION_REPORT: (
            "### 1.2 中心线索状态外推",
            "### 1.6 机间几何代价、图网络和目标簇",
            "### 3.1 中心交接过程与结果",
            "### 3.2 机间关联过程与结果",
        ),
    }
    for path, markers in required_markers.items():
        if not path.is_file():
            raise FileNotFoundError(path)
        text = path.read_text(encoding="utf-8")
        missing = [marker for marker in markers if marker not in text]
        if missing:
            raise RuntimeError(f"reviewed source changed: {path.name}; missing {missing}")
    if not DUAL_OPTICAL_EVIDENCE.is_file():
        raise FileNotFoundError(DUAL_OPTICAL_EVIDENCE)


def outside_target_hash(document: Document) -> str:
    """Hash body nodes outside 4.1, 4.4 and 4.5 for preservation checks."""

    boundaries = (
        (find_paragraph(document, "4.1 侦察的想法"), find_paragraph(document, "4.2 火指控的想法")),
        (find_paragraph(document, "4.4 拦截区域搜索"), find_paragraph(document, "4.5 群对群目标配准")),
        (find_paragraph(document, "4.5 群对群目标配准"), find_paragraph(document, "4.3.6 主动降级与分级目标分配")),
    )
    body = document._element.body
    excluded: set[int] = set()
    for start, end in boundaries:
        start_index = body.index(start._p)
        end_index = body.index(end._p)
        excluded.update(range(start_index + 1, end_index))

    digest = hashlib.sha256()
    for index, element in enumerate(body.iterchildren()):
        if index not in excluded and element.tag != qn("w:sectPr"):
            digest.update(element.xml.encode("utf-8"))
    return digest.hexdigest()


def set_a4_page_size(document: Document) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)


def find_paragraph(document: Document, prefix: str):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def remove_between(start_paragraph, end_paragraph) -> None:
    node = start_paragraph._p.getnext()
    while node is not None and node is not end_paragraph._p:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node
    if node is None:
        raise RuntimeError("section boundary was not found in the same document body")


def style_samples(document: Document):
    by_name = {}
    for paragraph in document.paragraphs:
        by_name.setdefault(paragraph.style.name, paragraph.style)
    required = ("Normal", "Heading 3", "Heading 4")
    missing = [name for name in required if name not in by_name]
    if missing:
        raise RuntimeError(f"missing paragraph styles: {missing}")
    return by_name


class SectionWriter:
    def __init__(self, document: Document, anchor, styles):
        self.document = document
        self.anchor = anchor
        self.styles = styles

    def _place(self, element):
        self.anchor._p.addprevious(element)

    def paragraph(
        self,
        text: str,
        *,
        style: str = "Normal",
        bold_prefix: str | None = None,
        align: WD_ALIGN_PARAGRAPH | None = None,
        first_line: bool = True,
        keep_with_next: bool = False,
    ):
        paragraph = self.document.add_paragraph()
        paragraph.style = self.styles[style]
        if bold_prefix and text.startswith(bold_prefix):
            run = paragraph.add_run(bold_prefix)
            run.bold = True
            paragraph.add_run(text[len(bold_prefix) :])
        else:
            paragraph.add_run(text)
        paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = Pt(28)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(28) if first_line and style == "Normal" else Pt(0)
        paragraph.paragraph_format.keep_with_next = keep_with_next
        for run in paragraph.runs:
            run.font.name = HEADING_FONT if style.startswith("Heading") else BODY_FONT
            run._element.get_or_add_rPr().rFonts.set(
                qn("w:eastAsia"), HEADING_FONT if style.startswith("Heading") else BODY_FONT
            )
            run.font.size = Pt(14)
        self._place(paragraph._p)
        return paragraph

    def heading(self, text: str, level: int = 3):
        style = "Heading 3" if level == 3 else "Heading 4"
        paragraph = self.paragraph(
            text,
            style=style,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_line=False,
            keep_with_next=True,
        )
        for run in paragraph.runs:
            run.bold = level == 3
        paragraph.paragraph_format.space_before = Pt(8 if level == 3 else 4)
        paragraph.paragraph_format.space_after = Pt(2)
        return paragraph

    def formula(self, text: str):
        paragraph = self.document.add_paragraph()
        paragraph.style = self.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = Pt(24)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.font.name = "Cambria Math"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Cambria Math")
        run.font.size = Pt(12.5)
        self._place(paragraph._p)
        return paragraph

    def image(self, filename: str, caption: str, *, width: float = 6.15):
        self.image_path(ASSET_DIR / filename, caption, width=width)

    def image_path(self, path: Path, caption: str, *, width: float = 6.15):
        if not path.is_file():
            raise FileNotFoundError(path)
        paragraph = self.document.add_paragraph()
        paragraph.style = self.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run().add_picture(str(path), width=Inches(width))
        self._place(paragraph._p)
        self.caption(caption)

    def cloned_image(self, paragraph_element, caption: str):
        clone = copy.deepcopy(paragraph_element)
        properties = clone.get_or_add_pPr()
        style = properties.find(qn("w:pStyle"))
        if style is None:
            style = OxmlElement("w:pStyle")
            properties.insert(0, style)
        style.set(qn("w:val"), self.styles["Normal"].style_id)
        justification = properties.find(qn("w:jc"))
        if justification is None:
            justification = OxmlElement("w:jc")
            properties.append(justification)
        justification.set(qn("w:val"), "center")
        self._place(clone)
        self.caption(caption)

    def caption(self, text: str):
        paragraph = self.document.add_paragraph()
        paragraph.style = self.styles["Normal"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = Pt(20)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run(text)
        run.font.name = CAPTION_FONT
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CAPTION_FONT)
        run.font.size = Pt(10.5)
        self._place(paragraph._p)

    def table(
        self,
        headers: Sequence[str],
        rows: Iterable[Sequence[str]],
        *,
        widths: Sequence[float] | None = None,
        font_size: float = 9.5,
    ):
        row_data = [list(map(str, row)) for row in rows]
        table = self.document.add_table(rows=1, cols=len(headers))
        table.autofit = False
        table.alignment = 1
        self._set_table_borders(table)
        for column, value in enumerate(headers):
            self._set_cell(table.rows[0].cells[column], value, header=True, font_size=font_size)
        for values in row_data:
            if len(values) != len(headers):
                raise ValueError(f"table row has {len(values)} cells, expected {len(headers)}")
            cells = table.add_row().cells
            for column, value in enumerate(values):
                self._set_cell(cells[column], value, header=False, font_size=font_size)
        if widths:
            for row in table.rows:
                for column, width in enumerate(widths):
                    row.cells[column].width = Inches(width)
        for row in table.rows:
            self._prevent_row_split(row)
        self._repeat_header(table.rows[0])
        self._place(table._tbl)
        spacer = self.document.add_paragraph()
        spacer.style = self.styles["Normal"]
        spacer.paragraph_format.space_after = Pt(2)
        spacer.paragraph_format.line_spacing = Pt(8)
        self._place(spacer._p)
        return table

    @staticmethod
    def _set_table_borders(table) -> None:
        properties = table._tbl.tblPr
        borders = properties.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            properties.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "7F8C9A")
            borders.append(element)

    @staticmethod
    def _set_cell(cell, value: str, *, header: bool, font_size: float) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = Pt(15)
        run = paragraph.add_run(value)
        run.font.name = "黑体" if header else "宋体"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体" if header else "宋体")
        run.font.size = Pt(font_size)
        run.bold = header
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "274C77")
            cell._tc.get_or_add_tcPr().append(shading)

    @staticmethod
    def _repeat_header(row) -> None:
        properties = row._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        properties.append(repeat)

    @staticmethod
    def _prevent_row_split(row) -> None:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))


def write_section_41(writer: SectionWriter) -> None:
    writer.heading("4.1.1 场景及存在的问题")
    writer.paragraph(
        "群目标侦察需要在较短时间内给出目标数量、空间位置和连续航迹。单台无源光电设备只能测得目标方位和俯仰，缺少可靠距离；逐个使用激光测距时，目标数量增加会直接拉长更新周期。按单目标测距0.2秒估算，100个目标完成一轮测距约需20秒，50米/秒目标在此期间可前进约1000米。该更新速度难以支撑后续区域搜索和拦截交接。"
    )
    writer.heading("4.1.2 存在的难点")
    writer.paragraph(
        "双站光电可以利用视线交会补足距离，但群目标条件下必须先把两站的局部航迹正确对上号。目标在远距离图像中只有少量像素，外观难以区分；两台云台按各自周期周扫，同一目标的观测时刻通常不同；目标交叉、漏检、虚警和云台姿态偏差还会造成多条候选关系。配错一条航迹后，交会位置、速度和后续目标编号都会随之错误。处理流程还要满足在线时限，不能把全部航迹组合直接交给学习模型。"
    )
    writer.heading("4.1.3 拟采用的方案")
    writer.paragraph(
        "采用“单站先成轨、双站再配准、配准后定位”的处理顺序。两站先把一次扫过中的连续检测合并成扫描片段，再把相邻扫描周期的片段连接为本站局部航迹。双站处理按拍摄时刻进行时间对齐，利用共面性、极线关系和运动连续性排除不可能组合。图神经网络只对保留下来的候选关系评分，匈牙利算法形成一一对应；最近三圈中至少两圈保持同一关系后，双站视线才进入位置和速度联合拟合。"
    )
    writer.image(
        "02_gnn_matching_process.png",
        "图4.1-1  双站光电配准与定位流程。几何筛选负责确定候选范围，图神经网络只在候选范围内排序。",
    )

    writer.heading("4.1.4 单站局部航迹")
    writer.paragraph(
        "光电云台连续扫过目标时，同一目标会在若干相邻图像中重复出现。系统先按拍摄时间、角位置和运动方向把这些检测合并成一个扫描片段，避免把一次扫过中的多帧检测重复计成多个目标。扫描片段保存检测框中心、拍摄时刻、云台姿态、目标单位视线和测量质量，不依赖检测框面积判断身份。"
    )
    writer.paragraph(
        "相邻扫描周期之间按角位置和角速度预测连接片段，形成本站局部航迹。航迹记录观测次数、持续时间、最近命中比例、方位和俯仰变化率、方向误差以及漏检情况。短时漏检时可在限定窗口内保持；超过保持时间后停止发布，防止长期外推形成虚假目标。两站的局部编号各自独立，不能直接作为跨站身份。"
    )

    writer.heading("4.1.5 时间对齐和几何筛选")
    writer.paragraph(
        "双站周扫通常不能在同一时刻看到同一目标。算法以图像拍摄时间为准，把两条局部航迹外推到共同参考时刻，再比较空间关系。消息到达时间单独用于评估通信延迟，不参与视线几何计算。像素坐标经相机内参、镜头畸变、云台角度和站址姿态换算为世界坐标系中的单位视线。"
    )
    writer.formula("r₍共面₎ = |bᵀ(d_A × d_B)| / (‖b‖·‖d_A × d_B‖ + ε)")
    writer.paragraph(
        "式中，b为两站基线，d_A和d_B为两站单位视线。同一目标对应的两条视线应与基线接近共面。残差明显超限的组合直接排除；接近门限的组合继续比较多时刻残差变化、观测时间重叠、角速度差、视线交会夹角、重投影误差和拟合速度。视线近似平行或目标沿不利方向排列时，系统保留多个候选，不强行确定身份。"
    )
    writer.image(
        "01_epipolar_geometry.png",
        "图4.1-2  共面性和极线筛选原理。正确候选的双站视线接近同一极平面，明显偏离的组合提前排除。",
    )

    writer.heading("4.1.6 候选关系图和图神经网络")
    writer.paragraph(
        "几何筛选后，把A站和B站局部航迹分别放在关系图两侧。每条航迹是一个节点，只有通过时间和几何条件的两条航迹之间才建立候选边。20条A站航迹和20条B站航迹理论上有400种组合，几何门控先删除大部分不可能组合，图神经网络只处理剩余的稀疏候选图。"
    )
    writer.image(
        "06_candidate_graph_assignment.png",
        "图4.1-3  候选关系图和一一配准。左侧保留有物理可能的候选边，右侧为整体分配后的对应关系。",
    )
    writer.paragraph(
        "节点信息反映单条航迹是否稳定，包括航迹长度、连续性、角运动、方向误差和漏检比例；边信息反映两条航迹能否由同一目标产生，包括共面性残差、时间重叠、重投影误差、视线夹角、拟合速度和运动一致性。图神经网络进行两轮信息交换，使每条候选边在评分时能够同时考虑周围竞争关系。例如一条B站航迹与两条A站航迹都接近时，网络会结合这两条A站航迹的其他候选关系调整排序。"
    )
    writer.image(
        "03_gnn_message_passing.png",
        "图4.1-4  图神经网络的信息交换过程。评分同时利用候选本身和相邻候选的竞争关系。",
    )
    writer.paragraph(
        "网络输出0至1之间的同目标评分。该评分不直接生成目标编号，也不参与三角定位。候选评分与几何代价共同形成代价矩阵，匈牙利算法从全部候选中选择总体代价较小的一一组合，并为证据不足的航迹保留“暂不匹配”选项。单圈结果只作为候选，连续多个扫描周期一致后才发布稳定关系。"
    )
    writer.image(
        "04_gnn_assignment_example.png",
        "图4.1-5  候选评分和匈牙利一一分配。证据不足的航迹保持未匹配，避免弱关系被强制确认。",
    )

    writer.heading("4.1.7 双站联合定位")
    writer.paragraph(
        "配准确认后，将同一目标在多个时刻的双站视线一起用于位置和速度拟合。设目标在参考时刻的位置为p₀、速度为v，第k条观测的相机位置和单位视线分别为cₖ、dₖ，按视线垂直方向误差进行加权最小二乘估计："
    )
    writer.formula("min Σₖ ‖(I − dₖdₖᵀ)[p₀ + v(tₖ − t₀) − cₖ]‖²_Wₖ")
    writer.paragraph(
        "权重由方向测量误差、云台姿态质量和时间误差确定。多时刻联合拟合允许使用不同扫描时刻的观测，不要求两站逐帧同步。输出包括参考时刻位置、速度、协方差和拟合条件数。拟合结果还要重新投影到A、B站图像；重投影残差持续偏大时撤销结果，并检查配准、时间同步和相机标定。"
    )
    writer.image(
        "05_joint_fit_reprojection.png",
        "图4.1-6  双站多时刻联合拟合与重投影检查。几何退化时扩大误差范围，不发布虚假的高精度坐标。",
    )

    writer.heading("4.1.8 AirSim仿真验证")
    writer.heading("试验条件", level=4)
    writer.paragraph(
        "试验采用AirSim计算机视觉模式。两台固定光电节点横向间隔2千米，目标走廊位于双站前方约2千米。目标采用长度约3米的无人机网格，以50米/秒飞行，空间位置前后错列并存在轨迹交叉。相机分辨率为1280×1024，焦距300毫米。360度试验按2秒一圈连续周扫并观察6圈；180度试验按1秒一次单程扫描并观察12轮。三种方法统一保留共面筛选、匈牙利一一分配和连续确认。"
    )
    writer.image_path(
        DUAL_OPTICAL_RAW_DIR / "01_scene_geometry_3d.png",
        "图4.1-7  双站光电和群目标场景原图。两站从不同方向观察前后错列、航迹交叉的来袭目标。",
        width=4.8,
    )

    writer.heading("360度理想单站条件", level=4)
    writer.paragraph(
        "理想单站条件用于单独核对双站算法：先保证两台光电内部的目标编号全部正确，再比较20、40、60目标下几何方法、图神经网络和增强型图神经网络的跨站配准。该九项矩阵尚未形成机器记录，不能用180度试验或单站正确子集替代。表中的待测试不代表结果为零。"
    )
    writer.table(
        ("目标数", "方法", "最后一圈关联精度", "最后一圈目标覆盖度", "处理耗时P95", "证据状态"),
        (
            ("20", "几何方法", "待测试", "待测试", "待测试", "无机器记录"),
            ("20", "图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
            ("20", "增强型图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
            ("40", "几何方法", "待测试", "待测试", "待测试", "无机器记录"),
            ("40", "图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
            ("40", "增强型图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
            ("60", "几何方法", "待测试", "待测试", "待测试", "无机器记录"),
            ("60", "图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
            ("60", "增强型图神经网络", "待测试", "待测试", "待测试", "无机器记录"),
        ),
        widths=(0.55, 1.2, 1.25, 1.25, 1.05, 1.2),
        font_size=8.4,
    )

    writer.heading("360度实际单站航迹", level=4)
    writer.paragraph(
        "第一组对照复用20目标、5个测试场景的同一批匿名单站航迹，只统计第6圈。无干扰时几何方法的精度和覆盖度最高；轻度干扰时增强型图神经网络的两项指标最高；基础图神经网络处理时间最短。几何方法在轻度干扰下有1个场景超过1000毫秒期限。"
    )
    writer.table(
        ("方法", "条件", "最后一圈精度", "最后一圈覆盖度", "处理耗时P95", "证据状态"),
        (
            ("几何方法", "无干扰", "98.8%", "85.0%", "867.9毫秒", "封存回放，诊断"),
            ("几何方法", "轻度干扰", "85.5%", "47.0%", "997.9毫秒", "诊断；1/5超时"),
            ("图神经网络", "无干扰", "91.5%", "75.0%", "75.2毫秒", "封存回放，诊断"),
            ("图神经网络", "轻度干扰", "77.2%", "61.0%", "88.5毫秒", "封存回放，诊断"),
            ("增强型图神经网络", "无干扰", "92.2%", "71.0%", "241.2毫秒", "封存回放，诊断"),
            ("增强型图神经网络", "轻度干扰", "89.0%", "65.0%", "288.5毫秒", "封存回放，诊断"),
        ),
        widths=(1.35, 0.8, 1.0, 1.0, 1.0, 1.35),
        font_size=8.4,
    )
    writer.image_path(
        DUAL_OPTICAL_REPORT_ASSET_DIR / "11_360_clean_light_route_comparison.png",
        "图4.1-8  二十目标360度周扫第6圈的无干扰与轻度干扰对照。",
    )

    writer.paragraph(
        "第二组冻结各规模模型并使用5个随机场景，加入无干扰、轻度、中度和重度四档条件。只有图神经网络形成20、40、60目标的完整分规模记录。20目标几何方法仅无干扰形成结果，另外三档均为5/5超时；40和60目标几何方法未开展，增强型图神经网络在本批次全部未开展。"
    )
    writer.table(
        ("目标数", "条件", "图网络精度", "图网络覆盖度", "处理耗时P95", "其他路线状态"),
        (
            ("20", "无干扰", "84.0%", "63.0%", "78.3毫秒", "几何85.7%/78.0%；增强型未开展"),
            ("20", "轻度", "82.2%", "60.0%", "96.7毫秒", "几何5/5超时；增强型未开展"),
            ("20", "中度", "76.6%", "49.0%", "121.8毫秒", "几何5/5超时；增强型未开展"),
            ("20", "重度", "69.4%", "43.0%", "171.8毫秒", "几何5/5超时；增强型未开展"),
            ("40", "无干扰", "91.7%", "60.5%", "193.3毫秒", "几何、增强型未开展"),
            ("40", "轻度", "80.0%", "50.0%", "204.2毫秒", "几何、增强型未开展"),
            ("40", "中度", "76.1%", "35.0%", "229.9毫秒", "几何、增强型未开展"),
            ("40", "重度", "67.1%", "27.5%", "270.3毫秒", "几何、增强型未开展"),
            ("60", "无干扰", "94.3%", "55.0%", "305.2毫秒", "几何、增强型未开展"),
            ("60", "轻度", "77.3%", "39.7%", "337.6毫秒", "几何、增强型未开展"),
            ("60", "中度", "77.6%", "40.3%", "366.2毫秒", "几何、增强型未开展"),
            ("60", "重度", "72.3%", "28.7%", "403.9毫秒", "几何、增强型未开展"),
        ),
        widths=(0.55, 0.65, 0.9, 0.95, 1.05, 2.05),
        font_size=8.0,
    )
    writer.image_path(
        DUAL_OPTICAL_REPORT_ASSET_DIR / "12_360_multiseed_cascade.png",
        "图4.1-9  三百六十度周扫在无干扰及轻、中、重度随机干扰下的结果。",
    )
    writer.paragraph(
        "随机干扰先造成单站断轨、错误重接和重复建轨，随后才表现为双站精度和覆盖度下降。两站没有形成对应的正确局部航迹时，跨站评分无法从后端补回；单站航迹基本完整但覆盖仍有损失的场景，还需继续校准跨站评分和连续确认。"
    )

    writer.heading("180度扫描", level=4)
    writer.paragraph(
        "180度扫描每秒形成一次关联结果，重访频率高于2秒一圈的360度周扫。下表每个单元格依次给出最后一轮精度、覆盖度和处理耗时P95。图神经网络六组均按时形成结果；几何方法六组全部超时；增强型图神经网络在20和40目标形成结果，60目标全部超时。"
    )
    writer.table(
        ("目标数", "条件", "几何方法", "图神经网络", "增强型图神经网络", "证据状态"),
        (
            ("20", "无干扰", "超时/2067.5毫秒", "97.8%/89.0%/100.4毫秒", "88.5%/77.0%/379.9毫秒", "正式"),
            ("20", "轻度", "超时/1950.6毫秒", "83.9%/73.0%/108.8毫秒", "88.8%/79.0%/401.3毫秒", "正式"),
            ("40", "无干扰", "超时/5174.2毫秒", "89.6%/81.5%/292.3毫秒", "77.5%/58.5%/972.7毫秒", "诊断"),
            ("40", "轻度", "超时/5022.3毫秒", "89.3%/71.0%/270.2毫秒", "79.1%/58.5%/955.0毫秒", "诊断"),
            ("60", "无干扰", "超时/9257.9毫秒", "95.8%/90.7%/509.5毫秒", "超时/1683.9毫秒", "诊断"),
            ("60", "轻度", "超时/9063.3毫秒", "94.6%/82.0%/504.9毫秒", "超时/1686.2毫秒", "诊断"),
        ),
        widths=(0.5, 0.65, 1.15, 1.55, 1.55, 0.65),
        font_size=7.8,
    )
    writer.image_path(
        DUAL_OPTICAL_REPORT_ASSET_DIR / "08_s180_selected_results.png",
        "图4.1-10  一百八十度扫描三种方法的最终一轮结果。",
    )
    writer.paragraph(
        "20目标轻度干扰下，180度图神经网络精度和覆盖度为83.9%和73.0%，360度同输入诊断批次为77.2%和61.0%。现有结果支持提高重访频率有利于维持关联，但两批数据不是同一种子、同一模型的单变量试验，差值不能全部归因于扫描范围。40和60目标180度结果还受单站航迹器验收状态限制。"
    )

    writer.heading("40目标交汇定位", level=4)
    writer.paragraph(
        "40目标理想演示形成37条双站关系，其中36条正确、1条错误。正确关系进入交汇定位后，平均位置误差为0.080米，95%位置误差不超过0.091米；平均速度误差为0.0081米/秒，95%速度误差不超过0.0197米/秒。"
    )
    writer.table(
        ("正确关系", "错误关系", "关联精度", "固定目标覆盖度", "平均/95%位置误差", "平均/95%速度误差"),
        (("36", "1", "97.3%", "90.0%", "0.080/0.091米", "0.0081/0.0197米/秒"),),
        widths=(0.75, 0.75, 0.8, 1.0, 1.35, 1.55),
        font_size=8.6,
    )
    writer.image_path(
        DUAL_OPTICAL_REPORT_ASSET_DIR / "10_ranging_reconstruction_and_error.png",
        "图4.1-11  正确配准关系的三维轨迹重建和误差分布。",
    )

    writer.heading("证据边界", level=4)
    writer.paragraph(
        "360度理想单站九项矩阵仍待测试；360度分规模随机干扰只有图神经网络形成完整记录，其他路线按未开展或超时列示；180度40和60目标结果属于诊断。40目标厘米级定位误差来自理想位姿、理想时间和仿真检测条件，不代表设备定位能力。识别输入来自AirSim检测接口，漏检和虚警按试验协议注入，后续还需加入安装测量误差、云台偏差、时间同步误差和真实检测中心偏差。"
    )
    writer.heading("4.1.9 输出内容")
    writer.paragraph(
        "双光电处理结果向后续环节提供目标临时编号、位置、速度、六维协方差、最近拍摄时间、消息到达时间、支持该结果的双站局部航迹、配准质量、定位质量、标定版本和有效期。结果过期、几何退化或配准冲突时，状态降为待确认，不发布精确点目标。"
    )


def write_section_44(writer: SectionWriter, original_region_image) -> None:
    writer.heading("4.4.1 场景及存在的问题")
    writer.paragraph(
        "中心双光电能够给出目标数量区间、主要来袭方向、粗略空间范围和信息时刻，部分目标可以形成带位置、速度和误差范围的源航迹，其余目标只能形成较宽的方位或概率区域。本节按中心关联精度80%、召回率80%的设计工况组织搜索。以100个真实目标为例，中心可能漏掉约20个目标，已发布线索也可能含有错误关联、重复航迹或虚假目标。拦截无人机只能依靠这些粗线索飞向目标附近，不能把一条中心线索直接视为一个已经确认的机载目标。"
    )
    writer.heading("4.4.2 存在的难点")
    writer.paragraph(
        "机载光电达到稳定识别像素后，可用于搜索和确认的时间只有数秒；距离继续缩短时目标像素增大，单个视场覆盖范围却迅速减小。中心漏检目标没有对应的指向线索，必须通过空档搜索发现。多架无人机若同时跟随同一条错误线索，会造成重复覆盖并留下未搜索区域。搜索还要处理云台转向、平台到达、视场重叠、连续复访和已锁定任务占用，资源数量不足时首先出现的是区域漏扫，后续配准无法补偿没有被观察的空域。"
    )
    writer.heading("4.4.3 拟采用的方案")
    writer.paragraph(
        "搜索任务面向目标概率区域和可能通道，不按中心航迹逐条固定分配无人机。中心向各机下发同一版本的概率区域、候选源航迹、目标数量区间和任务有效期。各机把区域划分为带方位、距离和高度层的搜索单元，交换已观察单元、未发现记录、候选短航迹、云台状态和信息时刻。线索搜索机检查高概率区域，空档搜索机覆盖中心未成轨通道，交叉复核机从第二方向观察疑似目标，机动预备机接续未完成单元。发现目标后转入连续跟踪，剩余搜索任务由其他无人机滚动接替。"
    )
    writer.image(
        "08_center_interceptor_search_architecture.png",
        "图4.4-1  中心粗线索下的协同搜索。多派出的无人机展开到不同区域和观察方向，角色随搜索结果滚动调整。",
    )
    writer.image_path(
        CENTER_TERMINAL_REPORT_ASSET_DIR / "01_search_flow.png",
        "图4.4-2  协同搜索计算流程。中心线索和空白走廊统一形成搜索单元，再进行滚动分配和连续确认。",
    )
    writer.cloned_image(
        original_region_image,
        "图4.4-3  区域重划示意。区域边界、重点通道和机动余量作为搜索单元划分依据。",
    )

    writer.heading("中心线索与搜索单元数量", level=4)
    writer.paragraph(
        "设真实目标数为N，正确中心线索数为T，中心发布线索总数为S。召回率T/N和精度T/S均按80%构造，因此T=0.8N、S=N；错误线索和中心漏掉的目标均为0.2N。正确线索和错误线索都生成线索搜索单元。中心漏检目标没有对应线索，另在来袭走廊铺设max(5，向上取整0.4N)个空白单元。"
    )
    writer.table(
        ("场景", "正确线索", "错误线索", "中心漏掉", "空白单元", "总搜索单元"),
        (
            ("20目标/8机", "16", "4", "4", "8", "28"),
            ("20目标/30机", "16", "4", "4", "8", "28"),
            ("40目标/50机", "32", "8", "8", "16", "56"),
        ),
        widths=(1.2, 0.9, 0.9, 0.9, 0.9, 1.0),
        font_size=9.0,
    )
    writer.paragraph(
        "每条线索先按p(t)=p₀+vΔt外推到规划时刻，再按三个方向的max(30米，3√Pᵢᵢ)确定搜索半宽。本轮位置标准差为1米，采用30米下限。空白单元分布在北向2500至3500米、东向负650至650米、高度负220至负70米的来袭走廊，概率取0.32，只表示需要观察，不表示已发现目标。"
    )

    writer.heading("4.4.4 交接位置和搜索窗口")
    writer.paragraph(
        "交接分四段进行。约5千米时下发概率区域、候选源航迹和预测交会点；约3千米时根据最新线索重划搜索责任，并将云台预置到预测方向；对3米级目标，约1.4千米是机载可见光达到10像素的建议稳定识别距离，此后形成机载局部航迹；约400至500米前完成连续确认，后续以稳定跟踪为主。上述距离是按当前方案参数计算的建议值，需通过多组仿真和真实光电试验标定。"
    )
    writer.table(
        ("距离区间", "50米/秒目标", "155米/秒目标", "主要工作"),
        (
            ("5千米至1.4千米", "约24.5秒", "约14.3秒", "飞向预测交会点、更新线索、云台预置"),
            ("3千米至1.4千米", "约10.9秒", "约6.3秒", "小范围预搜索和责任重划"),
            ("1.4千米至500米", "约6.1秒", "约3.6秒", "可见光搜索、局部成轨和多机复核"),
            ("500米至接近目标", "约3.4秒", "约2.0秒", "稳定跟踪和末段接管"),
        ),
        widths=(1.3, 1.1, 1.1, 2.7),
    )
    writer.paragraph(
        "扣除信息传输、云台转向和连续确认后，50米/秒目标约有4.5至5秒用于机载视觉搜索，155米/秒高速目标约有2至3秒。该窗口只支持在预测误差范围内执行有限观察动作。对更小的穿越机目标，达到10像素的距离更短，现有宽视场通道不适合承担大范围首次搜索。"
    )

    writer.heading("4.4.5 搜索区域和任务分配")
    writer.paragraph(
        "中心线索先按信息时刻外推到当前时刻，并根据位置、速度、协方差、目标机动能力和通信延迟形成三维概率区域。已成轨目标对应较小的预测椭圆，尚未稳定成轨的目标保留为较宽的方向扇区或空间块。概率区域按方位、距离和高度层划分为搜索单元，每个单元记录目标存在概率、预计目标数、最近观察时刻、观察质量、预计成像像素、可执行平台和任务有效期。"
    )
    writer.paragraph(
        "按视场尺寸划分扫描条带的计算继续作为搜索单元设计依据，执行平台采用拦截无人机自身双光云台。可见光全视场约19度乘11度，适合较远距离搜索；红外全视场约22度乘18度，进入有效成像距离后用于热目标复核和连续跟踪。可见光在1400米斜距的理想覆盖约469米乘270米，按20%重叠覆盖1平方千米单一高度层约需15个视场。距离缩短到500米后，同一区域约需104个视场，因此广域搜索应尽量前置，近距阶段只做预测区域凝视和局部补扫。"
    )
    writer.image(
        "09_interceptor_search_cell_allocation.png",
        "图4.4-4  拦截无人机搜索单元分配。单元根据目标概率、观察条件和平台任务状态滚动调整。",
    )
    writer.formula("Uᵢⱼ = 3pⱼ + 4Gᵢⱼ − 0.8C_转向 − C_到达 − 4C_重复，Gᵢⱼ = pⱼVᵢⱼQⱼ")
    writer.paragraph(
        "式中，p为单元概率，V为相机在700米观察距离上能够覆盖该单元的比例，Q=1/(1+n)反映既有覆盖次数n；其余三项分别扣除云台转向、到达观察位置和刚刚重复观察的代价。19度水平视场在700米处的名义覆盖宽度约234.3米，面积较大的空白单元一次不容易看全，需要后续复访。"
    )
    writer.paragraph(
        "有M台相机、K个有效单元时，先建立M×K收益矩阵，再增加M个空闲列，得到M×(K+M)矩阵，空闲收益为负0.05。匈牙利算法一次求解全部占用冲突，使每台相机每轮最多承担一个单元，每个真实单元最多分给一台相机。所有真实单元收益低于空闲项时，相机保持待命。"
    )

    writer.heading("4.4.6 未发现信息和连续复访")
    writer.paragraph(
        "一次扫过没有发现目标，只能降低该单元的目标概率。目标成像不足、云台未稳定、遮挡明显或扫描过快时，未发现记录的证据强度较低；连续多次高质量凝视仍未发现时，概率才明显下降。简化更新可写为："
    )
    writer.formula("p⁺ = p(1 − P_D) / (1 − pP_D)")
    writer.paragraph(
        "检测到候选后先形成局部短航迹，并把目标方向、拍摄时间和预测区域发给邻机。邻机形成第二视角或同机连续多帧确认后，再进入目标配准。该处理能够避免把一次低质量未发现直接写成空域清空，也能防止单帧虚警立即占用拦截资源。"
    )
    writer.image(
        "10_search_probability_update.png",
        "图4.4-5  未发现信息的概率更新。高质量凝视和快速扫过具有不同的证据强度。",
    )
    writer.paragraph(
        "试验中的局部确认采用确定性短航迹。当前检测与上一帧航迹的检测框中心距离超过180像素时排除，其余候选由匈牙利算法一一连接。每个单元连续观察3帧，帧间隔0.1秒；同一匿名航迹连续2帧达到10像素才生成确认记录。中间漏掉一帧时，前后两次检测不能直接拼成连续确认。"
    )

    writer.heading("4.4.7 AirSim仿真验证")
    writer.heading("试验条件", level=4)
    writer.paragraph(
        "试验设置20目标/8机、20目标/30机和40目标/50机三组场景。目标采用长度3米的无人机网格模型，以50米/秒移动；拦截无人机采用AirSim计算机视觉相机节点，不包含飞行动力学。机载相机分辨率1920×1080，水平视场19度，标称观察距离700米。检测框最长边达到10像素记为可识别，连续两帧满足条件才形成确认。中心线索精度和召回率固定为80%，三组均使用随机种子20260816，运行18秒，状态步长0.1秒，仿真时钟倍率0.1。"
    )
    writer.heading("试验过程", level=4)
    writer.paragraph(
        "20目标场景包含16条正确线索、4条错误线索、4个中心漏检目标和8个空白搜索单元，共28个搜索单元；40目标场景包含32条正确线索、8条错误线索、8个中心漏检目标和16个空白单元，共56个搜索单元。三组均执行三轮滚动分配。在线计算读取匿名检测框、拍摄时间、相机位姿和相机参数，真实目标编号只在试验结束后用于核对。"
    )
    writer.image(
        "13_airsim_validation_chain.png",
        "图4.4-6  三组AirSim试验计算链路。搜索、中心结果交接和机间配准在同一进程中分段运行。",
    )
    writer.heading("试验结果", level=4)
    writer.table(
        ("场景", "搜索单元", "三轮容量/实际分配", "唯一覆盖", "10像素/确认", "中心漏检补获", "平均规划时间"),
        (
            ("20目标/8机", "28", "24/24", "24", "20/19", "3/4", "2.758毫秒"),
            ("20目标/30机", "28", "90/44", "28", "20/20", "4/4", "11.739毫秒"),
            ("40目标/50机", "56", "150/88", "56", "40/40", "8/8", "35.753毫秒"),
        ),
        widths=(1.0, 0.75, 1.15, 0.8, 0.95, 0.95, 1.1),
        font_size=8.6,
    )
    writer.image(
        "14_airsim_search_capacity.png",
        "图4.4-7  搜索容量对比。20目标/8机场景的三轮任务槽少于搜索单元数量。",
    )
    writer.image(
        "15_airsim_search_results.png",
        "图4.4-8  目标可见和连续确认结果。达到10像素只表示目标曾经可见，连续确认才可进入后续交接。",
    )
    writer.paragraph(
        "20目标/8机三轮只有24个任务槽，无法覆盖28个搜索单元，最终留下4个单元未搜索。20个目标都曾达到10像素，但只有19个通过连续确认，中心漏掉的4个目标补获3个。20目标/30机首轮覆盖全部28个单元，40目标/50机在第二轮补齐剩余单元，两组均完成全部目标连续确认，并补获中心漏掉的全部目标。"
    )
    writer.paragraph(
        "本轮结果表明，搜索容量是首先需要满足的条件。无人机数量乘搜索轮次小于待覆盖单元数时，应增加搜索轮次、合并低价值单元或调整责任区。资源充足后，空白搜索单元和滚动复访能够补回部分中心漏检目标。表中计算时间不含通信排队、无人机飞行、云台稳定和图像识别处理，不能作为机载处理器指标。"
    )
    writer.heading("4.4.8 后续验证")
    writer.paragraph(
        "后续试验将逐项加入航迹过期、重复线索、方位偏差、导航误差、云台误差、漏检、虚警和通信延迟，每个规模至少运行10个独立随机种子。确定性基线稳定后，再比较强化学习搜索策略。学习策略只选择已有搜索单元、观察方式和停留时间，云台限位、任务边界、最低重访频率和已锁定目标保持由确定性规则控制。"
    )


def write_section_45(writer: SectionWriter, original_images: Sequence) -> None:
    writer.heading("4.5.1 场景及存在的问题")
    writer.paragraph(
        "拦截无人机进入目标附近后，每架相机只能看到群目标的一部分，且相邻视场存在重叠。可见光全视场约19度乘11度，按目标横向间隔100米估算，1400米距离通常可见4至5个横向目标，500米距离收窄到1至2个。不同无人机为同一目标建立的本地编号互不相同；中心源航迹在80%精度、80%召回率工况下也不能作为机载视觉真值。系统必须把中心源航迹、各机匿名局部航迹和机间共同观测整理成一致关系，才能进行末段任务核对。"
    )
    writer.table(
        ("交战距离", "可见光视场", "横向可见目标数", "纵向层数"),
        (
            ("1400米", "469米×270米", "4～5", "1～2"),
            ("1000米", "335米×193米", "3～4", "1～2"),
            ("500米", "167米×96米", "1～2", "1"),
            ("300米", "100米×58米", "1", "1"),
        ),
        widths=(1.2, 1.8, 1.5, 1.2),
    )
    writer.cloned_image(
        original_images[0],
        "图4.5-1  多架拦截无人机看到不同目标子集。各机本地编号不同，需要找出共同目标、独有目标和待复核关系。",
    )

    writer.heading("4.5.2 存在的难点")
    writer.paragraph(
        "同一目标在不同相机中的像素位置、运动方向和目标框大小通常不同，像素坐标不能直接比较。拍摄时刻不一致、相机位姿误差、通信延迟、遮挡和检测框抖动会产生多个候选；目标数量和相机数量增加后，全部相机两两比较会形成大量没有共同视场的无效组合。目标交叉时，两种身份排列在单帧内可能都成立。没有共同目标、可信源航迹或可靠时空交接时，几何方法和图神经网络都不能凭空建立关系。"
    )

    writer.heading("4.5.3 拟采用的方案")
    writer.paragraph(
        "配准分为中心结果交接和机间跨视角配准两条链路。中心结果先按拍摄时刻外推，并通过拦截机位置、机体姿态、云台角度和相机参数投影到图像平面，形成带误差范围的预测区域。机载检测在预测区域内形成匿名局部航迹，再用时间、几何和运动条件筛选。机间配准先按责任区和视场重叠选择需要比较的相机对，再比较两机局部航迹的单位视线、交会位置、重投影和运动连续性。少量候选使用最近邻或匈牙利算法；密集候选由图神经网络调整排序，最终仍由一一分配和连续多帧确认决定是否绑定。"
    )
    writer.image_path(
        CENTER_TERMINAL_REPORT_ASSET_DIR / "06_terminal_flow.png",
        "图4.5-2  多无人机视觉配准流程。各机先形成局部航迹，再进行跨视角候选筛选和整体分配。",
    )
    writer.cloned_image(
        original_images[2],
        "图4.5-3  匈牙利算法的多视角一一配准。候选门控后建立代价矩阵，避免多条航迹占用同一对象。",
    )
    writer.cloned_image(
        original_images[3],
        "图4.5-4  图神经网络辅助配准。网络调整候选边排序，几何门限和一一约束继续保留。",
    )

    writer.heading("4.5.4 中心源航迹投影")
    writer.heading("状态外推和误差增长", level=4)
    writer.paragraph(
        "中心源航迹包含北东地坐标系中的位置、速度、六维协方差、测量时刻、消息到达时刻、有效期和源编号。状态写成x=[p，v]。系统以图像拍摄时刻为基准，先按匀速模型外推位置和速度，再按白噪声加速度模型增加过程不确定度。线索时间越旧，预测位置范围越大；到达时间只用于检查通信延迟，不能替代拍摄时间参与几何计算。"
    )
    writer.formula("x(t) = Fx₀，P(t) = FP₀Fᵀ + Q，F = [[I，ΔtI]，[0，I]]")
    writer.paragraph(
        "试验中加速度标准差取0.5米/秒²。过程噪声的位置块随时间差四次方增长，位置与速度交叉块随三次方增长，速度块随二次方增长。该数值只用于本轮仿真标定，实际系统需要根据目标机动和中心测量误差重新确定。"
    )

    writer.heading("坐标转换和像面预测", level=4)
    writer.paragraph(
        "外推位置先减去拦截机位置，再依次经过机体、云台和相机安装旋转，得到AirSim相机坐标。相机x轴朝前、y轴朝图像右侧、z轴朝下；相机安装在机体前方0.5米。每帧读取最终相机位置和姿态，避免只使用初始设置值。针孔投影为："
    )
    writer.formula("û = cₓ + fₓy_c/x_c，v̂ = cᵧ + fᵧz_c/x_c")
    writer.paragraph(
        "投影雅可比矩阵把中心位置协方差转换到图像平面，再与投影噪声和本地检测中心噪声相加形成像面协方差S。中心定位误差、拦截机导航误差、机体姿态误差、云台误差、时间误差和检测误差因此统一表现为预测椭圆。云台先指向椭圆中心；椭圆超过一个视场时，再按概率从高到低执行局部扫描。"
    )
    writer.formula("S = JP_位置Jᵀ + R_投影 + R_本地，d_M² = (z − ẑ)ᵀS⁻¹(z − ẑ)")
    writer.image(
        "12_center_interceptor_direct_registration.png",
        "图4.5-5  中心源航迹与机载局部航迹直接配准。低质量或过期线索不进入绑定。",
    )

    writer.heading("候选门控和一一绑定", level=4)
    writer.table(
        ("检查项", "试验门限", "处理作用"),
        (
            ("机载识别", "检测框最长边不小于10像素", "排除没有稳定识别条件的检测"),
            ("线索时效", "已经到达且处于有效期", "阻止未来信息和过期信息参与"),
            ("像面范围", "预测点位于当前图像内", "排除当前相机不可见线索"),
            ("马氏距离", "d_M²不大于9.2103", "按预测椭圆自适应收紧或放宽"),
            ("像面运动", "速度差不大于80像素/秒", "排除位置接近但运动方向不符的候选"),
        ),
        widths=(1.2, 2.0, 3.0),
    )
    writer.paragraph(
        "若有S条中心线索和L条机载局部航迹，先形成S×L个真实候选，再为每条中心线索增加一个专用未匹配列，最终矩阵为S×(L+S)。几何代价由马氏距离和归一化像面运动残差组成；已确认源航迹若切换到其他局部航迹，增加4.0的切换代价；每条源航迹的未匹配项代价为12.0。匈牙利算法在整幅矩阵上同时选择，使一条源航迹最多绑定一条局部航迹，一条局部航迹也不会被多条源航迹重复占用。"
    )
    writer.formula("C_几何 = d_M² + (像面速度残差/80)²")
    writer.paragraph(
        "图神经网络对照只处理已经通过硬门控的候选，输出同目标概率P_图，并按C_最终=C_几何−2log(P_图)修正代价。网络不能放回已经被时间、像面或运动门限拒绝的关系。系统连续采集5帧，最近3帧中至少2帧保持同一对应后才正式绑定。机载侧只保存源编号与本地编号的关系，不改写中心源编号。"
    )
    writer.image_path(
        CENTER_TERMINAL_DIR
        / "airsim_m50_n40_scale_v2_20260816/center_handover/figures/projection_ellipse_matching.png",
        "图4.5-6  四十目标/五十机回放中的中心预测椭圆与机载局部检测原图。圆点为局部检测，椭圆表示中心线索投影后的不确定范围。",
    )
    writer.image_path(
        CENTER_TERMINAL_DIR
        / "airsim_m50_n40_scale_v2_20260816/center_handover/figures/matching_cost_matrix.png",
        "图4.5-7  四十目标/五十机回放中的候选代价矩阵原图。空白位置表示未通过门控，保留位置再进入整体一一分配。",
    )

    writer.heading("4.5.5 机间跨视角配准")
    writer.heading("本机局部航迹和视线交会", level=4)
    writer.paragraph(
        "每架无人机先在本相机内把匿名检测框连接为局部短航迹。检测框中心经过相机内参反投影为相机坐标中的单位视线，再利用每帧相机姿态转换到北东地坐标系。不同相机的检测不要求完全同步，两条局部航迹在0.16秒范围内插值或选取最近观测，至少取得3个有效交会样本。"
    )
    writer.formula("d_c = normalize([1，(u−cₓ)/fₓ，(v−cᵧ)/fᵧ])，d_n = normalize(R_nc d_c)")
    writer.paragraph(
        "两台相机的视线分别写成o_a+s d_a和o_b+t d_b。算法求两条视线的正深度最近点，取最近点中点作为该时刻的三角交会位置；两最近点距离反映视线分离误差，视线夹角反映三角定位的几何强度。交会中点随时间进行直线拟合，拟合误差和运动转角用于排除偶然交会或方向矛盾。"
    )
    writer.formula("δ=o_b−o_a，c=d_a·d_b，D=1−c²")
    writer.formula("s=[δ·d_a−c(δ·d_b)]/D，t=[c(δ·d_a)−δ·d_b]/D")
    writer.formula("q_a=o_a+s d_a，q_b=o_b+t d_b，q_中=(q_a+q_b)/2，e_分离=‖q_a−q_b‖")
    writer.image_path(
        CENTER_TERMINAL_DIR
        / "airsim_m30_n20_scale_20260816/crossview/figures/01_ned_top_and_height_views.png",
        "图4.5-8  二十目标/三十机回放的北东地平面和高度关系原图。灰线表示不同相机局部航迹之间形成的候选空间关系。",
    )

    writer.heading("几何门控和关系代价", level=4)
    writer.table(
        ("检查项", "试验门限", "作用"),
        (
            ("图像识别", "最长边不小于10像素", "排除尺寸不足的检测"),
            ("时间对齐", "不超过0.16秒", "限制插值和最近观测时间差"),
            ("航迹交接间隔", "不超过0.65秒", "排除长期未更新的航迹"),
            ("有效交会样本", "不少于3个", "避免单帧偶然交会"),
            ("视线夹角", "不小于0.35度", "排除近似平行视线"),
            ("视线分离", "不大于2米", "要求双视线在空间接近"),
            ("重投影误差", "不大于8像素", "核对交会点回到两幅图像的位置"),
            ("运动拟合误差", "不大于5米", "排除不连续三维运动"),
            ("运动转角", "不大于55度", "排除方向明显矛盾"),
            ("尺度对数差", "不大于0.28", "核对成像尺度和估计距离"),
        ),
        widths=(1.2, 1.7, 3.3),
        font_size=9,
    )
    writer.paragraph(
        "通过硬门控后，几何代价综合视线分离、重投影、时间差、运动拟合、转角、尺度和观测质量七项误差。各项按对应门限归一化并限制最大值，防止单项异常完全支配结果。图神经网络仍只在几何白名单内调整候选排序："
    )
    writer.formula(
        "C_几何 = 0.24C_分离 + 0.20C_重投影 + 0.10C_时间 + 0.18C_运动 + 0.12C_转角 + 0.08C_尺度 + 0.08C_质量"
    )
    writer.paragraph(
        "几何基线直接使用C_几何。图神经网络输出同目标概率P_图后，采用C_最终=0.55C_几何+0.45(1−P_图)融合。每个相机对建立包含未匹配项的代价矩阵，再用匈牙利算法执行一一配对；最近3帧中至少2次保持同一关系后确认。证据不足的局部航迹继续保留为未匹配。"
    )
    writer.image_path(
        CENTER_TERMINAL_DIR
        / "airsim_m30_n20_scale_20260816/crossview/figures/04_candidate_costs.png",
        "图4.5-9  二十目标/三十机回放的跨视角候选代价分布原图。低代价候选进入一一配对，高代价候选保持拒绝。",
    )

    writer.heading("稀疏相机图和目标簇合并", level=4)
    writer.paragraph(
        "相机数量增加后，先按责任区和视场重叠构建稀疏相机图。同一责任区的相机对保留；相邻责任区只有在共同观测时刻存在视锥重叠时才保留；不相邻责任区直接排除。若第i、j台相机分别有m_i、m_j条活动局部航迹，全量候选数近似为所有相机对m_i×m_j之和。8、30、50台相机分别对应28、435和1225个全量相机对，必须在精细几何计算前压缩。"
    )
    writer.formula("单帧候选边数 E = Σᵢ<ⱼ mᵢmⱼ，相机对数量 = M(M−1)/2")
    writer.paragraph(
        "确认关系按代价从小到大合并为跨相机目标簇。同一目标簇不允许出现同一相机的两条局部航迹，防止一台相机内部的两个目标被合并。两个已经成熟的目标簇准备合并时，至少需要两个不同相机对提供支持；只有2帧的短航迹进入成熟目标簇时，需要该簇内至少两台相机共同支持。该条件用于阻断一条偶然错误关系在多机网络中扩散。"
    )

    writer.heading("4.5.6 两目标交叉和部分可见")
    writer.paragraph(
        "两架拦截机同时看到两个接近目标时，系统同时比较两种完整排列。每种排列都计算两机视线交会位置、与中心预测的归一化残差、重投影误差、速度连续性和后续多帧一致性。两种排列的证据差距未达到门限时，不强行选择；一架无人机保持稳定跟踪，另一架可小幅侧向机动以增大观察夹角，待目标分离或取得新一轮中心预测后再确认。"
    )
    writer.paragraph(
        "A机和B机没有共同目标时，系统保留各自局部航迹，不建立跨机关系。目标从A机视场离开后，A机发布最后视线、角速度、拍摄时刻和预测走廊，B机在预计到达时间内搜索该走廊。中心没有对应源航迹的新目标可由两机通过时间对齐、视线交会和运动连续性形成区域临时候选编号，再上报中心去重和注册。临时候选编号只维持本地搜索连续性。"
    )

    writer.heading("4.5.7 AirSim仿真验证")
    writer.heading("试验条件", level=4)
    writer.paragraph(
        "中心交接和机间配准沿用4.4的三组AirSim场景：20目标/8机、20目标/30机和40目标/50机。目标长度3米、速度50米/秒，机载相机1920×1080、水平视场19度，检测框最长边达到10像素并连续两帧后进入局部航迹。中心线索精度和召回率固定为80%。三组使用同一随机种子20260816，每个规模正式采集一次。"
    )
    writer.table(
        ("项目", "试验设置"),
        (
            ("仿真模式", "AirSim计算机视觉模式"),
            ("中心相机", "2台，1280×1024，水平视场3.67度"),
            ("机载相机", "1920×1080，水平视场19度，相机前移0.5米"),
            ("目标", "长度3米的无人机网格，速度50米/秒"),
            ("中心线索", "精度80%、召回率80%，位置标准差1米、速度标准差0.2米/秒"),
            ("识别输入", "AirSim检测元数据；在线去除Actor名称；最长边不小于10像素"),
            ("图网络", "独立合成数据训练；AirSim种子20260816只作留出回放"),
        ),
        widths=(1.35, 4.85),
        font_size=9.0,
    )
    writer.table(
        ("批次", "中心正确/错误/漏检线索", "机载相机", "主要检查内容"),
        (
            ("20目标/8机", "16/4/4", "8", "资源受限时的中心交接和小规模跨视角关系"),
            ("20目标/30机", "16/4/4", "30", "相机数量增加后的候选膨胀和学习评分增益"),
            ("40目标/50机", "32/8/8", "50", "密集目标和大规模相机网络的稀疏化效果"),
        ),
        widths=(1.2, 1.65, 0.9, 2.6),
    )
    writer.heading("对照方案", level=4)
    writer.paragraph(
        "中心交接设置两条路线。几何路线使用时效、像面、马氏距离和运动门控，再执行匈牙利分配与多帧确认；图神经网络路线只在同一候选白名单内修正代价。机间配准设置全相机几何、全相机图网络、稀疏相机图几何和稀疏相机图图网络四种对照。全相机路线用于观察候选膨胀，稀疏路线按责任区和视场先筛相机对。"
    )
    writer.heading("试验过程", level=4)
    writer.paragraph(
        "三个专项在同一个AirSim Blocks进程中按重置分段运行。中心交接先按图像时刻外推源航迹并投影到各机图像，使用五项门控、未匹配项和匈牙利算法完成一一分配。机间配准把匿名检测框连接为局部航迹，完成时间对齐、视线交会、重投影、运动拟合和目标簇合并。图神经网络对照复用保存的匿名观测离线计算，没有重新运行场景。真实目标编号只在全部输出完成后用于离线评分。"
    )
    writer.heading("中心交接结果", level=4)
    writer.table(
        ("场景", "方法", "正确绑定", "错误绑定", "绑定精度", "绑定召回率", "中位复算时间"),
        (
            ("20目标/8机", "几何", "16", "0", "1.0000", "1.0000", "2.391秒"),
            ("20目标/8机", "图神经网络", "16", "0", "1.0000", "1.0000", "2.460秒"),
            ("20目标/30机", "几何", "14", "0", "1.0000", "0.8750", "2.751秒"),
            ("20目标/30机", "图神经网络", "14", "0", "1.0000", "0.8750", "2.882秒"),
            ("40目标/50机", "几何", "31", "1", "0.9688", "0.9688", "15.819秒"),
            ("40目标/50机", "图神经网络", "31", "0", "1.0000", "0.9688", "15.862秒"),
        ),
        widths=(0.9, 1.0, 0.7, 0.7, 0.8, 0.85, 1.0),
        font_size=8.5,
    )
    writer.image(
        "16_airsim_handover_results.png",
        "图4.5-10  中心结果交接对照。图神经网络在40目标/50机回放中拒绝了1条错误绑定。",
    )
    writer.paragraph(
        "三组正确绑定数分别为16、14和31。资源数量增加没有自动提高中心交接召回率。40目标/50机中，几何方法接受了1条连续落入预测区的错误线索，图神经网络对照拒绝了该关系，同时保留31条正确绑定。该结果只有一个AirSim随机种子，默认路径仍采用几何白名单、匈牙利分配和多帧确认。"
    )

    writer.heading("机间配准结果", level=4)
    writer.table(
        ("场景", "全部/稀疏相机对", "全量/稀疏候选边", "候选减少", "全部/稀疏几何时间"),
        (
            ("20目标/8机", "28/16", "5,778/3,296", "43.0%", "12.43/8.87秒"),
            ("20目标/30机", "435/267", "85,847/52,635", "38.7%", "139.78/97.82秒"),
            ("40目标/50机", "1,225/403", "1,104,646/375,236", "66.0%", "1842.79/770.99秒"),
        ),
        widths=(1.15, 1.25, 1.55, 0.85, 1.45),
        font_size=8.8,
    )
    writer.image(
        "17_airsim_camera_pair_sparsification.png",
        "图4.5-11  相机对和候选边压缩。40目标/50机场景在精细计算前排除了66%的候选边。",
    )
    writer.table(
        ("场景", "方法", "正确/错误/漏配", "关系精度/召回率", "身份混合"),
        (
            ("20目标/8机", "全相机几何", "30/0/2", "1.0000/0.9375", "0"),
            ("20目标/8机", "稀疏几何", "30/0/2", "1.0000/0.9375", "0"),
            ("20目标/8机", "稀疏图网络", "30/0/2", "1.0000/0.9375", "0"),
            ("20目标/30机", "全相机几何", "558/302/71", "0.6488/0.8871", "5"),
            ("20目标/30机", "稀疏几何", "564/198/65", "0.7402/0.8967", "4"),
            ("20目标/30机", "稀疏图网络", "571/142/58", "0.8008/0.9078", "2"),
            ("40目标/50机", "全相机几何", "3538/2537/794", "0.5824/0.8167", "18"),
            ("40目标/50机", "全相机图网络", "4031/2094/301", "0.6581/0.9305", "7"),
            ("40目标/50机", "稀疏几何", "4031/16/301", "0.9960/0.9305", "0"),
            ("40目标/50机", "稀疏图网络", "4031/16/301", "0.9960/0.9305", "0"),
        ),
        widths=(1.15, 1.25, 1.35, 1.4, 0.8),
        font_size=8.6,
    )
    writer.image(
        "18_airsim_crossview_comparison.png",
        "图4.5-12  跨视角配准方法对照。稀疏相机图先减少无共同视场关系，再进行几何和学习评分。",
    )
    writer.paragraph(
        "20目标/8机中，各种设置的关联质量相同，稀疏相机图只减少计算量。20目标/30机中，稀疏几何把错误关系从302条降到198条，图神经网络进一步降到142条，身份混合由4个降到2个。40目标/50机中，稀疏几何把关系精度从0.5824提高到0.9960，并将身份混合从18个降到0；图神经网络没有继续改善质量，复算时间反而增加。当前最稳定的改进是先限制需要比较的相机范围。"
    )

    writer.heading("4.5.8 证据边界和后续验证")
    writer.paragraph(
        "三组AirSim观测均使用随机种子20260816，每个规模只正式采集一次。识别输入为AirSim检测元数据，尚未接入真实可见光或红外探测器；试验未注入导航误差、云台姿态误差、时间同步偏差、相机标定漂移、真实漏检虚警、通信丢包和飞行动力学。图神经网络只使用合成数据训练，AirSim观测用于留出回放。40目标/50机复算时间包含关联、审计输出和制图，不能作为处理器部署指标。"
    )
    writer.paragraph(
        "后续按单项误差注入、10个以上独立随机种子和真实双光记录三个层次开展验证。默认路径保留责任区与视场稀疏相机图、几何门限、匈牙利一一分配和多帧确认。图神经网络只有在多组未见场景中稳定减少错误关系且满足处理时限后，才进入在线试验。仿真结果属于算法链路证据，不能写成装备实测或飞行试验能力。"
    )


def image_before_caption(document: Document, *caption_prefixes: str):
    """Capture an existing useful image by its visible caption, not package media name."""

    paragraphs = document.paragraphs
    captions = [
        index
        for index, paragraph in enumerate(paragraphs)
        if any(paragraph.text.strip().startswith(prefix) for prefix in caption_prefixes)
    ]
    if len(captions) != 1:
        raise RuntimeError(f"expected one caption matching {caption_prefixes}, found {len(captions)}")
    for index in range(captions[0] - 1, max(-1, captions[0] - 5), -1):
        if paragraphs[index]._p.xpath(".//a:blip/@r:embed"):
            return copy.deepcopy(paragraphs[index]._p)
    raise RuntimeError(f"image not found before caption {caption_prefixes}")


def capture_existing_images(document: Document):
    return {
        "region_repartition": image_before_caption(
            document,
            "图4.4-2  区域重划",
            "图4.4-3  区域重划",
        ),
        "multicamera_subset": image_before_caption(document, "图4.5-1  多架拦截无人机"),
        "multicamera_flow": image_before_caption(document, "图4.5-2  多无人机视觉配准流程"),
        "hungarian": image_before_caption(document, "图4.5-3  匈牙利算法"),
        "gnn": image_before_caption(document, "图4.5-4  图神经网络辅助配准"),
    }


def validate_target_headings(document: Document) -> None:
    expected = (
        *(f"4.1.{index}" for index in range(1, 10)),
        *(f"4.4.{index}" for index in range(1, 9)),
        *(f"4.5.{index}" for index in range(1, 9)),
    )
    for prefix in expected:
        find_paragraph(document, prefix)


def integrate(document_path: Path, output_path: Path, backup_path: Path | None) -> tuple[str, str]:
    if not document_path.is_file():
        raise FileNotFoundError(document_path)
    verify_reviewed_sources()
    for asset in (
        "01_epipolar_geometry.png",
        "02_gnn_matching_process.png",
        "03_gnn_message_passing.png",
        "04_gnn_assignment_example.png",
        "05_joint_fit_reprojection.png",
        "06_candidate_graph_assignment.png",
        "07_v4_20target_results_cn.png",
        "08_center_interceptor_search_architecture.png",
        "09_interceptor_search_cell_allocation.png",
        "10_search_probability_update.png",
        "12_center_interceptor_direct_registration.png",
        "13_airsim_validation_chain.png",
        "14_airsim_search_capacity.png",
        "15_airsim_search_results.png",
        "16_airsim_handover_results.png",
        "17_airsim_camera_pair_sparsification.png",
        "18_airsim_crossview_comparison.png",
    ):
        if not (ASSET_DIR / asset).is_file():
            raise FileNotFoundError(ASSET_DIR / asset)

    if backup_path is not None:
        backup_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(document_path, backup_path)

    document = Document(document_path)
    styles = style_samples(document)
    existing_images = capture_existing_images(document)
    outside_before = outside_target_hash(document)

    heading_41 = find_paragraph(document, "4.1 侦察的想法")
    heading_42 = find_paragraph(document, "4.2 火指控的想法")
    heading_44 = find_paragraph(document, "4.4 拦截区域搜索")
    heading_45 = find_paragraph(document, "4.5 群对群目标配准")
    heading_after_45 = find_paragraph(document, "4.3.6 主动降级与分级目标分配")

    remove_between(heading_41, heading_42)
    remove_between(heading_44, heading_45)
    remove_between(heading_45, heading_after_45)

    write_section_41(SectionWriter(document, heading_42, styles))
    write_section_44(SectionWriter(document, heading_45, styles), existing_images["region_repartition"])
    write_section_45(
        SectionWriter(document, heading_after_45, styles),
        (
            existing_images["multicamera_subset"],
            existing_images["multicamera_flow"],
            existing_images["hungarian"],
            existing_images["gnn"],
        ),
    )

    validate_target_headings(document)
    outside_after = outside_target_hash(document)
    if outside_before != outside_after:
        raise RuntimeError("content outside 4.1, 4.4 and 4.5 changed during integration")
    set_a4_page_size(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_name(output_path.name + ".tmp.docx")
    document.save(temporary)
    Document(temporary)
    os.replace(temporary, output_path)
    return outside_before, outside_after


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--document", type=Path, default=DEFAULT_DOCUMENT)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--backup", type=Path)
    args = parser.parse_args()

    document_path = args.document.resolve()
    output_path = (args.output or document_path).resolve()
    backup_path = args.backup.resolve() if args.backup else None
    before = sha256(document_path)
    outside_before, outside_after = integrate(document_path, output_path, backup_path)
    after = sha256(output_path)
    print(f"document={output_path}")
    print(f"sha256_before={before}")
    print(f"sha256_after={after}")
    print(f"outside_sha256_before={outside_before}")
    print(f"outside_sha256_after={outside_after}")
    if backup_path:
        print(f"backup={backup_path}")


if __name__ == "__main__":
    main()
