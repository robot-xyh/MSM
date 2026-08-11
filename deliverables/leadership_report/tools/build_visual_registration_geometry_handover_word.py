#!/usr/bin/env python3
"""Build the geometry-led visual-registration report as a formatted DOCX."""

from __future__ import annotations

import io
import re
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPORT_ROOT = Path(__file__).resolve().parents[1]
SOURCE = REPORT_ROOT / "VISUAL_REGISTRATION_GEOMETRY_AND_HANDOVER_SOLUTION_CN.md"
OUTPUT = REPORT_ROOT / "VISUAL_REGISTRATION_GEOMETRY_AND_HANDOVER_SOLUTION_CN.docx"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"

INK = "202833"
BLUE = "1F4E78"
TEAL = "176B73"
MUTED = "5F6B78"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|\$[^$\n]+\$)")


def _set_run_font(
    run,
    *,
    size: float = 11.0,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str = INK,
    east_asia: str = BODY_FONT,
    latin: str = LATIN_FONT,
) -> None:
    run.font.name = latin
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _latex_to_linear(source: str) -> str:
    text = _normalize_mathtext(source.strip())
    replacements = {
        r"\rightarrow": "→",
        r"\lambda": "λ",
        r"\sigma": "σ",
        r"\theta": "θ",
        r"\pi": "π",
        r"\Sigma": "Σ",
        r"\approx": "≈",
        r"\sim": "∼",
        r"\times": "×",
        r"\perp": "⊥",
        r"\qquad": "  ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    for command in ("mathbf", "mathrm", "hat", "tilde"):
        pattern = re.compile(rf"\\{command}\{{([^{{}}]+)\}}")
        while pattern.search(text):
            text = pattern.sub(r"\1", text)
    text = text.replace(r"\,", " ")
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text


def _normalize_mathtext(source: str) -> str:
    """Normalize compact LaTeX accepted by Markdown but rejected by mathtext."""
    return re.sub(r"\\mathbf\s+([A-Za-z])", r"\\mathbf{\1}", source)


def _add_inline(paragraph, text: str, *, size: float = 11.0, color: str = INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            _set_run_font(paragraph.add_run(text[position : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            _set_run_font(
                paragraph.add_run(token[2:-2]),
                size=size,
                bold=True,
                color=color,
                east_asia=HEADING_FONT,
            )
        elif token.startswith("*"):
            _set_run_font(
                paragraph.add_run(token[1:-1]),
                size=size,
                italic=True,
                color=color,
            )
        else:
            _set_run_font(
                paragraph.add_run(_latex_to_linear(token[1:-1])),
                size=size,
                italic=True,
                color=INK,
                east_asia=BODY_FONT,
                latin=MATH_FONT,
            )
        position = match.end()
    if position < len(text):
        _set_run_font(paragraph.add_run(text[position:]), size=size, color=color)


def _configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Heading 1": (16.0, BLUE, 14, 7),
        "Heading 2": (13.5, TEAL, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    _set_run_font(run, size=8.5, color=MUTED, east_asia=HEADING_FONT)


def _add_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        header.add_run("弱特征条件下多无人机视觉配准方案"),
        size=8.5,
        color=MUTED,
        east_asia=HEADING_FONT,
    )

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(footer.add_run("MSM 项目组  ·  "), size=8.5, color=MUTED)
    _add_page_number(footer)


def _add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Cm(0)
    _set_run_font(
        title.add_run("弱特征条件下多无人机视觉配准方案"),
        size=25,
        bold=True,
        color=INK,
        east_asia=HEADING_FONT,
    )

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = Cm(0)
    _set_run_font(
        subtitle.add_run("固定中心参考、扇区基准航迹与末端身份确认"),
        size=14.5,
        bold=True,
        color=BLUE,
        east_asia=HEADING_FONT,
    )

    for _ in range(9):
        document.add_paragraph()

    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    owner.paragraph_format.first_line_indent = Cm(0)
    _set_run_font(owner.add_run("MSM 项目组"), size=12, color=INK)

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.first_line_indent = Cm(0)
    _set_run_font(date.add_run("2026 年 8 月"), size=11, color=MUTED)

    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boundary.paragraph_format.first_line_indent = Cm(0)
    _set_run_font(boundary.add_run("科研仿真与技术论证材料"), size=9.5, color=TEAL)


def _add_image(document: Document, relative: str, caption: str) -> None:
    image_path = (SOURCE.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.first_line_indent = Cm(0)
    image_paragraph.paragraph_format.space_before = Pt(6)
    image_paragraph.paragraph_format.space_after = Pt(2)
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.add_run().add_picture(str(image_path), width=Cm(16.0))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)
    caption_paragraph.paragraph_format.space_after = Pt(6)
    caption_paragraph.paragraph_format.keep_together = True
    _set_run_font(
        caption_paragraph.add_run(caption),
        size=9,
        color=MUTED,
        east_asia=BODY_FONT,
    )


def _render_formula(formula_lines: list[str]) -> tuple[io.BytesIO, float]:
    source = " ".join(line.strip() for line in formula_lines if line.strip())
    lines = [source]
    if len(source) > 110:
        split_points = [index for index, char in enumerate(source) if char == "+"]
        if split_points:
            midpoint = len(source) / 2
            split_at = min(split_points, key=lambda index: abs(index - midpoint))
            lines = [source[:split_at].rstrip(), r"\quad " + source[split_at:].lstrip()]

    longest = max(len(line) for line in lines)
    font_size = 15 if longest <= 85 else 14
    rendered = "\n".join("$" + _normalize_mathtext(line) + "$" for line in lines)

    fig_height = max(0.45, 0.38 * len(lines) + 0.18)
    fig = plt.figure(figsize=(8.0, fig_height), facecolor="white")
    fig.text(
        0.5,
        0.5,
        rendered,
        ha="center",
        va="center",
        fontsize=font_size,
        color="#202833",
        linespacing=1.35,
    )
    stream = io.BytesIO()
    fig.savefig(
        stream,
        format="png",
        dpi=300,
        bbox_inches="tight",
        pad_inches=0.08,
        facecolor="white",
    )
    plt.close(fig)
    stream.seek(0)

    with Image.open(stream) as image:
        natural_width_cm = image.width / 300.0 * 2.54
    stream.seek(0)
    return stream, min(15.2, max(5.0, natural_width_cm))


def _add_formula(document: Document, formula_lines: list[str]) -> None:
    stream, width_cm = _render_formula(formula_lines)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(stream, width=Cm(width_cm))


def _add_list_item(
    document: Document,
    marker: str,
    content: str,
    *,
    compact: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(-0.58)
    paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    if compact:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    size = 9.8 if compact else 10.8
    _set_run_font(
        paragraph.add_run(f"{marker}  "),
        size=size,
        bold=True,
        color=BLUE,
        east_asia=HEADING_FONT,
    )
    _add_inline(paragraph, content, size=size)


def build_document() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    _configure_section(document.sections[0])
    _configure_styles(document)
    _add_cover(document)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(content_section)
    _add_header_footer(content_section)

    content_started = False
    in_references = False
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if not content_started and not line.startswith("## "):
            index += 1
            continue

        heading = re.match(r"^(#{2,3})\s+(.+)$", line)
        if heading:
            content_started = True
            title = heading.group(2).strip()
            level = 1 if len(heading.group(1)) == 2 else 2
            if title == "参考资料":
                in_references = True
                document.add_page_break()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline(
                paragraph,
                title,
                size=16 if level == 1 else 13.5,
                color=BLUE if level == 1 else TEAL,
            )
            index += 1
            continue

        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            caption = image_match.group(1).strip()
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            if lookahead < len(lines):
                candidate = lines[lookahead].strip()
                if candidate.startswith("*图") and candidate.endswith("*"):
                    caption = candidate[1:-1].strip()
                    lookahead += 1
            _add_image(document, image_match.group(2), caption)
            index = lookahead
            continue

        if line == "$$":
            formula_lines: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "$$":
                formula_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise RuntimeError("unterminated display formula")
            _add_formula(document, formula_lines)
            index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            _add_list_item(
                document,
                f"{numbered.group(1)}.",
                numbered.group(2),
                compact=in_references,
            )
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            _add_list_item(document, "•", bullet.group(1))
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate:
                break
            if candidate.startswith(("#", "![", "$$", "- ", "*")):
                break
            if re.match(r"^\d+\.\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        _add_inline(paragraph, " ".join(parts), size=11)
        index = lookahead

    properties = document.core_properties
    properties.title = "弱特征条件下多无人机视觉配准方案"
    properties.subject = "固定中心参考、扇区基准航迹与末端身份确认"
    properties.author = "MSM 项目组"
    properties.keywords = "视觉配准, 扇区基准航迹, 多视角几何, 目标关联"

    temporary = OUTPUT.with_name(f".{OUTPUT.stem}.tmp.docx")
    document.save(temporary)
    temporary.replace(OUTPUT)


def validate_document() -> dict[str, int]:
    document = Document(OUTPUT)
    with ZipFile(OUTPUT) as archive:
        damaged = archive.testzip()
        if damaged is not None:
            raise RuntimeError(f"damaged DOCX member: {damaged}")
        media = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]

    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = (
        "一、任务条件",
        "二、总体方案",
        "三、统一几何基础",
        "四、两级配准算法",
        "五、无共同视场",
        "六、误差控制和失效路径",
        "七、学习算法边界",
        "八、信息交换",
        "九、实施步骤",
        "十、验证方案",
        "十一、方案结论",
        "参考资料",
    )
    for heading in required:
        if heading not in text:
            raise RuntimeError(f"missing heading: {heading}")

    captions = [paragraph.text for paragraph in document.paragraphs if re.match(r"^图\d+\s", paragraph.text)]
    if len(captions) != 12:
        raise RuntimeError(f"expected 12 figure captions, found {len(captions)}")
    if len(document.inline_shapes) != 25:
        raise RuntimeError(
            f"expected 12 figures and 13 formula images, found {len(document.inline_shapes)} inline shapes"
        )
    if len(media) != 25:
        raise RuntimeError(f"expected 25 embedded media files, found {len(media)}")
    if "$$" in text or r"\mathbf" in text:
        raise RuntimeError("raw LaTeX leaked into DOCX paragraphs")
    if len(document.paragraphs) < 120:
        raise RuntimeError("generated document is unexpectedly short")

    return {
        "paragraphs": len(document.paragraphs),
        "figures": len(captions),
        "formulas": len(document.inline_shapes) - len(captions),
        "media": len(media),
        "bytes": OUTPUT.stat().st_size,
    }


def main() -> None:
    build_document()
    metrics = validate_document()
    print(
        f"{OUTPUT.name}: paragraphs={metrics['paragraphs']}, "
        f"figures={metrics['figures']}, formulas={metrics['formulas']}, "
        f"media={metrics['media']}, bytes={metrics['bytes']}"
    )


if __name__ == "__main__":
    main()
