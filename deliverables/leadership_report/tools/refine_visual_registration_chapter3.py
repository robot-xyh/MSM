#!/usr/bin/env python3
"""Refine chapter 3 of the existing visual-registration Word report in place.

The document contains user-authored additions that are not present in the Markdown
source. This script therefore edits the current DOCX directly and verifies that
embedded drawings and media are preserved.
"""

from __future__ import annotations

import os
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPORT_ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_PATH = REPORT_ROOT / "VISUAL_REGISTRATION_SECTION_CN.docx"
BACKUP_PATH = Path("/tmp/VISUAL_REGISTRATION_SECTION_CN_before_refine_20260804.docx")

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
INK = "202833"
TEAL = "176B73"
MUTED = "5F6B78"


REPLACEMENTS = {
    39: (
        "图匹配（Graph Matching）是在两个或多个图结构之间建立节点对应关系的方法。"
        "在计算机视觉中，它通常用于确定不同图像中关键点的对应关系，可用于图像检索、"
        "位姿估计和多视角配准。本方案按三个步骤处理：每架拦截无人机先从识别与跟踪结果中"
        "提取目标节点；再利用 Delaunay 三角剖分建立节点之间的邻接关系；最后由图神经网络"
        "综合节点特征和边连接关系，判断多架无人机所见目标是否对应，并形成统一的多视角"
        "目标关系，供后续任务分配使用。"
    ),
    43: (
        "在计算机科学中，图由顶点集合和边集合组成，边用于表示顶点之间的连接关系。"
        "对应到多视角目标匹配，顶点代表检测或跟踪得到的目标节点，边代表节点之间的空间"
        "邻接关系。给定节点及其连接关系后，即可形成完整的图结构。"
    ),
    45: "关键点可通过以下三类方法提取：",
    46: "1）人工标注；",
    47: "2）传统特征方法，如 ORB、AKAZE 和 SIFT；",
    48: "3）深度学习方法，如 SuperPoint。",
    49: (
        "本方案利用机载相机的目标检测和连续跟踪结果获得节点集合 P，再通过构图建立目标"
        "之间的局部空间关系。Delaunay 三角剖分是一种常用的构图方法。"
    ),
    51: "目标识别与三角剖分",
    52: (
        "Delaunay 三角剖分是一种广泛用于计算机图形学、计算几何和数值分析的三角剖分"
        "方法。它能够形成结构较规整的三角网格，便于表达离散点之间的邻接关系。"
    ),
    53: (
        "三角剖分（triangulation）是针对给定平面点集生成一组三角形的过程。设平面点集 "
        "P={P1，…，Pn}，希望得到三角形集合 T={t1，…，tm}，其结果应满足："
    ),
    54: "1）所有三角形的顶点共同构成点集 P；",
    55: "2）任意两个三角形的内部不相交，其边界只允许重合或无交点；",
    56: "3）全部三角形的并集构成点集 P 的凸包（convex hull）。",
    58: (
        "对一组离散点，Delaunay 三角剖分构建相互邻接且内部不重叠的三角形，并使任一"
        "三角形的外接圆内部不包含点集中的其他点。由此形成的三角网格可作为每个视角内的"
        "稀疏邻接关系；跨视角比较仍需结合相机几何、节点特征和多帧信息。"
    ),
    60: "Delaunay 三角剖分主要遵循以下两个准则：",
    61: (
        "1）空圆特性：任一 Delaunay 三角形的外接圆内部不包含点集中的其他点。"
        "当不存在四点共圆等退化情况时，剖分结果唯一，如下图所示。"
    ),
    62: (
        "2）最大化最小角特性：在同一散点集的不同三角剖分中，Delaunay 结果倾向于使"
        "最小角尽可能大，从而减少狭长三角形。对于两个相邻三角形构成的凸四边形，交换"
        "公共对角线后，六个内角中的最小值不会进一步增大。"
    ),
    64: (
        "Delaunay 三角剖分可由多种算法实现。在不存在四点共圆等退化条件时，剖分结果"
        "具有唯一性。以视场内 6 个目标为例，Watson 增量算法的基本步骤如下："
    ),
    65: "1）构造一个包含全部散点的超级三角形，并将其放入三角形链表；",
    66: (
        "2）依次插入点集中的散点。在当前三角形集合中找出外接圆包含该点的影响三角形，"
        "删除这些三角形并保留其边界边，再将新点与每条边界边连接，形成新的三角形；"
    ),
    67: (
        "3）检查局部新三角形是否满足 Delaunay 条件，必要时进行换边优化，并将结果写回"
        "三角形链表；"
    ),
    68: "4）重复步骤 2 和步骤 3，直至全部散点插入完毕。",
    73: (
        "图可采用邻接矩阵、邻接表、边集数组、十字链表和邻接多重表等方式存储。邻接矩阵、"
        "邻接表和边集数组既可表示有向图，也可表示无向图；十字链表主要用于有向图；邻接"
        "多重表适用于无向图。本方案使用邻接矩阵表达节点之间的连接和边权，便于后续批量计算。"
    ),
    74: (
        "对于含 n 个顶点的图，可用 n×n 矩阵 A 表示。元素 Aij 表示顶点 i 与顶点 j "
        "之间是否存在边：无连接时取 0；非加权图有连接时取 1；加权图则记录相应边权。"
    ),
    76: "对于加权无向图，若顶点 i 与顶点 j 相连，Aij 为该边权；若两点不相连，Aij=0。",
    78: (
        "无向图的邻接矩阵关于主对角线对称。因为顶点 i 与顶点 j 的连接没有方向，Aij "
        "与 Aji 相等；加权与非加权无向图均满足这一性质。"
    ),
    80: (
        "图匹配需要在两个图的节点之间建立对应关系，同时考虑节点特征相似度和边结构"
        "相似度。精确图匹配（Exact Graph Matching）要求节点与边完全对应，约束严格，"
        "难以适应检测误差、遮挡和部分可见等实际情况。非精确图匹配（Inexact Graph "
        "Matching）允许两个图存在结构差异，把问题转化为节点相似度和边相似度总体最大的"
        "优化问题。本方案采用非精确图匹配思路，并允许部分节点保持未匹配。"
    ),
    82: (
        "非精确图匹配经过适当放宽后仍属于 NP 难问题。工程上通常继续松弛置换矩阵的"
        "离散约束，求得可计算的近似解。常用策略包括："
    ),
    83: "1）谱松弛（Spectral Relaxation）；",
    84: "2）半正定规划松弛（Semidefinite Programming Relaxation）；",
    85: "3）双随机松弛（Doubly Stochastic Relaxation）。",
    87: (
        "传统松弛方法在节点数量较多时计算开销上升，且特征表达能力受人工设计限制。"
        "近年来，研究开始利用深度学习模型学习节点和边的联合表示。Zanfir 等较早将端到端"
        "深度学习框架用于图匹配；图卷积网络和图神经网络可以聚合邻接节点的信息，并与视觉"
        "描述子、几何约束和分配算法组合，形成可训练的匹配流程。"
    ),
    88: (
        "SuperGlue 是一种面向两组局部特征的学习型匹配方法，可同时估计对应点并拒绝无法"
        "匹配的点。该方法先由注意力图神经网络预测匹配得分，再通过可微最优传输求解部分"
        "分配。网络在图像对上训练，以学习局部外观、上下文关系和常见几何变化。公开研究"
        "表明，SuperGlue 可用于室内外姿态估计，并可接入三维重建（Structure from Motion，"
        "SfM）或同步定位与建图（Simultaneous Localization and Mapping，SLAM）流程。其"
        "在线速度取决于关键点数量和图形处理器能力，本项目仍需按机载算力实测。"
    ),
    89: (
        "当某个关键点存在多个外观相近的候选时，只比较局部描述子容易产生歧义。SuperGlue "
        "在两幅图像之间反复聚合上下文信息，比较候选点与周围结构的关系，使注意力逐步集中"
        "到更合理的位置。"
    ),
    91: (
        "SuperGlue 的第一部分是注意力图神经网络。输入两组初始局部特征后，网络通过图像"
        "内部和图像之间的信息传递更新描述子，使每个关键点同时利用本图上下文和另一幅图像"
        "的候选信息。"
    ),
    94: (
        "多重图神经网络：将两幅图像中的关键点作为一个图的节点，并设置两类无向边。图像"
        "内边（自边 E_self）把关键点 i 与同一图像中的其他关键点连接，用于提取单幅图像"
        "内部的结构关系；图像间边（交叉边 E_cross）把关键点 i 与另一幅图像中的候选"
        "关键点连接，用于交换跨图像匹配信息。网络沿两类边交替进行消息传递，每一层都聚合"
        "相关节点信息并更新节点表示。"
    ),
    96: (
        "SuperGlue 的第二部分是最优匹配层，用于生成部分分配矩阵。先计算由 M 行、N 列"
        "构成的候选得分矩阵 S，再在一对一约束和未匹配约束下求分配矩阵 P，使总得分 "
        "ΣSijPij 最大。该过程可视为带未匹配项的线性分配问题，并通过可微最优传输近似求解。"
    ),
    98: (
        "匹配得分预测：若为 M×N 个候选对分别建立高维表示，计算和存储开销较大。"
        "SuperGlue 以两端匹配描述子的内积表示成对得分，即 Sij=〈fiA，fjB〉。匹配描述子"
        "不强制归一化，其幅值可随特征和训练过程变化，用于表达预测置信程度。"
    ),
    100: (
        "遮挡和可见性处理：为允许关键点保持未匹配，SuperGlue 在两个集合中各增加一个"
        "“未匹配项”（dustbin）。没有可靠对应关系的关键点可分配到该项，避免被强制配给"
        "错误目标。实现时在得分矩阵中增加一行和一列，并用可学习参数 z 填充与未匹配项"
        "相关的得分，即 Si,N+1=SM+1,j=SM+1,N+1=z。"
    ),
}

SECTION_START = "三、主要研究内容"
SECTION_END = "四、当前进展"
SUBHEADINGS = {42, 72, 79}
MINOR_HEADING = 51
LIST_PARAGRAPHS = {46, 47, 48, 54, 55, 56, 61, 62, 65, 66, 67, 68, 83, 84, 85}
CAPTIONS = {35, 117}
FIGURE_NOTES = {36, 118}


def package_counts(path: Path) -> dict[str, int]:
    with ZipFile(path) as archive:
        names = archive.namelist()
        xml = archive.read("word/document.xml")
    return {
        "media": sum(
            name.startswith("word/media/") and not name.endswith("/") for name in names
        ),
        "embeddings": sum(
            name.startswith("word/embeddings/") and not name.endswith("/")
            for name in names
        ),
        "drawings": xml.count(b"<w:drawing"),
        "objects": xml.count(b"<w:object"),
        "pict": xml.count(b"<w:pict"),
    }


def set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    color: str = INK,
    east_asia: str = BODY_FONT,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_body(paragraph) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.widow_control = True
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, size=11)


def format_list(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.75)
    paragraph.paragraph_format.first_line_indent = Cm(-0.55)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.widow_control = True
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, size=11)


def format_heading(paragraph, *, level: int) -> None:
    style_id = f"Heading{level}"
    style = next(
        (candidate for candidate in paragraph.part.styles if candidate.style_id == style_id),
        None,
    )
    if style is None:
        raise RuntimeError(f"missing paragraph style: {style_id}")
    paragraph.style = style
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(9 if level == 3 else 6)
    paragraph.paragraph_format.space_after = Pt(4 if level == 3 else 3)
    paragraph.paragraph_format.keep_with_next = True
    size = 12 if level == 3 else 11.5
    color = TEAL if level == 3 else INK
    for run in paragraph.runs:
        if run.text:
            set_run_font(
                run,
                size=size,
                bold=True,
                color=color,
                east_asia=HEADING_FONT,
            )


def format_caption(paragraph, *, bold: bool) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3 if bold else 6)
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, size=9, bold=bold, color=MUTED)


def chapter_bounds(document: Document) -> tuple[int, int]:
    starts = [i for i, p in enumerate(document.paragraphs) if p.text.strip() == SECTION_START]
    ends = [i for i, p in enumerate(document.paragraphs) if p.text.strip() == SECTION_END]
    if len(starts) != 1 or len(ends) != 1 or starts[0] >= ends[0]:
        raise RuntimeError("unable to identify chapter 3 boundaries")
    return starts[0], ends[0]


def refine() -> None:
    if not DOCUMENT_PATH.exists():
        raise FileNotFoundError(DOCUMENT_PATH)

    before_counts = package_counts(DOCUMENT_PATH)
    before_document = Document(DOCUMENT_PATH)
    before_paragraphs = len(before_document.paragraphs)
    before_inline_shapes = len(before_document.inline_shapes)
    start, end = chapter_bounds(before_document)
    if start != 26 or end not in {106, 119}:
        raise RuntimeError(f"unexpected chapter 3 bounds: {(start, end)}")

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCUMENT_PATH, BACKUP_PATH)
    document = Document(DOCUMENT_PATH)

    for index, replacement in REPLACEMENTS.items():
        paragraph = document.paragraphs[index]
        if not paragraph.text.strip():
            raise RuntimeError(f"paragraph {index} is unexpectedly empty")
        if paragraph._p.xpath(".//w:drawing | .//w:object | .//w:pict"):
            raise RuntimeError(f"paragraph {index} contains an embedded object")
        paragraph.text = replacement

    for index in range(start + 1, end):
        paragraph = document.paragraphs[index]
        if not paragraph.text.strip() or paragraph._p.xpath(
            ".//w:drawing | .//w:object | .//w:pict"
        ):
            continue
        if index in SUBHEADINGS:
            format_heading(paragraph, level=3)
        elif index == MINOR_HEADING:
            format_heading(paragraph, level=4)
        elif index in LIST_PARAGRAPHS:
            format_list(paragraph)
        elif index in CAPTIONS:
            format_caption(paragraph, bold=True)
        elif index in FIGURE_NOTES:
            format_caption(paragraph, bold=False)
        elif paragraph.style.name not in {"Heading 1", "Heading 2"}:
            format_body(paragraph)

    summary_index = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.startswith("最终结果仍由明确规则收口")
    )
    preceding_text_index = max(
        i
        for i, paragraph in enumerate(document.paragraphs[:summary_index])
        if paragraph.text.strip()
    )
    removable_blanks = [
        document.paragraphs[i]
        for i in range(preceding_text_index + 1, summary_index)
        if not document.paragraphs[i].text.strip()
        and not document.paragraphs[i]._p.xpath(".//w:drawing | .//w:object | .//w:pict")
    ]
    for paragraph in removable_blanks:
        element = paragraph._element
        element.getparent().remove(element)

    temporary_path = DOCUMENT_PATH.with_name(f".{DOCUMENT_PATH.stem}.refined.docx")
    document.save(temporary_path)

    after_document = Document(temporary_path)
    after_counts = package_counts(temporary_path)
    checks = {
        "paragraphs": (
            before_paragraphs - len(removable_blanks),
            len(after_document.paragraphs),
        ),
        "inline_shapes": (before_inline_shapes, len(after_document.inline_shapes)),
        **{key: (value, after_counts[key]) for key, value in before_counts.items()},
    }
    mismatches = {key: values for key, values in checks.items() if values[0] != values[1]}
    if mismatches:
        temporary_path.unlink(missing_ok=True)
        raise RuntimeError(f"document package validation failed: {mismatches}")

    after_start, after_end = chapter_bounds(after_document)
    expected_after_end = end - len(removable_blanks)
    if (after_start, after_end) != (start, expected_after_end):
        temporary_path.unlink(missing_ok=True)
        raise RuntimeError("chapter boundaries changed during refinement")

    os.replace(temporary_path, DOCUMENT_PATH)
    print(f"updated: {DOCUMENT_PATH}")
    print(f"backup: {BACKUP_PATH}")
    print(f"chapter: paragraphs {start}-{end - 1}")
    print(f"revised text paragraphs: {len(REPLACEMENTS)}")
    print(f"removed redundant blank paragraphs: {len(removable_blanks)}")
    print(f"preserved package counts: {after_counts}")


if __name__ == "__main__":
    refine()
