#!/usr/bin/env python3
"""Build the cooperative-search and terminal-registration Chinese reports."""

from __future__ import annotations

from pathlib import Path
import re
import shutil
import sys
from typing import Any, Sequence
import warnings

warnings.filterwarnings(
    "ignore",
    message="Unable to import Axes3D.*",
    category=UserWarning,
    module="matplotlib.projections",
)

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[3]
EXPERIMENT_ROOT = (
    ROOT / "research_modules" / "independent_experiments" / "center_terminal_cv_campaign"
)
OUTPUTS_DIR = EXPERIMENT_ROOT / "outputs"
ASSET_DIR = ROOT / "deliverables" / "leadership_report" / "assets" / "center_terminal_split_reports"
REPORT_DIR = ROOT / "deliverables" / "leadership_report"
BENCHMARK_PATH = OUTPUTS_DIR / "gnn_offline_benchmark_20260816" / "benchmark_summary.json"

sys.path.insert(0, str(ROOT / "research_modules" / "independent_experiments"))

from center_terminal_cv_campaign.build_three_scale_report import (  # noqa: E402
    RUN_SPECS,
    BenchmarkEvidence,
    RunEvidence,
    fixture_counts,
    load_benchmark,
    load_evidence,
)


SEARCH_MD = REPORT_DIR / "协同搜索试验报告_CN.md"
SEARCH_DOCX = REPORT_DIR / "协同搜索试验报告_CN.docx"
TERMINAL_MD = REPORT_DIR / "末端目标配准试验报告_CN.md"
TERMINAL_DOCX = REPORT_DIR / "末端目标配准试验报告_CN.docx"


def _ratio(value: Any, digits: int = 4) -> str:
    return f"{float(value):.{digits}f}"


def _percent(value: Any, digits: int = 1) -> str:
    return f"{100.0 * float(value):.{digits}f}%"


def _table(headers: Sequence[Any], rows: Sequence[Sequence[Any]]) -> list[str]:
    lines = ["| " + " | ".join(str(value) for value in headers) + " |"]
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    lines.extend("| " + " | ".join(str(value) for value in row) + " |" for row in rows)
    return lines


def _figure(name: str, caption: str) -> str:
    return f"![{caption}](assets/center_terminal_split_reports/{name})"


def _load_evidence() -> tuple[tuple[RunEvidence, ...], BenchmarkEvidence]:
    runs = tuple(load_evidence(OUTPUTS_DIR, spec) for spec in RUN_SPECS)
    return runs, load_benchmark(BENCHMARK_PATH)


def _configure_plotting() -> None:
    plt.rcParams.update(
        {
            "font.family": "Noto Sans CJK JP",
            "axes.unicode_minus": False,
            "figure.facecolor": "white",
            "savefig.facecolor": "white",
        }
    )


def _flow_box(
    axis: plt.Axes,
    center: tuple[float, float],
    text: str,
    color: str,
    *,
    width: float = 0.15,
    height: float = 0.27,
) -> None:
    x, y = center
    axis.add_patch(
        FancyBboxPatch(
            (x - width / 2.0, y - height / 2.0),
            width,
            height,
            boxstyle="round,pad=0.015,rounding_size=0.018",
            facecolor=color,
            edgecolor="#39434b",
            linewidth=1.2,
        )
    )
    axis.text(x, y, text, ha="center", va="center", fontsize=10, linespacing=1.4)


def _flow_arrow(axis: plt.Axes, left: float, right: float, y: float = 0.58) -> None:
    axis.add_patch(
        FancyArrowPatch(
            (left, y),
            (right, y),
            arrowstyle="-|>",
            mutation_scale=13,
            linewidth=1.35,
            color="#46515a",
        )
    )


def _build_search_flow(path: Path) -> None:
    _configure_plotting()
    fig, axis = plt.subplots(figsize=(15.8, 5.0))
    axis.set_xlim(0.0, 1.0)
    axis.set_ylim(0.0, 1.0)
    axis.axis("off")
    centers = (0.09, 0.255, 0.42, 0.585, 0.75, 0.915)
    labels = (
        "80%精度、80%召回率\n中心线索",
        "线索单元\n空白走廊单元",
        "搜索收益矩阵\n概率、视场、转向、复访",
        "匈牙利一一分配\n相机与单元不冲突",
        "连续读取3帧\n检测框最长边≥10像素",
        "连续2帧确认\n未确认单元进入后续复访",
    )
    colors = ("#dfeaf1", "#e8f0df", "#f5e9cb", "#eedfd9", "#dfe8ef", "#e5ece4")
    for center, label, color in zip(centers, labels, colors, strict=True):
        _flow_box(axis, (center, 0.58), label, color, width=0.14)
    for left, right in zip(centers[:-1], centers[1:], strict=True):
        _flow_arrow(axis, left + 0.072, right - 0.072)
    axis.text(
        0.5,
        0.18,
        "没有形成连续确认时，只记录该单元已观察；系统不生成目标身份，也不把一次未发现写成目标不存在。",
        ha="center",
        va="center",
        fontsize=11,
        bbox={"boxstyle": "round,pad=0.35", "facecolor": "#f5f6f7", "edgecolor": "#9aa2a8"},
    )
    axis.set_title("协同搜索计算流程", fontsize=15, pad=10)
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)


def _build_terminal_flow(path: Path) -> None:
    _configure_plotting()
    fig, axes = plt.subplots(2, 1, figsize=(15.8, 7.8))
    for axis in axes:
        axis.set_xlim(0.0, 1.0)
        axis.set_ylim(0.0, 1.0)
        axis.axis("off")
    top_centers = (0.09, 0.255, 0.42, 0.585, 0.75, 0.915)
    top_labels = (
        "中心线索\n位置、速度、协方差",
        "外推到图像时刻\n协方差同步增长",
        "转入相机坐标\n预测像点和椭圆",
        "时间、像面、运动门控",
        "匈牙利一一匹配",
        "最近3帧至少2帧一致\n建立交接关系",
    )
    bottom_labels = (
        "各机匿名局部航迹",
        "0.16秒内时间对齐",
        "双视线交会\n恢复三维运动",
        "几何门控\n可选图网络修正代价",
        "相机对内匈牙利匹配",
        "稀疏相机图约束\n合并跨相机目标簇",
    )
    colors = ("#dfeaf1", "#e8f0df", "#f5e9cb", "#eedfd9", "#dfe8ef", "#e5ece4")
    for row, labels in enumerate((top_labels, bottom_labels)):
        axis = axes[row]
        for center, label, color in zip(top_centers, labels, colors, strict=True):
            _flow_box(axis, (center, 0.5), label, color, width=0.14, height=0.34)
        for left, right in zip(top_centers[:-1], top_centers[1:], strict=True):
            _flow_arrow(axis, left + 0.072, right - 0.072, y=0.5)
    axes[0].set_title("中心线索与机载局部航迹交接", fontsize=14, pad=8)
    axes[1].set_title("拦截无人机之间的跨视角关联", fontsize=14, pad=8)
    fig.suptitle("末端目标配准计算流程", fontsize=16, y=0.995)
    fig.tight_layout(h_pad=1.2)
    fig.savefig(path, dpi=190, bbox_inches="tight")
    plt.close(fig)


def build_assets() -> tuple[Path, ...]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    search_flow = ASSET_DIR / "01_search_flow.png"
    terminal_flow = ASSET_DIR / "06_terminal_flow.png"
    _build_search_flow(search_flow)
    _build_terminal_flow(terminal_flow)

    sources = {
        "02_search_capacity.png": OUTPUTS_DIR / "three_scale_report_figures" / "02_search_capacity.png",
        "03_search_results.png": OUTPUTS_DIR / "three_scale_report_figures" / "03_search_results.png",
        "04_search_coverage_20_8.png": OUTPUTS_DIR / "airsim_n20_formal_v3_20260816" / "search" / "figures" / "search_cell_coverage.png",
        "05_search_coverage_40_50.png": OUTPUTS_DIR / "airsim_m50_n40_scale_v2_20260816" / "search" / "figures" / "search_cell_coverage.png",
        "07_handover_geometry.png": OUTPUTS_DIR / "three_scale_report_figures" / "04_handover_geometry.png",
        "08_handover_results.png": OUTPUTS_DIR / "three_scale_report_figures" / "05_handover_results.png",
        "09_projection_ellipse_matching.png": OUTPUTS_DIR / "airsim_n20_formal_v3_20260816" / "center_handover" / "figures" / "projection_ellipse_matching.png",
        "10_crossview_rays.png": OUTPUTS_DIR / "three_scale_report_figures" / "06_crossview_rays.png",
        "11_crossview_funnel.png": OUTPUTS_DIR / "three_scale_report_figures" / "07_crossview_funnel.png",
        "12_crossview_results.png": OUTPUTS_DIR / "three_scale_report_figures" / "08_scaling_results.png",
        "13_local_pixel_tracks.png": OUTPUTS_DIR / "airsim_n20_formal_v3_20260816" / "crossview" / "figures" / "02_local_pixel_tracks.png",
        "14_crossview_relation_graph.png": OUTPUTS_DIR / "airsim_n20_formal_v3_20260816" / "crossview" / "figures" / "03_crossview_relation_graph.png",
    }
    for destination_name, source in sources.items():
        if not source.is_file():
            raise FileNotFoundError(source)
        shutil.copy2(source, ASSET_DIR / destination_name)
    return tuple(sorted(ASSET_DIR.glob("*.png")))


def build_search_markdown(runs: Sequence[RunEvidence]) -> str:
    lines = [
        "# 协同搜索试验报告",
        "",
        "本报告说明中心线索不完整时，拦截无人机如何把有限相机资源分配到搜索区域，并把一次检测逐步确认成可交接的匿名局部航迹。结论来自20目标/8机、20目标/30机和40目标/50机三组AirSim单次试验，只用于说明算法流程和规模变化。",
        "",
        "## 一、问题、难点与方法",
        "",
        "### 1.1 当前问题与难点",
        "",
        "中心线索按精度80%、召回率80%构造。以20个真实目标为例，中心给出20条线索，其中16条指向真实目标，4条是错误线索，同时还有4个真实目标没有线索。拦截无人机无法事先判断哪条线索错误，也不能只沿中心线索搜索，否则中心漏掉的目标不会进入机载视场。",
        "",
        "搜索还受到相机容量限制。窄视场相机一次只能观察一个有限区域，同一轮内多架无人机若指向同一位置，会浪费搜索容量；全部无人机只追逐高概率线索，又会放弃空白走廊。搜索结果还要经过10像素门限和连续两帧确认，一次短暂看见不能直接形成稳定航迹。",
        "",
        "### 1.2 搜索单元",
        "",
        "算法先把中心线索和来袭走廊统一表示为搜索单元。这样，正确线索、错误线索和无中心线索覆盖的区域可以进入同一套分配计算。设真实目标数为N，正确中心线索数为T，全部中心线索数为S：",
        "",
        "```text",
        "T = 0.8N",
        "S = T / 0.8 = N",
        "错误线索数 = S - T = 0.2N",
        "中心漏掉目标数 = N - T = 0.2N",
        "```",
        "",
        "每条中心线索按位置和速度外推到规划时刻。外推的用途是把旧线索移到目标当前可能出现的位置，关系式为 `p(t)=p0+vΔt`。搜索范围根据位置协方差确定，三个方向的半宽取 `max(30米, 3sqrt(P_ii))`。本试验位置标准差为1米，三倍标准差小于30米，因此使用30米下限。",
        "",
        "中心漏检目标没有对应线索。算法在北向2500至3500米、东向-650至650米、高度-220至-70米的来袭走廊内增加空白单元，数量取 `max(5, ceil(0.4N))`，单元概率为0.32。空白单元只代表需要查看的区域，不代表已经存在目标。",
        "",
    ]
    cue_rows = []
    for run in runs:
        counts = fixture_counts(run.spec.target_count)
        cue_rows.append(
            (
                run.spec.label,
                counts["true_cues"],
                counts["false_cues"],
                counts["missed_targets"],
                counts["gap_cells"],
                run.search["search_cell_count"],
            )
        )
    lines.extend(_table(("场景", "正确线索", "错误线索", "中心漏掉", "空白单元", "总搜索单元"), cue_rows))
    lines.extend(
        [
            "",
            "### 1.3 搜索收益和一一分配",
            "",
            "每轮先建立相机到搜索单元的收益矩阵。若有M台可用相机、K个有效单元，矩阵共有M行、K列。第i行第j列表示第i台相机本轮查看第j个单元的收益。收益需要兼顾目标概率、相机能否看全、转动和到达代价以及是否刚刚看过：",
            "",
            "```text",
            "U = 3p + 4G - 0.8C_slew - 1.0C_arrival - 4C_repeat",
            "G = p × V × Q",
            "```",
            "",
            "式中，p是单元概率；V是相机在700米观察距离上能够覆盖该单元的比例；Q=1/(1+n)，n为既有覆盖次数。C_slew衡量云台和机体需要转动多少，C_arrival衡量到观察位置的距离，C_repeat压低刚刚看过的单元。19度水平视场角在700米处的名义覆盖宽度约234.3米。范围较大的空白单元一次不容易看全，会通过后续复访补齐。",
            "",
            "矩阵右侧再增加M个空闲列，形成M×(K+M)矩阵，空闲收益为-0.05。匈牙利算法一次选出全局组合，使每台相机最多承担一个单元，每个单元在同一轮最多分给一台相机。某台相机若没有收益高于空闲项的单元，就保持空闲。",
            "",
            _figure("01_search_flow.png", "协同搜索计算流程"),
            "",
            "### 1.4 像素门限、连续确认和未发现处理",
            "",
            "相机被放到单元中心前方700米处，并指向单元中心。每个分配单元连续读取3帧，帧间隔0.1秒。1920×1080、水平视场角19度的机载相机像素焦距约5736.7像素。3米目标在700米处的理想成像最长边约24.6像素，超过10像素接口门限；实际判断仍以AirSim检测框为准。",
            "",
            "同一相机、同一单元内，当前检测与上一帧局部航迹先按检测框中心距离计算代价，超过180像素的组合排除，其余组合使用匈牙利算法一一连接。检测框最长边达到10像素记为可识别，同一局部航迹连续2帧满足门限后才生成确认记录。三帧观察提供三次检测机会，但中间漏掉一帧不能把前后检测直接算成连续两帧。",
            "",
            "本轮未发现或未连续确认时，系统只记录该单元已完成一次观察，不生成目标身份。下一轮重复代价会优先让资源转向尚未覆盖的单元；随时间增加，原单元的重复代价逐步下降，仍可再次分配。当前试验没有根据一次未发现动态降低区域概率，也没有实现长时间休眠航迹重接。",
            "",
            "## 二、试验配置",
            "",
            "### 2.1 场景与设备",
            "",
        ]
    )
    lines.extend(
        _table(
            ("项目", "设置"),
            (
                ("仿真模式", "AirSim ComputerVision"),
                ("规模", "20目标/8机、20目标/30机、40目标/50机"),
                ("目标", "无人机静态网格Actor，速度50米/秒，最长尺寸3米"),
                ("场景", "18秒，状态步长0.1秒，ClockSpeed=0.1"),
                ("机载相机", "1920×1080，水平视场角19度，标称观察距离700米"),
                ("中心线索", "固定构造精度80%、召回率80%；位置标准差1米，速度标准差0.2米/秒"),
                ("检测输入", "AirSim simGetDetections元数据，在线入口去除Actor名称"),
                ("识别条件", "检测框最长边不小于10像素，连续2帧确认"),
                ("分配轮数", "3轮，每个分配单元读取3帧，帧间隔0.1秒"),
                ("随机种子", "20260816；每个规模单次运行"),
            ),
        )
    )
    lines.extend(
        [
            "",
            "### 2.2 评价口径",
            "",
            "搜索单元覆盖表示至少执行过一次观察的单元数；达到10像素表示目标至少有一条检测达到接口门限；连续确认表示同一匿名局部航迹连续两帧通过门限；中心漏检补获表示没有正确中心线索的目标最终被空白单元搜索发现。Actor名称只在运行结束后计算这些指标。",
            "",
            "ComputerVision节点按指令改变位姿，不含飞行动力学。AirSim检测元数据不等同于真实可见光或红外探测器。本试验没有注入导航误差、云台误差、通信丢包和真实检测误差。",
            "",
            "## 三、试验过程、结果与分析",
            "",
            "### 3.1 试验过程",
            "",
            "三组试验使用同一计算流程。每轮先筛除过期线索，建立搜索收益矩阵，完成一一分配，再把ComputerVision节点放到对应观察位置。节点连续读取3帧匿名检测，建立短航迹并执行10像素和连续两帧确认。该轮结束后更新覆盖次数和相机姿态，进入下一轮。三轮结束后，离线使用Actor标签统计目标是否被发现。",
            "",
            _figure("02_search_capacity.png", "三组场景的搜索容量与单元覆盖"),
            "",
            "### 3.2 结果",
            "",
        ]
    )
    capacity_rows = []
    result_rows = []
    for run in runs:
        matrix = f"{run.spec.resource_count}×({run.search['search_cell_count']}+{run.spec.resource_count})"
        capacity_rows.append(
            (
                run.spec.label,
                matrix,
                run.search["assignment_capacity"],
                run.search["covered_cell_count"],
                run.search["unassigned_cell_count"],
                f"{float(run.search['planner_compute_mean_ms']):.3f}毫秒",
            )
        )
        result_rows.append(
            (
                run.spec.label,
                f"{run.search['covered_cell_count']}/{run.search['search_cell_count']}",
                f"{run.search['recognized_target_count']}/{run.spec.target_count}",
                f"{run.search['discovered_target_count']}/{run.spec.target_count}",
                f"{run.search['center_missed_recovered_count']}/{run.search['center_missed_target_count']}",
                run.search["online_detection_count"],
            )
        )
    lines.extend(_table(("场景", "分配矩阵", "三轮容量", "唯一覆盖", "未覆盖", "规划平均耗时"), capacity_rows))
    lines.extend(["", _figure("03_search_results.png", "三组场景的目标发现结果"), ""])
    lines.extend(_table(("场景", "单元覆盖", "达到10像素", "连续确认", "中心漏检补获", "匿名检测记录"), result_rows))
    lines.extend(
        [
            "",
            "### 3.3 结果分析",
            "",
            "20目标/8机共有28个搜索单元，三轮最多提供24个分配槽，因此至少有4个单元无法观察。试验实际覆盖24个单元。20个目标都曾达到10像素，但只有19个形成连续确认；中心漏掉的4个目标补获3个。该场景表明，资源容量不足首先造成区域覆盖缺口，随后压缩重访和连续确认机会。",
            "",
            _figure("04_search_coverage_20_8.png", "20目标/8机场景的搜索单元覆盖"),
            "",
            "20目标/30机首轮容量已经超过28个搜索单元。三轮实际执行44次分配，覆盖全部28个单元，其余16次用于复访。全部20个目标连续确认，中心漏掉的4个目标全部补获。规划平均耗时11.739毫秒。继续增加资源的主要作用已经从首次覆盖转向复访。",
            "",
            "40目标/50机共有56个搜索单元，首轮最多覆盖50个，第二轮补齐其余6个。三轮实际执行88次分配，全部40个目标连续确认，中心漏掉的8个目标全部补获。规划矩阵为50×106，平均计算35.753毫秒。该数字只包含确定性分配计算，不包含相机运动、接口等待和通信排队。",
            "",
            _figure("05_search_coverage_40_50.png", "40目标/50机场景的搜索单元覆盖"),
            "",
            "三组结果支持两个判断。搜索单元数量必须与可用相机和滚动轮数共同核算，资源少于单元时无法依靠分配算法消除物理容量缺口。资源达到全覆盖条件后，空白走廊单元能够补获中心漏检目标。当前证据只有单个种子，不能给出稳定成功概率；真实探测器、机动约束和通信时延仍需另行验证。",
            "",
        ]
    )
    return "\n".join(lines)


def build_terminal_markdown(
    runs: Sequence[RunEvidence], benchmark: BenchmarkEvidence
) -> str:
    center_geometry = [
        benchmark.result(run.spec.scenario_id, "center_handover", "geometry")
        for run in runs
    ]
    center_gnn = [
        benchmark.result(run.spec.scenario_id, "center_handover", "gnn") for run in runs
    ]
    sparse_geometry = [
        benchmark.result(run.spec.scenario_id, "crossview", "geometry", "sector_fov")
        for run in runs
    ]
    sparse_gnn = [
        benchmark.result(run.spec.scenario_id, "crossview", "gnn", "sector_fov")
        for run in runs
    ]
    full_geometry = [
        benchmark.result(run.spec.scenario_id, "crossview", "geometry", "full")
        for run in runs
    ]
    full_gnn = [
        benchmark.result(run.spec.scenario_id, "crossview", "gnn", "full")
        for run in runs
    ]

    lines = [
        "# 末端目标配准试验报告",
        "",
        "本报告说明两类关系如何建立：中心线索与拦截无人机局部航迹的交接，以及多架拦截无人机局部航迹之间的跨视角关联。两类计算都保留未匹配状态，并通过一一匹配和多帧确认限制错误绑定。结果来自三组AirSim匿名观测和保存观测上的图神经网络离线复算。",
        "",
        "## 一、问题、难点与方法",
        "",
        "### 1.1 当前问题与难点",
        "",
        "中心线索给出北东地坐标中的位置、速度和协方差，机载相机给出图像中的匿名检测框，两者坐标不同、时间不同、编号也不同。若直接按最近像点绑定，旧线索、目标交叉和相机姿态变化都会引起错配。中心本身还有20%的错误线索和20%的漏检目标，因此算法必须允许中心线索不匹配，也必须允许机载新发现目标保持未注册。",
        "",
        "多架拦截无人机看到同一目标时，各机只掌握自己的局部编号。目标在不同视角中的像素位置和检测框大小差异较大；相机数增加后，全部相机两两比较会产生大量没有共同视场的候选。机间关联需要先恢复共同几何关系，再解决一条航迹被多个候选同时占用以及错误关系在相机网络中扩散的问题。",
        "",
        _figure("06_terminal_flow.png", "中心交接与机间关联的统一计算流程"),
        "",
        "### 1.2 中心线索状态外推",
        "",
        "第一步把中心线索推到机载图像的测量时刻。这样比较的是同一时刻的预测位置和实际检测，而不是拿旧坐标直接对当前图像。状态记为位置和速度 `x=[p,v]`，图像时刻与线索测量时刻相差Δt：",
        "",
        "```text",
        "x(t) = F x0",
        "F = [[I, Δt·I], [0, I]]",
        "P(t) = F P0 F^T + Q",
        "```",
        "",
        "P表示位置和速度的不确定范围，Q表示外推期间可能发生的机动。本试验加速度标准差为0.5米/秒²，Q的位置块、交叉块和速度块分别按 `qΔt⁴/4`、`qΔt³/2` 和 `qΔt²` 增长，其中q=0.5²。线索越旧，预测范围越大，后续像面门限会自动放宽，但错误候选也会增多。",
        "",
        "### 1.3 坐标转换与像面预测",
        "",
        "外推后的目标位置先减去相机位置，再经过机体、云台和相机安装旋转，得到相机坐标 `(x_c,y_c,z_c)`。AirSim相机x轴朝前、y轴朝图像右侧、z轴朝下。预测点落入图像的位置为：",
        "",
        "```text",
        "u = c_x + f_x · y_c / x_c",
        "v = c_y + f_y · z_c / x_c",
        "```",
        "",
        "式中f是像素焦距，c是图像中心。若目标位于相机后方或预测点在图像外，候选直接排除。线索位置协方差通过投影雅可比矩阵J转换到图像，再与投影误差和检测中心误差相加，得到预测椭圆S。实际检测中心z与预测中心z_hat的偏差按预测椭圆归一化：",
        "",
        "```text",
        "S = J P_pos J^T + R_projection + R_local",
        "d² = (z - z_hat)^T S^-1 (z - z_hat)",
        "```",
        "",
        "d²不只看相差多少像素，还看这条线索本来有多不确定。候选需同时满足检测框不小于10像素、线索已经到达且未过期、预测点在图像内、d²不大于9.2103，以及有历史速度时像面速度差不大于80像素/秒。",
        "",
        _figure("07_handover_geometry.png", "中心线索投影到机载图像并形成预测椭圆"),
        "",
        "### 1.4 中心交接代价与确认",
        "",
        "通过门控的候选进入代价矩阵。若有S条中心线索、L条机载局部航迹，先建立S×L个真实候选，再为每条中心线索增加一个专用未匹配列，矩阵为S×(L+S)。几何代价同时考虑位置和像面运动：",
        "",
        "```text",
        "C_geo = d² + (运动残差 / 80)²",
        "```",
        "",
        "已经确认的中心线索若切换到其他局部航迹，代价增加4.0；未匹配项代价为12.0。匈牙利算法一次求解整个矩阵，使一条中心线索最多占用一条局部航迹，一条局部航迹也不会被多条中心线索重复使用。关系需要在最近3帧中至少2帧被选中，才转为正式交接。未达到条件的中心线索保持未匹配，机载航迹保持未注册。",
        "",
        "图神经网络对照只处理已经通过硬门控的候选，不扩大候选范围。网络输出同一目标概率P_gnn后，将代价修正为 `C_final=C_geo-2log(P_gnn)`，再使用相同的匈牙利算法和多帧确认。网络不能绕过图像范围、时间有效性和马氏距离门限。",
        "",
        "### 1.5 机间时间对齐与双视线交会",
        "",
        "各机先把匿名检测框串成局部航迹。机间关联以局部航迹为单位，不比较不同相机的局部编号。检测框中心 `(u,v)` 先反投影成相机坐标单位视线，再根据相机姿态转到北东地坐标：",
        "",
        "```text",
        "d_c = normalize([1, (u-c_x)/f_x, (v-c_y)/f_y])",
        "d_n = normalize(R_n_c d_c)",
        "```",
        "",
        "两台相机的观测时间允许相差0.16秒。算法对两条局部航迹插值或选择最近观测，形成多组时间对齐样本。对每组样本，分别从相机位置沿两条单位视线延伸，求两条空间直线的最近点，最近点中点作为交会位置，最近点距离作为视线分离误差。至少3组有效交会样本才能进入后续运动拟合。",
        "",
        _figure("10_crossview_rays.png", "两台机载相机通过双视线交会核对同一目标"),
        "",
        "### 1.6 机间几何代价、图网络和目标簇",
        "",
        "候选首先经过硬门控：航迹更新时间差不大于0.65秒、视线夹角不小于0.35度、视线分离不大于2米、重投影误差不大于8像素、运动拟合误差不大于5米、运动转角不大于55度、检测框尺度对数差不大于0.28。通过后，将各项误差除以门限并封顶为3，再计算几何代价：",
        "",
        "```text",
        "C_geo = 0.24C_sep + 0.20C_reproj + 0.10C_time",
        "      + 0.18C_motion + 0.12C_turn + 0.08C_scale + 0.08C_conf",
        "```",
        "",
        "几何代价越小，两条局部航迹越可能属于同一目标。图神经网络对照在同一硬门控白名单内输出同目标概率，再按 `C_final=0.55C_geo+0.45(1-P_gnn)` 修正代价。每个相机对分别建立航迹代价矩阵，未匹配代价为1.05，匈牙利算法完成一一匹配；最近3帧至少2次选中同一关系后确认。",
        "",
        "确认关系最终合并为跨相机目标簇。一个目标簇不能包含同一相机的两条航迹；两个成熟目标簇至少需要两个不同相机对共同支持才能合并；只有2帧的短航迹需要成熟簇内至少两台相机支持。该约束用于阻止一条错误关系把多个真实目标串成一个簇。",
        "",
        "相机数量较多时，不再比较所有相机对。同一搜索责任区的相机对保留；相邻责任区只有在共同帧内视场重叠时保留，并增加5度视场余量；其余相机对排除。保留相机作为节点、可能共同观测的相机对作为边，形成稀疏相机图。这个图只由责任区、相机参数和观测姿态生成，不读取真实目标编号。",
        "",
        _figure("11_crossview_funnel.png", "相机对筛选与跨视角候选压缩"),
        "",
        "## 二、试验配置",
        "",
        "### 2.1 场景与输入",
        "",
    ]
    lines.extend(
        _table(
            ("项目", "设置"),
            (
                ("仿真模式", "AirSim ComputerVision"),
                ("规模", "20目标/8机、20目标/30机、40目标/50机"),
                ("中心相机", "2台，1280×1024，水平视场角3.67度"),
                ("机载相机", "1920×1080，水平视场角19度，相机前移0.5米"),
                ("目标", "无人机静态网格Actor，速度50米/秒，最长尺寸3米"),
                ("中心线索", "精度80%、召回率80%，位置标准差1米，速度标准差0.2米/秒"),
                ("识别接口", "AirSim simGetDetections；检测框最长边不小于10像素"),
                ("在线身份", "只使用匿名检测和局部航迹；不读取Actor名称和真实目标编号"),
                ("图网络", "独立合成数据训练，seed 20260816留出；只做保存观测离线复算"),
                ("随机种子", "20260816；每个规模单次采集"),
            ),
        )
    )
    lines.extend(
        [
            "",
            "### 2.2 评价口径与边界",
            "",
            "中心交接精度按正确绑定数除以全部绑定数计算，召回率按正确绑定数除以正确中心线索数计算。机间关系精度按正确跨相机关系数除以全部输出关系数计算，关系召回率按正确关系数除以应建立关系数计算。身份混合表示同一目标簇内出现多个离线真实目标。上述真值只在运行结束后评分。",
            "",
            "中心交接和跨视角几何基线读取AirSim匿名观测。图神经网络使用同一批保存观测离线复算，没有重新运行Blocks。复算时间包含关联、审计文件和绘图，不能当作机载实时延迟。ComputerVision节点不含飞行动力学，本试验没有验证导航误差、云台姿态误差、通信丢包、真实检测器和物理拦截。",
            "",
            "## 三、试验过程、结果与分析",
            "",
            "### 3.1 中心交接过程与结果",
            "",
            "每组中心交接读取5帧观测。每帧先把全部有效中心线索外推到图像时刻，按相机位姿投影到各机图像，再对10像素、时间、图像范围、马氏距离和像面运动逐项门控。通过门控的候选组成代价矩阵，完成一一匹配并累计最近3帧确认次数。图网络对照复用相同白名单、匹配和确认逻辑。",
            "",
            _figure("09_projection_ellipse_matching.png", "中心预测椭圆与机载局部航迹匹配示例"),
            "",
        ]
    )
    center_rows = []
    for index, run in enumerate(runs):
        for method, result in (("几何", center_geometry[index]), ("图神经网络", center_gnn[index])):
            metrics = result["metrics"]
            center_rows.append(
                (
                    run.spec.label,
                    method,
                    metrics["true_binding_count"],
                    metrics["false_binding_count"],
                    _ratio(metrics["binding_precision"]),
                    _ratio(metrics["binding_recall"]),
                    f"{float(result['timing']['median_wall_duration_s']):.3f}秒",
                )
            )
    lines.extend(_table(("场景", "方法", "正确绑定", "错误绑定", "精度", "召回率", "中位复算时间"), center_rows))
    lines.extend(
        [
            "",
            _figure("08_handover_results.png", "三组中心交接结果"),
            "",
            "20目标/8机中，几何和图网络都正确绑定16条有效中心线索，没有错误绑定。20目标/30机中，两种方法均正确绑定14条，另有2条正确线索没有完成绑定，召回率为0.8750。资源数量增加没有自动提高交接结果，因为三组专项在reset后独立采样，局部航迹数量和连续性并不完全相同。",
            "",
            "40目标/50机中，几何方法正确绑定31条，同时误接1条错误线索，精度和召回率均为0.9688。图网络保留31条正确绑定并拒绝该错误关系，精度提高到1.0000，召回率保持0.9688。该差异只出现在单个留出回放中，支持保留学习对照，不能据此替换几何门控和一一匹配。",
            "",
            "### 3.2 机间关联过程与结果",
            "",
            "跨视角复算先按责任区和视场形成相机图，再对每条保留边上的局部航迹执行时间对齐、双视线交会、运动拟合和几何门控。相机对内完成匈牙利匹配和多帧确认后，再按相机唯一性和多边支持规则合并目标簇。全相机策略作为规模压力对照，责任区/视场稀疏策略作为当前默认路径。",
            "",
            _figure("13_local_pixel_tracks.png", "20目标/8机场景中的多相机局部像面航迹"),
            "",
        ]
    )
    cross_quality_rows = []
    cross_scale_rows = []
    for index, run in enumerate(runs):
        groups = (
            ("全相机", "几何", full_geometry[index]),
            ("全相机", "图神经网络", full_gnn[index]),
            ("责任区/视场稀疏", "几何", sparse_geometry[index]),
            ("责任区/视场稀疏", "图神经网络", sparse_gnn[index]),
        )
        for graph_label, method, result in groups:
            metrics = result["metrics"]
            audit = result["candidate_audit"]
            cross_quality_rows.append(
                (
                    run.spec.label,
                    graph_label,
                    method,
                    metrics["true_positive_relations"],
                    metrics["false_positive_relations"],
                    _ratio(metrics["association_precision"]),
                    _ratio(metrics["association_recall"]),
                    metrics["id_switch_count"],
                )
            )
            cross_scale_rows.append(
                (
                    run.spec.label,
                    graph_label,
                    method,
                    audit["camera_pair_retained_count"],
                    metrics["candidate_edge_count"],
                    f"{float(result['timing']['median_wall_duration_s']):.2f}秒",
                )
            )
    lines.extend(
        _table(
            ("场景", "相机图", "方法", "正确", "错误", "精度", "召回率", "身份混合"),
            cross_quality_rows,
        )
    )
    lines.extend(["", "候选规模和复算时间单独列示，避免将质量指标与计算量混在一张宽表中。", ""])
    lines.extend(
        _table(
            ("场景", "相机图", "方法", "保留相机对", "候选边", "复算时间"),
            cross_scale_rows,
        )
    )
    lines.extend(
        [
            "",
            _figure("14_crossview_relation_graph.png", "20目标/8机场景的跨相机关联关系"),
            "",
            "20目标/8机中，四种组合都得到30条正确关系、0条错误关系和2条漏配，精度1.0000、召回率0.9375。稀疏相机图把相机对从28组降到16组，候选边从5778条降到3296条，质量不变，复算时间从全相机几何的12.43秒降到8.87秒。",
            "",
            "20目标/30机中，全相机候选过密，几何方法精度只有0.6488。责任区和视场筛选后，几何精度提高到0.7402。稀疏图网络进一步把错误关系从198条降到142条，精度达到0.8008、召回率达到0.9078，身份混合从4个降到2个。该规模下，先缩小相机比较范围，再由学习模型处理剩余歧义，产生了可测增益。",
            "",
            "40目标/50机中，全相机几何需要处理1,104,646条候选边，输出2537条错误关系并形成18个身份混合。稀疏策略把相机对从1225组降到403组，候选边降到375,236条。稀疏几何得到4031条正确关系、16条错误关系，精度0.9960、召回率0.9305，身份混合为0。稀疏图网络结果相同，复算时间由770.99秒增加到812.96秒，因此该规模没有继续使用图网络的质量收益。",
            "",
            _figure("12_crossview_results.png", "全相机、稀疏相机图与图网络对照"),
            "",
            "三组结果表明，中心交接和机间关联都需要保留几何白名单、未匹配状态、一一匹配和多帧确认。大规模机间关联的主要改进来自责任区和视场形成的稀疏相机图。图神经网络适合处理硬门控后仍有歧义的候选，在20目标/30机场景产生增益，在40目标/50机稀疏几何已经接近满精度时没有继续改善。当前默认路径仍为稀疏相机图、几何门控和匈牙利匹配。",
            "",
            "以上数字来自seed 20260816的单次AirSim观测和对应离线复算，不具有多seed统计意义。真实相机标定漂移、时间同步偏差、导航误差和真实漏检虚警会直接影响投影、视线交会和单机局部航迹，仍需在后续试验中单独注入并标定。",
            "",
        ]
    )
    return "\n".join(lines)


def _set_font(run: Any, chinese: str, western: str, size: float, *, bold: bool | None = None) -> None:
    run.font.name = western
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), chinese)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.349)
    section.right_margin = Cm(2.349)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name, size in (("Heading 1", 16.5), ("Heading 2", 14.0), ("Heading 3", 12.5)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
    document.styles["Heading 1"].paragraph_format.space_after = Pt(8)


def _add_cover(document: Document, title: str, subtitle: str) -> None:
    for _ in range(5):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(paragraph.add_run(title), "黑体", "Times New Roman", 26, bold=True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(paragraph.add_run(subtitle), "黑体", "Times New Roman", 15, bold=True)
    for _ in range(7):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(paragraph.add_run("MSM 项目组"), "宋体", "Times New Roman", 12)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(paragraph.add_run("2026 年 8 月"), "宋体", "Times New Roman", 11)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(paragraph.add_run("科研仿真与技术论证材料"), "宋体", "Times New Roman", 9.5)
    document.add_page_break()


def _plain_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.replace("\\|", "|")


def _parse_table_row(line: str) -> list[str]:
    return [_plain_markdown(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_table(document: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = True
    font_size = 7.5 if column_count >= 10 else (8.5 if column_count >= 7 else 9.5)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_font(run, "宋体", "Times New Roman", font_size, bold=row_index == 0)
            if row_index == 0:
                _shade_cell(cell, "D9E2F3")
    _repeat_table_header(table.rows[0])
    document.add_paragraph()


def _add_picture(document: Document, image_path: Path, caption: str, figure_number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.0))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(5)
    _set_font(
        caption_paragraph.add_run(f"图 {figure_number}  {caption}"),
        "宋体",
        "Times New Roman",
        9.5,
    )


def build_docx(markdown_path: Path, docx_path: Path, *, title: str, subtitle: str) -> Path:
    document = Document()
    _configure_document(document)
    _add_cover(document, title, subtitle)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    figure_number = 0
    first_h1_skipped = False
    first_section = True
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("# ") and not first_h1_skipped:
            first_h1_skipped = True
            index += 1
            continue
        if line.startswith("## "):
            paragraph = document.add_paragraph(
                _plain_markdown(line[3:]), style="Heading 1"
            )
            if not first_section:
                paragraph.paragraph_format.page_break_before = True
            first_section = False
            index += 1
            continue
        if line.startswith("### "):
            document.add_paragraph(_plain_markdown(line[4:]), style="Heading 2")
            index += 1
            continue
        if line.startswith("#### "):
            document.add_paragraph(_plain_markdown(line[5:]), style="Heading 3")
            index += 1
            continue
        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            figure_number += 1
            image_path = (markdown_path.parent / image_match.group(2)).resolve()
            if not image_path.is_file():
                raise FileNotFoundError(image_path)
            _add_picture(document, image_path, image_match.group(1), figure_number)
            index += 1
            continue
        if line.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(5)
            _shade_cell_like_paragraph(paragraph, "F2F2F2")
            _set_font(paragraph.add_run("\n".join(code_lines)), "宋体", "Courier New", 10)
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [_parse_table_row(value) for value in table_lines]
            if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in parsed[1]):
                parsed.pop(1)
            _add_table(document, parsed)
            continue
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            paragraph = document.add_paragraph(style="List Number")
            _set_font(paragraph.add_run(_plain_markdown(text)), "宋体", "Times New Roman", 12)
            index += 1
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _set_font(paragraph.add_run(_plain_markdown(line[2:])), "宋体", "Times New Roman", 12)
            index += 1
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        _set_font(paragraph.add_run(_plain_markdown(line)), "宋体", "Times New Roman", 12)
        index += 1
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)
    return docx_path


def _shade_cell_like_paragraph(paragraph: Any, fill: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    paragraph_properties.append(shading)


def generate() -> tuple[Path, ...]:
    runs, benchmark = _load_evidence()
    assets = build_assets()
    SEARCH_MD.write_text(build_search_markdown(runs), encoding="utf-8")
    TERMINAL_MD.write_text(build_terminal_markdown(runs, benchmark), encoding="utf-8")
    build_docx(
        SEARCH_MD,
        SEARCH_DOCX,
        title="协同搜索试验报告",
        subtitle="算法原理、试验配置与结果分析",
    )
    build_docx(
        TERMINAL_MD,
        TERMINAL_DOCX,
        title="末端目标配准试验报告",
        subtitle="中心交接、机间关联与试验结果",
    )
    return (SEARCH_MD, SEARCH_DOCX, TERMINAL_MD, TERMINAL_DOCX, *assets)


def main() -> int:
    for path in generate():
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
