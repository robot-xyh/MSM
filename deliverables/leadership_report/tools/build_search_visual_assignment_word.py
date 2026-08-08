#!/usr/bin/env python3
"""Build the four coarse-cue leadership reports as A4 Word documents."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPORT_ROOT = Path(__file__).resolve().parents[1]
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
BODY_SIZE = 12.0

BLUE = "1F4E78"
TEAL = "176B73"
INK = "202833"
MUTED = "5F6B78"
RED = "9C2F2F"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|`.+?`)")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
FORMULA_RE = re.compile(r"^\$\$(.+)\$\$$")


@dataclass(frozen=True)
class ReportSpec:
    source_name: str
    output_name: str
    title: str
    subtitle: str
    keywords: str
    expected_images: int
    required_sections: tuple[str, ...]


REPORTS = (
    ReportSpec(
        source_name="INTERCEPTION_REGION_REPARTITION_SOLUTION_CN.md",
        output_name="INTERCEPTION_REGION_REPARTITION_SOLUTION_CN.docx",
        title="拦截区域重划方案",
        subtitle="八区域空中二级节点条件下的智能搜索责任区设计",
        keywords="八区域, 空中二级节点, 区域重划, 强化学习, 三维搜索单元",
        expected_images=11,
        required_sections=(
            "背景",
            "存在的难点",
            "提出的方案",
            "传统方案",
            "人工智能方案",
            "关键技术",
            "训练区域重划策略",
            "实施方案",
            "验证方法",
        ),
    ),
    ReportSpec(
        source_name="INTERCEPTOR_COOPERATIVE_GIMBAL_SEARCH_SOLUTION_CN.md",
        output_name="INTERCEPTOR_COOPERATIVE_GIMBAL_SEARCH_SOLUTION_CN.docx",
        title="拦截无人机协同搜索方案",
        subtitle="盘旋二级节点与拦截无人机的智能主动搜索",
        keywords="空中二级节点, 协同搜索, 强化学习, 主动视觉, 概率搜索",
        expected_images=11,
        required_sections=(
            "背景",
            "存在的难点",
            "提出的方案",
            "传统方案",
            "人工智能方案",
            "关键技术",
            "连续决策模型",
            "实施方案",
            "验证方法",
        ),
    ),
    ReportSpec(
        source_name="VISUAL_REGISTRATION_SECTION_CN.md",
        output_name="VISUAL_REGISTRATION_SECTION_CN.docx",
        title="多无人机目标配准方案",
        subtitle="不同相机观察不同目标子集条件下的智能身份对应",
        keywords="多无人机, 目标配准, 图神经网络, 多视角几何, 轨迹关联",
        expected_images=11,
        required_sections=(
            "背景",
            "存在的难点",
            "提出的方案",
            "传统方案",
            "人工智能方案",
            "关键技术",
            "稀疏轨迹图",
            "实施方案",
            "验证方法",
        ),
    ),
    ReportSpec(
        source_name="LOCAL_TARGET_ASSIGNMENT_SOLUTION_CN.md",
        output_name="LOCAL_TARGET_ASSIGNMENT_SOLUTION_CN.docx",
        title="任务组本地目标分配方案",
        subtitle="空中二级节点主持的学习增强滚动任务组织",
        keywords="空中二级节点, 目标分配, 强化学习, 匈牙利算法, 失效接替",
        expected_images=11,
        required_sections=(
            "背景",
            "存在的难点",
            "提出的方案",
            "传统方案",
            "人工智能方案",
            "关键技术",
            "学习代价修正",
            "实施方案",
            "验证方法",
        ),
    ),
)


def _set_run_font(
    run,
    *,
    size: float = 11,
    bold: bool | None = None,
    color: str = INK,
    east_asia: str = BODY_FONT,
    italic: bool | None = None,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_inline(paragraph, text: str, *, size: float = BODY_SIZE, color: str = INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            _set_run_font(paragraph.add_run(text[position : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            _set_run_font(
                paragraph.add_run(token[2:-2]),
                size=size,
                bold=True,
                color=color,
                east_asia=HEADING_FONT,
            )
        elif token.startswith("*"):
            _set_run_font(
                paragraph.add_run(token[1:-1]),
                size=size,
                color=color,
                east_asia=BODY_FONT,
                italic=True,
            )
        else:
            _set_run_font(
                paragraph.add_run(token[1:-1]),
                size=size,
                color=BLUE,
                east_asia=BODY_FONT,
            )
        position = match.end()
    if position < len(text):
        _set_run_font(paragraph.add_run(text[position:]), size=size, color=color)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    _set_run_font(run, size=8.5, color=MUTED, east_asia=HEADING_FONT)


def _configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    specs = {
        "Heading 1": (16, BLUE, 14, 8),
        "Heading 2": (13.5, TEAL, 11, 6),
        "Heading 3": (12, INK, 9, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)


def _add_header_footer(section, title: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_inline(header, title, size=8.5, color=MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline(footer, "MSM 项目组  ·  ", size=8.5, color=MUTED)
    _add_page_number(footer)


def _add_cover(document: Document, spec: ReportSpec) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(title.add_run(spec.title), size=27, bold=True, color=INK, east_asia=HEADING_FONT)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(subtitle.add_run(spec.subtitle), size=15, bold=True, color=BLUE, east_asia=HEADING_FONT)
    for _ in range(9):
        document.add_paragraph()
    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(owner.add_run("MSM 项目组"), size=12, color=INK)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(date.add_run("2026 年 8 月"), size=11, color=MUTED)
    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(boundary.add_run("科研仿真与技术论证材料"), size=9.5, color=TEAL)


def _add_image(document: Document, source: Path, alt: str, relative: str, number: int) -> None:
    image_path = (source.parent / relative).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(16.1))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(4)
    caption_text = re.sub(r"^图\s*\d+\s*", "", alt).strip()
    _set_run_font(caption.add_run(f"图 {number}  {caption_text}"), size=9, color=MUTED)


def _add_formula(document: Document, formula: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.right_indent = Cm(0.45)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_together = True

    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F7F9")
    properties.append(shading)

    run = paragraph.add_run(formula.strip())
    run.font.name = MATH_FONT
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), MATH_FONT)
    fonts.set(qn("w:hAnsi"), MATH_FONT)
    fonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(INK)


def _add_list_paragraph(
    document: Document,
    marker: str,
    content: str,
    *,
    compact: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(-0.58)
    paragraph.paragraph_format.space_after = Pt(2 if compact else 4)
    if compact:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    size = 10.5 if compact else BODY_SIZE
    _set_run_font(paragraph.add_run(f"{marker}  "), size=size, bold=True, color=BLUE, east_asia=HEADING_FONT)
    _add_inline(paragraph, content, size=size)


def build_document(spec: ReportSpec) -> Path:
    source = REPORT_ROOT / spec.source_name
    output = REPORT_ROOT / spec.output_name
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    _configure_section(document.sections[0])
    _configure_styles(document)
    _add_cover(document, spec)

    content_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(content_section)
    _add_header_footer(content_section, spec.title)

    image_number = 0
    in_references = False
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            if heading_match.group(2).strip() == "参考资料":
                in_references = True
                document.add_page_break()
            level = min(3, len(heading_match.group(1)) - 1)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline(paragraph, heading_match.group(2), size=(16, 13.5, 12)[level - 1], color=(BLUE, TEAL, INK)[level - 1])
            index += 1
            continue
        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            image_number += 1
            _add_image(document, source, image_match.group(1), image_match.group(2), image_number)
            index += 1
            continue
        formula_match = FORMULA_RE.fullmatch(line)
        if formula_match:
            _add_formula(document, formula_match.group(1))
            index += 1
            continue
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            note = document.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note.paragraph_format.first_line_indent = Cm(0)
            note.paragraph_format.space_after = Pt(6)
            note_text = re.sub(r"^图\s*\d+\s*", "说明：", line[1:-1].strip())
            _set_run_font(note.add_run(note_text), size=9, color=MUTED, italic=True)
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            _add_list_paragraph(
                document,
                f"{numbered.group(1)}.",
                numbered.group(2),
                compact=in_references,
            )
            index += 1
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if bullet:
            _add_list_paragraph(document, "•", bullet.group(1))
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate or candidate.startswith(("#", "![", "*", "- ")):
                break
            if re.match(r"^\d+\.\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        _add_inline(paragraph, " ".join(parts), size=BODY_SIZE)
        index = lookahead

    properties = document.core_properties
    properties.title = spec.title
    properties.subject = spec.subtitle
    properties.author = "MSM 项目组"
    properties.keywords = spec.keywords
    document.save(output)
    return output


def validate_document(spec: ReportSpec, output: Path) -> dict[str, int]:
    document = Document(output)
    with ZipFile(output) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
    if len(images) != spec.expected_images:
        raise RuntimeError(f"{output.name}: expected {spec.expected_images} images, found {len(images)}")
    if len(document.paragraphs) < 30:
        raise RuntimeError(f"{output.name}: generated document is unexpectedly short")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for required in spec.required_sections:
        if required not in text:
            raise RuntimeError(f"{output.name}: missing section {required}")
    return {"paragraphs": len(document.paragraphs), "images": len(images), "bytes": output.stat().st_size}


def build_all() -> tuple[Path, ...]:
    outputs = []
    for spec in REPORTS:
        output = build_document(spec)
        metrics = validate_document(spec, output)
        print(
            f"{output.name}: paragraphs={metrics['paragraphs']}, "
            f"images={metrics['images']}, bytes={metrics['bytes']}"
        )
        outputs.append(output)
    return tuple(outputs)


def main() -> None:
    build_all()


if __name__ == "__main__":
    main()
