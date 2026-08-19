#!/usr/bin/env python3
"""Build figures, evidence manifest, and Word output for the dual-optical report."""

from __future__ import annotations

import ast
import hashlib
import json
import math
import re
import warnings
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
warnings.filterwarnings("ignore", message="Unable to import Axes3D.*")

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Polygon, Rectangle
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[3]
REPORT_ROOT = REPO_ROOT / "deliverables" / "leadership_report"
REPORT_MD = REPORT_ROOT / "双光电多目标轨迹配准与交汇定位试验报告_CN.md"
REPORT_DOCX = REPORT_MD.with_suffix(".docx")
ASSET_DIR = REPORT_ROOT / "assets" / "dual_optical_registration_report"
EVIDENCE_MANIFEST = REPORT_ROOT / "双光电多目标轨迹配准与交汇定位试验报告_EVIDENCE.json"

S180_ROOT = (
    REPO_ROOT
    / "research_modules"
    / "independent_experiments"
    / "dual_optical_online_benchmark"
    / "outputs"
    / "s180_1s_sector_v1"
)
S180_METRICS = S180_ROOT / "s180_combined_metrics.json"
S180_REPRODUCTION = S180_ROOT / "reproduction_manifest.json"

CLEAN_LIGHT_ROOT = (
    REPO_ROOT
    / "research_modules"
    / "independent_experiments"
    / "dual_optical_online_benchmark"
    / "outputs"
    / "clean_light_sealed_rescore_20260817_v2"
)
CLEAN_LIGHT_METRICS = CLEAN_LIGHT_ROOT / "clean_light_metrics.json"

CONTINUOUS_360_ROOT = (
    REPO_ROOT
    / "research_modules"
    / "independent_experiments"
    / "dual_optical_online_benchmark"
    / "outputs"
    / "scale_funnel_v3"
)
CONTINUOUS_360_SUMMARY = CONTINUOUS_360_ROOT / "summary" / "scale_funnel_summary.json"
CONTINUOUS_360_TARGET_COUNTS = (20, 40, 60)

RANGING_ROOT = (
    REPO_ROOT
    / "research_modules"
    / "independent_experiments"
    / "dual_optical_40target"
    / "outputs"
    / "airsim_seed_20260810_run11"
)
RANGING_METRICS = RANGING_ROOT / "metrics.json"
RANGING_SCENARIO = RANGING_ROOT / "scenario.json"
RANGING_TRACKS = RANGING_ROOT / "truth" / "target_trajectories.csv"
RANGING_MATCHES = RANGING_ROOT / "online" / "cross_camera_matches.csv"
RANGING_MATCH_SCORES = RANGING_ROOT / "truth" / "match_scoring.csv"
RANGING_DETECTIONS = RANGING_ROOT / "online" / "anonymous_detections.csv"

FONT = "Noto Sans CJK SC"
SERIF_FONT = "Noto Serif CJK SC"
FONT_PATH = Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc")
BLUE = "#245B78"
ORANGE = "#C96D2D"
GREEN = "#2E7D5A"
RED = "#A9473E"
INK = "#1F2933"
MUTED = "#65727E"
GRID = "#D7DEE5"
LIGHT_BLUE = "#E8F0F5"
LIGHT_ORANGE = "#F8EDE4"
LIGHT_GREEN = "#E8F3ED"
LIGHT_GRAY = "#F3F5F7"

BODY_FONT = "宋体"
HEADING_FONT = "黑体"
LATIN_FONT = "Times New Roman"
WORD_BLUE = "1F4E78"
WORD_TEAL = "176B73"
WORD_INK = "202833"
WORD_MUTED = "5F6B78"

IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
TABLE_DIVIDER_RE = re.compile(r"^\|(?:\s*:?-+:?\s*\|)+$")

S180_TARGET_COUNTS = (20, 40, 60)
S180_CONDITIONS = ("clean", "light")
ROUTE_ORDER = ("epipolar_mht", "gnn", "track_superglue")
ROUTE_LABELS_CN = {
    "epipolar_mht": "几何方法",
    "gnn": "图神经网络",
    "track_superglue": "增强型图神经网络（航迹级注意力）",
}
ROUTE_FIGURE_LABELS_CN = {
    "epipolar_mht": "几何方法",
    "gnn": "图神经网络",
    "track_superglue": "增强型图网络\n（航迹级注意力）",
}

EXPECTED_S180_FINAL = {
    (20, "clean", "epipolar_mht"): (0.0, 0.0, 5),
    (20, "clean", "gnn"): (0.978021978021978, 0.89, 0),
    (20, "clean", "track_superglue"): (0.8850574712643678, 0.77, 0),
    (20, "light", "epipolar_mht"): (0.0, 0.0, 5),
    (20, "light", "gnn"): (0.8390804597701149, 0.73, 0),
    (20, "light", "track_superglue"): (0.8876404494382022, 0.79, 0),
    (40, "clean", "epipolar_mht"): (0.0, 0.0, 5),
    (40, "clean", "gnn"): (0.8956043956043956, 0.8150000000000001, 0),
    (40, "clean", "track_superglue"): (0.7748344370860927, 0.585, 0),
    (40, "light", "epipolar_mht"): (0.0, 0.0, 5),
    (40, "light", "gnn"): (0.8930817610062893, 0.71, 0),
    (40, "light", "track_superglue"): (0.7905405405405406, 0.585, 0),
    (60, "clean", "epipolar_mht"): (0.0, 0.0, 5),
    (60, "clean", "gnn"): (0.9577464788732394, 0.9066666666666666, 0),
    (60, "clean", "track_superglue"): (0.0, 0.0, 5),
    (60, "light", "epipolar_mht"): (0.0, 0.0, 5),
    (60, "light", "gnn"): (0.9461538461538461, 0.82, 0),
    (60, "light", "track_superglue"): (0.0, 0.0, 5),
}

EXPECTED_RANGING = {
    "correct_match_count": 36,
    "false_match_count": 1,
    "association_precision": 0.972972972972973,
    "association_full_target_recall": 0.9,
    "position_error_mean_m": 0.08030544908018379,
    "position_error_p95_m": 0.09140807343853961,
    "velocity_error_mean_mps": 0.008106441752602696,
    "velocity_error_p95_mps": 0.019719089744373565,
}

EXPECTED_CLEAN_LIGHT_FINAL = {
    ("epipolar_mht", "clean"): (0.9883720930232558, 0.85),
    ("epipolar_mht", "light"): (0.8545454545454545, 0.47),
    ("gnn", "clean"): (0.9146341463414634, 0.75),
    ("gnn", "light"): (0.7721518987341772, 0.61),
    ("track_superglue", "clean"): (0.922077922077922, 0.71),
    ("track_superglue", "light"): (0.8904109589041096, 0.65),
}

EXPECTED_CONTINUOUS_360_FINAL = {
    (20, "clean"): (0.97, 1.0, 0.84, 0.63, 0.20, 0.95),
    (20, "light"): (0.96, 0.9595959595959596, 0.821917808219178, 0.60, 0.25, 0.85),
    (20, "medium"): (0.87, 0.9680851063829787, 0.765625, 0.49, 0.20, 0.80),
    (20, "heavy"): (0.83, 0.8695652173913043, 0.6935483870967742, 0.43, 0.10, 0.80),
    (40, "clean"): (0.885, 0.9417989417989417, 0.9166666666666666, 0.605, 0.40, 0.80),
    (40, "light"): (0.90, 0.900523560209424, 0.80, 0.50, 0.325, 0.70),
    (40, "medium"): (0.77, 0.8034682080924855, 0.7608695652173914, 0.35, 0.20, 0.55),
    (40, "heavy"): (0.78, 0.7457627118644068, 0.6707317073170732, 0.275, 0.175, 0.35),
    (60, "clean"): (0.82, 0.832089552238806, 0.9428571428571428, 0.55, 0.4666666666666667, 0.6166666666666667),
    (60, "light"): (0.79, 0.8068181818181818, 0.7727272727272727, 0.39666666666666667, 0.35, 0.43333333333333335),
    (60, "medium"): (0.7566666666666667, 0.84765625, 0.7756410256410257, 0.4033333333333333, 0.31666666666666665, 0.48333333333333334),
    (60, "heavy"): (0.7333333333333333, 0.7086614173228346, 0.7226890756302521, 0.2866666666666667, 0.25, 0.31666666666666665),
}


def configure_matplotlib() -> None:
    if FONT_PATH.exists():
        font_manager.fontManager.addfont(str(FONT_PATH))
        configured_font = font_manager.FontProperties(fname=str(FONT_PATH)).get_name()
    else:
        configured_font = FONT
    matplotlib.rcParams.update(
        {
            "font.family": configured_font,
            "font.sans-serif": [configured_font, FONT, "Noto Sans CJK SC"],
            "axes.unicode_minus": False,
            "axes.edgecolor": MUTED,
            "axes.labelcolor": INK,
            "xtick.color": MUTED,
            "ytick.color": MUTED,
            "text.color": INK,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def final_window_latency_p95_ms(rows: list[dict], *, expected_count: int = 5) -> float:
    """Return P95 of end-to-end latency for one final-window, five-scene batch."""
    if len(rows) != expected_count:
        raise RuntimeError(f"expected {expected_count} final-window latency rows, found {len(rows)}")
    values = []
    for row in rows:
        value = row.get("end_to_end_ms")
        if value is None or not math.isfinite(float(value)):
            raise RuntimeError("final-window row has no finite end_to_end_ms")
        values.append(float(value))
    return float(np.percentile(np.asarray(values, dtype=float), 95))


def relative(path: Path) -> str:
    return path.resolve().relative_to(REPO_ROOT.resolve()).as_posix()


def save_figure(figure: plt.Figure, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    figure.savefig(path, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close(figure)
    return path


def rounded_box(
    axis,
    xy: tuple[float, float],
    width: float,
    height: float,
    text: str,
    *,
    facecolor: str,
    edgecolor: str,
    fontsize: float = 13,
) -> None:
    patch = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.012,rounding_size=0.015",
        facecolor=facecolor,
        edgecolor=edgecolor,
        linewidth=1.8,
    )
    axis.add_patch(patch)
    axis.text(
        xy[0] + width / 2,
        xy[1] + height / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        color=INK,
        linespacing=1.35,
    )


def arrow(axis, start: tuple[float, float], end: tuple[float, float], *, color: str = INK) -> None:
    axis.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=15,
            linewidth=1.7,
            color=color,
            connectionstyle="arc3,rad=0",
        )
    )


def project_points(
    points: np.ndarray,
    *,
    center: np.ndarray,
    scale: np.ndarray,
    azimuth_deg: float = -55.0,
    elevation_deg: float = 25.0,
) -> np.ndarray:
    """Orthographically project 3-D points without depending on mpl_toolkits."""
    normalized = (np.asarray(points, dtype=float) - center) / scale
    azimuth = math.radians(azimuth_deg)
    elevation = math.radians(elevation_deg)
    horizontal = math.cos(azimuth) * normalized[:, 0] - math.sin(azimuth) * normalized[:, 1]
    depth = math.sin(azimuth) * normalized[:, 0] + math.cos(azimuth) * normalized[:, 1]
    vertical = -math.sin(elevation) * depth + math.cos(elevation) * normalized[:, 2]
    return np.column_stack([horizontal, vertical])


def draw_projected_axes(
    axis,
    *,
    origin: np.ndarray,
    vectors: tuple[np.ndarray, np.ndarray, np.ndarray],
    center: np.ndarray,
    scale: np.ndarray,
    labels: tuple[str, str, str],
) -> None:
    projected_origin = project_points(origin[None, :], center=center, scale=scale)[0]
    for vector, label, color in zip(vectors, labels, (BLUE, ORANGE, GREEN)):
        endpoint = project_points((origin + vector)[None, :], center=center, scale=scale)[0]
        axis.annotate(
            "",
            xy=endpoint,
            xytext=projected_origin,
            arrowprops={"arrowstyle": "-|>", "color": color, "linewidth": 1.5},
        )
        axis.text(endpoint[0], endpoint[1], f"  {label}", fontsize=10.5, color=color)


def continuous_360_metrics_path(target_count: int) -> Path:
    return (
        CONTINUOUS_360_ROOT
        / f"targets_{target_count:03d}"
        / "results"
        / "comparison_metrics.json"
    )


def continuous_360_diagnostics_path(target_count: int) -> Path:
    return (
        CONTINUOUS_360_ROOT
        / f"targets_{target_count:03d}"
        / "results"
        / "failure_diagnostics.json"
    )


def load_continuous_360_evidence() -> dict:
    clean_light = json.loads(CLEAN_LIGHT_METRICS.read_text(encoding="utf-8"))
    if clean_light.get("truth_used_online") is not False:
        raise RuntimeError("clean/light rescore does not record online truth isolation")
    if clean_light.get("source_mode") != "sealed_publication_rescore":
        raise RuntimeError("clean/light evidence is not a sealed publication rescore")
    protocol = clean_light["protocol"]
    if (
        protocol.get("scan_span_deg") != 360.0
        or protocol.get("scan_period_s") != 2.0
        or protocol.get("duration_s") != 12.0
        or protocol.get("target_count") != 20
        or len(protocol.get("test_seeds", [])) != 5
    ):
        raise RuntimeError("clean/light 360-degree protocol changed")

    clean_light_final_rows = []
    for route_name, condition in EXPECTED_CLEAN_LIGHT_FINAL:
        matches = [
            row
            for row in clean_light["rows"]
            if row["route_name"] == route_name
            and row["corruption_level"] == condition
            and row["revolution_index"] == 6
        ]
        if len(matches) != 5:
            raise RuntimeError(f"expected five clean/light final-round rows: {(route_name, condition)}")
        correct_count = sum(int(row["correct_match_count"]) for row in matches)
        false_count = sum(int(row["false_association_count"]) for row in matches)
        deadline_miss_count = sum(not bool(row["deadline_met"]) for row in matches)
        row = {
            "route_name": route_name,
            "route_label_cn": ROUTE_LABELS_CN[route_name],
            "corruption_level": condition,
            "revolution_index": 6,
            "sample_count": len(matches),
            "correct_match_count": correct_count,
            "false_association_count": false_count,
            "association_precision": correct_count / (correct_count + false_count),
            "fixed_target_coverage": correct_count / (protocol["target_count"] * len(matches)),
            "latency_p95_ms": final_window_latency_p95_ms(matches),
            "deadline_miss_count": deadline_miss_count,
            "evidence_status": "diagnostic_sealed_rescore",
            "seeds": [int(item["seed"]) for item in matches],
        }
        expected = EXPECTED_CLEAN_LIGHT_FINAL[(route_name, condition)]
        actual = (
            float(row["association_precision"]),
            float(row["fixed_target_coverage"]),
        )
        if any(
            not math.isclose(value, reference, rel_tol=0.0, abs_tol=1e-12)
            for value, reference in zip(actual, expected)
        ):
            raise RuntimeError(f"clean/light metric changed: {(route_name, condition)}")
        clean_light_final_rows.append(row)

    scale_summary = json.loads(CONTINUOUS_360_SUMMARY.read_text(encoding="utf-8"))
    if scale_summary.get("completed_target_counts") != [20, 40, 60]:
        raise RuntimeError("continuous 360-degree completed target counts changed")

    final_rows = []
    scale_matrix_rows = []
    for target_count in CONTINUOUS_360_TARGET_COUNTS:
        metrics = json.loads(continuous_360_metrics_path(target_count).read_text(encoding="utf-8"))
        diagnostics = json.loads(continuous_360_diagnostics_path(target_count).read_text(encoding="utf-8"))
        protocol = metrics["protocol"]
        if (
            metrics.get("truth_used_online") is not False
            or protocol.get("scan_span_deg") != 360.0
            or protocol.get("scan_period_s") != 2.0
            or protocol.get("duration_s") != 12.0
            or protocol.get("target_count") != target_count
            or len(protocol.get("test_seeds", [])) != 5
        ):
            raise RuntimeError(f"continuous 360-degree protocol changed for {target_count} targets")

        active_routes = tuple(metrics.get("active_routes", []))
        expected_active_routes = ("epipolar_mht", "gnn") if target_count == 20 else ("gnn",)
        if active_routes != expected_active_routes:
            raise RuntimeError(
                f"continuous 360-degree active routes changed for {target_count}: {active_routes}"
            )

        for condition in ("clean", "light", "medium", "heavy"):
            for route_name in ROUTE_ORDER:
                if route_name not in active_routes:
                    scale_matrix_rows.append(
                        {
                            "target_count": target_count,
                            "corruption_level": condition,
                            "route_name": route_name,
                            "route_label_cn": ROUTE_LABELS_CN[route_name],
                            "test_seed_count": 0,
                            "association_precision": None,
                            "fixed_target_coverage": None,
                            "latency_p95_ms": None,
                            "deadline_miss_count": None,
                            "confirmed_output_available": False,
                            "evidence_status": "not_run",
                        }
                    )
                    continue

                route_rows = [
                    row
                    for row in metrics["rows"]
                    if row["route_name"] == route_name
                    and row["corruption_level"] == condition
                    and row["revolution_index"] == 6
                ]
                if len(route_rows) != 5:
                    raise RuntimeError(
                        f"expected five scale final-round rows: "
                        f"{(target_count, condition, route_name)}"
                    )
                correct_count = sum(int(row["correct_match_count"]) for row in route_rows)
                false_count = sum(int(row["false_association_count"]) for row in route_rows)
                output_count = correct_count + false_count
                deadline_miss_count = sum(not bool(row["deadline_met"]) for row in route_rows)
                confirmed_output_available = output_count > 0
                scale_matrix_rows.append(
                    {
                        "target_count": target_count,
                        "corruption_level": condition,
                        "route_name": route_name,
                        "route_label_cn": ROUTE_LABELS_CN[route_name],
                        "test_seed_count": len(route_rows),
                        "correct_match_count": correct_count,
                        "false_association_count": false_count,
                        "association_precision": (
                            correct_count / output_count if confirmed_output_available else None
                        ),
                        "fixed_target_coverage": (
                            correct_count / (target_count * len(route_rows))
                            if confirmed_output_available
                            else None
                        ),
                        "latency_p95_ms": final_window_latency_p95_ms(route_rows),
                        "deadline_miss_count": deadline_miss_count,
                        "confirmed_output_available": confirmed_output_available,
                        "evidence_status": (
                            "timeout" if deadline_miss_count == len(route_rows) else "diagnostic"
                        ),
                        "seeds": [int(row["seed"]) for row in route_rows],
                    }
                )

        diagnostic_index = {
            (row["corruption_level"], row["revolution_index"]): row
            for row in diagnostics["by_corruption_and_revolution"]
        }
        for condition in ("clean", "light", "medium", "heavy"):
            rows = [
                row
                for row in metrics["rows"]
                if row["route_name"] == "gnn"
                and row["corruption_level"] == condition
                and row["revolution_index"] == 6
            ]
            if len(rows) != 5 or any(not row["deadline_met"] for row in rows):
                raise RuntimeError(f"expected five on-time final-round rows: {(target_count, condition)}")
            correct_count = sum(int(row["correct_match_count"]) for row in rows)
            false_count = sum(int(row["false_association_count"]) for row in rows)
            retained_count = sum(int(row["candidate_true_retained_count"]) for row in rows)
            opportunity_count = sum(int(row["candidate_true_opportunity_count"]) for row in rows)
            coverages = [float(row["correct_match_count"]) / target_count for row in rows]
            stable_common_coverage = (
                float(diagnostic_index[(condition, 6)]["stable_common_truth_count"])
                / target_count
            )
            final = {
                "target_count": target_count,
                "corruption_level": condition,
                "test_seed_count": len(rows),
                "correct_match_count": correct_count,
                "false_association_count": false_count,
                "stable_common_coverage": stable_common_coverage,
                "candidate_true_retention_rate": retained_count / opportunity_count,
                "association_precision": correct_count / (correct_count + false_count),
                "fixed_target_coverage": correct_count / (target_count * len(rows)),
                "latency_p95_ms": final_window_latency_p95_ms(rows),
                "deadline_miss_count": 0,
                "seed_coverage_min": min(coverages),
                "seed_coverage_max": max(coverages),
                "seed_coverages": coverages,
                "seeds": [int(row["seed"]) for row in rows],
            }
            expected = EXPECTED_CONTINUOUS_360_FINAL[(target_count, condition)]
            actual = tuple(
                final[key]
                for key in (
                    "stable_common_coverage",
                    "candidate_true_retention_rate",
                    "association_precision",
                    "fixed_target_coverage",
                    "seed_coverage_min",
                    "seed_coverage_max",
                )
            )
            if any(
                not math.isclose(value, reference, rel_tol=0.0, abs_tol=1e-12)
                for value, reference in zip(actual, expected)
            ):
                raise RuntimeError(f"continuous 360-degree metric changed: {(target_count, condition)}")
            final_rows.append(final)

    return {
        "clean_light_metrics": clean_light,
        "clean_light_final_rows": clean_light_final_rows,
        "scale_summary": scale_summary,
        "final_rows": final_rows,
        "scale_matrix_rows": scale_matrix_rows,
    }


def load_and_validate_evidence() -> tuple[dict, list[dict], dict, dict, dict]:
    s180 = json.loads(S180_METRICS.read_text(encoding="utf-8"))
    reproduction = json.loads(S180_REPRODUCTION.read_text(encoding="utf-8"))
    ranging = json.loads(RANGING_METRICS.read_text(encoding="utf-8"))
    scenario = json.loads(RANGING_SCENARIO.read_text(encoding="utf-8"))

    if s180.get("coverage_denominator") != "fixed_target_count":
        raise RuntimeError("S180 coverage denominator changed")
    if s180.get("truth_used_online") is not False:
        raise RuntimeError("S180 online truth isolation is not recorded")

    s180_final_rows: list[dict] = []
    evidence_by_count = {item["target_count"]: item for item in s180["evidence"]}
    for target_count in S180_TARGET_COUNTS:
        evidence = evidence_by_count[target_count]
        status = "正式" if evidence["formal_use_allowed"] else "诊断"
        for corruption in S180_CONDITIONS:
            for route_name in ROUTE_ORDER:
                matches = [
                    row
                    for row in s180["summary"]
                    if row["target_count"] == target_count
                    and row["route_name"] == route_name
                    and row["corruption_level"] == corruption
                    and row["window"] == "final_round"
                ]
                if len(matches) != 1:
                    raise RuntimeError(
                        f"expected one S180 final row for {(target_count, corruption, route_name)}, "
                        f"found {len(matches)}"
                    )
                row = dict(matches[0])
                expected = EXPECTED_S180_FINAL[(target_count, corruption, route_name)]
                actual = (
                    float(row["association_precision"]),
                    float(row["fixed_denominator_coverage"]),
                    int(row["no_confirmed_output_round_count"]),
                )
                if any(
                    not math.isclose(value, reference, rel_tol=0.0, abs_tol=1e-12)
                    for value, reference in zip(actual, expected)
                ):
                    raise RuntimeError(
                        f"S180 final metric changed for {(target_count, corruption, route_name)}"
                    )
                raw_final_rows = [
                    raw
                    for raw in s180["rows"]
                    if raw["target_count"] == target_count
                    and raw["route_name"] == route_name
                    and raw["corruption_level"] == corruption
                    and raw["revolution_index"] == 12
                ]
                latency_p95_ms = final_window_latency_p95_ms(raw_final_rows)
                if not math.isclose(
                    latency_p95_ms,
                    float(row["latency_p95_ms"]),
                    rel_tol=0.0,
                    abs_tol=1e-9,
                ):
                    raise RuntimeError(
                        f"S180 final latency changed for {(target_count, corruption, route_name)}"
                    )
                row["route_label_cn"] = ROUTE_LABELS_CN[route_name]
                row["evidence_status_cn"] = status
                row["confirmed_output_available"] = row["no_confirmed_output_round_count"] == 0
                row["latency_p95_ms"] = latency_p95_ms
                row["deadline_miss_count"] = sum(
                    not bool(raw["deadline_met"]) for raw in raw_final_rows
                )
                s180_final_rows.append(row)

    for key, value in EXPECTED_RANGING.items():
        actual = ranging[key]
        if isinstance(value, int):
            if actual != value:
                raise RuntimeError(f"ranging metric changed: {key}")
        elif not math.isclose(float(actual), value, rel_tol=0.0, abs_tol=1e-12):
            raise RuntimeError(f"ranging metric changed: {key}")
    if ranging.get("online_truth_leakage_count") != 0:
        raise RuntimeError("ranging evidence reports online truth leakage")

    continuous_360 = load_continuous_360_evidence()
    return s180, s180_final_rows, ranging, scenario, continuous_360


def figure_algorithm_flow() -> Path:
    figure, axis = plt.subplots(figsize=(14, 7.4))
    axis.set_xlim(0, 1)
    axis.set_ylim(0, 1)
    axis.axis("off")
    axis.text(0.5, 0.94, "双光电多目标轨迹配准与交汇定位", ha="center", fontsize=22, fontweight="bold")

    boxes = [
        ((0.04, 0.62), "两站图像检测\n分别形成单站航迹", LIGHT_BLUE, BLUE),
        ((0.28, 0.62), "时刻与设备姿态对齐\n像素点转换为空间视线", LIGHT_GRAY, MUTED),
        ((0.52, 0.62), "共面关系粗筛\n剔除明显不可能组合", LIGHT_ORANGE, ORANGE),
        ((0.76, 0.62), "候选关系评分\n图神经网络为主", LIGHT_GREEN, GREEN),
        ((0.52, 0.25), "匈牙利算法一一分配\n连续多轮确认", LIGHT_BLUE, BLUE),
        ((0.76, 0.25), "双射线交汇定位\n短时位置与速度拟合", LIGHT_GREEN, GREEN),
    ]
    for xy, label, fill, edge in boxes:
        rounded_box(axis, xy, 0.2, 0.18, label, facecolor=fill, edgecolor=edge, fontsize=13)

    arrow(axis, (0.24, 0.71), (0.28, 0.71))
    arrow(axis, (0.48, 0.71), (0.52, 0.71))
    arrow(axis, (0.72, 0.71), (0.76, 0.71))
    arrow(axis, (0.86, 0.62), (0.66, 0.43), color=GREEN)
    arrow(axis, (0.72, 0.34), (0.76, 0.34))

    rounded_box(
        axis,
        (0.04, 0.25),
        0.38,
        0.18,
        "确定性安全边界\n几何筛选限制候选范围；一一分配防止重复占用；\n低分关系保持空匹配或待确认",
        facecolor="#FAFAFA",
        edgecolor=MUTED,
        fontsize=12,
    )
    axis.text(
        0.5,
        0.08,
        "先确认两站看到的是同一目标，再计算距离和位置。真实目标编号只用于试验结束后的离线核对。",
        ha="center",
        fontsize=12.5,
        color=MUTED,
    )
    return save_figure(figure, "01_algorithm_flow.png")


def figure_single_station_tracking() -> Path:
    rng = np.random.default_rng(20260817)
    figure, axes = plt.subplots(1, 2, figsize=(14, 6.6), gridspec_kw={"width_ratios": [1.25, 1]})
    time = np.linspace(0, 12, 25)
    tracks = {
        "目标1": 18 + 1.25 * time + 0.05 * time**2,
        "目标2": 34 - 0.55 * time + 0.035 * time**2,
        "目标3": 23 + 0.45 * time,
    }
    colors = [BLUE, ORANGE, GREEN]
    masks = [np.ones_like(time, dtype=bool), ~((time > 4.0) & (time < 6.5)), ~((time > 8.0) & (time < 9.5))]
    for (label, values), color, mask in zip(tracks.items(), colors, masks):
        observed = values + rng.normal(0, 0.18, len(time))
        axes[0].plot(time, values, color=color, linewidth=1.5, alpha=0.35)
        axes[0].scatter(time[mask], observed[mask], s=28, color=color, label=label, zorder=3)
        axes[0].plot(time[mask], observed[mask], color=color, linewidth=2.0)
    axes[0].axvspan(4.0, 6.5, color=RED, alpha=0.08)
    axes[0].annotate(
        "短时漏检\n保留预测，不立即删除",
        xy=(5.2, 28.5),
        xytext=(2.0, 39),
        arrowprops={"arrowstyle": "->", "color": RED},
        fontsize=11,
        color=RED,
    )
    axes[0].set_title("一次扫过与重访形成单站航迹", fontsize=16, fontweight="bold")
    axes[0].set_xlabel("时间 / 秒")
    axes[0].set_ylabel("观测方位 / 度")
    axes[0].grid(True, color=GRID, linewidth=0.7)
    axes[0].legend(frameon=False, ncol=3, loc="upper center")

    axes[1].axis("off")
    axes[1].set_title("目标规模增大后的主要风险", fontsize=16, fontweight="bold", pad=12)
    states = [
        (0.73, "连续航迹", "同一目标跨扫描重访\n保持原编号", LIGHT_GREEN, GREEN),
        (0.44, "航迹碎片", "漏检后另建新编号\n正确目标没有完整节点", LIGHT_ORANGE, ORANGE),
        (0.15, "错误重接", "交叉目标被接入同一航迹\n后续配准输入已混合", "#F8E8E7", RED),
    ]
    for y, title, detail, fill, edge in states:
        rounded_box(axes[1], (0.08, y), 0.84, 0.19, f"{title}\n{detail}", facecolor=fill, edgecolor=edge, fontsize=12)
    axes[1].text(
        0.5,
        0.04,
        "双站算法只能比较已经形成的局部航迹，不能补回从未成轨的目标。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    figure.tight_layout(w_pad=3.0)
    return save_figure(figure, "02_single_station_tracking.png")


def figure_coplanarity_screening() -> Path:
    figure, axis = plt.subplots(figsize=(12, 8))
    station_a = np.array([-1.0, 0.0, 0.0])
    station_b = np.array([1.0, 0.0, 0.0])
    target = np.array([0.25, 2.3, 1.05])
    wrong_target = np.array([-0.45, 2.1, 1.9])
    plane = np.array(
        [
            station_a,
            station_b,
            station_b + 1.12 * (target - station_a),
            station_a + 1.12 * (target - station_a),
        ]
    )
    center = np.array([0.0, 1.25, 0.9])
    scale = np.array([1.5, 1.55, 1.3])
    projected_plane = project_points(plane, center=center, scale=scale)
    axis.add_patch(Polygon(projected_plane, closed=True, facecolor="#DCEAF2", edgecolor=BLUE, linewidth=1.2, alpha=0.42))

    def draw_segment(start: np.ndarray, end: np.ndarray, *, color: str, label: str, style: str = "-") -> None:
        projected = project_points(np.vstack([start, end]), center=center, scale=scale)
        axis.plot(projected[:, 0], projected[:, 1], color=color, linewidth=2.6, linestyle=style, label=label)

    draw_segment(station_a, station_b, color=INK, label="两站基线")
    draw_segment(station_a, target, color=BLUE, label="左站候选视线")
    draw_segment(station_b, target, color=GREEN, label="同目标视线")
    draw_segment(station_b, wrong_target, color=RED, label="错误候选视线", style="--")
    projected_points = project_points(np.vstack([station_a, station_b, target, wrong_target]), center=center, scale=scale)
    axis.scatter(*projected_points[0], s=150, marker="^", color=BLUE, zorder=5)
    axis.scatter(*projected_points[1], s=150, marker="^", color=ORANGE, zorder=5)
    axis.scatter(*projected_points[2], s=110, color=GREEN, zorder=5)
    axis.scatter(*projected_points[3], s=110, color=RED, marker="x", zorder=5)
    for projected, label, color in zip(
        projected_points,
        ("左站", "右站", "同一目标", "其他目标"),
        (INK, INK, GREEN, RED),
    ):
        axis.text(projected[0] + 0.025, projected[1] + 0.025, label, fontsize=12, color=color)
    plane_label = project_points(np.array([[0.0, 1.2, 0.55]]), center=center, scale=scale)[0]
    axis.text(plane_label[0], plane_label[1], "共面候选区域", fontsize=13, color=BLUE, ha="center")
    draw_projected_axes(
        axis,
        origin=np.array([-1.25, -0.05, 0.0]),
        vectors=(np.array([0.65, 0.0, 0.0]), np.array([0.0, 0.75, 0.0]), np.array([0.0, 0.0, 0.6])),
        center=center,
        scale=scale,
        labels=("横向", "纵向", "高度"),
    )
    axis.set_title("共面关系先排除明显不可能的航迹组合", fontsize=19, fontweight="bold", pad=18)
    axis.set_aspect("equal", adjustable="datalim")
    axis.axis("off")
    axis.legend(loc="upper left", frameon=False)
    figure.text(
        0.5,
        0.04,
        "三维正交投影。设备姿态、观测时间和航迹协方差共同决定门限；通过共面条件不等于已经确认身份。",
        ha="center",
        fontsize=12,
        color=MUTED,
    )
    return save_figure(figure, "03_coplanarity_screening_3d.png")


def _draw_bipartite(axis, *, scored: bool) -> None:
    count = 8
    y_values = np.linspace(0.86, 0.12, count)
    left_x, right_x = 0.13, 0.87
    axis.set_xlim(0, 1)
    axis.set_ylim(0, 1)
    axis.axis("off")
    if not scored:
        for index, y in enumerate(y_values):
            candidates = {index, max(0, index - 1), min(count - 1, index + 1)}
            if index in {1, 4, 6}:
                candidates.add((index + 3) % count)
            for candidate in sorted(candidates):
                axis.plot(
                    [left_x, right_x],
                    [y, y_values[candidate]],
                    color=GREEN if candidate == index else "#C9D2D9",
                    linewidth=2.2 if candidate == index else 0.9,
                    alpha=0.85 if candidate == index else 0.75,
                    zorder=1,
                )
        axis.set_title("共面筛选后的候选关系", fontsize=16, fontweight="bold")
    else:
        for index, y in enumerate(y_values):
            axis.plot([left_x, right_x], [y, y], color=colors_for_index(index), linewidth=2.8, zorder=1)
        axis.set_title("图网络评分与一一分配后", fontsize=16, fontweight="bold")
    for index, y in enumerate(y_values):
        axis.scatter(left_x, y, s=440, color=BLUE, edgecolor="white", linewidth=1.2, zorder=3)
        axis.scatter(right_x, y, s=440, color=ORANGE, edgecolor="white", linewidth=1.2, zorder=3)
        axis.text(left_x, y, str(index + 1), ha="center", va="center", color="white", fontsize=10.5, zorder=4)
        axis.text(right_x, y, str(index + 1), ha="center", va="center", color="white", fontsize=10.5, zorder=4)
    axis.text(left_x, 0.98, "左站航迹", ha="center", fontsize=12, fontweight="bold")
    axis.text(right_x, 0.98, "右站航迹", ha="center", fontsize=12, fontweight="bold")


def colors_for_index(index: int) -> str:
    palette = [BLUE, ORANGE, GREEN, "#7B5EA7", "#B88A2E", "#3F8F9C", "#A64F4A", "#65727E"]
    return palette[index % len(palette)]


def figure_candidate_graph() -> Path:
    figure, axes = plt.subplots(1, 2, figsize=(14, 7))
    _draw_bipartite(axes[0], scored=False)
    _draw_bipartite(axes[1], scored=True)
    figure.suptitle("候选关系从多对多收敛为一一对应", fontsize=20, fontweight="bold", y=0.99)
    figure.text(
        0.5,
        0.025,
        "图神经网络比较候选关系的相对可信度；匈牙利算法负责最终一一选择；低分候选允许空匹配。",
        ha="center",
        fontsize=12,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.06, 1, 0.94), w_pad=3.0)
    return save_figure(figure, "04_candidate_graph_gnn_assignment.png")


def figure_multitime_triangulation() -> Path:
    figure, axis = plt.subplots(figsize=(12, 8))
    station_a = np.array([0.0, -1.0, 0.0])
    station_b = np.array([0.0, 1.0, 0.0])
    times = [0.0, 0.55, 1.1]
    points = [np.array([1.6 + 0.58 * t, -0.1 + 0.17 * t, 0.82 + 0.04 * t]) for t in times]
    colors = [BLUE, GREEN, ORANGE]
    center = np.array([1.1, 0.0, 0.55])
    scale = np.array([1.55, 1.35, 0.85])
    projected_track = project_points(np.array(points), center=center, scale=scale)
    axis.plot(projected_track[:, 0], projected_track[:, 1], color=RED, linewidth=3.0, marker="o", label="目标短时轨迹")
    for index, (point, color) in enumerate(zip(points, colors), start=1):
        for station in (station_a, station_b):
            projected = project_points(np.vstack([station, point]), center=center, scale=scale)
            axis.plot(projected[:, 0], projected[:, 1], color=color, linewidth=1.9, alpha=0.86)
        projected_point = project_points(point[None, :], center=center, scale=scale)[0]
        axis.text(projected_point[0] + 0.025, projected_point[1] + 0.02, f"时刻{index}", fontsize=11, color=color)
    stations = project_points(np.vstack([station_a, station_b]), center=center, scale=scale)
    axis.scatter(*stations[0], s=160, marker="^", color=BLUE, zorder=5)
    axis.scatter(*stations[1], s=160, marker="^", color=ORANGE, zorder=5)
    axis.text(stations[0, 0] + 0.03, stations[0, 1], "左站", fontsize=12)
    axis.text(stations[1, 0] + 0.03, stations[1, 1], "右站", fontsize=12)
    axis.plot(stations[:, 0], stations[:, 1], color=INK, linewidth=3.0, label="两站基线")
    draw_projected_axes(
        axis,
        origin=np.array([-0.05, -1.1, 0.0]),
        vectors=(np.array([0.7, 0.0, 0.0]), np.array([0.0, 0.65, 0.0]), np.array([0.0, 0.0, 0.45])),
        center=center,
        scale=scale,
        labels=("纵向", "横向", "高度"),
    )
    axis.set_title("稳定配准后使用多时刻双射线交汇定位", fontsize=19, fontweight="bold", pad=18)
    axis.set_aspect("equal", adjustable="datalim")
    axis.axis("off")
    axis.legend(loc="upper left", frameon=False)
    figure.text(
        0.5,
        0.04,
        "三维正交投影。每个时刻取两条视线最近点的中点，再用多个时刻拟合位置和速度。",
        ha="center",
        fontsize=12,
        color=MUTED,
    )
    return save_figure(figure, "05_multitime_triangulation_3d.png")


def figure_airsim_scene(scenario: dict) -> Path:
    tracks = pd.read_csv(RANGING_TRACKS)
    scene = scenario["scenario"]
    figure, axis = plt.subplots(figsize=(12.5, 8.4))
    cmap = plt.get_cmap("turbo")
    center = np.array([1250.0, 0.0, 100.0])
    scale = np.array([1500.0, 1200.0, 42.0])
    for index, (_, group) in enumerate(tracks.groupby("truth_id")):
        sampled = group.iloc[::25]
        points = np.column_stack(
            [sampled["px_ned_m"].to_numpy(), sampled["py_ned_m"].to_numpy(), -sampled["pz_ned_m"].to_numpy()]
        )
        projected = project_points(points, center=center, scale=scale)
        axis.plot(
            projected[:, 0],
            projected[:, 1],
            color=cmap(index / 40),
            linewidth=1.2,
            alpha=0.76,
        )
    station_a = np.asarray(scene["camera_a_position_ned"], dtype=float)
    station_b = np.asarray(scene["camera_b_position_ned"], dtype=float)
    station_points = np.vstack(
        [
            [station_a[0], station_a[1], -station_a[2]],
            [station_b[0], station_b[1], -station_b[2]],
        ]
    )
    projected_stations = project_points(station_points, center=center, scale=scale)
    axis.scatter(*projected_stations[0], s=190, marker="^", color=BLUE, label="左站光电", zorder=5)
    axis.scatter(*projected_stations[1], s=190, marker="^", color=ORANGE, label="右站光电", zorder=5)
    axis.plot(projected_stations[:, 0], projected_stations[:, 1], color=INK, linewidth=2.5)
    axis.text(projected_stations[0, 0] + 0.025, projected_stations[0, 1], "左站", fontsize=11)
    axis.text(projected_stations[1, 0] + 0.025, projected_stations[1, 1], "右站", fontsize=11)
    draw_projected_axes(
        axis,
        origin=np.array([0.0, -1000.0, 78.0]),
        vectors=(np.array([650.0, 0.0, 0.0]), np.array([0.0, 450.0, 0.0]), np.array([0.0, 0.0, 18.0])),
        center=center,
        scale=scale,
        labels=("前向", "横向", "高度"),
    )
    axis.set_title("两台固定光电与40个三维运动目标", fontsize=20, fontweight="bold", pad=18)
    axis.set_aspect("equal", adjustable="datalim")
    axis.axis("off")
    axis.legend(loc="upper left", frameon=False)
    figure.text(
        0.5,
        0.03,
        "三维正交投影。目标以50米/秒运动；两站横向间隔2千米；目标最近间隔约27.1米。",
        ha="center",
        fontsize=12,
        color=MUTED,
    )
    return save_figure(figure, "06_airsim_scene_40_targets_cn.png")


def figure_optical_observations() -> Path:
    frame_index = 425
    detections = pd.read_csv(RANGING_DETECTIONS)
    figure, axes = plt.subplots(1, 2, figsize=(14, 6.6))
    for axis, camera, label, color in zip(
        axes,
        ("Optical_A", "Optical_B"),
        ("左站视图", "右站视图"),
        (BLUE, ORANGE),
    ):
        image_path = RANGING_ROOT / "keyframes" / camera / f"frame_{frame_index:05d}.png"
        image = Image.open(image_path).convert("RGB")
        axis.imshow(image)
        rows = detections[(detections["camera_id"] == camera) & (detections["frame_index"] == frame_index)]
        for _, row in rows.iterrows():
            x1, y1, x2, y2 = ast.literal_eval(row["bbox_xyxy"])
            axis.add_patch(Rectangle((x1, y1), x2 - x1, y2 - y1, fill=False, edgecolor="#E5B83B", linewidth=2.0))
        axis.set_title(f"{label}（{len(rows)}个检测）", fontsize=16, fontweight="bold", color=color)
        axis.axis("off")
    figure.suptitle("同一扫描时刻两台光电看到不同目标子集", fontsize=20, fontweight="bold")
    figure.text(
        0.5,
        0.025,
        "黄色框为匿名检测结果。两站视图不能直接按像素位置比较，需先转换为空间视线并形成局部航迹。",
        ha="center",
        fontsize=12,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.055, 1, 0.93), w_pad=1.2)
    return save_figure(figure, "07_airsim_optical_observations_cn.png")


def figure_s180_results(s180_final_rows: list[dict]) -> Path:
    indexed = {
        (row["target_count"], row["corruption_level"], row["route_name"]): row
        for row in s180_final_rows
    }
    condition_labels = {"clean": "无干扰", "light": "轻度干扰"}
    method_labels = [ROUTE_FIGURE_LABELS_CN[route] for route in ROUTE_ORDER]
    x = np.arange(len(ROUTE_ORDER))
    width = 0.34
    figure, axes = plt.subplots(2, 3, figsize=(15, 9.2), sharey=True)
    for row_index, condition in enumerate(S180_CONDITIONS):
        for column_index, target_count in enumerate(S180_TARGET_COUNTS):
            axis = axes[row_index, column_index]
            rows = [indexed[(target_count, condition, route)] for route in ROUTE_ORDER]
            precision = np.array([row["association_precision"] for row in rows]) * 100
            coverage = np.array([row["fixed_denominator_coverage"] for row in rows]) * 100
            bars_precision = axis.bar(x - width / 2, precision, width, color=BLUE, label="关联精度")
            bars_coverage = axis.bar(x + width / 2, coverage, width, color=GREEN, label="目标覆盖度")
            precision_labels = [f"{value:.1f}" if row["confirmed_output_available"] else "" for value, row in zip(precision, rows)]
            coverage_labels = [f"{value:.1f}" if row["confirmed_output_available"] else "" for value, row in zip(coverage, rows)]
            axis.bar_label(bars_precision, labels=precision_labels, padding=2, fontsize=8.2)
            axis.bar_label(bars_coverage, labels=coverage_labels, padding=2, fontsize=8.2)
            for method_index, row in enumerate(rows):
                if not row["confirmed_output_available"]:
                    axis.text(method_index, 4.5, "未形成结果", ha="center", fontsize=8.2, color=RED)
            status = rows[0]["evidence_status_cn"]
            axis.set_title(
                f"{target_count}目标  {condition_labels[condition]}（{status}）",
                fontsize=13,
                fontweight="bold",
            )
            axis.set_xticks(x, method_labels)
            axis.set_ylim(0, 106)
            axis.set_ylabel("比例 / %")
            axis.grid(axis="y", color=GRID, linewidth=0.7)
            axis.spines[["top", "right"]].set_visible(False)
    handles, labels = axes[0, 0].get_legend_handles_labels()
    figure.legend(handles, labels, frameon=False, ncol=2, loc="upper center", bbox_to_anchor=(0.5, 0.93))
    figure.suptitle("180度扫描三种方法的最终一轮结果", fontsize=20, fontweight="bold")
    figure.text(
        0.5,
        0.018,
        "每项为5个测试场景第12轮的合并结果；未形成结果表示最终一轮没有可用于统计的确认关系。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.055, 1, 0.87), h_pad=2.0, w_pad=1.4)
    return save_figure(figure, "08_s180_selected_results.png")


def figure_stage_metrics(s180_final_rows: list[dict]) -> Path:
    gnn_rows = [row for row in s180_final_rows if row["route_name"] == "gnn"]
    condition_labels = {"clean": "无干扰", "light": "轻度干扰"}
    labels = [f"{row['target_count']}目标\n{condition_labels[row['corruption_level']]}" for row in gnn_rows]
    precision = np.array([row["association_precision"] for row in gnn_rows]) * 100
    coverage = np.array([row["fixed_denominator_coverage"] for row in gnn_rows]) * 100
    x = np.arange(len(labels))
    width = 0.34
    figure, axis = plt.subplots(figsize=(13, 7))
    bars_precision = axis.bar(x - width / 2, precision, width, label="关联精度", color=BLUE)
    bars_coverage = axis.bar(x + width / 2, coverage, width, label="目标覆盖度", color=GREEN)
    axis.bar_label(bars_precision, fmt="%.1f", padding=3, fontsize=9.2)
    axis.bar_label(bars_coverage, fmt="%.1f", padding=3, fontsize=9.2)
    axis.set_ylim(60, 103)
    axis.set_ylabel("比例 / %")
    axis.set_xticks(x, labels)
    axis.set_title("图神经网络最终一轮精度与覆盖度", fontsize=20, fontweight="bold", pad=16)
    axis.grid(axis="y", color=GRID, linewidth=0.7)
    axis.spines[["top", "right"]].set_visible(False)
    axis.legend(frameon=False, ncol=2, loc="lower center")
    figure.text(
        0.5,
        0.02,
        "20目标为正式结果；40和60目标为诊断结果。所有数值均取第12轮。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.055, 1, 0.93))
    return save_figure(figure, "09_tracking_and_registration_loss.png")


def figure_360_clean_light(clean_light_rows: list[dict]) -> Path:
    indexed = {
        (row["route_name"], row["corruption_level"]): row
        for row in clean_light_rows
    }
    x = np.arange(len(ROUTE_ORDER))
    width = 0.34
    figure, axes = plt.subplots(1, 2, figsize=(14, 6.8), sharey=True)
    panels = (
        ("association_precision", "已输出关系正确比例"),
        ("fixed_target_coverage", "全部目标覆盖比例"),
    )
    for axis, (metric, title) in zip(axes, panels):
        clean = np.array([indexed[(route, "clean")][metric] for route in ROUTE_ORDER]) * 100
        light = np.array([indexed[(route, "light")][metric] for route in ROUTE_ORDER]) * 100
        clean_bars = axis.bar(x - width / 2, clean, width, color=BLUE, label="无干扰")
        light_bars = axis.bar(x + width / 2, light, width, color=ORANGE, label="轻干扰")
        axis.set_title(title, fontsize=16, fontweight="bold", pad=12)
        axis.set_xticks(x, [ROUTE_FIGURE_LABELS_CN[route] for route in ROUTE_ORDER])
        axis.set_ylim(0, 105)
        axis.set_ylabel("比例 / %")
        axis.grid(axis="y", color=GRID, linewidth=0.7)
        axis.spines[["top", "right"]].set_visible(False)
        axis.bar_label(clean_bars, fmt="%.1f", padding=3, fontsize=9.5)
        axis.bar_label(light_bars, fmt="%.1f", padding=3, fontsize=9.5)
    handles, legend_labels = axes[0].get_legend_handles_labels()
    figure.legend(handles, legend_labels, frameon=False, ncol=2, loc="upper center", bbox_to_anchor=(0.5, 0.90))
    figure.suptitle("20目标连续360度周扫的最终一圈结果", fontsize=20, fontweight="bold")
    figure.text(
        0.5,
        0.02,
        "每项为5个随机场景第6圈的合并结果；轻度干扰为3%随机漏检和每台相机每秒2个瞬时虚警。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.06, 1, 0.84), w_pad=2.4)
    return save_figure(figure, "11_360_clean_light_route_comparison.png")


def figure_360_multiseed_cascade(final_rows: list[dict]) -> Path:
    condition_order = ("clean", "light", "medium", "heavy")
    condition_labels = ("无干扰", "轻度", "中度", "重度")
    target_counts = (20, 40, 60)
    colors = {20: BLUE, 40: ORANGE, 60: GREEN}
    indexed = {
        (row["target_count"], row["corruption_level"]): row
        for row in final_rows
    }
    x = np.arange(len(condition_order))
    figure, axes = plt.subplots(1, 2, figsize=(14, 6.8), sharex=True, sharey=True)
    panels = (
        ("association_precision", "双站关联精度"),
        ("fixed_target_coverage", "最终正确覆盖"),
    )
    for axis, (metric, title) in zip(axes, panels):
        for series_index, target_count in enumerate(target_counts):
            values = np.array(
                [indexed[(target_count, condition)][metric] for condition in condition_order]
            ) * 100
            plot_x = x + (series_index - 1) * 0.045
            axis.plot(
                plot_x,
                values,
                color=colors[target_count],
                marker="o",
                markersize=7,
                linewidth=2.2,
                label=f"{target_count}目标",
            )
            label_offset = (7, 0, -9)[series_index]
            for x_value, value in zip(plot_x, values):
                axis.annotate(
                    f"{value:.1f}",
                    xy=(x_value, value),
                    xytext=(0, label_offset),
                    textcoords="offset points",
                    ha="center",
                    va="bottom" if label_offset >= 0 else "top",
                    fontsize=8.5,
                )
        axis.axhline(80, color=MUTED, linestyle="--", linewidth=1.0, alpha=0.7)
        axis.set_title(title, fontsize=15, fontweight="bold", pad=10)
        axis.set_xticks(x, condition_labels)
        axis.set_ylim(20, 106)
        axis.set_ylabel("比例 / %")
        axis.grid(axis="y", color=GRID, linewidth=0.7)
        axis.spines[["top", "right"]].set_visible(False)
    handles, legend_labels = axes[0].get_legend_handles_labels()
    figure.legend(handles, legend_labels, frameon=False, ncol=3, loc="upper center", bbox_to_anchor=(0.5, 0.925))
    figure.suptitle("连续360度周扫的无干扰及轻、中、重度干扰结果", fontsize=20, fontweight="bold")
    figure.text(
        0.5,
        0.018,
        "每个点为5个随机场景第6圈的合并结果；虚线为80%参考线。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    figure.tight_layout(rect=(0, 0.06, 1, 0.84), w_pad=2.0)
    return save_figure(figure, "12_360_multiseed_cascade.png")


def figure_ranging_reconstruction() -> Path:
    matches = pd.read_csv(RANGING_MATCHES)
    scores = pd.read_csv(RANGING_MATCH_SCORES)
    scored = matches.merge(scores, on=["match_id", "track_a_id", "track_b_id"], how="inner")
    correct = scored[scored["correct"]].copy()
    tracks = pd.read_csv(RANGING_TRACKS)

    figure = plt.figure(figsize=(14, 8))
    grid = figure.add_gridspec(2, 2, width_ratios=[1.35, 1], hspace=0.34, wspace=0.26)
    axis_3d = figure.add_subplot(grid[:, 0])
    center = np.array([1750.0, 0.0, 100.0])
    scale = np.array([1000.0, 600.0, 35.0])
    for index, (_, row) in enumerate(correct.iloc[:10].iterrows()):
        estimate_position = np.asarray(ast.literal_eval(row["position_ned"]), dtype=float)
        estimate_velocity = np.asarray(ast.literal_eval(row["velocity_ned"]), dtype=float)
        reference = float(row["reference_timestamp"])
        timeline = np.linspace(1.0, 11.0, 60)
        estimate = estimate_position[None, :] + (timeline - reference)[:, None] * estimate_velocity[None, :]
        truth = tracks[tracks["truth_id"] == row["truth_a"]].sort_values("simulation_timestamp")
        truth_window = truth[
            (truth["simulation_timestamp"] >= 1.0)
            & (truth["simulation_timestamp"] <= 11.0)
        ].iloc[::10]
        color = colors_for_index(index)
        truth_points = np.column_stack(
            [
                truth_window["px_ned_m"].to_numpy(),
                truth_window["py_ned_m"].to_numpy(),
                -truth_window["pz_ned_m"].to_numpy(),
            ]
        )
        estimate_points = np.column_stack([estimate[:, 0], estimate[:, 1], -estimate[:, 2]])
        truth_projected = project_points(truth_points, center=center, scale=scale)
        estimate_projected = project_points(estimate_points, center=center, scale=scale)
        axis_3d.plot(
            truth_projected[:, 0],
            truth_projected[:, 1],
            color="#AEB8C1",
            linewidth=2.0,
        )
        axis_3d.plot(
            estimate_projected[:, 0],
            estimate_projected[:, 1],
            color=color,
            linewidth=1.8,
            linestyle="--",
        )
    axis_3d.set_title("正确关系的短时轨迹重建", fontsize=16, fontweight="bold")
    draw_projected_axes(
        axis_3d,
        origin=np.array([900.0, -450.0, 78.0]),
        vectors=(np.array([350.0, 0.0, 0.0]), np.array([0.0, 250.0, 0.0]), np.array([0.0, 0.0, 15.0])),
        center=center,
        scale=scale,
        labels=("前向", "横向", "高度"),
    )
    axis_3d.axis("off")
    axis_3d.set_aspect("equal", adjustable="datalim")
    axis_3d.text(
        0.04,
        0.04,
        "三维正交投影\n灰色实线：真实轨迹\n彩色虚线：交汇定位后拟合",
        transform=axis_3d.transAxes,
        fontsize=10.5,
    )

    axis_pos = figure.add_subplot(grid[0, 1])
    axis_pos.hist(correct["position_error_m"], bins=8, color=BLUE, alpha=0.86, edgecolor="white")
    axis_pos.axvline(correct["position_error_m"].mean(), color=RED, linestyle="--", linewidth=1.8)
    axis_pos.set_title("位置误差分布", fontsize=15, fontweight="bold")
    axis_pos.set_xlabel("位置误差 / 米")
    axis_pos.set_ylabel("关系数量")
    axis_pos.grid(axis="y", color=GRID, linewidth=0.7)
    axis_pos.spines[["top", "right"]].set_visible(False)
    axis_pos.text(0.98, 0.88, f"平均 {correct['position_error_m'].mean():.3f} 米", ha="right", transform=axis_pos.transAxes)

    axis_vel = figure.add_subplot(grid[1, 1])
    axis_vel.hist(correct["velocity_error_mps"], bins=8, color=GREEN, alpha=0.86, edgecolor="white")
    axis_vel.axvline(correct["velocity_error_mps"].mean(), color=RED, linestyle="--", linewidth=1.8)
    axis_vel.set_title("速度误差分布", fontsize=15, fontweight="bold")
    axis_vel.set_xlabel("速度误差 / 米/秒")
    axis_vel.set_ylabel("关系数量")
    axis_vel.grid(axis="y", color=GRID, linewidth=0.7)
    axis_vel.spines[["top", "right"]].set_visible(False)
    axis_vel.text(0.98, 0.88, f"平均 {correct['velocity_error_mps'].mean():.4f} 米/秒", ha="right", transform=axis_vel.transAxes)

    figure.suptitle("40目标理想条件下的交汇定位结果", fontsize=20, fontweight="bold", y=0.98)
    figure.text(
        0.5,
        0.02,
        "本图用于验证计算链路；设备位姿、时间同步和检测结果均为理想条件。",
        ha="center",
        fontsize=11.5,
        color=MUTED,
    )
    return save_figure(figure, "10_ranging_reconstruction_and_error.png")


def generate_figures(s180_final_rows: list[dict], scenario: dict, continuous_360: dict) -> list[Path]:
    configure_matplotlib()
    curated_figures = [
        ASSET_DIR / "01_algorithm_flow.png",
        ASSET_DIR / "04a_local_tracks_before_registration.png",
        ASSET_DIR / "04b_candidate_graph_gnn_assignment.png",
    ]
    missing = [path for path in curated_figures if not path.is_file()]
    if missing:
        raise FileNotFoundError(f"missing curated report figures: {missing}")
    return [
        curated_figures[0],
        figure_single_station_tracking(),
        figure_coplanarity_screening(),
        curated_figures[1],
        curated_figures[2],
        figure_multitime_triangulation(),
        figure_airsim_scene(scenario),
        figure_optical_observations(),
        figure_s180_results(s180_final_rows),
        figure_stage_metrics(s180_final_rows),
        figure_ranging_reconstruction(),
        figure_360_clean_light(continuous_360["clean_light_final_rows"]),
        figure_360_multiseed_cascade(continuous_360["final_rows"]),
    ]


CONDITION_LABELS_CN = {
    "clean": "无干扰",
    "light": "轻度干扰",
    "medium": "中度干扰",
    "heavy": "重度干扰",
}


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join("---" for _ in headers) + "|"]
    lines.extend("| " + " | ".join(str(value) for value in row) + " |" for row in rows)
    return "\n".join(lines)


def percent_text(value: float | None) -> str:
    return "未记录" if value is None else f"{100.0 * float(value):.1f}%"


def latency_text(value: float | None) -> str:
    return "未记录" if value is None else f"{float(value):.1f}毫秒"


def ideal_360_pending_rows() -> list[dict]:
    return [
        {
            "target_count": target_count,
            "route_name": route_name,
            "route_label_cn": ROUTE_LABELS_CN[route_name],
            "association_precision": None,
            "fixed_target_coverage": None,
            "latency_p95_ms": None,
            "evidence_status": "pending_test_no_machine_evidence",
        }
        for target_count in CONTINUOUS_360_TARGET_COUNTS
        for route_name in ROUTE_ORDER
    ]


def build_markdown(
    s180_final_rows: list[dict],
    ranging: dict,
    continuous_360: dict,
) -> None:
    ideal_rows = ideal_360_pending_rows()
    clean_light_rows = continuous_360["clean_light_final_rows"]
    scale_rows = continuous_360["scale_matrix_rows"]

    ideal_table = markdown_table(
        ["目标数", "方法", "最后一圈关联精度", "最后一圈目标覆盖度", "处理耗时P95", "证据状态"],
        [
            [
                str(row["target_count"]),
                row["route_label_cn"],
                "待测试",
                "待测试",
                "待测试",
                "待测试；无机器记录",
            ]
            for row in ideal_rows
        ],
    )

    clean_light_table = markdown_table(
        ["方法", "条件", "最后一圈关联精度", "最后一圈目标覆盖度", "处理耗时P95", "证据状态"],
        [
            [
                row["route_label_cn"],
                CONDITION_LABELS_CN[row["corruption_level"]],
                percent_text(row["association_precision"]),
                percent_text(row["fixed_target_coverage"]),
                latency_text(row["latency_p95_ms"]),
                (
                    f"封存回放（诊断）；{row['deadline_miss_count']}/5超时"
                    if row["deadline_miss_count"]
                    else "封存回放（诊断）"
                ),
            ]
            for row in clean_light_rows
        ],
    )

    scale_table_rows = []
    for row in scale_rows:
        if row["evidence_status"] == "not_run":
            precision = coverage = "未开展"
            latency = "未记录"
            status = "未开展"
        elif row["evidence_status"] == "timeout":
            precision = coverage = "超时"
            latency = latency_text(row["latency_p95_ms"])
            status = f"诊断；{row['deadline_miss_count']}/5超时"
        else:
            precision = percent_text(row["association_precision"])
            coverage = percent_text(row["fixed_target_coverage"])
            latency = latency_text(row["latency_p95_ms"])
            status = "诊断"
        scale_table_rows.append(
            [
                str(row["target_count"]),
                CONDITION_LABELS_CN[row["corruption_level"]],
                row["route_label_cn"],
                precision,
                coverage,
                latency,
                status,
            ]
        )
    scale_table = markdown_table(
        ["目标数", "条件", "方法", "最后一圈关联精度", "最后一圈目标覆盖度", "处理耗时P95", "证据状态"],
        scale_table_rows,
    )

    s180_table_rows = []
    for row in s180_final_rows:
        if row["confirmed_output_available"]:
            precision = percent_text(row["association_precision"])
            coverage = percent_text(row["fixed_denominator_coverage"])
        else:
            precision = coverage = "超时"
        status = row["evidence_status_cn"]
        if row["deadline_miss_count"]:
            status += f"；{row['deadline_miss_count']}/5超时"
        s180_table_rows.append(
            [
                str(row["target_count"]),
                CONDITION_LABELS_CN[row["corruption_level"]],
                row["route_label_cn"],
                precision,
                coverage,
                latency_text(row["latency_p95_ms"]),
                status,
            ]
        )
    s180_table = markdown_table(
        ["目标数", "条件", "方法", "最后一轮关联精度", "最后一轮目标覆盖度", "处理耗时P95", "证据状态"],
        s180_table_rows,
    )

    ranging_table = markdown_table(
        ["项目", "结果"],
        [
            ["正确关系", f"{ranging['correct_match_count']}条"],
            ["错误关系", f"{ranging['false_match_count']}条"],
            ["双站关联精度", percent_text(ranging["association_precision"])],
            ["固定目标覆盖度", percent_text(ranging["association_full_target_recall"])],
            ["平均位置误差", f"{ranging['position_error_mean_m']:.3f}米"],
            ["95%位置误差", f"{ranging['position_error_p95_m']:.3f}米"],
            ["平均速度误差", f"{ranging['velocity_error_mean_mps']:.4f}米/秒"],
            ["95%速度误差", f"{ranging['velocity_error_p95_mps']:.4f}米/秒"],
        ],
    )

    report = f"""# 双光电多目标轨迹配准与交汇定位试验报告

本报告说明双光电形成单站航迹、建立双站对应关系并完成交汇定位的处理方法。试验结果先列出360度扫描且单站航迹完全正确时需要完成的验证矩阵，再分析360度实际单站航迹下的无干扰和轻、中、重度干扰，随后给出180度扫描结果，最后说明40目标理想条件下的交汇定位演示。各批数据的输入、随机场景和算法冻结状态不同，结果分别列示，不合并计算。

双光电配准在几何上可行，现有40目标理想演示也已贯通配准和交汇定位链路。20、40、60目标在360度扫描且单站关联完全正确条件下的三算法完整矩阵尚未试验，因此本报告不填写推算值。实际单站航迹进入链路后，断轨、错误重接和重复建轨会直接减少可供双站比较的正确对象。当前主要问题仍在单站航迹连续性，跨站评分和连续确认还存在少量独立损失。

## 一、算法原理

### 1.1 处理流程

两台光电分别处理图像，把同一目标在连续画面中的检测结果连成单站航迹，并输出观测时间、方位、俯仰、运动趋势和航迹质量。系统随后校正两站观测时刻和设备姿态，用共面关系排除明显不可能的航迹组合。候选关系再由几何方法、图神经网络或增强型图神经网络评分，经过一一分配和连续确认后形成双站对应关系。稳定关系进入双射线交汇定位。

![双光电多目标轨迹配准与交汇定位流程](assets/dual_optical_registration_report/01_algorithm_flow.png)

配准和定位分两步完成。前一步判断两台光电看到的是否为同一目标，后一步计算该目标的位置和速度。错误配准会把两条不相关视线送入定位环节，因此系统先确认对应关系，再进行距离和位置解算。

### 1.2 单站航迹

窄视场光电扫描时，同一目标只在短时间内进入画面。一次扫过形成若干连续检测点，下一次重访还要判断新片段是否属于原目标。单站航迹器使用方位和俯仰变化、运动方向、时间间隔和预测误差完成连接。短时漏检时保留休眠航迹；目标交叉时暂时保留少量候选，等待后续观测消解。

![单站检测点形成连续航迹及断轨风险](assets/dual_optical_registration_report/02_single_station_tracking.png)

单站成轨决定双站配准的输入质量。一个目标被拆成多条航迹，或两个目标被错误合成一条航迹，都会改变跨站算法实际处理的对象。双站评分可以拒绝部分错误关系，但无法恢复没有形成的正确局部航迹。

### 1.3 共面筛选

两台光电的位置和姿态已知时，每个检测框中心可由针孔相机模型转换为空间视线。对于同一时刻的同一目标，A站视线、B站视线和两站基线应近似位于同一平面。系统计算归一化共面性残差，并结合姿态误差、时间差和航迹预测不确定度设置门限。残差过大的组合直接剔除。

![双站视线的三维共面筛选](assets/dual_optical_registration_report/03_coplanarity_screening_3d.png)

共面筛选用于缩小候选范围，不直接判定身份。目标密集、航迹交叉或设备存在小幅姿态误差时，多组航迹可能同时满足共面条件，还需比较一段时间内的运动一致性和候选竞争关系。

### 1.4 图神经网络与航迹级注意力

候选关系可表示为二部图。左侧节点为A站航迹，右侧节点为B站航迹，通过共面筛选的组合形成连线。节点记录方向、角速度、航迹年龄和不确定度；连线记录多时刻共面残差、同步后的运动差、交汇稳定性和历史确认状态。图神经网络在候选图上交换信息，比较一条关系与周边竞争关系的相对合理性。

![双站局部航迹对应关系待确定](assets/dual_optical_registration_report/04a_local_tracks_before_registration.png)

![候选关系经图神经网络评分和一一分配后收敛](assets/dual_optical_registration_report/04b_candidate_graph_gnn_assignment.png)

本报告所称“增强型图神经网络”对应实验记录中的 `track_superglue` 路线。该路线借鉴SuperGlue的自注意力、交叉注意力和部分匹配思想，在两组航迹之间反复交换信息，再用归一化分配求出可匹配关系。它的输入是单站航迹及其几何和运动特征，不是原始图像关键点，也不是原始图像版SuperGlue。低分关系允许保持空匹配，最终仍由一一分配和连续确认约束输出。

### 1.5 几何方法和安全边界

几何方法把多时刻共面残差、运动方向差、角速度差和交汇稳定性按固定规则合成为代价，再执行匈牙利一一分配和时间确认。这条路线容易解释，适合候选较少和扫描规律稳定的情况，也可作为学习方法不可用时的回退。目标数量增加后，候选组合和多时刻拟合的计算量上升，固定权重和门限需要随扫描协议重新标定。

三条路线共用两道确定性边界。共面筛选阻止明显违反成像几何的关系进入评分，一一分配和连续确认阻止一条航迹被多个目标重复占用。学习算法负责比较复杂候选，几何条件负责限制物理上不合理的输出。

### 1.6 交汇定位

双站关系稳定后，系统在相邻时刻取出两台光电对同一目标的视线。理想情况下两条视线相交；存在离散采样和数值误差时，取两条视线最近点的中点作为位置。多个时刻的位置再进行短时运动拟合，得到位置、速度和残差。交会角过小、最近点距离过大或结果跳变时，定位结果保持待确认。

![多时刻双射线交汇定位原理](assets/dual_optical_registration_report/05_multitime_triangulation_3d.png)

交汇角决定距离解算条件。两台设备与目标近似共线时，小幅角度误差会被放大为较大的距离误差。实际部署需要保证足够基线和侧向观察角，并把设备位置、姿态和时间同步误差纳入不确定度。本文40目标测距结果采用理想位姿和时间条件，只用于验证计算链路。

### 1.7 评价口径

{markdown_table(
        ["指标", "计算口径", "说明"],
        [
            ["最后一圈或最后一轮关联精度", "正确双站关系数除以全部已输出关系数", "判断已输出关系中有多少正确"],
            ["最后一圈或最后一轮目标覆盖度", "正确配准目标数除以目标总数", "判断全部目标中有多少完成正确配准"],
            ["处理耗时P95", "五个场景最后一圈或最后一轮端到端耗时的95%分位值", "超时场景仍纳入耗时统计"],
            ["证据状态", "区分正式、诊断、超时、未开展和待测试", "无记录不写成零值"],
        ],
    )}

180度扫描统计第12轮，360度周扫统计第6圈。每项可用结果合并5个测试场景，未输出和超时场景仍保留在覆盖度分母中。若五个场景均超时且没有确认关系，精度和覆盖度写“超时”，不写0。耗时只读取机器记录中的 `end_to_end_ms`，并在对应最终窗口内计算P95。没有机器记录的组合写“未记录”或“待测试”。真实身份只用于试验结束后的离线评分。

## 二、试验配置

### 2.1 360度扫描及理想单站条件

360度连续周扫的基础设置如下。理想单站条件要求每个真实目标在每台光电内各形成一条连续且身份正确的航迹，不出现断轨、错误重接和重复建轨。该条件只隔离检查双站配准，不代表实际单站航迹器已经达到完全正确。

{markdown_table(
        ["项目", "设置"],
        [
            ["光电数量与基线", "2台固定光电，横向间隔2千米"],
            ["图像分辨率", "1280×1024"],
            ["等效焦距", "300毫米"],
            ["水平视场角", "2.93度"],
            ["仿真等效垂直视场角", "2.344度"],
            ["目标尺寸与速度", "长度3米，速度50米/秒"],
            ["扫描方式", "2秒连续旋转360度"],
            ["场景时长", "12秒，共6圈"],
            ["目标规模", "20、40、60"],
            ["统计窗口", "5个测试场景的第6圈"],
            ["在线期限", "每圈1000毫秒"],
        ],
    )}

20、40、60目标与三种算法的理想单站完整矩阵目前没有现成机器记录。本报告在结果表中保留九个待测试条目，不使用180度理想数据、实际航迹中的正确子集或40目标交汇演示回填。

### 2.2 360度实际单站航迹

实际航迹试验包含两批封存数据。第一批为20目标无干扰和轻度干扰，同一批匿名单站航迹分别输入几何方法、图神经网络和增强型图神经网络，用于比较评分方法。第二批为20、40、60目标的随机场景，完整保留了图神经网络结果；20目标还运行了几何方法，40和60目标没有运行几何方法，增强型图神经网络也没有进入该分规模批次。两批数据使用不同冻结版本，结果不能拼接。

每个场景运行12秒，共6圈。目标初始位置、前后间隔和交叉关系随种子变化，一半目标沿零度航向飞行，另一半沿负30度航向飞行。四档条件均保留0.4毫弧度固定云台偏差和0.3毫弧度逐帧随机抖动。

{markdown_table(
        ["干扰条件", "随机漏检率", "每台每秒瞬时虚警", "每台持续虚警"],
        [
            ["无干扰", "0", "0", "0"],
            ["轻度干扰", "3%", "2个", "0"],
            ["中度干扰", "7%", "4个", "1个"],
            ["重度干扰", "12%", "8个", "2个"],
        ],
    )}

### 2.3 180度扫描

180度试验把单程扫描时间缩短为1秒，机械往返周期仍为2秒。场景持续12秒，每秒形成一次双站关联结果。20、40和60目标分别使用8个训练种子、2个验证种子和5个保留测试种子，结果只取5个测试场景的第12轮。

{markdown_table(
        ["项目", "设置"],
        [
            ["光电数量与基线", "2台固定光电，横向间隔2千米"],
            ["目标尺寸与速度", "长度3米，速度50米/秒"],
            ["扫描方式", "1秒单程180度，2秒机械往返"],
            ["场景时长", "12秒，共12轮"],
            ["目标规模", "20、40、60"],
            ["固定姿态偏差", "均方根0.4毫弧度"],
            ["随机姿态抖动", "均方根0.3毫弧度"],
            ["轻度干扰", "3%漏检，每台相机每秒2个虚警"],
            ["在线期限", "每轮1000毫秒"],
        ],
    )}

三条路线共享同一批匿名单站航迹和候选关系。20目标共享航迹器通过验收；40目标错误重接率超过门限，60目标同时存在错误重接率和航迹器耗时问题，因此40和60目标只作为诊断。

### 2.4 40目标交汇定位演示

交汇定位演示采用AirSim计算机视觉模式。两台光电横向间隔2千米，40个无人机网格目标在三维空间以50米/秒运动，目标之间最近间隔约27.1米。扫描范围为正负45度，0.5秒完成单程，1秒完成往返。检测使用仿真元数据，真实编号只用于离线核对。

![40目标三维AirSim场景](assets/dual_optical_registration_report/06_airsim_scene_40_targets_cn.png)

![两台光电在同一扫描时刻看到的目标子集](assets/dual_optical_registration_report/07_airsim_optical_observations_cn.png)

该演示采用单个种子，设备位姿、时间同步和检测结果为理想条件，没有注入姿态误差、漏检和虚警。它用于确认“单站航迹、双站配准、交汇定位、速度拟合、离线评分”能够贯通，不能替代360度理想单站规模矩阵，也不能作为实际装备精度指标。

### 2.5 证据边界

三组关联试验只报告最后一圈或最后一轮的关联精度、目标覆盖度和端到端耗时P95。360度20目标三路线同输入对照是封存回放诊断；分规模批次中没有运行的算法写“未开展”。180度20目标结果使用通过验收的共享航迹器，40和60目标结果因航迹器未通过验收而标为诊断。不同目标规模采用不同随机场景和分别冻结的模型，不能据此推导目标数量与精度之间的单变量关系。

## 三、试验结果

### 3.1 360度理想单站条件

理想单站条件用于回答一个独立问题：单站航迹身份全部正确时，双站几何和关系评分能否在20、40、60目标下稳定完成配准。所需九个组合尚未形成完整试验记录，表中全部保留为待测试。

{ideal_table}

“待测试”表示没有可引用的精度、覆盖度和耗时记录，不表示结果为0。共面约束、一一分配和双射线交汇给出了理论处理路径，40目标理想演示也形成36条正确关系，但该演示采用正负45度扫描、单个种子和独立处理链，不能证明上述360度规模矩阵已经完成验证。

### 3.2 360度实际单站航迹

#### 3.2.1 20目标三种方法同输入对照

本组使用同一批20目标匿名单站航迹和5个测试场景，分别运行三种方法。表中只统计第6圈。

{clean_light_table}

无干扰条件下，几何方法的精度和覆盖度最高，分别为98.8%和85.0%，端到端耗时P95为867.9毫秒。轻度干扰下，增强型图神经网络的精度和覆盖度最高，分别为89.0%和65.0%；图神经网络耗时最低，为88.5毫秒。几何方法在轻度干扰下有1个场景超过1000毫秒期限，表中保留该超时事实。

![20目标360度周扫第6圈的无干扰与轻度干扰对比](assets/dual_optical_registration_report/11_360_clean_light_route_comparison.png)

本组没有一种方法在两种条件下同时占优。几何方法适合干扰较轻、候选较少的场景；增强型图神经网络在本组轻度干扰下更稳；基础图神经网络处理速度更快。三者需要在相同输入和时限下继续对照，不能只依据一个干扰等级确定最终方案。

#### 3.2.2 20、40、60目标随机干扰

本组按目标规模分别冻结模型并使用5个随机测试场景。表中列出无干扰、轻度、中度和重度干扰的第6圈结果。未进入该批次的算法明确写“未开展”；已经运行但五个场景全部超时的组合写“超时”。

{scale_table}

图神经网络是该批次唯一覆盖20、40和60目标四档条件的方法。无干扰时，20、40、60目标的覆盖度分别为63.0%、60.5%和55.0%；重度干扰时分别降至43.0%、27.5%和28.7%。目标数量和干扰增加后，端到端耗时P95从20目标无干扰的78.3毫秒增加到60目标重度干扰的403.9毫秒，仍低于本批次1000毫秒期限。

![360度周扫无干扰及轻、中、重度随机干扰结果](assets/dual_optical_registration_report/12_360_multiseed_cascade.png)

故障记录显示，目标交叉、姿态抖动、漏检和虚警先造成单站断轨、错误重接和重复建轨，随后才表现为双站精度和覆盖度下降。两站没有形成对应的正确局部航迹时，跨站图网络无法从后端补回。个别场景在单站航迹基本完整时仍出现双站覆盖损失，说明跨站评分和连续确认也需要单独校准。

### 3.3 180度扫描

180度扫描每秒形成一次关联结果，重访频率高于2秒周扫360度。每个规模和条件使用5个测试场景，表中只统计第12轮。

{s180_table}

图神经网络在六组场景均按时形成结果，精度为83.9%至97.8%，覆盖度为71.0%至90.7%，耗时P95为100.4至509.5毫秒。几何方法六组均超过1000毫秒期限。增强型图神经网络在20和40目标形成结果，60目标五个场景全部超时。

![180度扫描三种方法的最终一轮结果](assets/dual_optical_registration_report/08_s180_selected_results.png)

![图神经网络在180度扫描中的最终一轮精度与覆盖度](assets/dual_optical_registration_report/09_tracking_and_registration_loss.png)

20目标轻度干扰下，180度图神经网络的最后一轮精度和覆盖度为83.9%和73.0%；360度同输入诊断批次相应数值为77.2%和61.0%。现有对照支持“提高重访频率有利于维持关联”的判断，但两组数据并非同一随机种子、同一冻结模型的严格消融试验，不能把差值全部归因于扫描范围变化。40和60目标180度结果还受单站航迹器未通过验收的限制。

60目标无干扰精度高于40目标，主要由不同随机场景造成。40目标有一个困难场景贡献了该规模第12轮19个错误关系中的12个；去除该场景后，其余4个场景合并精度为95.3%，与60目标95.8%接近。现有数据不支持“目标越多，关联越准”的结论。

### 3.4 交汇定位结果

40目标理想演示形成37条双站关系，其中36条正确、1条错误。正确关系进入交汇定位后，平均位置误差为0.080米，95%位置误差不超过0.091米；平均速度误差为0.0081米/秒，95%速度误差不超过0.0197米/秒。

{ranging_table}

![正确配准关系的三维轨迹重建和误差分布](assets/dual_optical_registration_report/10_ranging_reconstruction_and_error.png)

上述误差来自理想位姿、理想时间和仿真检测条件，主要反映计算链路的一致性。真实系统还需加入安装测量误差、云台偏差、时间同步误差、检测中心偏差和大气条件。完成这些误差标定前，不能把厘米级仿真结果写成设备定位能力。

### 3.5 结论

1. **双光电配准具有明确的理论处理路径，完整规模矩阵仍待试验。** 共面筛选、候选评分、一一分配和交汇定位构成闭合链路，40目标理想演示完成了单个场景验证。360度理想单站条件下20、40、60目标的三算法矩阵没有现成记录，当前不能给出规模结论和耗时结论。

2. **实际360度试验表明三种方法各有边界。** 20目标同输入对照中，几何方法在无干扰条件下数值最高，增强型图神经网络在轻度干扰下数值最高，基础图神经网络耗时最低。分规模随机干扰批次只有图神经网络形成20、40、60目标的完整结果，其他组合按实际情况标为超时或未开展。

3. **提高重访频率能够改善现有图神经网络结果，但仍需严格同条件复核。** 180度扫描在现有20目标轻度干扰对照中比360度周扫取得更高精度和覆盖度。两批数据不是同种子的单变量试验，后续仍需使用相同场景、相同模型和相同干扰条件完成扫描频率消融。

4. **当前主要卡点是单站航迹关联。** 40目标航迹器未通过错误重接率门限，60目标还存在航迹器耗时问题。360度随机干扰下，断轨、错误重接和重复建轨先减少正确航迹，双站算法只能处理剩余候选。下一步应先稳定单站重访航迹，再补齐360度理想矩阵和跨站困难场景标定。
"""
    REPORT_MD.write_text(report, encoding="utf-8")


def set_run_font(
    run,
    *,
    size: float = 12,
    bold: bool | None = None,
    color: str = WORD_INK,
    east_asia: str = BODY_FONT,
    italic: bool | None = None,
) -> None:
    run.font.name = LATIN_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str, *, size: float = 12, color: str = WORD_INK) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(
                paragraph.add_run(token[2:-2]),
                size=size,
                bold=True,
                color=color,
                east_asia=HEADING_FONT,
            )
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size=size, color=WORD_BLUE)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size=size, color=color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    set_run_font(run, size=8.5, color=WORD_MUTED, east_asia=HEADING_FONT)


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.35)
    section.header_distance = Cm(0.72)
    section.footer_distance = Cm(0.72)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Cm(0.74)

    styles = {
        "Heading 1": (16.5, WORD_BLUE, 14, 8),
        "Heading 2": (14, WORD_TEAL, 11, 6),
        "Heading 3": (12.5, WORD_INK, 9, 4),
    }
    for name, (size, color, before, after) in styles.items():
        style = document.styles[name]
        style.font.name = LATIN_FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), HEADING_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)


def add_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline(header, "双光电多目标轨迹配准与交汇定位试验报告", size=8.5, color=WORD_MUTED)
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(footer, "科研仿真与技术论证材料  ·  ", size=8.5, color=WORD_MUTED)
    add_page_number(footer)


def add_cover(document: Document) -> None:
    for _ in range(5):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run("双光电多目标轨迹配准与交汇定位"), size=26, bold=True, east_asia=HEADING_FONT)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        subtitle.add_run("算法原理、试验配置与结果分析"),
        size=15,
        bold=True,
        color=WORD_BLUE,
        east_asia=HEADING_FONT,
    )
    for _ in range(9):
        document.add_paragraph()
    owner = document.add_paragraph()
    owner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(owner.add_run("MSM 项目组"), size=12)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("2026 年 8 月"), size=11, color=WORD_MUTED)
    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(boundary.add_run("科研仿真与技术论证材料"), size=9.5, color=WORD_TEAL)


def table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_table(document: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    font_size = 8.3 if columns >= 7 else 9.4
    for row in table.rows:
        row_properties = row._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_properties.append(cannot_split)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, values in enumerate(rows):
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(0.5)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = values[column_index] if column_index < len(values) else ""
            run = paragraph.add_run(text)
            set_run_font(
                run,
                size=font_size,
                bold=row_index == 0,
                color="FFFFFF" if row_index == 0 else WORD_INK,
                east_asia=HEADING_FONT if row_index == 0 else BODY_FONT,
            )
            shade_cell(cell, "1F5F99" if row_index == 0 else ("F1F5F8" if row_index % 2 == 0 else "FFFFFF"))
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_image(document: Document, alt: str, path_text: str, number: int) -> None:
    image_path = (REPORT_MD.parent / path_text).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    with Image.open(image_path) as source:
        width_px, height_px = source.size
    max_width_cm = 16.0
    max_height_cm = 11.8
    ratio = width_px / height_px
    width_cm = min(max_width_cm, max_height_cm * ratio)
    height_cm = width_cm / ratio
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm * ratio

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_after = Pt(5)
    set_run_font(caption.add_run(f"图 {number}  {alt}"), size=9, color=WORD_MUTED)


def add_list(document: Document, marker: str, content: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.78)
    paragraph.paragraph_format.first_line_indent = Cm(-0.52)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(f"{marker}  "), size=11, bold=True, color=WORD_BLUE, east_asia=HEADING_FONT)
    add_inline(paragraph, content, size=11)


def build_word() -> None:
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_section(document.sections[0])
    configure_styles(document)
    add_cover(document)
    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    add_header_footer(body_section)

    image_number = 0
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading:
            level = min(3, len(heading.group(1)) - 1)
            paragraph = document.add_paragraph(style=f"Heading {level}")
            if level == 1 or heading.group(2).strip() == "3.5 结论":
                paragraph.paragraph_format.page_break_before = True
            add_inline(
                paragraph,
                heading.group(2),
                size=(16.5, 14, 12.5)[level - 1],
                color=(WORD_BLUE, WORD_TEAL, WORD_INK)[level - 1],
            )
            index += 1
            continue
        image = IMAGE_RE.fullmatch(line)
        if image:
            image_number += 1
            add_image(document, image.group(1), image.group(2), image_number)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and TABLE_DIVIDER_RE.fullmatch(lines[index + 1].strip()):
            rows = [table_cells(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(table_cells(lines[index]))
                index += 1
            if len(rows[0]) >= 8:
                landscape = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(landscape, landscape=True)
                add_header_footer(landscape)
                add_table(document, rows)
                portrait = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(portrait)
                add_header_footer(portrait)
            else:
                add_table(document, rows)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        if numbered or bullet:
            marker = f"{numbered.group(1)}." if numbered else "•"
            content = numbered.group(2) if numbered else bullet.group(1)
            add_list(document, marker, content)
            index += 1
            continue

        parts = [line]
        lookahead = index + 1
        while lookahead < len(lines):
            candidate = lines[lookahead].strip()
            if not candidate or candidate.startswith(("#", "![", "|")):
                break
            if re.match(r"^(?:\d+\.|[-*])\s+", candidate):
                break
            parts.append(candidate)
            lookahead += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline(paragraph, " ".join(parts), size=12)
        index = lookahead

    properties = document.core_properties
    properties.title = "双光电多目标轨迹配准与交汇定位试验报告"
    properties.subject = "算法原理、试验配置与结果分析"
    properties.author = "MSM 项目组"
    properties.keywords = "双光电, 多目标轨迹配准, 图神经网络, 交汇定位, AirSim"
    document.save(REPORT_DOCX)


def validate_word() -> dict[str, int]:
    document = Document(REPORT_DOCX)
    with ZipFile(REPORT_DOCX) as archive:
        damaged = archive.testzip()
        if damaged is not None:
            raise RuntimeError(f"generated DOCX archive is damaged: {damaged}")
        images = [name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/")]
    if len(images) != 13:
        raise RuntimeError(f"expected 13 embedded figures, found {len(images)}")
    if len(document.tables) != 9:
        raise RuntimeError(f"expected 9 tables, found {len(document.tables)}")
    text_parts = [paragraph.text for paragraph in document.paragraphs]
    text_parts.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    text = "\n".join(text_parts)
    for required in (
        "算法原理",
        "试验配置",
        "试验结果",
        "3.1 360度理想单站条件",
        "3.2 360度实际单站航迹",
        "3.3 180度扫描",
        "待测试；无机器记录",
        "增强型图神经网络（航迹级注意力）",
        "封存回放（诊断）",
        "未开展",
        "处理耗时P95",
        "403.9毫秒",
        "1/5超时",
        "97.8%",
        "81.5%",
        "90.7%",
        "82.0%",
        "360度连续周扫",
        "中度干扰",
        "重度干扰",
        "28.7%",
        "39.7%",
        "83.9%",
        "73.0%",
        "98.8%",
        "85.0%",
        "超时",
        "最后一圈关联精度",
        "当前主要卡点是单站航迹关联",
        "现有数据不支持“目标越多，关联越准”的结论",
        "40和60目标结果因航迹器未通过验收而标为诊断",
    ):
        if required not in text:
            raise RuntimeError(f"Word output is missing required content: {required}")
    for forbidden in (
        "单站正确时双站覆盖度",
        "条件双站覆盖度",
        "全6圈覆盖度",
        "第3至6圈覆盖度",
        "两站共同稳定成轨",
        "正确候选保留",
    ):
        if forbidden in text:
            raise RuntimeError(f"Word output still contains a removed metric: {forbidden}")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(images),
        "sections": len(document.sections),
        "bytes": REPORT_DOCX.stat().st_size,
    }


def build_evidence_manifest(
    s180: dict,
    s180_final_rows: list[dict],
    ranging: dict,
    continuous_360: dict,
    figures: list[Path],
) -> None:
    evidence_by_count = {item["target_count"]: item for item in s180["evidence"]}
    reported_s180_rows = []
    for row in s180_final_rows:
        reported_s180_rows.append(
            {
                "target_count": row["target_count"],
                "corruption_level": row["corruption_level"],
                "window": row["window"],
                "route": row["route_name"],
                "evidence_status": row["evidence_status_cn"],
                "test_seed_count": row["sample_count"],
                "confirmed_output_available": row["confirmed_output_available"],
                "association_precision": row["association_precision"],
                "fixed_target_coverage": row["fixed_denominator_coverage"],
                "latency_p95_ms": row["latency_p95_ms"],
                "deadline_miss_count": row["deadline_miss_count"],
                "tracker_acceptance_passed": evidence_by_count[row["target_count"]]["tracker_acceptance_passed"],
                "tracker_failure_reasons": evidence_by_count[row["target_count"]]["tracker_failure_reasons"],
            }
        )
    source_paths = [
        S180_METRICS,
        S180_REPRODUCTION,
        CLEAN_LIGHT_METRICS,
        Path(continuous_360["clean_light_metrics"]["test_manifest"]),
        *(
            Path(item["path"])
            for item in continuous_360["clean_light_metrics"]["route_manifests"].values()
        ),
        CONTINUOUS_360_SUMMARY,
        *(continuous_360_metrics_path(count) for count in CONTINUOUS_360_TARGET_COUNTS),
        *(continuous_360_diagnostics_path(count) for count in CONTINUOUS_360_TARGET_COUNTS),
        RANGING_METRICS,
        RANGING_SCENARIO,
        RANGING_TRACKS,
        RANGING_MATCHES,
        RANGING_MATCH_SCORES,
        RANGING_DETECTIONS,
    ]
    generated_paths = [REPORT_MD, REPORT_DOCX, *figures]
    manifest = {
        "schema_version": "dual-optical-leadership-report-evidence-v5",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "report": relative(REPORT_MD),
        "reporting_policy": {
            "window": "final_round",
            "s180_round": 12,
            "continuous_360_revolution": 6,
            "reported_metrics": [
                "association_precision",
                "fixed_target_coverage",
                "final_window_end_to_end_latency_p95_ms",
            ],
            "latency_policy": (
                "P95 of end_to_end_ms across exactly five scenes in the reported final "
                "round or revolution; timed-out scenes remain included"
            ),
            "s180_routes": list(ROUTE_ORDER),
        },
        "ideal_360_single_station_matrix": {
            "status": "pending_test_no_machine_evidence",
            "scan_span_deg": 360.0,
            "target_counts": list(CONTINUOUS_360_TARGET_COUNTS),
            "routes": list(ROUTE_ORDER),
            "rows": ideal_360_pending_rows(),
            "prohibited_substitutions": [
                "s180 ideal-local-track evidence",
                "correct-local-track conditional subsets",
                "the single-seed 40-target ranging demonstration",
            ],
        },
        "s180_final_round_rows": reported_s180_rows,
        "continuous_360_diagnostics": {
            "status": "diagnostic_only",
            "scan_span_deg": 360.0,
            "scan_period_s": 2.0,
            "duration_s": 12.0,
            "clean_light_final_round_rows": [
                {
                    key: row[key]
                    for key in (
                        "route_name",
                        "route_label_cn",
                        "corruption_level",
                        "revolution_index",
                        "sample_count",
                        "association_precision",
                        "fixed_target_coverage",
                        "latency_p95_ms",
                        "deadline_miss_count",
                    )
                }
                for row in continuous_360["clean_light_final_rows"]
            ],
            "multi_scale_route_matrix_rows": [
                {
                    key: row.get(key)
                    for key in (
                        "target_count",
                        "corruption_level",
                        "route_name",
                        "route_label_cn",
                        "test_seed_count",
                        "association_precision",
                        "fixed_target_coverage",
                        "latency_p95_ms",
                        "deadline_miss_count",
                        "confirmed_output_available",
                        "evidence_status",
                    )
                }
                for row in continuous_360["scale_matrix_rows"]
            ],
            "multi_scale_final_round_rows": [
                {
                    key: row[key]
                    for key in (
                        "target_count",
                        "corruption_level",
                        "test_seed_count",
                        "association_precision",
                        "fixed_target_coverage",
                        "latency_p95_ms",
                        "deadline_miss_count",
                    )
                }
                for row in continuous_360["final_rows"]
            ],
        },
        "ranging_demonstration": {
            "status": "ideal_pose_time_single_seed_chain_demonstration",
            "seed": ranging["seed"],
            "target_count": ranging["target_count"],
            "correct_match_count": ranging["correct_match_count"],
            "false_match_count": ranging["false_match_count"],
            "association_precision": ranging["association_precision"],
            "fixed_target_coverage": ranging["association_full_target_recall"],
            "position_error_mean_m": ranging["position_error_mean_m"],
            "position_error_p95_m": ranging["position_error_p95_m"],
            "velocity_error_mean_mps": ranging["velocity_error_mean_mps"],
            "velocity_error_p95_mps": ranging["velocity_error_p95_mps"],
        },
        "source_artifacts": [{"path": relative(path), "sha256": sha256(path)} for path in source_paths],
        "generated_artifacts": [{"path": relative(path), "sha256": sha256(path)} for path in generated_paths],
        "limitations": [
            "40- and 60-target S180 rows are diagnostic because the shared local tracker failed acceptance.",
            "The 40- and 60-target rows use different test scenes and separately trained GNN models; they are not a target-count-only controlled comparison.",
            "One difficult 40-target seed contributes 12 of the 19 final-round false associations; excluding it yields 95.27% precision.",
            "The geometry comparison is limited to the unretuned S180 enhanced-geometry baseline.",
            "The requested 360-degree perfect-local-track 20/40/60 by three-route matrix has no machine evidence; all nine cells remain pending test.",
            "The 360-degree clean/light route comparison and scale diagnostic use different frozen campaigns and must not be pooled.",
            "The 20-, 40-, and 60-target 360-degree rows use separately trained GNN models and five test seeds per scale.",
            "In the multi-scale 360-degree campaign, geometry ran only at 20 targets and track-level attention did not run; absent combinations are recorded as not_run, not zero.",
            "The 360-degree random-interference table reports final-round metrics for clean, light, medium, and heavy corruption; it does not pool all six revolutions.",
            "The random 360-degree scenes vary initial geometry and crossings but retain two heading classes (0 and -30 degrees); continuously randomized headings are not validated.",
            "All association result tables report final-round precision and fixed-target coverage only; intermediate diagnostic metrics remain source evidence but are not shown.",
            "The 40-target ranging result uses ideal pose/time and one seed and is not an equipment claim.",
        ],
    }
    EVIDENCE_MANIFEST.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def main() -> None:
    s180, s180_final_rows, ranging, scenario, continuous_360 = load_and_validate_evidence()
    figures = generate_figures(s180_final_rows, scenario, continuous_360)
    build_markdown(s180_final_rows, ranging, continuous_360)
    build_word()
    word_metrics = validate_word()
    build_evidence_manifest(s180, s180_final_rows, ranging, continuous_360, figures)
    print(
        f"generated {REPORT_DOCX.name}: figures={len(figures)}, tables={word_metrics['tables']}, "
        f"sections={word_metrics['sections']}, bytes={word_metrics['bytes']}"
    )
    print(f"evidence: {EVIDENCE_MANIFEST.name}")


if __name__ == "__main__":
    main()
