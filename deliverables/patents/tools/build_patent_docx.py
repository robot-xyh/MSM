#!/usr/bin/env python3
"""Build review-ready Word drafts from the two patent Markdown files."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
PATENT_DIR = ROOT / "deliverables" / "patents"
PATENTS = (
    PATENT_DIR / "基于多智能体决策的集群无人机协同反无人机系统_CN.md",
    PATENT_DIR / "基于图神经网络的双光电多目标轨迹配准与交会定位方法_CN.md",
)


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12) -> None:
    run.font.name = latin
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(24)
    normal.paragraph_format.space_after = Pt(0)

    for name, east_asia, size in (
        ("Title", "黑体", 18),
        ("Heading 1", "黑体", 16),
        ("Heading 2", "黑体", 14),
        ("Heading 3", "黑体", 12),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def configure_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.3)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9)


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def latex_to_plain(text: str) -> str:
    """Convert the limited TeX used in the drafts to readable Word text."""
    replacements = (
        (r"\operatorname{normalize}", "normalize"),
        (r"\operatorname{sigmoid}", "sigmoid"),
        (r"\operatorname{atan2}", "atan2"),
        (r"\operatorname{AGG}", "AGG"),
        (r"\begin{bmatrix}", "["),
        (r"\end{bmatrix}", "]"),
        (r"\begin{cases}", "{"),
        (r"\end{cases}", ""),
        (r"\boldsymbol", ""),
        (r"\widetilde", "~"),
        (r"\widehat", "^"),
        (r"\hat", "^"),
        (r"\overline", "¯"),
        (r"\mathrm", ""),
        (r"\text", ""),
        (r"\left", ""),
        (r"\right", ""),
        (r"\lVert", "||"),
        (r"\rVert", "||"),
        (r"\|", "||"),
        (r"\times", "×"),
        (r"\cdot", "·"),
        (r"\leq", "≤"),
        (r"\geq", "≥"),
        (r"\neq", "≠"),
        (r"\infty", "∞"),
        (r"\in", "∈"),
        (r"\sum", "Σ"),
        (r"\sqrt", "√"),
        (r"\ln", "ln"),
        (r"\max", "max"),
        (r"\min", "min"),
        (r"\tanh", "tanh"),
        (r"\arcsin", "arcsin"),
        (r"\varepsilon", "ε"),
        (r"\lambda", "λ"),
        (r"\alpha", "α"),
        (r"\beta", "β"),
        (r"\gamma", "γ"),
        (r"\Delta", "Δ"),
        (r"\delta", "δ"),
        (r"\sigma", "σ"),
        (r"\Sigma", "Σ"),
        (r"\phi", "φ"),
        (r"\psi", "ψ"),
        (r"\eta", "η"),
        (r"\nu", "ν"),
        (r"\theta", "θ"),
        (r"\rho", "ρ"),
        (r"\partial", "∂"),
    )
    value = text.strip().strip("$").strip()
    for source, target in replacements:
        value = value.replace(source, target)
    value = value.replace(r"\\", "; ").replace(r"\,", " ").replace(r"\;", " ")
    value = value.replace(r"\qquad", "    ").replace(r"\quad", "  ")
    value = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", value)
    value = re.sub(r"\\(?:frac|sqrt)\{", "", value)
    value = value.replace("{", "").replace("}", "")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def normalize_inline_math(text: str) -> str:
    text = re.sub(r"\$\$([^$]+)\$\$", lambda match: latex_to_plain(match.group(1)), text)
    return re.sub(r"\$([^$]+)\$", lambda match: latex_to_plain(match.group(1)), text)


def add_inline_runs(paragraph, text: str, size: float = 12) -> None:
    text = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", text)
    text = normalize_inline_math(text)
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        code = part.startswith("`") and part.endswith("`")
        content = part[2:-2] if bold else part[1:-1] if code else part
        run = paragraph.add_run(content)
        set_run_font(run, east_asia="宋体", latin="Consolas" if code else "Times New Roman", size=size)
        run.bold = bold


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        if not is_table_separator(lines[index]):
            rows.append([strip_inline_markdown(c) for c in lines[index].strip().strip("|").split("|")])
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(paragraph, row[col_index] if col_index < len(row) else "", size=9.5)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def add_body_paragraph(document: Document, text: str, *, claim: bool = False, list_item: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(24)
    paragraph.paragraph_format.space_after = Pt(0)
    if not claim and not list_item:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    elif list_item:
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.37)
    add_inline_runs(paragraph, text)


def add_equation(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(latex_to_plain(text))
    set_run_font(run, east_asia="宋体", latin="Cambria Math", size=11.5)


def add_figure(document: Document, md_path: Path, alt: str, relative_path: str) -> None:
    image_path = (md_path.parent / relative_path).resolve()
    if not image_path.exists():
        add_body_paragraph(document, f"[附图文件缺失：{relative_path}]", claim=True)
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(6)
    add_inline_runs(caption, alt, size=10.5)


def section_heading(document: Document, title: str, level: int) -> None:
    if title in {"摘要", "权利要求书", "说明书"} and len(document.paragraphs) > 1:
        document.add_page_break()
    paragraph = document.add_heading(title, level=min(level, 3))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT


def render_markdown(document: Document, md_path: Path, lines: Iterable[str]) -> None:
    source = list(lines)
    in_code = False
    in_equation = False
    equation_lines: list[str] = []
    in_claims = False
    index = 0
    while index < len(source):
        raw = source[index].rstrip()
        stripped = raw.strip()
        if stripped == "$$":
            if in_equation:
                add_equation(document, " ".join(equation_lines))
                equation_lines = []
                in_equation = False
            else:
                in_equation = True
            index += 1
            continue
        if in_equation:
            equation_lines.append(stripped)
            index += 1
            continue
        if stripped.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("# "):
            title = strip_inline_markdown(stripped[2:])
            if index == 0:
                p = document.add_paragraph()
                p.style = document.styles["Title"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(p, title, size=18)
            else:
                section_heading(document, title, 1)
            in_claims = title == "权利要求书"
            index += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            title = strip_inline_markdown(heading.group(2))
            section_heading(document, title, len(heading.group(1)) - 1)
            in_claims = title == "权利要求书" or in_claims and title not in {"技术领域", "背景技术", "发明内容", "附图说明", "具体实施方式"}
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^]]*)]\(([^)]+)\)", stripped)
        if image_match:
            add_figure(document, md_path, image_match.group(1), image_match.group(2))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(source) and is_table_separator(source[index + 1]):
            rows, index = parse_table(source, index)
            add_table(document, rows)
            continue
        if stripped.startswith(">"):
            text = stripped.lstrip("> ")
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.right_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(22)
            add_inline_runs(paragraph, text, size=10.5)
            index += 1
            continue
        if in_code:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.74)
            add_inline_runs(paragraph, stripped, size=10)
            index += 1
            continue
        list_match = re.match(r"^(?:[-*]|\d+[.)])\s+(.+)$", stripped)
        if list_match and not (in_claims and re.match(r"^\d+\.\s*", stripped)):
            marker = stripped[: stripped.find(list_match.group(1))].strip()
            add_body_paragraph(document, f"{marker} {list_match.group(1)}".strip(), list_item=True)
            index += 1
            continue
        add_body_paragraph(document, stripped, claim=in_claims)
        index += 1


def build(md_path: Path) -> Path:
    if not md_path.exists():
        raise FileNotFoundError(md_path)
    document = Document()
    configure_styles(document)
    configure_page(document.sections[0])
    for section in document.sections:
        add_page_number(section.footer.paragraphs[0])

    render_markdown(document, md_path, md_path.read_text(encoding="utf-8").splitlines())

    properties = document.core_properties
    properties.title = md_path.stem.replace("_CN", "")
    properties.subject = "中国发明专利申请初稿"
    properties.author = "【待填写】"
    properties.keywords = "发明专利; MSM; 无人机; 光电"
    output = md_path.with_suffix(".docx")
    document.save(output)
    return output


def main() -> None:
    for md_path in PATENTS:
        output = build(md_path)
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
